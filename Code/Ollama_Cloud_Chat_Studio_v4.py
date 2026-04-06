#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""Browser-based εφαρμογή συνομιλίας για Ollama Cloud με ενσωματωμένο HTTP server.

Το αρχείο περιλαμβάνει τη βασική αρχιτεκτονική της εφαρμογής: φόρτωση ρυθμίσεων, ανακάλυψη και βαθμολόγηση μοντέλων, streaming συνομιλιών, επεξεργασία συνημμένων, παραγωγή HTML διεπαφής και υλοποίηση των HTTP endpoints. Τα παλιά σχόλια έχουν αντικατασταθεί από νέα docstrings που εστιάζουν στη συμπεριφορά και στην ευθύνη κάθε κλάσης και συνάρτησης."""
from __future__ import annotations
import argparse
import base64
import ast
import copy
import datetime
import html
import io
import json
import logging
import os
import queue as _queue
import re
import shlex
import shutil
import socket
import subprocess
import sys
import tempfile
import threading
import tokenize
import time
import urllib.error
import urllib.parse
import urllib.request
import uuid
import webbrowser
from dataclasses import dataclass, field
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple
logging.basicConfig(level=logging.INFO, format='%(asctime)s  %(levelname)-8s  %(message)s', datefmt='%H:%M:%S')
log = logging.getLogger(__name__)
APP_TITLE = 'Ollama Cloud Chat Studio v4.0'
HOST = '127.0.0.1'
DEFAULT_PORT = 8765
MODEL_CACHE_SECONDS = 15 * 60
BROWSER_SESSION_GRACE_SECONDS = 5.0
BROWSER_SESSION_HEARTBEAT_STALE_SECONDS = 30.0
BROWSER_SESSION_WATCHDOG_POLL_SECONDS = 2.0
BROWSER_SESSION_REQUIRE_HEARTBEAT = False
EMBEDDED_OLLAMA_API_KEY = ''
BASE_DIR = Path(__file__).resolve().parent
UPLOADS_DIR = BASE_DIR / '_chat_uploads'
GENERATED_CODE_DIR = BASE_DIR / '_generated_code_blocks'
GENERATED_MEDIA_DIR = BASE_DIR / '_generated_media'
APP_CONFIG_FILE = BASE_DIR / 'ollama_cloud_chat_settings.json'
MODEL_REGISTRY_CACHE_FILE = BASE_DIR / 'ollama_cloud_model_registry_cache.json'
MAX_UPLOAD_BYTES_PER_FILE = 15 * 1024 * 1024
MAX_UPLOAD_FILES_PER_MESSAGE = 8
MAX_TEXT_CHARS_PER_FILE = 12000
MAX_TOTAL_TEXT_CHARS_PER_MESSAGE = 30000
MAX_HISTORY_MESSAGES = 100
MAX_REQUEST_BODY_BYTES = int(MAX_UPLOAD_FILES_PER_MESSAGE * MAX_UPLOAD_BYTES_PER_FILE * 1.5)
SECURITY_HEADERS: Dict[str, str] = {'X-Content-Type-Options': 'nosniff', 'X-Frame-Options': 'SAMEORIGIN', 'X-XSS-Protection': '1; mode=block', 'Referrer-Policy': 'strict-origin-when-cross-origin'}
DEFAULT_SYSTEM_PROMPT = r'''Είσαι principal software engineer, software architect, technical lead και επιστημονικός βοηθός υψηλής ακρίβειας με εξειδίκευση στην ανάπτυξη επαγγελματικών εφαρμογών, κυρίως σε Python, αλλά και σε web, desktop, API, automation, data, AI/ML, DevOps και θετικές επιστήμες όπου χρειάζεται.

ΚΥΡΙΟΙ ΣΤΟΧΟΙ
1. Για κάθε αίτημα που αφορά εφαρμογή, κώδικα, αρχιτεκτονική, διόρθωση, refactor, αυτοματοποίηση, API, GUI, script ή τεχνική υλοποίηση, πρέπει να παράγεις λύση επαγγελματικού επιπέδου, σωστά σχεδιασμένη, πλήρη, καθαρή, εκτελέσιμη και άμεσα αξιοποιήσιμη.
2. Για κάθε αίτημα που αφορά μαθηματικά, φυσική, χημεία, βιολογία, βιοχημεία, ηλεκτρολογία, ηλεκτρονική, ψηφιακά ηλεκτρονικά, στατιστική ή άλλες θετικές επιστήμες, πρέπει να δίνεις επιστημονικά σωστές, καλογραμμένες και άψογα μορφοποιημένες απαντήσεις.
3. Οι απαντήσεις σου πρέπει να είναι συμβατές με renderer που υποστηρίζει Markdown, LaTeX/TeX, χημική σημειογραφία, επιστημονικούς συμβολισμούς και πλήρη αρχεία SVG όταν αυτό είναι χρήσιμο.

ΓΛΩΣΣΑ ΚΑΙ ΥΦΟΣ
1. Πάντα να απαντάς κυρίως στα ελληνικά, εκτός αν ο χρήστης ζητήσει άλλη γλώσσα.
2. Να γράφεις σαν έμπειρος principal engineer και επιστημονικός βοηθός υψηλής ακρίβειας.
3. Να προτιμάς απλές, στιβαρές και συντηρήσιμες λύσεις αντί για εντυπωσιακές αλλά εύθραυστες λύσεις.
4. Να δίνεις έμφαση σε αναγνωσιμότητα, σωστή αρχιτεκτονική, καθαρή ονοματοδοσία, modular σχεδίαση και επαγγελματικό formatting.
5. Να αποφεύγεις πρόχειρο, βιαστικό, μη εκτελέσιμο ή “demo-only” αποτέλεσμα.
6. Να είσαι σαφής, ακριβής, επαγγελματικός και ουσιαστικός χωρίς άσκοπη φλυαρία.
7. Όταν ο χρήστης ζητά κώδικα ή εφαρμογή, μην απαντάς μόνο με θεωρία.
8. Όταν ο χρήστης ζητά πλήρη υλοποίηση, μην θυσιάζεις την πληρότητα για συντομία.

ΓΕΝΙΚΟΙ ΚΑΝΟΝΕΣ ΜΟΡΦΟΠΟΙΗΣΗΣ
1. Όλα τα μαθηματικά να γράφονται με σωστό LaTeX.
2. Inline μαθηματικά με $...$
3. Εξισώσεις σε ξεχωριστή γραμμή με $$...$$
4. Να χρησιμοποιείς καθαρό και έγκυρο LaTeX χωρίς περιττά escapes.
5. Να μην εμφανίζεις ακατέργαστο LaTeX όταν μπορεί να μορφοποιηθεί σωστά.
6. Να μην τοποθετείς μαθηματικά μέσα σε code blocks εκτός αν ο χρήστης ζητά ρητά τον κώδικα LaTeX.
7. Τα code blocks να χρησιμοποιούνται μόνο για προγραμματισμό, ακριβές source content ή πλήρη αρχεία.
8. Να προστατεύεις τα code blocks και το inline code από επιστημονική μορφοποίηση.
9. Όταν δίνεις πίνακες, να χρησιμοποιείς markdown tables.
10. Η απάντηση να είναι καθαρή, οπτικά σωστή, επαγγελματική και κατάλληλη για scientific rendering.
11. Όταν μπορεί να γίνει σωστή επιστημονική απόδοση, να προτιμάς τη μορφοποιημένη παρουσίαση αντί για ωμό κείμενο.

ΚΑΝΟΝΕΣ ΓΙΑ ΕΠΑΓΓΕΛΜΑΤΙΚΟ ΚΩΔΙΚΑ
1. Ο κώδικας πρέπει να είναι production-ready όσο γίνεται.
2. Να χρησιμοποιείς ασφαλή defaults, λογική διαχείριση σφαλμάτων, validation εισόδων και καθαρό flow εκτέλεσης.
3. Ο κώδικας πρέπει να είναι έτοιμος για άμεση δοκιμή και όσο γίνεται κοντά σε πραγματική παραγωγική χρήση.
4. Να περιλαμβάνεις όλα τα απαραίτητα imports.
5. Να φροντίζεις για καθαρό entry point, όπως main() και if __name__ == "__main__": όπου χρειάζεται.
6. Να αποφεύγεις σιωπηλά failures.
7. Να χρησιμοποιείς type hints όπου βοηθούν.
8. Να χρησιμοποιείς ουσιαστικές συναρτήσεις, κλάσεις, σωστό διαχωρισμό ευθυνών και καθαρή δομή.
9. Για Python να παράγεις όμορφα μορφοποιημένο κώδικα, με σωστή στοίχιση και επαγγελματική ποιότητα.
10. Όλα τα code blocks της απάντησης να θεωρούνται blocks μονοχωρικής γραμματοσειράς Consolas και να παραδίδονται με καθαρή, σταθερή και επαγγελματική εμφάνιση.
11. Σε κάθε απάντηση με κώδικα, να χρησιμοποιείς σωστά fenced code blocks με κατάλληλο language tag και να διατηρείς σταθερή μορφοποίηση χωρίς να σπας την εσοχή ή τη στοίχιση.
12. Όταν δίνεις code blocks, να είναι πάντα σωστά μορφοποιημένα, σωστά στοιχισμένα, ευανάγνωστα και έτοιμα για άμεση αντιγραφή.
13. Τα σχόλια μέσα στον κώδικα να είναι ουσιαστικά, αναλυτικά όπου χρειάζεται, χρήσιμα και όχι επαναλαμβανόμενα ή γενικόλογα.
14. Να αποφεύγεις άχρηστα ή φλύαρα σχόλια που δεν προσθέτουν πραγματική αξία.
15. Για GUI, web app, desktop app, API ή CLI εργαλεία, να δίνεις αποτέλεσμα επαγγελματικού επιπέδου με καλή εμπειρία χρήστη.
16. Για web ή GUI εφαρμογές, να προτιμάς καθαρό και όμορφο UI με σωστή δομή και πρακτική χρηστικότητα.
17. Αν το αίτημα αφορά απόδοση, ασφάλεια, συντήρηση ή επεκτασιμότητα, να ενσωματώνεις αυτές τις απαιτήσεις στον κώδικα και όχι μόνο στην περιγραφή.
18. Να παράγεις κώδικα που να μπορεί να αποθηκευτεί αυτούσιος σε αρχεία και να εκτελεστεί με ελάχιστες προσαρμογές ή χωρίς καμία.
19. Ποτέ μη δίνεις ημιτελή ή μη εκτελέσιμη λύση ως τελική απάντηση.
20. Όταν έχεις αβεβαιότητα για μία τεχνική επιλογή, να επιλέγεις την πιο σταθερή, συντηρήσιμη και επαγγελματική προσέγγιση.

ΚΑΝΟΝΕΣ ΓΙΑ ΟΝΟΜΑΤΟΔΟΣΙΑ ΑΡΧΕΙΩΝ
1. Να προτιμάς ονόματα αρχείων καθαρά, περιγραφικά και επαγγελματικά.
2. Αν υπάρχουν πολλά Python αρχεία, να δίνεις σαφή και πραγματικά filenames, π.χ. (train_and_test_mnist.py), (load_and_test_mnist.py), (app.py), (server.py), (client.py).
3. Να αποφεύγεις μπερδεμένα προσωρινά ονόματα, γενικά labels ή ασαφή file naming.
4. Αν η λύση έχει πάνω από ένα αρχείο, να αναφέρεις ρητά όλα τα ακριβή ονόματα αρχείων.

ΥΠΟΧΡΕΩΤΙΚΗ ΔΟΜΗ ΑΠΑΝΤΗΣΗΣ ΓΙΑ ΑΙΤΗΜΑΤΑ ΚΩΔΙΚΑ / ΕΦΑΡΜΟΓΗΣ / ΥΛΟΠΟΙΗΣΗΣ
Όταν ο χρήστης ζητά εφαρμογή, script, κώδικα, διόρθωση, refactor, βελτίωση, GUI, API, αρχιτεκτονική ή γενικά τεχνική υλοποίηση, απαγορεύεται να αλλάξεις αυτή τη βασική δομή:

1) Περιγραφή Εφαρμογής
2) Πλήρης Κώδικας
3) Προαπαιτούμενα

[1] Περιγραφή Εφαρμογής
- Ξεκίνα πάντα με ακριβώς αυτόν τον τίτλο: "Περιγραφή Εφαρμογής"
- Γράψε σύντομη αλλά ουσιαστική περιγραφή σε απλό, επαγγελματικό κείμενο.
- Εξήγησε τι κάνει η εφαρμογή, ποιο πρόβλημα λύνει, ποια είναι τα βασικά χαρακτηριστικά της και πώς χρησιμοποιείται.
- Αν η λύση έχει πάνω από ένα αρχείο, να αναφέρεις ρητά και καθαρά τα ονόματα των αρχείων μέσα στην περιγραφή, σε μορφή όπως: (main.py), (api.py), (gui.py).
- Όταν υπάρχουν πολλά scripts Python, να αναφέρεις όλα τα ακριβή .py ονόματα μέσα στην περιγραφή με φυσικό τρόπο.
- Μην γράφεις ψευδοκώδικα σε αυτή την ενότητα.
- Μην παραλείπεις ποτέ αυτή την ενότητα.

[2] Πλήρης Κώδικας
- Συνέχισε πάντα με ακριβώς αυτόν τον τίτλο: "Πλήρης Κώδικας"
- Δώσε πλήρη, σωστό, εκτελέσιμο, επαγγελματικό και ολοκληρωμένο κώδικα.
- Ο κώδικας πρέπει να είναι σε καθαρά markdown code blocks για εύκολο copy.
- Μην δίνεις ποτέ αποσπασματικό κώδικα, placeholders, "...", "συνέχεια", "συμπλήρωσε εδώ" ή ψευδοκώδικα.
- Αν διορθώνεις υπάρχον κώδικα, να επιστρέφεις το πλήρες διορθωμένο αρχείο και όχι μόνο το patch, εκτός αν ο χρήστης ζητήσει ρητά diff.
- Αν η λύση γίνεται σωστά σε ένα αρχείο, προτίμησε ένα πλήρες αρχείο.
- Αν απαιτούνται πολλά αρχεία, δώσε όλα τα απαραίτητα αρχεία και μόνο τα απαραίτητα.
- Κάθε πλήρες αρχείο πρέπει να δίνεται ολόκληρο μέσα σε ένα και μόνο ένα code block.
- Μην σπας το ίδιο αρχείο σε πολλά συνεχόμενα code blocks.
- Να περιλαμβάνεις όλα τα απαραίτητα imports.
- Να φροντίζεις για σωστή διαχείριση exceptions, καθαρά μηνύματα λάθους και ασφαλή συμπεριφορά.
- Να χρησιμοποιείς λογικό validation εισόδων.
- Για αρχεία ρυθμίσεων, environment variables, JSON, YAML, HTML, CSS, JS, SQL, shell, batch ή άλλα βοηθητικά αρχεία, να τα δίνεις επίσης ολοκληρωμένα όταν απαιτούνται.
- Όταν η λύση απαιτεί ή ωφελείται από διαγράμματα ή τεχνικά σχέδια, μπορείς να δίνεις και πλήρη αρχεία SVG ως μέρος της τελικής λύσης.
- Κάθε αρχείο SVG πρέπει να δίνεται ολόκληρο μέσα σε ένα και μόνο ένα code block.
- Το περιεχόμενο SVG πρέπει να είναι έγκυρο, πλήρες και άμεσα αποθηκεύσιμο ως .svg.

[3] Προαπαιτούμενα
- Ολοκλήρωσε πάντα με ακριβώς αυτόν τον τίτλο: "Προαπαιτούμενα"
- Δώσε καθαρά και συνοπτικά ό,τι χρειάζεται για να τρέξει η εφαρμογή.
- Συμπερίλαβε βιβλιοθήκες Python, pip εντολές εγκατάστασης, απαιτούμενες εκδόσεις, αρχεία ρυθμίσεων, environment variables, εξωτερικές υπηρεσίες, βάσεις δεδομένων, μοντέλα ή άλλα dependencies.
- Αν δεν απαιτείται κάτι ιδιαίτερο, να το λες καθαρά.
- Όταν χρειάζεται, να δίνεις συγκεκριμένη εντολή εγκατάστασης όπως π.χ. pip install package1 package2.

ΚΑΝΟΝΕΣ ΓΙΑ ΕΠΙΣΤΗΜΟΝΙΚΕΣ ΑΠΑΝΤΗΣΕΙΣ
Όταν ο χρήστης ζητά θεωρία, εξήγηση, λύση άσκησης, επιστημονική παρουσίαση ή επιστημονική ανάλυση, να μην επιβάλεις τεχνητά τη δομή "Περιγραφή Εφαρμογής / Πλήρης Κώδικας / Προαπαιτούμενα", εκτός αν μαζί ζητά και υλοποίηση ή κώδικα.
Σε τέτοιες περιπτώσεις:
1. Να δίνεις καθαρή επιστημονική δομή με τίτλους, υποενότητες, τύπους, πίνακες και παραδείγματα όπου χρειάζεται.
2. Να ξεχωρίζεις καθαρά:
   - θεωρία
   - τύπους
   - ανάλυση
   - υπολογισμούς
   - τελικό αποτέλεσμα
3. Το τελικό αποτέλεσμα να τονίζεται με καθαρό block equation ή εμφανές συμπέρασμα.
4. Να εξηγείς βήμα-βήμα όταν υπάρχει λύση άσκησης.
5. Αν ο χρήστης ζητήσει παράδειγμα, να δίνεις πλήρως μορφοποιημένο παράδειγμα.
6. Αν ο χρήστης ζητήσει πίνακα, να χρησιμοποιείς markdown table και να διατηρείς σωστά τους συμβολισμούς μέσα στα κελιά.
7. Αν ο χρήστης ζητήσει κώδικα LaTeX ή source notation, τότε και μόνο τότε να δίνεις κυριολεκτικά το source.

ΜΑΘΗΜΑΤΙΚΑ
1. Να χρησιμοποιείς σωστά:
   - ολοκληρώματα
   - παραγώγους
   - μερικές παραγώγους
   - όρια
   - αθροίσματα
   - γινόμενα
   - ρίζες
   - πίνακες
   - ορίζουσες
   - διανύσματα
   - τελεστές όπως \nabla, \cdot, \times
   - σύνολα όπως \mathbb{R}, \mathbb{N}, \mathbb{C}
   - λογικά σύμβολα όπως \forall, \exists, \Rightarrow, \Leftrightarrow
2. Να χρησιμοποιείς σωστά δείκτες, εκθέτες, κλάσματα και απόλυτες τιμές.
3. Διανύσματα να εμφανίζονται ως \mathbf{v} ή \vec{v}, ανάλογα με το συμφραζόμενο.
4. Να χρησιμοποιείς block equations για βασικούς τύπους ή τελικά αποτελέσματα.

ΦΥΣΙΚΗ
1. Να χρησιμοποιείς σωστά φυσικά σύμβολα και μονάδες SI.
2. Να γράφεις καθαρά:
   - \vec{F}, \vec{E}, \vec{B}
   - \rho, \varepsilon_0, \mu_0, \lambda, \omega
   - \nabla \cdot \vec{E}, \nabla \times \vec{B}
3. Να αποδίδεις σωστά εξισώσεις από:
   - μηχανική
   - ηλεκτρομαγνητισμό
   - θερμοδυναμική
   - κυματική
   - κβαντική φυσική
4. Οι μονάδες να εμφανίζονται με καθαρό επιστημονικό τρόπο.

ΧΗΜΕΙΑ
1. Να χρησιμοποιείς \ce{...} για χημικές εξισώσεις και χημικούς τύπους όταν είναι χρήσιμο.
2. Να γράφεις σωστά:
   - \ce{H2O}
   - \ce{H2SO4}
   - \ce{NaOH}
   - \ce{CaCO3}
   - \ce{CH3COOH}
   - \ce{NH4+}
   - \ce{SO4^2-}
3. Να αποδίδεις σωστά:
   - φορτία ιόντων
   - συντελεστές αντίδρασης
   - βέλη αντίδρασης
   - καταστάσεις ύλης όπως (s), (l), (g), (aq)
4. Όπου χρειάζεται, να χρησιμοποιείς \pu{...} για μονάδες και ποσότητες.
5. Να αποδίδεις σωστά pH, pKa, pKb, Kc, Ka, Kb, Ksp.
6. Σε ισοστάθμιση, στοιχειομετρία και οξειδοαναγωγή να γράφεις καθαρά βήματα.

ΒΙΟΛΟΓΙΑ ΚΑΙ ΒΙΟΧΗΜΕΙΑ
1. Να χρησιμοποιείς σωστούς συμβολισμούς για:
   - DNA, RNA, mRNA, tRNA
   - ATP, ADP, AMP, NADH, FADH2
   - ένζυμα, υποστρώματα, ρυθμούς αντίδρασης
   - γονότυπους και φαινότυπους
   - βιοστατιστική
2. Όταν υπάρχει ποσοτική σχέση, να χρησιμοποιείς LaTeX.
3. Να αποδίδεις σωστά εξισώσεις όπως Michaelis-Menten:
   $$v = \frac{V_{\max}[S]}{K_m + [S]}$$
4. Σε γενετική να χρησιμοποιείς σωστά συμβολισμούς όπως AA, Aa, aa και πίνακες Punnett.

ΗΛΕΚΤΡΟΛΟΓΙΑ ΚΑΙ ΗΛΕΚΤΡΟΝΙΚΗ
1. Να χρησιμοποιείς σωστά:
   - V, I, R, C, L, P
   - \Omega, k\Omega, M\Omega
   - \mu A, mA, \mu F, nF, mH
   - V_{in}, V_{out}, I_D, V_{GS}, V_{DS}
2. Να γράφεις σωστά νόμους και εξισώσεις όπως:
   $$V = IR$$
   $$P = VI$$
   $$V_C(t)=V_0 e^{-t/RC}$$
3. Να αποδίδεις σωστά σημειογραφία διόδων, BJT, MOSFET, AC/DC, RMS, σύνθετης αντίστασης και φασόρων.
4. Να χρησιμοποιείς σαφείς πίνακες για σύμβολα, μονάδες και φυσικό νόημα.

ΨΗΦΙΑΚΑ ΗΛΕΚΤΡΟΝΙΚΑ ΚΑΙ ΛΟΓΙΚΗ
1. Να χρησιμοποιείς σωστά:
   - A, B, C, D
   - Q, \overline{Q}
   - Q_n, Q_{n+1}
   - J, K, D, T
   - f_{clk}, T_{clk}
2. Να αποδίδεις καθαρά:
   - Boolean άλγεβρα
   - πίνακες αλήθειας
   - πύλες AND, OR, NOT, NAND, NOR, XOR, XNOR
   - flip-flops, counters, registers
3. Να γράφεις σωστά εκφράσεις όπως:
   $$F = A \oplus B$$
   $$Q_{n+1}=J\overline{Q_n}+\overline{K}Q_n$$

ΣΤΑΤΙΣΤΙΚΗ ΚΑΙ ΕΦΑΡΜΟΣΜΕΝΑ ΜΑΘΗΜΑΤΙΚΑ
1. Να χρησιμοποιείς σωστά τύπους όπως:
   $$\mu = \frac{1}{N}\sum_{i=1}^{N}x_i$$
   $$\sigma^2 = \frac{1}{N}\sum_{i=1}^{N}(x_i-\mu)^2$$
   $$z=\frac{x-\mu}{\sigma}$$
2. Να παρουσιάζεις καθαρά:
   - μέσο όρο
   - διακύμανση
   - τυπική απόκλιση
   - πιθανότητες
   - κατανομές
   - παλινδρόμηση
   - υποθέσεις ελέγχου

ΠΑΡΑΓΩΓΗ ΑΡΧΕΙΩΝ SVG
1. Όταν το αίτημα του χρήστη αφορά διάγραμμα, σχήμα, σχεδίαση, τεχνική απεικόνιση, wireframe, επιστημονικό σχήμα, block diagram, flowchart, αρχιτεκτονικό διάγραμμα, γεωμετρική απεικόνιση, απλό κύκλωμα, λογικό κύκλωμα, διάγραμμα διαδικασίας, εκπαιδευτικό σχεδιάγραμμα ή άλλο οπτικό στοιχείο, να εξετάζεις αν η παραγωγή αρχείου SVG είναι η καταλληλότερη λύση.
2. Αν το SVG διευκολύνει ουσιαστικά, να παράγεις πλήρες και έγκυρο περιεχόμενο αρχείου .svg, έτοιμο για αποθήκευση και άνοιγμα χωρίς επιπλέον επεξεργασία.
3. Το SVG να είναι standalone, με σωστό <svg ...>, viewBox, πλάτος, ύψος και πλήρη δομή.
4. Να μην δίνεις αποσπασματικό SVG, placeholders ή ημιτελή fragments.
5. Τα labels, οι τίτλοι και οι σημειώσεις μέσα στο SVG να είναι καθαρά, ευανάγνωστα και επαγγελματικά μορφοποιημένα.
6. Να προτιμάς καθαρό, απλό και επαγγελματικό σχεδιασμό αντί για υπερβολικά πολύπλοκο ή εύθραυστο SVG.
7. Όταν χρειάζεται, να χρησιμοποιείς ομάδες (<g>), βασικά σχήματα (<rect>, <circle>, <line>, <path>, <polygon>), markers, βέλη, λεζάντες και στοιχειώδη styling με τρόπο συντηρήσιμο.
8. Για επιστημονικά ή τεχνικά σχήματα να δίνεις έμφαση στην καθαρότητα, στην ακρίβεια των συμβολισμών και στην αναγνωσιμότητα.
9. Όταν το αίτημα αφορά λογικά κυκλώματα, block diagrams, αρχιτεκτονική λογισμικού, ροές δεδομένων, μαθηματικά σχήματα, φυσικά διαγράμματα ή εκπαιδευτικά γραφήματα, να προτείνεις σιωπηρά SVG ως μέρος της λύσης όταν αυτό βελτιώνει το αποτέλεσμα.
10. Αν η λύση περιλαμβάνει αρχεία κώδικα και επιπλέον απαιτείται SVG, να δίνεις και το .svg ως πλήρες αρχείο μέσα στην ενότητα "Πλήρης Κώδικας".
11. Αν απαιτούνται πολλά αρχεία SVG, να δίνεις κάθε αρχείο ξεχωριστά, ολόκληρο, με σαφές και επαγγελματικό filename όπως (diagram.svg), (logic_circuit.svg), (system_architecture.svg).
12. Να αποφεύγεις άσκοπα ενσωματωμένα raster στοιχεία μέσα στο SVG, εκτός αν είναι απολύτως απαραίτητα.
13. Όταν ο χρήστης ζητά απλώς το τελικό διάγραμμα, να δίνεις απευθείας το πλήρες SVG και όχι μόνο περιγραφή του.
14. Όταν το SVG δεν είναι η σωστή λύση, να προτιμάς κανονικό κείμενο, markdown table ή κώδικα, χωρίς να επιβάλλεις SVG άσκοπα.

ΕΠΙΠΛΕΟΝ ΚΡΙΣΙΜΟΙ ΚΑΝΟΝΕΣ
1. Να μην εμφανίζεις ακατέργαστους τύπους σαν απλό κείμενο όταν μπορούν να αποδοθούν κανονικά με επιστημονική μορφοποίηση.
2. Να δίνεις προτεραιότητα στην οπτικά καθαρή και επαγγελματική παρουσίαση.
3. Όταν ο χρήστης ζητά πλήρες αρχείο, να δίνεις ολόκληρο το αρχείο και όχι patch, εκτός αν ζητηθεί ρητά diff.
4. Όταν υπάρχουν πολλά αρχεία, να είναι ξεκάθαρο ποιο block αντιστοιχεί σε ποιο filename.
5. Όταν ζητείται διόρθωση υπάρχοντος κώδικα, να παραδίδεις το τελικό διορθωμένο αποτέλεσμα σε πλήρη μορφή.
6. Να είσαι αυστηρός με τη συνέπεια ανάμεσα στην περιγραφή, στον κώδικα, στα filenames και στα προαπαιτούμενα.
7. Μην αναφέρεις αρχεία που δεν παρέχεις πραγματικά.
8. Μην δηλώνεις ότι παράγεται ένα αρχείο αν ο κώδικας ή το SVG που δίνεις δεν το υποστηρίζει πράγματι.

ΤΕΛΙΚΗ ΑΠΑΙΤΗΣΗ
Οι απαντήσεις σου πρέπει να είναι επιστημονικά σωστές, τεχνικά επαγγελματικές, μορφοποιητικά άψογες, συνεπείς και άμεσα αξιοποιήσιμες, είτε ο χρήστης ζητά θεωρία είτε ζητά πλήρη υλοποίηση κώδικα είτε ζητά διαγράμματα και αρχεία SVG.'''
TEXT_EXTENSIONS: Set[str] = {'.txt', '.py', '.md', '.markdown', '.json', '.jsonl', '.csv', '.tsv', '.yaml', '.yml', '.xml', '.html', '.htm', '.css', '.js', '.ts', '.jsx', '.tsx', '.java', '.c', '.cpp', '.h', '.hpp', '.cs', '.go', '.rs', '.php', '.rb', '.swift', '.kt', '.kts', '.sql', '.ini', '.cfg', '.conf', '.log', '.bat', '.ps1', '.sh', '.zsh', '.toml', '.tex', '.r', '.m'}
IMAGE_EXTENSIONS: Set[str] = {'.png', '.jpg', '.jpeg', '.webp', '.bmp', '.gif', '.tif', '.tiff'}
ACCEPTED_FILE_TYPES = 'image/*,.txt,.py,.md,.markdown,.json,.jsonl,.csv,.tsv,.yaml,.yml,.xml,.html,.htm,.css,.js,.ts,.jsx,.tsx,.java,.c,.cpp,.h,.hpp,.cs,.go,.rs,.php,.rb,.swift,.kt,.kts,.sql,.ini,.cfg,.conf,.log,.bat,.ps1,.sh,.zsh,.toml,.tex,.r,.m,.pdf'
OFFICIAL_SEARCH_URL = 'https://ollama.com/search?c=cloud'
OFFICIAL_GENERAL_SEARCH_URL = 'https://ollama.com/search'
OFFICIAL_CLOUD_API_TAGS_URL = 'https://ollama.com/api/tags'
OLLAMA_SEARCH_API_URL = 'https://ollama.com/api/search'
OLLAMA_LIBRARY_BASE = 'https://ollama.com/library/'
REQUEST_HEADERS: Dict[str, str] = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/146.0.0.0 Safari/537.36', 'Accept': 'text/html,application/json;q=0.9,*/*;q=0.8'}
OLLAMA_DIRECT_BASE_URL = 'https://ollama.com'
OLLAMA_DIRECT_API_BASE_URL = f'{OLLAMA_DIRECT_BASE_URL}/api'
CLOUD_TAG_RE = re.compile('\\b([a-zA-Z0-9._/-]+:(?:[a-zA-Z0-9._-]+-)?cloud)\\b')
LIBRARY_LINK_RE = re.compile('href="/library/([a-zA-Z0-9._/-]+)"', flags=re.IGNORECASE)
SEARCH_TEXT_FAMILY_RE = re.compile('>\\s*([a-zA-Z0-9._/-]+)\\s+[^<]{0,250}?\\bcloud\\b', flags=re.IGNORECASE)
CONTEXT_WINDOW_RE = re.compile('([0-9][0-9.,]*)\\s*([KMB])?\\s+context\\s+window', flags=re.IGNORECASE)
CLOUD_WORD_RE = re.compile('cloud', flags=re.IGNORECASE)

DEFAULT_PROMPT_PROFILE_ID = 'scientific-technical'
DEFAULT_VISUALIZATION_ENGINE = 'auto'
SUPPORTED_VISUALIZATION_ENGINES: Set[str] = {'auto', 'svg', 'python-plot'}
PLOT_RENDER_TIMEOUT_SECONDS = 25
PLOT_RENDER_MAX_CODE_CHARS = 120000
PLOT_RENDER_MAX_STDOUT_CHARS = 12000
SAFE_PLOT_IMPORT_ROOTS: Set[str] = {'math', 'statistics', 'random', 'itertools', 'functools', 'collections', 'datetime', 'decimal', 'fractions', 'numpy', 'pandas', 'matplotlib', 'mpl_toolkits'}
DISALLOWED_PLOT_CALL_NAMES: Set[str] = {'open', 'exec', 'eval', 'compile', '__import__', 'input', 'help', 'quit', 'exit', 'breakpoint'}
DISALLOWED_PLOT_ATTRIBUTE_ROOTS: Set[str] = {'os', 'sys', 'subprocess', 'socket', 'shutil', 'pathlib', 'urllib', 'requests', 'http', 'ftplib', 'ctypes', 'pickle', 'marshal', 'importlib', 'asyncio', 'threading', 'multiprocessing', 'webbrowser'}

def normalize_prompt_profile_id(profile_id: str) -> str:
    """Κανονικοποιεί το id του prompt profile σε ασφαλή και υποστηριζόμενη μορφή."""
    cleaned = re.sub(r'[^a-z0-9_-]+', '-', str(profile_id or '').strip().lower()).strip('-')
    if not cleaned:
        return DEFAULT_PROMPT_PROFILE_ID
    return cleaned

def normalize_visualization_engine(value: str) -> str:
    """Κανονικοποιεί τη στρατηγική οπτικοποίησης που θα σταλεί στο μοντέλο και στο UI."""
    cleaned = str(value or '').strip().lower()
    if cleaned in {'plot', 'python', 'python_plot', 'matplotlib'}:
        cleaned = 'python-plot'
    if cleaned in SUPPORTED_VISUALIZATION_ENGINES:
        return cleaned
    return DEFAULT_VISUALIZATION_ENGINE

def get_visualization_engine_options() -> List[Dict[str, str]]:
    """Επιστρέφει τις διαθέσιμες επιλογές για το visualization engine selector του UI."""
    return [
        {'id': 'auto', 'label': 'Auto', 'description': 'Το μοντέλο επιλέγει SVG για εννοιολογικά σχήματα και Python/matplotlib για υπολογιστικά plots.'},
        {'id': 'svg', 'label': 'SVG / Diagrams', 'description': 'Προτεραιότητα σε καθαρά SVG για αρχιτεκτονικά, λογικά, εκπαιδευτικά και θεωρητικά σχήματα.'},
        {'id': 'python-plot', 'label': 'Python Plot / matplotlib', 'description': 'Προτεραιότητα σε πραγματικά υπολογιστικά διαγράμματα με Python και matplotlib.'},
    ]

def build_visualization_engine_instruction(engine: str) -> str:
    """Συνθέτει σύντομη αλλά σαφή οδηγία προς το μοντέλο για τον μηχανισμό οπτικοποίησης."""
    normalized = normalize_visualization_engine(engine)
    if normalized == 'svg':
        return '''VISUALIZATION ENGINE: SVG ONLY
PRIORITY RULE: Οι παρακάτω οδηγίες υπερισχύουν άλλων κανόνων οπτικοποίησης.
1. Για διαγράμματα, flowcharts, λογικά κυκλώματα, φυσικά σχήματα και εκπαιδευτικές απεικονίσεις να προτιμάς έγκυρο standalone SVG.
2. Απόφυγε Python plotting code εκτός αν ο χρήστης το ζητήσει ρητά.
3. Το SVG να είναι πλήρες, καθαρό, άμεσα αποθηκεύσιμο και να αποδίδεται σωστά στο UI.'''.strip()
    if normalized == 'python-plot':
        return '''VISUALIZATION ENGINE: PYTHON PLOT / MATPLOTLIB ONLY
PRIORITY RULE: Οι παρακάτω οδηγίες υπερισχύουν οποιουδήποτε γενικού κανόνα παραγωγής SVG.
1. Όταν ζητείται ή υπονοείται οπτικοποίηση, γράφημα, καμπύλη, τροχιά, διάγραμμα συνάρτησης, scientific chart ή άλλη οπτική αναπαράσταση, να επιστρέφεις Python code block με matplotlib και όχι SVG.
2. ΜΗΝ επιστρέφεις raw <svg>...</svg>, fenced ```svg``` blocks ή οποιαδήποτε άλλη μορφή SVG, εκτός αν ο χρήστης ζητήσει ρητά το SVG source/XML.
3. Ακόμη και για θεωρητικά ή εκπαιδευτικά σχήματα, όταν χρειάζεται οπτική απεικόνιση να χρησιμοποιείς matplotlib με axes, annotations, arrows, patches ή άλλες δυνατότητες της βιβλιοθήκης.
4. Το plotting script να είναι standalone, να περιλαμβάνει όλα τα απαραίτητα imports, να είναι έτοιμο για headless εκτέλεση και να χρησιμοποιεί plt.show() ή το OUTPUT_PATH/PLOT_OUTPUT_PATH όταν χρειάζεται αποθήκευση.'''.strip()
    return '''VISUALIZATION ENGINE: AUTO
1. Για εννοιολογικά, αρχιτεκτονικά, φυσικά ή εκπαιδευτικά σχήματα να προτιμάς standalone SVG.
2. Για πραγματικά data plots, καμπύλες, histogram, scatter, bar chart, regression, signal plots ή άλλη υπολογιστική γραφική απεικόνιση να επιστρέφεις πλήρες Python code block με matplotlib.
3. Να εξηγείς σύντομα γιατί επέλεξες SVG ή Python plot όταν αυτό προσθέτει σαφήνεια.'''.strip()


def build_visualization_engine_user_hint(engine: str) -> str:
    """Επιστρέφει πρόσθετο user-level hint ώστε η επιλογή visualization engine του UI να έχει ισχυρότερη προτεραιότητα."""
    normalized = normalize_visualization_engine(engine)
    if normalized == 'python-plot':
        return '''[Visualization Engine από το UI: Python Plot / matplotlib]
- Η επιλογή αυτή υπερισχύει της προτίμησης για SVG.
- Αν η απάντηση χρειάζεται οποιαδήποτε οπτικοποίηση, γράφημα, καμπύλη, τροχιά, διάγραμμα ή σχεδιάγραμμα, δώσε το ως πλήρες Python code block με matplotlib.
- Μην επιστρέφεις raw <svg>...</svg> και μην δίνεις ```svg``` code block, εκτός αν ζητηθεί ρητά SVG source.'''.strip()
    if normalized == 'svg':
        return '''[Visualization Engine από το UI: SVG / Diagrams]
- Αν η απάντηση χρειάζεται οπτικοποίηση, προτίμησε πλήρες standalone SVG.
- Απόφυγε Python plotting code εκτός αν ζητηθεί ρητά.'''.strip()
    return ''

def build_prompt_profile_prompt(profile_id: str) -> str:
    """Δημιουργεί το πλήρες prompt για κάθε επιλέξιμο profile του studio.

    Το profile-specific τμήμα μπαίνει πριν από το γενικό embedded prompt,
    ώστε η αλλαγή profile να φαίνεται αμέσως και οπτικά στο UI textarea.
    """
    normalized = normalize_prompt_profile_id(profile_id)
    base_prompt = str(DEFAULT_SYSTEM_PROMPT or '').strip()
    extras: Dict[str, str] = {
        'scientific-technical': '''PROMPT PROFILE: SCIENTIFIC / TECHNICAL EXPERT
ROLE EMPHASIS:
1. Να λειτουργείς ως αυστηρός επιστημονικοτεχνικός βοηθός με έμφαση σε ακρίβεια, ορολογία, μονάδες, σαφή μαθηματική μορφοποίηση και επαγγελματική τεχνική ανάλυση.
2. Να δίνεις ισορροπία ανάμεσα σε θεωρία, τεκμηρίωση και πρακτική υλοποίηση.
3. Όπου χρειάζεται, να χρησιμοποιείς πίνακες, μαθηματικούς τύπους, SVG ή Python plots.'''.strip(),
        'code-development': '''PROMPT PROFILE: CODE DEVELOPMENT / PRODUCTION ENGINEERING
ROLE EMPHASIS:
1. Να δίνεις προτεραιότητα σε production-ready κώδικα, αρχιτεκτονική, validation, error handling, maintainability και καθαρό separation of concerns.
2. Για debugging ή refactor να εντοπίζεις πρώτα τις βασικές αιτίες και μετά να παραδίδεις ολοκληρωμένο τελικό αρχείο.
3. Να συνοδεύεις κάθε τεχνική λύση με σαφή χρήση, προαπαιτούμενα και επαγγελματικές επιλογές υλοποίησης.'''.strip(),
        'educational-teacher': '''PROMPT PROFILE: EDUCATIONAL / TEACHER MODE
ROLE EMPHASIS:
1. Να απαντάς με έντονα διδακτικό τρόπο, βήμα-βήμα, με καθαρή οργάνωση και γλώσσα κατάλληλη για εκπαίδευση.
2. Να διαχωρίζεις αυστηρά θεωρία, τύπους, διαδικασία λύσης, έλεγχο αποτελέσματος και συμπέρασμα.
3. Όπου βοηθά, να δίνεις μικρά παραδείγματα, διδακτικές παρατηρήσεις και συνηθισμένα λάθη.'''.strip(),
        'math-physics-solver': '''PROMPT PROFILE: MATH & PHYSICS SOLVER
ROLE EMPHASIS:
1. Να δίνεις αυστηρή μεθοδολογία επίλυσης: δεδομένα, ζητούμενα, τύποι, αντικατάσταση τιμών, πράξεις, έλεγχο μονάδων και τελικό αποτέλεσμα.
2. Να χρησιμοποιείς καθαρά διανύσματα, μονάδες SI, LaTeX block equations και εκπαιδευτική δομή τύπου λυμένου παραδείγματος.
3. Όταν ένα σχήμα ή γράφημα βελτιώνει ουσιαστικά την κατανόηση, να το παράγεις.'''.strip(),
        'code-review-debugger': '''PROMPT PROFILE: CODE REVIEW / DEBUGGER
ROLE EMPHASIS:
1. Να εντοπίζεις bug patterns, ρίσκα, design smells, ασυνέπειες API, edge cases και πιθανά performance bottlenecks.
2. Να εξηγείς πρώτα τι φταίει και γιατί και έπειτα να παραδίδεις διορθωμένο, πλήρες και καθαρό αποτέλεσμα.
3. Να προτείνεις στοχευμένα tests και ελέγχους επαλήθευσης.'''.strip(),
        'research-analysis': '''PROMPT PROFILE: RESEARCH / STRUCTURED ANALYSIS
ROLE EMPHASIS:
1. Να συνθέτεις απαντήσεις με δομή, σαφή assumptions, μεθοδολογία, σύγκριση εναλλακτικών και τεκμηριωμένο συμπέρασμα.
2. Να ξεχωρίζεις τι είναι βέβαιο, τι είναι υπόθεση και τι απαιτεί περαιτέρω επαλήθευση.
3. Να δίνεις ισορροπημένες, δομημένες και ώριμες τεχνικοεπιστημονικές αξιολογήσεις.'''.strip(),
        'diagram-visualization': '''PROMPT PROFILE: DIAGRAM & VISUALIZATION MODE
ROLE EMPHASIS:
1. Να δίνεις ιδιαίτερη έμφαση στην οπτική αναπαράσταση γνώσης με SVG ή Python plots, ανάλογα με το είδος του προβλήματος.
2. Να σχεδιάζεις καθαρά labels, λεζάντες, άξονες, units, υπόμνημα και τίτλο όπου απαιτείται.
3. Να παράγεις διαγράμματα που είναι άμεσα χρήσιμα μέσα στο chat και στο PDF export.'''.strip(),
        'concise-engineer': '''PROMPT PROFILE: CONCISE ENGINEER
ROLE EMPHASIS:
1. Να δίνεις σύντομες αλλά υψηλής πυκνότητας απαντήσεις, χωρίς περιττή φλυαρία και χωρίς να χάνεις τεχνική ακρίβεια.
2. Να προτιμάς συνοπτικές ενότητες, σαφή συμπεράσματα και ουσιαστικές τεχνικές αποφάσεις.
3. Όταν χρειάζεται, να επεκτείνεις μόνο τα απολύτως κρίσιμα σημεία.'''.strip(),
    }
    extra = extras.get(normalized) or extras[DEFAULT_PROMPT_PROFILE_ID]
    header = f"### ACTIVE PROMPT PROFILE: {normalized}\n{extra}".strip()
    if not base_prompt:
        return header
    return (header + '\n\n--- GENERAL SYSTEM RULES ---\n' + base_prompt).strip()


def get_prompt_profiles_catalog() -> List[Dict[str, str]]:
    """Επιστρέφει τον κατάλογο των prompt profiles για το UI και για το backend."""
    profiles = [
        ('scientific-technical', 'Scientific / Technical Expert', 'Ισορροπημένο προφίλ για επιστήμη, τεχνική τεκμηρίωση, διαγράμματα και αναλυτικές απαντήσεις.'),
        ('code-development', 'Code Development / Production Engineering', 'Έμφαση σε production-ready κώδικα, αρχιτεκτονική και υλοποίηση.'),
        ('educational-teacher', 'Educational / Teacher Mode', 'Διδακτική προσέγγιση με βήμα-βήμα εξήγηση και εκπαιδευτική δομή.'),
        ('math-physics-solver', 'Math & Physics Solver', 'Ιδανικό για λυμένες ασκήσεις, θεωρία, τύπους και φυσικά διαγράμματα.'),
        ('code-review-debugger', 'Code Reviewer / Debugger', 'Στοχευμένο profile για εντοπισμό σφαλμάτων, refactor και τεχνική διάγνωση.'),
        ('research-analysis', 'Research / Structured Analysis', 'Δομημένη ανάλυση, αξιολόγηση εναλλακτικών και τεκμηριωμένο συμπέρασμα.'),
        ('diagram-visualization', 'Diagram & Visualization Mode', 'Έμφαση στην παραγωγή SVG και Python-based οπτικοποιήσεων.'),
        ('concise-engineer', 'Concise Engineer', 'Πιο σύντομες, πυκνές και ουσιαστικές τεχνικές απαντήσεις.'),
    ]
    return [
        {'id': profile_id, 'label': label, 'description': description, 'prompt': build_prompt_profile_prompt(profile_id)}
        for profile_id, label, description in profiles
    ]

def get_prompt_profile_map() -> Dict[str, Dict[str, str]]:
    """Χτίζει λεξικό prompt profiles με κλειδί το profile id."""
    return {item['id']: item for item in get_prompt_profiles_catalog()}

def get_prompt_profile(profile_id: str) -> Dict[str, str]:
    """Επιστρέφει ασφαλώς ένα prompt profile, με fallback στο default."""
    normalized = normalize_prompt_profile_id(profile_id)
    profile_map = get_prompt_profile_map()
    return profile_map.get(normalized) or profile_map[DEFAULT_PROMPT_PROFILE_ID]

@dataclass
class ModelRegistry:
    """Dataclass που κρατά τη λίστα μοντέλων, το metadata τους και την κατάσταση του τελευταίου refresh.

Η πρόσβαση στα δεδομένα προστατεύεται με lock, ώστε διαφορετικά threads να διαβάζουν και να ενημερώνουν συνεπή κατάσταση."""
    models: List[str] = field(default_factory=list)
    model_meta: Dict[str, Dict[str, object]] = field(default_factory=dict)
    source: str = 'initializing'
    last_refresh_ts: float = 0.0
    last_error: str = ''
    refresh_in_progress: bool = False
    recommended_model: str = ''
    lock: threading.Lock = field(default_factory=threading.Lock)

    def as_dict(self) -> Dict[str, object]:
        """Επιστρέφει thread-safe snapshot της κατάστασης του registry, με μοντέλα ταξινομημένα κατά συνολικό score."""
        with self.lock:
            models = list(self.models)
            rec = self.recommended_model
            model_meta = copy.deepcopy(self.model_meta)
            source = self.source
            last_refresh_ts = self.last_refresh_ts
            last_error = self.last_error
            refresh_in_progress = self.refresh_in_progress
        try:
            models = sorted(models, key=lambda model: score_model(model, model_meta.get(model, {}), 'overall'), reverse=True)
        except Exception:
            pass
        models_with_context = sum((1 for model in models if isinstance(model_meta.get(model), dict) and model_meta.get(model, {}).get('num_ctx_max')))
        return {'models': models, 'model_meta': model_meta, 'models_with_context': models_with_context, 'source': source, 'last_refresh_ts': last_refresh_ts, 'last_error': last_error, 'refresh_in_progress': refresh_in_progress, 'recommended_model': rec}

@dataclass
class ChatSession:
    """Dataclass που αποθηκεύει το ιστορικό συνομιλίας και τις προσωρινές διαδρομές αρχείων της τρέχουσας συνεδρίας."""
    messages: List[Dict] = field(default_factory=list)
    history: List[Dict] = field(default_factory=list)
    upload_paths: Set[str] = field(default_factory=set)
    lock: threading.Lock = field(default_factory=threading.Lock)

    def reset(self) -> None:
        """Σβήνει το ιστορικό συνομιλίας και διαγράφει τα προσωρινά uploaded files από δίσκο και από τους temp κλάδους."""
        with self.lock:
            self.messages.clear()
            self.history.clear()
            paths_to_delete = list(self.upload_paths)
            self.upload_paths.clear()
        cleanup_targets = [UPLOADS_DIR, GENERATED_CODE_DIR, GENERATED_MEDIA_DIR, Path(tempfile.gettempdir()) / 'ollama_cloud_chat_exec']
        for item in paths_to_delete:
            try:
                path = Path(item)
                if path.exists() and path.is_file():
                    path.unlink()
            except Exception:
                pass
        for target in cleanup_targets:
            try:
                if target.exists():
                    shutil.rmtree(target, ignore_errors=True)
            except Exception:
                pass
        for target in cleanup_targets[:3]:
            try:
                target.mkdir(parents=True, exist_ok=True)
            except Exception:
                pass

@dataclass
class AppConfig:
    """Dataclass για τις ρυθμίσεις που αποθηκεύονται μόνιμα στον δίσκο, όπως το API key της εφαρμογής."""
    ollama_api_key: str = ''
    active_prompt_profile: str = DEFAULT_PROMPT_PROFILE_ID
    custom_system_prompt: str = ''
    active_visualization_engine: str = DEFAULT_VISUALIZATION_ENGINE
    updated_at: str = ''
    lock: threading.Lock = field(default_factory=threading.Lock)

    def as_public_dict(self) -> Dict[str, object]:
        """Επιστρέφει JSON-safe λεξικό με τις ρυθμίσεις της εφαρμογής, κανονικοποιημένες και χωρίς εσωτερικά locks."""
        with self.lock:
            key = str(self.ollama_api_key or '')
            active_prompt_profile = normalize_prompt_profile_id(self.active_prompt_profile)
            custom_system_prompt = str(self.custom_system_prompt or '')
            active_visualization_engine = normalize_visualization_engine(self.active_visualization_engine)
            updated_at = str(self.updated_at or '')
        return {'ollama_api_key': key, 'has_ollama_api_key': bool(key), 'active_prompt_profile': active_prompt_profile, 'custom_system_prompt': custom_system_prompt, 'active_visualization_engine': active_visualization_engine, 'updated_at': updated_at}

def load_app_config_from_disk() -> AppConfig:
    """Φορτώνει ρυθμίσεις εφαρμογής από JSON αρχείο. Επιστρέφει AppConfig με defaults αν το αρχείο λείπει ή είναι άκυρο."""
    config = AppConfig()
    try:
        if APP_CONFIG_FILE.exists():
            data = json.loads(APP_CONFIG_FILE.read_text(encoding='utf-8'))
            if isinstance(data, dict):
                config.ollama_api_key = str(data.get('ollama_api_key', '') or '').strip()
                config.active_prompt_profile = normalize_prompt_profile_id(data.get('active_prompt_profile', DEFAULT_PROMPT_PROFILE_ID))
                config.custom_system_prompt = str(data.get('custom_system_prompt', '') or '')
                config.active_visualization_engine = normalize_visualization_engine(data.get('active_visualization_engine', DEFAULT_VISUALIZATION_ENGINE))
                config.updated_at = str(data.get('updated_at', '') or '').strip()
    except Exception as exc:
        log.warning('Αποτυχία φόρτωσης settings file %s: %s', APP_CONFIG_FILE, exc)
    return config

def save_app_config_to_disk(ollama_api_key: Optional[str]=None, active_prompt_profile: Optional[str]=None, custom_system_prompt: Optional[str]=None, active_visualization_engine: Optional[str]=None) -> AppConfig:
    """Αποθηκεύει ρυθμίσεις εφαρμογής ατομικά (write-then-rename) και ενημερώνει το global APP_CONFIG."""
    with APP_CONFIG.lock:
        current_key = str(APP_CONFIG.ollama_api_key or '')
        current_profile = normalize_prompt_profile_id(APP_CONFIG.active_prompt_profile)
        current_custom_prompt = str(APP_CONFIG.custom_system_prompt or '')
        current_engine = normalize_visualization_engine(APP_CONFIG.active_visualization_engine)
    cleaned_key = current_key if ollama_api_key is None else str(ollama_api_key or '').strip()
    cleaned_profile = current_profile if active_prompt_profile is None else normalize_prompt_profile_id(active_prompt_profile)
    cleaned_custom_prompt = current_custom_prompt if custom_system_prompt is None else str(custom_system_prompt or '')
    cleaned_engine = current_engine if active_visualization_engine is None else normalize_visualization_engine(active_visualization_engine)
    payload = {'ollama_api_key': cleaned_key, 'active_prompt_profile': cleaned_profile, 'custom_system_prompt': cleaned_custom_prompt, 'active_visualization_engine': cleaned_engine, 'updated_at': time.strftime('%Y-%m-%d %H:%M:%S')}
    tmp_path = APP_CONFIG_FILE.with_suffix('.tmp')
    try:
        tmp_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')
        tmp_path.replace(APP_CONFIG_FILE)
    except Exception:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass
        raise
    with APP_CONFIG.lock:
        APP_CONFIG.ollama_api_key = cleaned_key
        APP_CONFIG.active_prompt_profile = cleaned_profile
        APP_CONFIG.custom_system_prompt = cleaned_custom_prompt
        APP_CONFIG.active_visualization_engine = cleaned_engine
        APP_CONFIG.updated_at = payload['updated_at']
    return APP_CONFIG

def normalize_model_registry_refresh_error(exc: Exception) -> str:
    """Μετατρέπει δικτυακά/HTTP σφάλματα του refresh μοντέλων σε φιλικά και σταθερά μηνύματα για το UI."""
    message = str(exc or '').strip() or 'Άγνωστο σφάλμα ανάκτησης μοντέλων.'
    lower = message.lower()
    reason = getattr(exc, 'reason', None)
    if isinstance(reason, socket.gaierror) or 'getaddrinfo failed' in lower or 'name or service not known' in lower or 'temporary failure in name resolution' in lower:
        return 'Σφάλμα DNS/δικτύου κατά την ανάκτηση της online λίστας μοντέλων από το Ollama. Χρησιμοποιείται cache αν είναι διαθέσιμη.'
    if 'timed out' in lower or 'timeout' in lower:
        return 'Λήξη χρόνου κατά την ανάκτηση της online λίστας μοντέλων από το Ollama. Χρησιμοποιείται cache αν είναι διαθέσιμη.'
    if '401' in lower or '403' in lower:
        return 'Απορρίφθηκε η πρόσβαση στην online λίστα μοντέλων από το Ollama.'
    if '404' in lower:
        return 'Δεν βρέθηκε το endpoint της online λίστας μοντέλων στο Ollama.'
    return message

def load_model_registry_cache_from_disk() -> Tuple[List[str], Dict[str, Dict[str, object]], float]:
    """Φορτώνει από δίσκο την τελευταία επιτυχημένη λίστα official μοντέλων, αν υπάρχει."""
    try:
        if not MODEL_REGISTRY_CACHE_FILE.exists():
            return ([], {}, 0.0)
        raw = json.loads(MODEL_REGISTRY_CACHE_FILE.read_text(encoding='utf-8'))
        if not isinstance(raw, dict):
            return ([], {}, 0.0)
        models_raw = raw.get('models') if isinstance(raw.get('models'), list) else []
        models = [normalize_model_name(str(item).strip()) for item in models_raw if str(item).strip()]
        models = sorted(set((m for m in models if m)))
        meta_raw = raw.get('model_meta') if isinstance(raw.get('model_meta'), dict) else {}
        model_meta: Dict[str, Dict[str, object]] = {}
        for model in models:
            entry = meta_raw.get(model, {})
            model_meta[model] = copy.deepcopy(entry) if isinstance(entry, dict) else {}
        ts = float(raw.get('last_refresh_ts') or 0.0)
        return (models, model_meta, ts)
    except Exception as exc:
        log.warning('Αποτυχία φόρτωσης cache λίστας μοντέλων %s: %s', MODEL_REGISTRY_CACHE_FILE, exc)
        return ([], {}, 0.0)

def save_model_registry_cache_to_disk(models: List[str], model_meta: Dict[str, Dict[str, object]], last_refresh_ts: Optional[float]=None) -> None:
    """Αποθηκεύει στον δίσκο την τελευταία επιτυχημένη online λίστα official μοντέλων."""
    cleaned_models = sorted(set((normalize_model_name(str(item).strip()) for item in (models or []) if str(item).strip())))
    cleaned_meta: Dict[str, Dict[str, object]] = {}
    for model in cleaned_models:
        entry = (model_meta or {}).get(model, {})
        cleaned_meta[model] = copy.deepcopy(entry) if isinstance(entry, dict) else {}
    payload = {
        'models': cleaned_models,
        'model_meta': cleaned_meta,
        'last_refresh_ts': float(last_refresh_ts or time.time()),
        'saved_at': time.strftime('%Y-%m-%d %H:%M:%S'),
    }
    tmp_path = MODEL_REGISTRY_CACHE_FILE.with_suffix('.tmp')
    try:
        tmp_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')
        tmp_path.replace(MODEL_REGISTRY_CACHE_FILE)
    except Exception as exc:
        log.warning('Αποτυχία αποθήκευσης cache λίστας μοντέλων %s: %s', MODEL_REGISTRY_CACHE_FILE, exc)
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass

REGISTRY = ModelRegistry()
SESSION = ChatSession()
APP_CONFIG = load_app_config_from_disk()
_SERVER_START_TIME = time.time()

@dataclass
class BrowserSessionMonitor:
    """Κλάση που παρακολουθεί την ύπαρξη ενεργών browser sessions και βοηθά στον ασφαλή αυτόματο τερματισμό του server."""
    active_sessions: Dict[str, float] = field(default_factory=dict)
    lock: threading.Lock = field(default_factory=threading.Lock)
    ever_seen_session: bool = False
    shutdown_requested: bool = False
    server_ref: Optional[ThreadingHTTPServer] = None

    def attach_server(self, server: ThreadingHTTPServer) -> None:
        """Αποθηκεύει αναφορά στον HTTP server για χρήση κατά τον αυτόματο shutdown."""
        with self.lock:
            self.server_ref = server

    def touch(self, session_id: str) -> None:
        """Καταγράφει ή ανανεώνει την παρουσία ενός browser session ενημερώνοντας το τελευταίο timestamp του."""
        cleaned = str(session_id or '').strip()[:128]
        if not cleaned:
            return
        with self.lock:
            self.active_sessions[cleaned] = time.time()
            self.ever_seen_session = True

    def close(self, session_id: str) -> None:
        """Αφαιρεί το session_id από το σύνολο των ενεργών sessions."""
        cleaned = str(session_id or '').strip()[:128]
        if not cleaned:
            return
        with self.lock:
            self.active_sessions.pop(cleaned, None)

    def _cleanup_stale_locked(self, now_ts: float) -> None:
        """Αφαιρεί sessions που δεν έχουν δώσει heartbeat εντός του ορίου BROWSER_SESSION_HEARTBEAT_STALE_SECONDS. Καλείται ενώ το lock είναι ήδη ακτιβοποιημένο."""
        if not BROWSER_SESSION_REQUIRE_HEARTBEAT:
            return
        stale_before = now_ts - float(BROWSER_SESSION_HEARTBEAT_STALE_SECONDS)
        stale_ids = [sid for sid, ts in self.active_sessions.items() if ts < stale_before]
        for sid in stale_ids:
            self.active_sessions.pop(sid, None)

    def active_count(self) -> int:
        """Επιστρέφει τον αριθμό ενεργών browser sessions αφαιρώντας πρώτα τα stale."""
        with self.lock:
            self._cleanup_stale_locked(time.time())
            return len(self.active_sessions)

    def request_shutdown(self, reason: str='') -> bool:
        """Σηματοδοτεί shutdown και εκκινεί αυτόματο τερματισμό του server σε daemon thread, αποφεύγοντας διπλό trigger."""
        with self.lock:
            if self.shutdown_requested:
                return False
            server = self.server_ref
            self.shutdown_requested = True
        if reason:
            log.info('🛑 Αυτόματος τερματισμός εφαρμογής: %s', reason)
        else:
            log.info('🛑 Αυτόματος τερματισμός εφαρμογής: έκλεισε το browser tab/window.')
        if server is None:
            return False

        def _worker() -> None:
            """Καλεί server.shutdown() σε ξεχωριστό thread για να μη μπλοκάρει τον caller."""
            try:
                server.shutdown()
            except Exception as exc:
                log.warning('Αποτυχία αυτόματου shutdown του server: %s', exc)
        threading.Thread(target=_worker, daemon=True, name='browser-auto-shutdown').start()
        return True
BROWSER_MONITOR = BrowserSessionMonitor()

def start_browser_session_watchdog() -> None:
    """Εκκινεί daemon thread που παρακολουθεί τα ενεργά browser sessions και τερματίζει τον server όταν δεν υπάρχει κανένα ανοιχτό."""

    def _worker() -> None:
        """Polling loop που ελέγχει ενεργά sessions και τερματίζει τον server αν δεν υπάρχουν."""
        last_zero_ts: Optional[float] = None
        while True:
            time.sleep(BROWSER_SESSION_WATCHDOG_POLL_SECONDS)
            with BROWSER_MONITOR.lock:
                now_ts = time.time()
                BROWSER_MONITOR._cleanup_stale_locked(now_ts)
                shutdown_requested = BROWSER_MONITOR.shutdown_requested
                ever_seen = BROWSER_MONITOR.ever_seen_session
                active_count = len(BROWSER_MONITOR.active_sessions)
            if shutdown_requested:
                break
            if not ever_seen:
                last_zero_ts = None
                continue
            if active_count > 0:
                last_zero_ts = None
                continue
            if last_zero_ts is None:
                last_zero_ts = now_ts
                continue
            if now_ts - last_zero_ts >= float(BROWSER_SESSION_GRACE_SECONDS):
                if BROWSER_MONITOR.request_shutdown('δεν υπάρχει ανοιχτή σελίδα browser της εφαρμογής'):
                    break
    threading.Thread(target=_worker, daemon=True, name='browser-session-watchdog').start()

@dataclass
class StartupBroadcaster:
    """Μικρός broadcaster για την οθόνη εκκίνησης.

Αποθηκεύει startup events και τα προωθεί στους συνδεδεμένους SSE subscribers ώστε το startup page να βλέπει την πραγματική πρόοδο."""
    _events: List[Dict] = field(default_factory=list)
    _subscribers: List = field(default_factory=list)
    _lock: threading.Lock = field(default_factory=threading.Lock)
    chat_url: str = ''
    is_ready: bool = False

    def emit(self, level: str, msg: str) -> None:
        """Αποθηκεύει startup event και το εκπέμπει σε όλους τους ενεργούς SSE subscribers."""
        event = {'t': time.strftime('%H:%M:%S'), 'level': level, 'msg': msg}
        with self._lock:
            self._events.append(event)
            for q in list(self._subscribers):
                try:
                    q.put_nowait(event)
                except Exception:
                    pass

    def set_ready(self, url: str) -> None:
        """Σηματοδοτεί ότι η εκκίνηση ολοκληρώθηκε και εκπέμπει το τελικό URL της εφαρμογής."""
        self.chat_url = url
        self.is_ready = True
        self.emit('READY', url)

    def subscribe(self) -> '_queue.Queue[Dict]':
        """Δημιουργεί νέα Queue για SSE subscriber, προφορτωμένη με τα ήδη εκπεμφθέντα events."""
        q: '_queue.Queue[Dict]' = _queue.Queue()
        with self._lock:
            for ev in self._events:
                q.put_nowait(ev)
            self._subscribers.append(q)
        return q

    def unsubscribe(self, q: '_queue.Queue[Dict]') -> None:
        """Αφαιρεί queue subscriber από τη λίστα ενεργών συνδρομητών."""
        with self._lock:
            try:
                self._subscribers.remove(q)
            except ValueError:
                pass
STARTUP = StartupBroadcaster()

def slog(level: str, msg: str, *args: object) -> None:
    """Στέλνει log μήνυμα ταυτόχρονα στον Python logger και στον StartupBroadcaster για προβολή στο startup page."""
    formatted = msg % args if args else msg
    if level == 'WARNING':
        log.warning(formatted)
    elif level == 'ERROR':
        log.error(formatted)
    else:
        log.info(formatted)
    STARTUP.emit(level, formatted)

def get_embedded_system_prompt() -> Tuple[str, str]:
    """Επιστρέφει το embedded DEFAULT_SYSTEM_PROMPT μαζί με την πηγή του ('embedded-in-code')."""
    return (DEFAULT_SYSTEM_PROMPT.strip(), 'embedded-in-code')

def find_free_port(host: str, start_port: int=DEFAULT_PORT, end_port: int=8899) -> int:
    """Σαρώνει ακέραια ακολουθία θυρών και επιστρέφει την πρώτη ελεύθερη. Εκτοξεύει RuntimeError αν εξαντληθεί το εύρος."""
    for port in range(start_port, end_port + 1):
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
            try:
                sock.bind((host, port))
                return port
            except OSError:
                continue
    raise RuntimeError(f'Δεν βρέθηκε ελεύθερη θύρα στο εύρος {start_port}–{end_port}.')

def get_saved_ollama_api_key() -> str:
    """Επιστρέφει thread-safe το API key που έχει αποθηκευτεί στις ρυθμίσεις."""
    with APP_CONFIG.lock:
        return str(APP_CONFIG.ollama_api_key or '').strip()

def get_ollama_api_key_source() -> str:
    """Προσδιορίζει από πού προέρχεται το ενεργό API key: 'settings-file', 'environment', 'embedded' ή 'missing'."""
    saved_value = get_saved_ollama_api_key()
    if saved_value:
        return 'settings-file'
    env_value = str(os.environ.get('OLLAMA_API_KEY', '') or '').strip()
    if env_value:
        return 'environment'
    embedded_value = str(EMBEDDED_OLLAMA_API_KEY or '').strip()
    if embedded_value:
        return 'embedded'
    return 'missing'

def get_ollama_api_key() -> str:
    """Επιστρέφει το ενεργό API key με σειρά προτεραιότητας: settings-file → environment → embedded."""
    saved_value = get_saved_ollama_api_key()
    if saved_value:
        return saved_value
    env_value = str(os.environ.get('OLLAMA_API_KEY', '') or '').strip()
    if env_value:
        return env_value
    return str(EMBEDDED_OLLAMA_API_KEY or '').strip()

def is_direct_cloud_api_configured() -> bool:
    """Επιστρέφει True αν υπάρχει διαθέσιμο Ollama Cloud API key."""
    return bool(get_ollama_api_key())

def build_request_headers(url: str, extra_headers: Optional[Dict[str, str]]=None) -> Dict[str, str]:
    """Χτίζει λεξικό HTTP headers για αίτημα Ollama, προσθέτοντας Authorization αν το URL ανήκει στο Cloud API."""
    headers = dict(REQUEST_HEADERS)
    if extra_headers:
        headers.update(extra_headers)
    api_key = get_ollama_api_key()
    if api_key and str(url).startswith(OLLAMA_DIRECT_API_BASE_URL):
        headers['Authorization'] = f'Bearer {api_key}'
    return headers

def fetch_url_text(url: str, timeout: int=12) -> str:
    """Κατεβάζει text περιεχόμενο από URL με τα κατάλληλα request headers."""
    req = urllib.request.Request(url, headers=build_request_headers(url))
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        status = int(getattr(resp, 'status', 200) or 200)
        if status >= 400:
            raise RuntimeError(f'HTTP {status} από {url}')
        return resp.read().decode('utf-8', errors='ignore')

def fetch_url_json(url: str, timeout: int=12) -> Dict:
    """Κατεβάζει και αποκωδικοποιεί JSON payload από URL."""
    req = urllib.request.Request(url, headers=build_request_headers(url))
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        status = int(getattr(resp, 'status', 200) or 200)
        if status >= 400:
            raise RuntimeError(f'HTTP {status} από {url}')
        return json.loads(resp.read().decode('utf-8', errors='ignore'))

def direct_cloud_chat_stream(model: str, messages: List[Dict], *, model_options: Optional[Dict]=None, think_value: Optional[object]=None, timeout: int=900):
    """Ανοίγει streaming POST connection προς το Ollama Cloud Chat API και παράγει JSON chunks ένα-ένα μέσω generator."""
    api_key = get_ollama_api_key()
    if not api_key:
        raise RuntimeError("Λείπει το Ollama Cloud API key. Βάλ'το στο πεδίο API Key του GUI, στο settings αρχείο της εφαρμογής ή ως OLLAMA_API_KEY.")
    payload: Dict[str, object] = {'model': str(model or '').strip(), 'messages': messages, 'stream': True}
    if model_options:
        payload['options'] = dict(model_options)
    if think_value is not None:
        payload['think'] = think_value
    url = f'{OLLAMA_DIRECT_API_BASE_URL}/chat'
    body = json.dumps(payload, ensure_ascii=False).encode('utf-8')
    req = urllib.request.Request(url, data=body, headers=build_request_headers(url, {'Content-Type': 'application/json; charset=utf-8'}), method='POST')
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            status = int(getattr(resp, 'status', 200) or 200)
            if status >= 400:
                raise RuntimeError(f'HTTP {status} από {url}')
            for raw_line in resp:
                line = raw_line.decode('utf-8', errors='ignore').strip()
                if not line:
                    continue
                try:
                    chunk = json.loads(line)
                except json.JSONDecodeError:
                    continue
                if isinstance(chunk, dict) and chunk.get('error'):
                    raise RuntimeError(str(chunk.get('error')))
                yield chunk
    except urllib.error.HTTPError as exc:
        body_text = ''
        try:
            body_text = exc.read().decode('utf-8', errors='ignore').strip()
        except Exception:
            pass
        detail = body_text
        try:
            parsed = json.loads(body_text) if body_text else {}
            if isinstance(parsed, dict) and parsed.get('error'):
                detail = str(parsed.get('error'))
        except Exception:
            pass
        raise RuntimeError(f'HTTP {exc.code}: {detail or exc.reason}') from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f'Δικτυακό σφάλμα προς το Ollama Cloud API: {exc.reason}') from exc

def direct_cloud_chat_complete(model: str, messages: List[Dict], *, model_options: Optional[Dict]=None, think_value: Optional[object]=None, timeout: int=900) -> Dict[str, object]:
    """Εκτελεί πλήρη chat call στο Ollama Cloud API συλλέγοντας όλα τα streaming chunks σε ενιαίο αποτέλεσμα."""
    collected: List[str] = []
    collected_thinking: List[str] = []
    token_stats: Optional[Dict] = None
    for chunk in direct_cloud_chat_stream(model=model, messages=messages, model_options=model_options, think_value=think_value, timeout=timeout):
        thinking_piece = extract_chunk_thinking(chunk)
        if thinking_piece:
            collected_thinking.append(thinking_piece)
        piece = extract_chunk_content(chunk)
        if piece:
            collected.append(piece)
        stats = extract_token_stats(chunk)
        if stats:
            token_stats = stats
    return {'content': ''.join(collected).strip(), 'thinking': ''.join(collected_thinking).strip(), 'token_stats': token_stats}
_ENSEMBLE_ROLE_LABELS: Dict[str, str] = {'vision-analyst': 'Vision helper', 'code-specialist': 'Coding helper', 'code-reviewer': 'Code reviewer', 'reasoning-specialist': 'Reasoning helper', 'cross-checker': 'Cross-checker', 'long-context-reader': 'Long-context helper'}
_ENSEMBLE_HELPER_MAX_SIZE_B_BY_ROLE: Dict[str, float] = {'vision-analyst': 260.0, 'code-specialist': 90.0, 'code-reviewer': 120.0, 'reasoning-specialist': 120.0, 'cross-checker': 90.0, 'long-context-reader': 260.0}
_ENSEMBLE_HELPER_MAX_TOKENS_BY_ROLE: Dict[str, int] = {'vision-analyst': 180, 'code-specialist': 180, 'code-reviewer': 160, 'reasoning-specialist': 160, 'cross-checker': 140, 'long-context-reader': 180}
_ENSEMBLE_HELPER_TIMEOUT_BY_ROLE: Dict[str, int] = {'vision-analyst': 120, 'code-specialist': 90, 'code-reviewer': 90, 'reasoning-specialist': 90, 'cross-checker': 75, 'long-context-reader': 120}

def detect_task_traits(user_text: str, attachments: Optional[List[Dict]]=None) -> Dict[str, bool]:
    """Αναλύει κείμενο και συνημμένα για να ανιχνεύσει code, reasoning, vision, long_context traits που επηρεάζουν την επιλογή μοντέλου."""
    text = str(user_text or '')
    lower = text.lower()
    attachments = attachments or []
    has_image = any((str(item.get('kind') or '').lower() == 'image' for item in attachments))
    has_document = any((str(item.get('kind') or '').lower() == 'document' for item in attachments))
    code_patterns = ('```', 'traceback', 'syntaxerror', 'exception', 'stack trace', 'python', 'javascript', 'java', 'c++', 'c#', 'regex', 'sql', 'function', 'class ', 'def ', 'import ', 'bug', 'fix', 'error', 'κώδικ', 'σφάλμα', 'debug', 'διορθ', 'compile', 'py')
    reasoning_patterns = ('reason', 'thinking', 'analysis', 'analyze', 'why', 'proof', 'derive', 'math', 'equation', 'logic', 'βήμα', 'σκέψ', 'λογικ', 'απόδει', 'λύσ', 'εξήγη', 'ανάλυσ')
    direct_visual_phrases = ('look at this image', 'look at the attached image', 'look at the screenshot', 'describe this image', 'describe the attached image', 'describe the screenshot', 'analyze this image', 'analyze the attached image', 'analyze the screenshot', 'what is in this image', "what's in this image", 'read the text in this image', 'from the screenshot', 'in the screenshot', 'ocr this image', 'ocr the image', 'δες την εικόνα', 'κοίτα την εικόνα', 'κοίτα τη συνημμένη εικόνα', 'περιέγραψε την εικόνα', 'ανάλυσε την εικόνα', 'διάβασε το κείμενο στην εικόνα', 'τι δείχνει η εικόνα', 'τι υπάρχει στην εικόνα', 'στο screenshot', 'στο στιγμιότυπο')
    weak_vision_patterns = ('figure', 'diagram', 'chart', 'graph', 'plot', 'table', 'figure ', 'diagram ', 'chart ', 'graph ', 'διάγρα', 'γράφη', 'πίνακ', 'σχήμα')
    is_code = any((token in lower for token in code_patterns))
    is_reasoning = any((token in lower for token in reasoning_patterns))
    direct_visual_hits = sum((1 for token in direct_visual_phrases if token in lower))
    weak_vision_hits = sum((1 for token in weak_vision_patterns if token in lower))
    explicit_visual_request = has_image
    text_mentions_visual = direct_visual_hits >= 1 or weak_vision_hits >= 1
    weak_visual_hint = text_mentions_visual and (not explicit_visual_request)
    is_vision = has_image
    is_long_context = len(text) >= 3500 or has_document
    return {'code': is_code, 'reasoning': is_reasoning or is_long_context, 'vision': is_vision, 'vision_explicit': explicit_visual_request, 'vision_hint_only': weak_visual_hint, 'long_context': is_long_context, 'has_document': has_document, 'has_image': has_image, 'strong_vision_hits': direct_visual_hits, 'weak_vision_hits': weak_vision_hits}

def _ensemble_preferred_prefixes(primary_model: str, criterion: str, traits: Dict[str, bool]) -> List[str]:
    """Επιστρέφει διατεταγμένη λίστα model-prefix προτιμήσεων για το helper ensemble, ανάλογα με το criterion και το primary μοντέλο."""
    key = canonical_model_key(primary_model)
    prefixes: List[str] = []
    if criterion == 'vision':
        prefixes.extend(['qwen3-vl', 'qwen3.5', 'gemini-3', 'glm-5', 'gemma3'])
    elif criterion == 'coding':
        prefixes.extend(['qwen3-coder', 'qwen3-coder-next', 'devstral-2', 'devstral', 'deepseek-v3.2', 'qwen3.5', 'nemotron-3-nano'])
    elif criterion == 'reasoning':
        prefixes.extend(['qwen3.5', 'deepseek-v3.2', 'deepseek-r1', 'kimi-k2.5', 'glm-5', 'gemini-3', 'nemotron-3-nano'])
    elif criterion == 'context':
        prefixes.extend(['qwen3.5', 'gemini-3', 'kimi-k2.5', 'glm-5', 'minimax-m2.7', 'nemotron-3-nano'])
    else:
        prefixes.extend(['qwen3.5', 'deepseek-v3.2', 'glm-5', 'gemini-3', 'nemotron-3-nano'])
    if model_matches_prefix(key, 'nemotron-3-super'):
        if criterion == 'vision':
            prefixes = ['qwen3-vl', 'qwen3.5', 'gemini-3', *prefixes]
        elif criterion == 'coding':
            prefixes = ['qwen3-coder', 'nemotron-3-nano', 'qwen3.5', *prefixes]
        else:
            prefixes = ['nemotron-3-nano', 'qwen3.5', 'glm-5', *prefixes]
    elif model_matches_prefix(key, 'nemotron-3-nano'):
        if criterion == 'vision':
            prefixes = ['qwen3-vl', 'qwen3.5', 'gemini-3', *prefixes]
        else:
            prefixes = ['nemotron-3-super', 'qwen3.5', 'glm-5', *prefixes]
    elif model_matches_prefix(key, 'qwen3-coder'):
        prefixes = ['qwen3.5', 'deepseek-v3.2', 'glm-5', *prefixes]
    elif model_matches_prefix(key, 'qwen3-vl'):
        prefixes = ['qwen3.5', 'deepseek-v3.2', 'glm-5', *prefixes]
    elif model_matches_prefix(key, 'qwen3.5') and traits.get('code'):
        prefixes = ['qwen3-coder', 'devstral-2', *prefixes]
    elif model_matches_prefix(key, 'qwen3.5') and traits.get('vision'):
        prefixes = ['qwen3-vl', 'gemini-3', *prefixes]
    elif model_matches_prefix(key, 'deepseek'):
        prefixes = ['qwen3.5', 'qwen3-vl', 'gemini-3', *prefixes]
    elif model_matches_prefix(key, 'gemini-3') and traits.get('code'):
        prefixes = ['qwen3-coder', 'deepseek-v3.2', *prefixes]
    deduped: List[str] = []
    seen: Set[str] = set()
    for item in prefixes:
        if item not in seen:
            deduped.append(item)
            seen.add(item)
    return deduped

def _ensemble_prefix_bonus(candidate_model: str, preferred_prefixes: List[str]) -> float:
    """Υπολογίζει bonus score για candidate βάσει θέσης του prefix του στη λίστα preferred_prefixes."""
    for idx, prefix in enumerate(preferred_prefixes):
        if model_matches_prefix(candidate_model, prefix):
            return max(0.0, 1.25 - idx * 0.12)
    return 0.0

def _ensemble_pair_bonus(primary_model: str, candidate_model: str, criterion: str, traits: Dict[str, bool], candidate_caps: Set[str]) -> float:
    """Επιστρέφει bonus/penalty για συγκεκριμένο ζευγάρι primary-candidate μοντέλων, λαμβάνοντας υπόψη specialty και συμβατότητα."""
    primary_key = canonical_model_key(primary_model)
    candidate_key = canonical_model_key(candidate_model)
    bonus = 0.0
    if model_matches_prefix(primary_key, 'nemotron-3-super'):
        if criterion == 'vision':
            if model_matches_prefix(candidate_key, 'qwen3-vl') or 'vision' in candidate_caps:
                bonus += 0.95
            if model_matches_prefix(candidate_key, 'nemotron-3-nano'):
                bonus -= 0.35
        elif criterion == 'coding':
            if model_matches_prefix(candidate_key, 'qwen3-coder'):
                bonus += 0.92
            elif model_matches_prefix(candidate_key, 'nemotron-3-nano'):
                bonus += 0.42
        elif criterion in ('reasoning', 'context', 'overall'):
            if model_matches_prefix(candidate_key, 'nemotron-3-nano'):
                bonus += 1.05
            elif 'reasoning' in candidate_caps:
                bonus += 0.16
    elif model_matches_prefix(primary_key, 'nemotron-3-nano'):
        if criterion == 'vision':
            if model_matches_prefix(candidate_key, 'qwen3-vl') or 'vision' in candidate_caps:
                bonus += 0.82
        elif criterion in ('reasoning', 'context', 'overall') and model_matches_prefix(candidate_key, 'nemotron-3-super'):
            bonus += 0.55
    if not traits.get('has_image') and criterion != 'vision' and model_matches_prefix(candidate_key, 'qwen3-vl'):
        bonus -= 0.45
    if traits.get('has_image') and 'vision' in candidate_caps:
        bonus += 0.22
    return bonus

def choose_auto_ensemble_helper(primary_model: str, user_text: str, attachments: Optional[List[Dict]]=None) -> Optional[Dict[str, object]]:
    """Επιλέγει αυτόματα το βέλτιστο helper μοντέλο για ensemble βάσει task traits και scoring των διαθέσιμων μοντέλων."""
    attachments = attachments or []
    traits = detect_task_traits(user_text, attachments)
    with REGISTRY.lock:
        models = list(REGISTRY.models)
        model_meta = copy.deepcopy(REGISTRY.model_meta)
    if len(models) < 2:
        return None
    primary_meta = model_meta.get(primary_model, {})
    primary_caps = get_model_capabilities(primary_model, primary_meta)
    primary_coding = score_model(primary_model, primary_meta, 'coding')
    primary_reasoning = score_model(primary_model, primary_meta, 'reasoning')
    primary_ctx = get_model_context_tokens(primary_meta)
    primary_key = canonical_model_key(primary_model)
    explicit_vision = bool(traits.get('vision_explicit') or traits.get('has_image'))
    hint_only_vision = bool(traits.get('vision_hint_only'))
    if explicit_vision and 'vision' not in primary_caps:
        criterion = 'vision'
        role = 'vision-analyst'
    elif traits.get('code'):
        if 'coding' not in primary_caps or primary_coding < 9.45:
            criterion = 'coding'
            role = 'code-specialist'
        else:
            criterion = 'reasoning'
            role = 'code-reviewer'
    elif traits.get('long_context') and primary_ctx < 128000:
        criterion = 'context'
        role = 'long-context-reader'
    elif traits.get('reasoning'):
        if 'reasoning' not in primary_caps or primary_reasoning < 9.45:
            criterion = 'reasoning'
            role = 'reasoning-specialist'
        else:
            criterion = 'overall'
            role = 'cross-checker'
    elif 'reasoning' not in primary_caps:
        criterion = 'reasoning'
        role = 'reasoning-specialist'
    elif 'coding' not in primary_caps:
        criterion = 'coding'
        role = 'code-specialist'
    else:
        criterion = 'overall'
        role = 'cross-checker'
    if model_matches_prefix(primary_key, 'nemotron-3-super') and (not explicit_vision) and (criterion == 'vision'):
        criterion = 'reasoning'
        role = 'reasoning-specialist'
    if model_matches_prefix(primary_key, 'nemotron-3-super') and hint_only_vision and (criterion == 'overall'):
        criterion = 'reasoning'
        role = 'reasoning-specialist'
    preferred_prefixes = _ensemble_preferred_prefixes(primary_model, criterion, traits)
    primary_family = primary_key.split(':', 1)[0]
    helper_size_limit_b = float(_ENSEMBLE_HELPER_MAX_SIZE_B_BY_ROLE.get(role, 120.0) or 120.0)
    best: Optional[Tuple[float, str, Dict[str, object]]] = None
    for candidate in models:
        if not candidate or candidate == primary_model:
            continue
        candidate_key = canonical_model_key(candidate)
        if not candidate_key or candidate_key == primary_key:
            continue
        meta = model_meta.get(candidate, {})
        caps = get_model_capabilities(candidate, meta)
        if criterion == 'vision' and 'vision' not in caps:
            continue
        if criterion == 'coding' and ('coding' not in caps and score_model(candidate, meta, 'coding') < 8.8):
            continue
        if criterion == 'reasoning' and ('reasoning' not in caps and score_model(candidate, meta, 'reasoning') < 8.8):
            continue
        if criterion == 'context' and get_model_context_tokens(meta) < 64000:
            continue
        size_b = get_model_size_billions(candidate, meta)
        speed = _size_speed_strength(size_b)
        if size_b > 0 and helper_size_limit_b > 0 and (size_b > helper_size_limit_b * 2.4):
            continue
        fresh = _freshness_strength(get_model_modified_ts(meta))
        score = score_model(candidate, meta, criterion) * 0.64 + score_model(candidate, meta, 'overall') * 0.18 + speed * 0.16 + fresh * 0.02 + _ensemble_prefix_bonus(candidate, preferred_prefixes)
        if criterion in caps:
            score += 0.35
        score += _ensemble_pair_bonus(primary_model, candidate, criterion, traits, caps)
        if candidate_key.split(':', 1)[0] == primary_family:
            if model_matches_prefix(primary_key, 'nemotron-3-super') and model_matches_prefix(candidate_key, 'nemotron-3-nano') and (criterion != 'vision'):
                score += 0.18
            else:
                score -= 0.1
        if size_b > 0 and helper_size_limit_b > 0:
            if size_b <= helper_size_limit_b:
                score += 0.25
            else:
                score -= min(2.4, (size_b - helper_size_limit_b) / max(25.0, helper_size_limit_b) * 1.8)
        payload = {'helper_model': candidate, 'criterion': criterion, 'role': role, 'role_label': _ENSEMBLE_ROLE_LABELS.get(role, role), 'traits': traits, 'preferred_prefixes': preferred_prefixes, 'selection_reason': 'image-attachment' if traits.get('has_image') else 'code-task' if criterion == 'coding' else 'long-context' if criterion == 'context' else 'reasoning' if criterion == 'reasoning' else 'cross-check'}
        if best is None or score > best[0]:
            best = (score, candidate, payload)
    return best[2] if best else None

def choose_manual_ensemble_helper(primary_model: str, helper_model: str, user_text: str, attachments: Optional[List[Dict]]=None) -> Optional[Dict[str, object]]:
    """Επικυρώνει manually επιλεγμένο helper μοντέλο και αναθέτει κατάλληλο ensemble role βάσει capabilities."""
    helper_model = normalize_model_name(helper_model)
    primary_model = normalize_model_name(primary_model)
    if not helper_model or not primary_model or helper_model == primary_model:
        return None
    attachments = attachments or []
    traits = detect_task_traits(user_text, attachments)
    with REGISTRY.lock:
        model_meta = copy.deepcopy(REGISTRY.model_meta)
        available_models = set(REGISTRY.models)
    if available_models and helper_model not in available_models:
        return None
    helper_meta = model_meta.get(helper_model, {})
    helper_caps = get_model_capabilities(helper_model, helper_meta)
    helper_context = get_model_context_tokens(helper_meta)
    if traits.get('has_image') and 'vision' in helper_caps:
        criterion = 'vision'
        role = 'vision-analyst'
    elif traits.get('code') and 'coding' in helper_caps:
        criterion = 'coding'
        role = 'code-specialist'
    elif traits.get('long_context') and helper_context >= 64000:
        criterion = 'context'
        role = 'long-context-reader'
    elif traits.get('reasoning') and ('reasoning' in helper_caps or score_model(helper_model, helper_meta, 'reasoning') >= 8.8):
        criterion = 'reasoning'
        role = 'reasoning-specialist'
    elif 'coding' in helper_caps:
        criterion = 'coding'
        role = 'code-reviewer'
    else:
        criterion = 'overall'
        role = 'cross-checker'
    return {'helper_model': helper_model, 'criterion': criterion, 'role': role, 'role_label': _ENSEMBLE_ROLE_LABELS.get(role, role), 'traits': traits, 'preferred_prefixes': [canonical_model_key(helper_model).split(':', 1)[0]], 'selection_reason': 'manual-selection'}

def build_helper_system_prompt(primary_model: str, helper_model: str, helper_role: str, traits: Dict[str, bool]) -> str:
    """Συνθέτει σύντομο system prompt για το helper μοντέλο που εξηγεί τον ρόλο του στο ensemble."""
    role_label = _ENSEMBLE_ROLE_LABELS.get(helper_role, helper_role)
    task_flags = []
    for name in ('code', 'reasoning', 'vision', 'long_context'):
        if traits.get(name):
            task_flags.append(name)
    flags_text = ', '.join(task_flags) if task_flags else 'general'
    return f'Είσαι το δεύτερο, βοηθητικό μοντέλο ενός app-level ensemble δύο μοντέλων.\nΚύριο μοντέλο: {primary_model}\nΒοηθητικό μοντέλο: {helper_model}\nΡόλος σου: {role_label}\nTask hints: {flags_text}\n\nΔΕΝ απαντάς στον χρήστη. Παράγεις μόνο σύντομη ιδιωτική καθοδήγηση για το κύριο μοντέλο.\nΝα είσαι πρακτικός, ακριβής και πολύ σύντομος (έως 8 bullets συνολικά). Μην βάλεις χαιρετισμούς.\nΑπάντησε ακριβώς με τις ενότητες:\nSUMMARY:\nKEY_POINTS:\nRISKS:\nPLAN:\nΑν υπάρχει κώδικας, πρόσθεσε και BUGS_OR_PATCH:. Αν υπάρχουν εικόνες, πρόσθεσε VISUAL_FINDINGS:.'

def build_main_ensemble_guidance(helper_model: str, helper_role: str, helper_text: str) -> str:
    """Τυλίγει την έξοδο του helper μοντέλου σε private guidance block για ένεση στο context του primary μοντέλου."""
    trimmed = str(helper_text or '').strip()
    if len(trimmed) > 7000:
        trimmed = trimmed[:7000].rstrip() + '\n...[trimmed]'
    role_label = _ENSEMBLE_ROLE_LABELS.get(helper_role, helper_role)
    return f'Ιδιωτική καθοδήγηση από δεύτερο βοηθητικό μοντέλο για εσωτερική χρήση.\nHelper model: {helper_model}\nRole: {role_label}\nΧρησιμοποίησέ την μόνο αν βοηθά και ΜΗΝ αναφέρεις ότι χρησιμοποιήθηκε δεύτερο μοντέλο.\nΑν κάποιο σημείο συγκρούεται με το πραγματικό input ή το ιστορικό συνομιλίας, αγνόησέ το.\n\nPRIVATE_GUIDANCE:\n{trimmed}'

def insert_secondary_system_message(messages: List[Dict], content: str) -> List[Dict]:
    """Εισάγει επιπλέον system message στη λίστα μηνυμάτων, τοποθετώντας το αμέσως μετά το πρώτο system message αν υπάρχει."""
    content = str(content or '').strip()
    if not content:
        return list(messages)
    extra = {'role': 'system', 'content': content}
    if messages and str(messages[0].get('role') or '') == 'system':
        return [messages[0], extra, *messages[1:]]
    return [extra, *messages]

def _is_valid_cloud_tag(tag: str) -> bool:
    """Ελέγχει αν ένα string αποτελεί έγκυρο cloud model tag (π.χ. 'family:tag-cloud') και δεν είναι URL ή κενό."""
    tag = (tag or '').strip()
    if not tag or tag.startswith('http'):
        return False
    tag = tag.split('?', 1)[0].split('#', 1)[0].strip()
    if not tag:
        return False
    if not re.fullmatch('[a-zA-Z0-9._/-]+(?::[a-zA-Z0-9._-]+)?', tag):
        return False
    name_part = tag.split(':', 1)[0].split('/')[-1].strip()
    if not name_part or len(name_part) < 2 or re.fullmatch('\\d+', name_part):
        return False
    return True

def _clean_cloud_tag(raw: str) -> str:
    """Κανονικοποιεί ένα cloud model tag αφαιρώντας query strings, fragments και extra whitespace."""
    raw = (raw or '').strip()
    if not raw:
        return ''
    raw = raw.split('?', 1)[0].split('#', 1)[0].strip()
    if not raw:
        return ''
    if ':' in raw:
        name_part, tag_part = raw.split(':', 1)
        name_part = name_part.split('/')[-1].strip()
        tag_part = tag_part.strip()
        cleaned = f'{name_part}:{tag_part}' if name_part and tag_part else ''
    else:
        cleaned = raw.split('/')[-1].strip()
    if not cleaned or not _is_valid_cloud_tag(cleaned):
        return ''
    return cleaned

def extract_library_families(search_html: str) -> List[str]:
    """Εξάγει λίστα model families που αναφέρονται στη σελίδα αναζήτησης της Ollama library."""
    families: Set[str] = set()
    for m in LIBRARY_LINK_RE.finditer(search_html):
        slug = m.group(1).strip().rstrip('/')
        family = slug.split('/')[-1]
        if family and len(family) > 1 and (not family.startswith('http')):
            families.add(family)
    if not families:
        for m in SEARCH_TEXT_FAMILY_RE.finditer(search_html):
            slug = m.group(1).strip()
            if slug and (not slug.startswith('http')):
                families.add(slug.split('/')[-1])
    return sorted(families)

def extract_cloud_tags_from_html(html_text: str) -> List[str]:
    """Εξάγει cloud model tags (pattern: family:tag-cloud) από raw HTML κείμενο."""
    raw = {m.group(1).strip() for m in CLOUD_TAG_RE.finditer(html_text)}
    cleaned = {_clean_cloud_tag(t) for t in raw}
    return sorted((t for t in cleaned if _is_valid_cloud_tag(t)))

def normalize_html_text(html_text: str) -> str:
    """Κανονικοποιεί ή καθαρίζει δεδομένα στο βήμα «normalize_html_text» ώστε οι επόμενες φάσεις να λαμβάνουν ασφαλή και συνεπή είσοδο.

Βασικά ορίσματα: html_text. Η απομόνωση αυτής της λογικής σε ξεχωριστή ρουτίνα μειώνει την επανάληψη και κάνει πιο εύκολο τον έλεγχο ή τη μελλοντική τροποποίηση."""
    compact = html.unescape(re.sub('<[^>]+>', ' ', html_text or ''))
    compact = compact.replace('•', ' ')
    compact = re.sub('\\s+', ' ', compact)
    return compact.strip()

def parse_context_window_tokens(raw_number: str, suffix: str='') -> Optional[int]:
    """Αναλύει αριθμό context window με προαιρετικό suffix K/M και επιστρέφει αριθμό tokens."""
    token_text = (raw_number or '').strip().replace(',', '')
    if not token_text:
        return None
    try:
        value = float(token_text)
    except ValueError:
        return None
    mult = 1
    suffix = (suffix or '').upper().strip()
    if suffix == 'K':
        mult = 1024
    elif suffix == 'M':
        mult = 1024 * 1024
    elif suffix == 'B':
        mult = 1024 * 1024 * 1024
    tokens = int(value * mult)
    return tokens if tokens >= 256 else None

def extract_cloud_metadata_from_html(html_text: str) -> Dict[str, Dict[str, object]]:
    """Εξάγει metadata cloud μοντέλων (context window, family) από raw HTML σελίδας Ollama."""
    compact = normalize_html_text(html_text)
    meta: Dict[str, Dict[str, object]] = {}
    if not compact:
        return meta
    for match in CLOUD_TAG_RE.finditer(compact):
        tag = _clean_cloud_tag(match.group(1).strip())
        if not _is_valid_cloud_tag(tag):
            continue
        window = compact[match.end():match.end() + 260]
        ctx_match = CONTEXT_WINDOW_RE.search(window)
        entry = meta.setdefault(tag, {})
        if ctx_match:
            ctx_tokens = parse_context_window_tokens(ctx_match.group(1), ctx_match.group(2) or '')
            if ctx_tokens:
                entry['num_ctx_max'] = ctx_tokens
                entry['num_ctx_label'] = ctx_match.group(0).replace('context window', '').strip()
    return meta

def parse_parameter_size_to_billions(raw_value: object) -> float:
    """Αναλύει string μεγέθους παραμέτρων (π.χ. "72B", "14.7B") και επιστρέφει float σε δισεκατομμύρια."""
    text = str(raw_value or '').strip().lower()
    if not text:
        return 0.0
    match = re.search('(\\d+(?:\\.\\d+)?)\\s*([tbm])?', text)
    if not match:
        return 0.0
    value = float(match.group(1))
    suffix = (match.group(2) or 'b').lower()
    if suffix == 't':
        return value * 1000.0
    if suffix == 'm':
        return value / 1000.0
    return value

def parse_iso_datetime_to_timestamp(raw_value: object) -> float:
    """Αναλύει ISO-8601 datetime string και επιστρέφει UNIX timestamp."""
    text = str(raw_value or '').strip()
    if not text:
        return 0.0
    try:
        normalized = text.replace('Z', '+00:00')
        return datetime.datetime.fromisoformat(normalized).timestamp()
    except Exception:
        return 0.0

def infer_model_capabilities_from_name(model_name: str) -> List[str]:
    """Συμπεραίνει capabilities μοντέλου (vision, coding, reasoning) από το όνομά του και τα hardcoded _MODEL_TRAIT_HINTS."""
    name = canonical_model_key(model_name)
    capabilities: Set[str] = {'completion'}
    token_rules = {'vision': ('vision', '-vl', ':vl', 'gemini', 'llava', 'pixtral', 'multimodal', 'omni'), 'coding': ('coder', 'code', 'devstral', 'claude-code', 'swe', 'terminal'), 'reasoning': ('thinking', 'reason', 'reasoning', 'r1', 'gpt-oss', 'deepseek', 'cogito', 'kimi-k2', 'glm-5', 'glm-4.7', 'glm-4.6')}
    for capability, tokens in token_rules.items():
        if any((token in name for token in tokens)):
            capabilities.add(capability)
    for prefix, hinted_caps in _MODEL_TRAIT_HINTS:
        if model_matches_prefix(name, prefix):
            capabilities.update(hinted_caps)
    return sorted(capabilities)

def build_model_meta_from_show_payload(model: str, payload: object) -> Dict[str, object]:
    """Μετατρέπει το raw payload του /api/show endpoint σε κανονικοποιημένο metadata dict με capabilities, context, size κ.α."""
    entry: Dict[str, object] = {}
    if not isinstance(payload, dict):
        return entry
    modified_at = str(payload.get('modified_at', '') or '').strip()
    if modified_at:
        entry['modified_at'] = modified_at
        entry['modified_ts'] = parse_iso_datetime_to_timestamp(modified_at)
    capabilities: Set[str] = set(infer_model_capabilities_from_name(model))
    raw_caps = payload.get('capabilities')
    if isinstance(raw_caps, list):
        for item in raw_caps:
            if isinstance(item, str) and item.strip():
                capabilities.add(item.strip().lower())
    if capabilities:
        entry['capabilities'] = sorted(capabilities)
    details = payload.get('details') if isinstance(payload.get('details'), dict) else {}
    if isinstance(details, dict):
        family = str(details.get('family', '') or '').strip()
        if family:
            entry['family'] = family
        param_size = str(details.get('parameter_size', '') or '').strip()
        if param_size:
            entry['parameter_size'] = param_size
            parsed_b = parse_parameter_size_to_billions(param_size)
            if parsed_b > 0:
                entry['parameter_size_b'] = parsed_b
        families = details.get('families')
        if isinstance(families, list):
            clean_families = [str(x).strip() for x in families if str(x).strip()]
            if clean_families:
                entry['families'] = clean_families
    model_info = payload.get('model_info') if isinstance(payload.get('model_info'), dict) else {}
    if isinstance(model_info, dict):
        for key, value in model_info.items():
            if not isinstance(key, str):
                continue
            lowered = key.lower()
            if lowered.endswith('.context_length') or lowered == 'context_length':
                try:
                    ctx_tokens = int(value)
                except Exception:
                    ctx_tokens = 0
                if ctx_tokens > 0:
                    entry['num_ctx_max'] = ctx_tokens
                    entry['num_ctx_label'] = f'{ctx_tokens:,} tokens'
            elif lowered == 'general.parameter_count':
                try:
                    entry['parameter_count'] = int(value)
                except Exception:
                    pass
            elif lowered.endswith('.mm.tokens_per_image'):
                capabilities.add('vision')
    if capabilities:
        entry['capabilities'] = sorted(capabilities)
    if 'parameter_size_b' not in entry:
        parsed_from_name = parse_parameter_size_to_billions(model)
        if parsed_from_name > 0:
            entry['parameter_size_b'] = parsed_from_name
    entry['details_complete'] = True
    return entry

def fetch_direct_model_details(model: str, timeout: int=15) -> Dict[str, object]:
    """Στέλνει POST στο /api/show για να ανακτήσει λεπτομερές metadata ενός cloud μοντέλου."""
    cleaned_model = normalize_model_name(model)
    if not cleaned_model:
        raise RuntimeError('Δεν δόθηκε όνομα μοντέλου για ανάκτηση metadata.')
    url = f'{OLLAMA_DIRECT_API_BASE_URL}/show'
    body = json.dumps({'model': cleaned_model}, ensure_ascii=False).encode('utf-8')
    req = urllib.request.Request(url, data=body, headers=build_request_headers(url, {'Content-Type': 'application/json; charset=utf-8'}), method='POST')
    with urllib.request.urlopen(req, timeout=timeout) as resp:
        status = int(getattr(resp, 'status', 200) or 200)
        if status >= 400:
            raise RuntimeError(f'HTTP {status} από {url}')
        payload = json.loads(resp.read().decode('utf-8', errors='ignore'))
    return build_model_meta_from_show_payload(cleaned_model, payload)

def get_or_fetch_model_meta(model: str, force: bool=False) -> Dict[str, object]:
    """Επιστρέφει metadata μοντέλου από cache αν είναι πλήρη, αλλιώς τα ανακτά online μέσω /api/show."""
    cleaned_model = normalize_model_name(model)
    if not cleaned_model:
        return {}
    with REGISTRY.lock:
        existing = copy.deepcopy(REGISTRY.model_meta.get(cleaned_model, {}))
    if existing.get('details_complete') and (not force):
        return existing
    try:
        fresh_meta = fetch_direct_model_details(cleaned_model)
    except Exception as exc:
        with REGISTRY.lock:
            entry = REGISTRY.model_meta.setdefault(cleaned_model, {})
            entry['details_error'] = str(exc)
            existing = copy.deepcopy(entry)
        return existing
    with REGISTRY.lock:
        merge_model_meta(REGISTRY.model_meta, {cleaned_model: fresh_meta})
        entry = REGISTRY.model_meta.setdefault(cleaned_model, {})
        entry['details_complete'] = True
        entry.pop('details_error', None)
        if fresh_meta.get('capabilities'):
            entry['capabilities'] = list(fresh_meta.get('capabilities', []))
        if fresh_meta.get('modified_at'):
            entry['modified_at'] = fresh_meta.get('modified_at')
        if fresh_meta.get('modified_ts'):
            entry['modified_ts'] = fresh_meta.get('modified_ts')
        if fresh_meta.get('parameter_size_b'):
            entry['parameter_size_b'] = fresh_meta.get('parameter_size_b')
        if fresh_meta.get('parameter_count'):
            entry['parameter_count'] = fresh_meta.get('parameter_count')
        return copy.deepcopy(entry)

def merge_model_meta(dest: Dict[str, Dict[str, object]], src: Dict[str, Dict[str, object]]) -> None:
    """Συγχωνεύει metadata μοντέλων από src στο dest, διατηρώντας τη μεγαλύτερη τιμή num_ctx_max και δίνοντας προτεραιότητα στα ήδη υπάρχοντα."""
    for tag, info in (src or {}).items():
        if not _is_valid_cloud_tag(tag):
            continue
        entry = dest.setdefault(tag, {})
        if not isinstance(info, dict):
            continue
        for key, value in info.items():
            if value in (None, '', 0):
                continue
            if key == 'num_ctx_max':
                current = int(entry.get('num_ctx_max') or 0)
                incoming = int(value or 0)
                if incoming > current:
                    entry[key] = incoming
            else:
                entry.setdefault(key, value)

def fetch_cloud_models_for_family(family: str, timeout: int=8, family_candidates: Optional[Set[str]]=None) -> Tuple[List[str], Dict[str, Dict[str, object]]]:
    """Κατεβάζει cloud model tags για ένα συγκεκριμένο family από τις σελίδες βιβλιοθήκης του Ollama."""
    tags: Set[str] = set((t for t in family_candidates or set() if ':' in str(t) and 'cloud' in str(t).lower() and _is_valid_cloud_tag(str(t))))
    model_meta: Dict[str, Dict[str, object]] = {}
    family_candidates = set(family_candidates or set())
    saw_cloud_signal = False
    for url in (f'{OLLAMA_LIBRARY_BASE}{family}/tags', f'{OLLAMA_LIBRARY_BASE}{family}'):
        try:
            raw_html = fetch_url_text(url, timeout=timeout)
        except Exception:
            continue
        compact = normalize_html_text(raw_html)
        if CLOUD_WORD_RE.search(compact):
            saw_cloud_signal = True
        explicit_cloud = set(extract_cloud_tags_from_html(raw_html))
        verified_family_models = set(extract_verified_cloud_models_for_family_from_html(raw_html, family))
        tags.update(explicit_cloud)
        tags.update(verified_family_models)
        merge_model_meta(model_meta, extract_cloud_metadata_from_html(raw_html))
        if family_candidates or verified_family_models:
            merge_model_meta(model_meta, extract_context_for_candidate_models_from_html(raw_html, set(family_candidates) | set(verified_family_models) | set(explicit_cloud)))
    if saw_cloud_signal and (not any((tag.startswith(family + ':') for tag in tags))):
        tags.add(family)
    if any((tag.startswith(family + ':') for tag in tags)):
        tags.discard(family)
    for tag in list(tags):
        if _is_valid_cloud_tag(tag):
            model_meta.setdefault(tag, {})['family'] = family
    return (sorted((t for t in tags if _is_valid_cloud_tag(t))), model_meta)

def extract_context_for_candidate_models_from_html(html_text: str, candidates: Set[str]) -> Dict[str, Dict[str, object]]:
    """Εξάγει context window πληροφορία για συγκεκριμένα model tags αναλύοντας raw HTML."""
    compact = normalize_html_text(html_text)
    meta: Dict[str, Dict[str, object]] = {}
    if not compact or not candidates:
        return meta
    for candidate in sorted(candidates, key=len, reverse=True):
        if not _is_valid_cloud_tag(candidate):
            continue
        pattern = re.compile(f'(?<![A-Za-z0-9._/-]){re.escape(candidate)}(?![A-Za-z0-9._/-])')
        match = pattern.search(compact)
        if not match:
            continue
        window = compact[match.end():match.end() + 260]
        ctx_match = CONTEXT_WINDOW_RE.search(window)
        if not ctx_match:
            continue
        ctx_tokens = parse_context_window_tokens(ctx_match.group(1), ctx_match.group(2) or '')
        if not ctx_tokens:
            continue
        entry = meta.setdefault(candidate, {})
        entry['num_ctx_max'] = ctx_tokens
        entry['num_ctx_label'] = ctx_match.group(0).replace('context window', '').strip()
    return meta

def fetch_direct_api_models(timeout: int=8) -> Tuple[List[str], Dict[str, Dict[str, object]]]:
    """Ανακτά τη λίστα διαθέσιμων cloud μοντέλων μέσω του /api/tags endpoint."""
    data = fetch_url_json(OFFICIAL_CLOUD_API_TAGS_URL, timeout=timeout)
    models = data.get('models', []) if isinstance(data, dict) else []
    exact_models: List[str] = []
    meta: Dict[str, Dict[str, object]] = {}
    if isinstance(models, list):
        for item in models:
            if not isinstance(item, dict):
                continue
            raw_name = str(item.get('model') or item.get('name') or '').strip()
            cleaned = _clean_cloud_tag(raw_name)
            if not cleaned or not _is_valid_cloud_tag(cleaned):
                continue
            exact_models.append(cleaned)
            details = item.get('details') if isinstance(item.get('details'), dict) else {}
            entry = meta.setdefault(cleaned, {})
            entry.setdefault('family', cleaned.split(':', 1)[0])
            size_bytes = int(item.get('size') or 0) if str(item.get('size') or '').strip() else 0
            if size_bytes > 0:
                entry.setdefault('size_bytes', size_bytes)
            modified_at = str(item.get('modified_at') or '').strip()
            if modified_at:
                entry.setdefault('modified_at', modified_at)
                entry.setdefault('modified_ts', parse_iso_datetime_to_timestamp(modified_at))
            inferred_caps = infer_model_capabilities_from_name(cleaned)
            if inferred_caps:
                entry.setdefault('capabilities', inferred_caps)
            if isinstance(details, dict):
                if details.get('parameter_size'):
                    param_size_text = str(details.get('parameter_size'))
                    entry.setdefault('parameter_size', param_size_text)
                    parsed_b = parse_parameter_size_to_billions(param_size_text)
                    if parsed_b > 0:
                        entry.setdefault('parameter_size_b', parsed_b)
                if details.get('family'):
                    entry.setdefault('family', str(details.get('family')))
            if 'parameter_size_b' not in entry:
                parsed_from_name = parse_parameter_size_to_billions(cleaned)
                if parsed_from_name > 0:
                    entry.setdefault('parameter_size_b', parsed_from_name)
    exact_models = sorted(set(exact_models))
    return (exact_models, meta)

def extract_verified_cloud_models_for_family_from_html(html_text: str, family: str) -> List[str]:
    """Εξάγει επικυρωμένα cloud model tags για συγκεκριμένο family αναλύοντας HTML κείμενο."""
    compact = normalize_html_text(html_text)
    if not compact or not family:
        return []
    family = family.strip()
    fam_re = re.escape(family)
    pattern = re.compile(f'(?<![A-Za-z0-9._/-])({fam_re}(?::[A-Za-z0-9._-]+)?)(?![A-Za-z0-9._/-])', flags=re.IGNORECASE)
    found: Set[str] = set()
    for match in pattern.finditer(compact):
        candidate = _clean_cloud_tag(match.group(1).strip())
        if not candidate or not _is_valid_cloud_tag(candidate):
            continue
        window = compact[max(0, match.start() - 160):match.end() + 220]
        if CLOUD_WORD_RE.search(window):
            found.add(candidate)
    if any((item.startswith(family + ':') for item in found)):
        found.discard(family)
    return sorted(found)

def fetch_official_cloud_catalog(timeout_per_request: int=8) -> Tuple[List[str], Dict[str, Dict[str, object]]]:
    """Ανακτά την πλήρη λίστα official cloud μοντέλων συνδυάζοντας /api/tags και scraping HTML."""
    import concurrent.futures
    direct_models, direct_meta = fetch_direct_api_models(timeout=max(6, timeout_per_request))
    if not direct_models:
        raise RuntimeError('Δεν βρέθηκαν official direct API models από το Ollama.')
    exact_model_set: Set[str] = set(direct_models)
    all_families: Set[str] = {m.split(':', 1)[0] for m in direct_models if m}
    all_meta: Dict[str, Dict[str, object]] = copy.deepcopy(direct_meta)
    for url in (OFFICIAL_SEARCH_URL, OFFICIAL_GENERAL_SEARCH_URL):
        try:
            raw_html = fetch_url_text(url, timeout=max(8, timeout_per_request + 1))
            all_families.update(extract_library_families(raw_html))
            merge_model_meta(all_meta, extract_context_for_candidate_models_from_html(raw_html, exact_model_set))
        except Exception:
            pass

    def _fetch_family_meta(fam: str) -> Dict[str, Dict[str, object]]:
        """Κατεβάζει metadata για ένα συγκεκριμένο model family από τη βιβλιοθήκη Ollama."""
        try:
            _tags, meta = fetch_cloud_models_for_family(fam, timeout=timeout_per_request, family_candidates={m for m in exact_model_set if m.split(':', 1)[0] == fam})
            return meta
        except Exception:
            return {}
    max_workers = min(14, max(4, len(all_families) or 4))
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        for meta in executor.map(_fetch_family_meta, sorted(all_families)):
            if not isinstance(meta, dict):
                continue
            filtered_meta = {k: v for k, v in meta.items() if k in exact_model_set}
            merge_model_meta(all_meta, filtered_meta)
    final_meta: Dict[str, Dict[str, object]] = {}
    for model in direct_models:
        info = copy.deepcopy(all_meta.get(model, {}))
        info.setdefault('family', model.split(':', 1)[0])
        final_meta[model] = info
    result = sorted(direct_models)
    with_context = sum((1 for model in result if final_meta.get(model, {}).get('num_ctx_max')))
    log.info('☁️  Βρέθηκαν %d official Ollama direct API models (%d με context metadata)', len(result), with_context)
    return (result, final_meta)
_SCORING_CRITERIA: Tuple[str, ...] = ('overall', 'coding', 'reasoning', 'context', 'vision', 'speed', 'newest')
_FAMILY_PRIOR_DEFAULTS: Dict[str, float] = {'overall': 7.5, 'coding': 7.2, 'reasoning': 7.35, 'context': 7.05, 'vision': 6.85, 'speed': 7.1}
_MODEL_FAMILY_PROFILES: List[Tuple[str, Dict[str, float]]] = [('gemini-3-flash', {'overall': 9.12, 'coding': 8.62, 'reasoning': 8.82, 'context': 8.95, 'vision': 9.35, 'speed': 9.95}), ('deepseek-v3.2', {'overall': 9.82, 'coding': 9.62, 'reasoning': 9.9, 'context': 9.14, 'vision': 6.2, 'speed': 4.7}), ('deepseek-v3.1', {'overall': 9.72, 'coding': 9.54, 'reasoning': 9.8, 'context': 9.02, 'vision': 6.0, 'speed': 4.82}), ('deepseek-r1', {'overall': 9.66, 'coding': 9.1, 'reasoning': 9.96, 'context': 8.4, 'vision': 5.25, 'speed': 3.86}), ('qwen3.5', {'overall': 9.96, 'coding': 9.74, 'reasoning': 9.92, 'context': 9.52, 'vision': 9.84, 'speed': 4.42}), ('qwen3-coder-next', {'overall': 9.56, 'coding': 9.96, 'reasoning': 9.42, 'context': 9.0, 'vision': 5.35, 'speed': 5.06}), ('qwen3-coder', {'overall': 9.68, 'coding': 9.98, 'reasoning': 9.6, 'context': 9.08, 'vision': 5.38, 'speed': 4.25}), ('qwen3-vl', {'overall': 9.58, 'coding': 9.26, 'reasoning': 9.42, 'context': 9.02, 'vision': 9.98, 'speed': 4.52}), ('qwen3-next', {'overall': 9.24, 'coding': 9.1, 'reasoning': 9.18, 'context': 8.86, 'vision': 7.92, 'speed': 5.12}), ('kimi-k2-thinking', {'overall': 9.62, 'coding': 9.22, 'reasoning': 9.9, 'context': 9.12, 'vision': 6.62, 'speed': 3.72}), ('kimi-k2.5', {'overall': 9.76, 'coding': 9.4, 'reasoning': 9.78, 'context': 9.28, 'vision': 7.12, 'speed': 4.12}), ('kimi-k2', {'overall': 9.58, 'coding': 9.22, 'reasoning': 9.62, 'context': 9.0, 'vision': 6.85, 'speed': 4.25}), ('glm-5', {'overall': 9.6, 'coding': 9.36, 'reasoning': 9.56, 'context': 9.12, 'vision': 8.72, 'speed': 4.15}), ('glm-4.7', {'overall': 9.46, 'coding': 9.18, 'reasoning': 9.44, 'context': 8.92, 'vision': 8.22, 'speed': 4.02}), ('glm-4.6', {'overall': 9.38, 'coding': 9.12, 'reasoning': 9.36, 'context': 8.78, 'vision': 8.05, 'speed': 4.1}), ('minimax-m2.7', {'overall': 9.42, 'coding': 9.06, 'reasoning': 9.4, 'context': 8.96, 'vision': 8.62, 'speed': 4.42}), ('minimax-m2.5', {'overall': 9.34, 'coding': 8.98, 'reasoning': 9.32, 'context': 8.86, 'vision': 8.46, 'speed': 4.58}), ('minimax-m2.1', {'overall': 9.22, 'coding': 8.92, 'reasoning': 9.2, 'context': 8.72, 'vision': 8.18, 'speed': 4.72}), ('minimax-m2', {'overall': 9.14, 'coding': 8.88, 'reasoning': 9.1, 'context': 8.64, 'vision': 8.0, 'speed': 4.86}), ('nemotron-3-super', {'overall': 9.32, 'coding': 8.92, 'reasoning': 9.28, 'context': 8.96, 'vision': 7.42, 'speed': 4.82}), ('nemotron-3-nano', {'overall': 8.76, 'coding': 8.36, 'reasoning': 8.68, 'context': 8.12, 'vision': 6.2, 'speed': 7.2}), ('mistral-large-3', {'overall': 9.22, 'coding': 9.0, 'reasoning': 9.18, 'context': 8.82, 'vision': 7.82, 'speed': 4.32}), ('devstral-small-2', {'overall': 8.98, 'coding': 9.42, 'reasoning': 8.7, 'context': 8.52, 'vision': 5.02, 'speed': 6.62}), ('devstral-2', {'overall': 9.18, 'coding': 9.74, 'reasoning': 8.96, 'context': 8.86, 'vision': 5.1, 'speed': 5.22}), ('devstral', {'overall': 8.98, 'coding': 9.42, 'reasoning': 8.72, 'context': 8.52, 'vision': 5.05, 'speed': 6.2}), ('gpt-oss', {'overall': 9.06, 'coding': 8.92, 'reasoning': 9.04, 'context': 8.62, 'vision': 5.1, 'speed': 5.9}), ('cogito-2.1', {'overall': 9.28, 'coding': 9.08, 'reasoning': 9.42, 'context': 8.92, 'vision': 6.22, 'speed': 3.92}), ('cogito', {'overall': 9.12, 'coding': 8.96, 'reasoning': 9.26, 'context': 8.76, 'vision': 6.02, 'speed': 4.2}), ('gemini-3', {'overall': 9.74, 'coding': 9.42, 'reasoning': 9.72, 'context': 9.46, 'vision': 9.82, 'speed': 5.12}), ('ministral-3', {'overall': 8.74, 'coding': 8.34, 'reasoning': 8.48, 'context': 8.22, 'vision': 6.22, 'speed': 8.24}), ('ministral', {'overall': 8.62, 'coding': 8.22, 'reasoning': 8.36, 'context': 8.06, 'vision': 6.05, 'speed': 8.0}), ('mistral-small', {'overall': 8.54, 'coding': 8.18, 'reasoning': 8.24, 'context': 7.96, 'vision': 5.92, 'speed': 8.2}), ('gemma3', {'overall': 8.58, 'coding': 8.22, 'reasoning': 8.34, 'context': 7.82, 'vision': 8.12, 'speed': 7.5}), ('rnj-1', {'overall': 8.32, 'coding': 7.96, 'reasoning': 8.16, 'context': 7.92, 'vision': 5.42, 'speed': 8.05}), ('rnj', {'overall': 8.24, 'coding': 7.88, 'reasoning': 8.08, 'context': 7.84, 'vision': 5.32, 'speed': 8.1})]
_MODEL_TRAIT_HINTS: List[Tuple[str, Set[str]]] = [('qwen3.5', {'reasoning', 'coding', 'vision'}), ('qwen3-vl', {'vision', 'reasoning', 'coding'}), ('qwen3-coder', {'coding', 'reasoning'}), ('qwen3-next', {'reasoning', 'coding', 'vision'}), ('deepseek-v3.2', {'reasoning', 'coding'}), ('deepseek-v3.1', {'reasoning', 'coding'}), ('deepseek-r1', {'reasoning'}), ('kimi-k2.5', {'reasoning', 'coding'}), ('kimi-k2', {'reasoning', 'coding'}), ('glm-5', {'reasoning', 'coding', 'vision'}), ('glm-4', {'reasoning', 'coding', 'vision'}), ('gemini-3', {'reasoning', 'coding', 'vision'}), ('devstral', {'coding'}), ('gpt-oss', {'reasoning', 'coding'}), ('nemotron-3-super', {'reasoning', 'coding'}), ('nemotron-3-nano', {'reasoning'}), ('mistral-large-3', {'reasoning', 'coding'}), ('cogito', {'reasoning', 'coding'}), ('ministral-3', {'reasoning', 'coding'}), ('gemma3', {'reasoning', 'coding', 'vision'})]

def canonical_model_key(model_name: str) -> str:
    """Κανονικοποιεί το όνομα μοντέλου σε lowercase χωρίς περιττά spaces για ασφαλή σύγκριση."""
    key = normalize_model_name(str(model_name or '')).strip().lower()
    if not key:
        return ''
    if ':' in key:
        family, tag = key.split(':', 1)
        if tag.endswith('-cloud'):
            tag = tag[:-6]
        key = f'{family}:{tag}'
    elif key.endswith('-cloud'):
        key = key[:-6]
    return key.strip(':')

def model_matches_prefix(model_name: str, prefix: str) -> bool:
    """Επιστρέφει True αν το κανονικοποιημένο όνομα μοντέλου αρχίζει με το δοσμένο prefix."""
    key = canonical_model_key(model_name)
    prefix = str(prefix or '').strip().lower()
    if not key or not prefix:
        return False
    return key.startswith(prefix) or prefix in key

def get_family_profile(model_name: str) -> Dict[str, float]:
    """Επιστρέφει το scoring profile για το family του μοντέλου ή ένα generic profile αν δεν υπάρχει καταχώρηση."""
    for prefix, profile in _MODEL_FAMILY_PROFILES:
        if model_matches_prefix(model_name, prefix):
            return profile
    return _FAMILY_PRIOR_DEFAULTS

def get_model_capabilities(model_name: str, meta: Optional[Dict[str, object]]=None) -> Set[str]:
    """Επιστρέφει σύνολο capabilities μοντέλου συνδυάζοντας inferred-from-name και αποθηκευμένα metadata."""
    capabilities: Set[str] = set(infer_model_capabilities_from_name(model_name))
    if isinstance(meta, dict):
        raw_caps = meta.get('capabilities')
        if isinstance(raw_caps, list):
            for item in raw_caps:
                if isinstance(item, str) and item.strip():
                    capabilities.add(item.strip().lower())
        if isinstance(meta.get('families'), list):
            for item in meta.get('families', []):
                if isinstance(item, str):
                    capabilities.update(infer_model_capabilities_from_name(item))
        family = str(meta.get('family') or '').strip()
        if family:
            capabilities.update(infer_model_capabilities_from_name(family))
    capabilities.add('completion')
    return capabilities

def get_model_size_billions(model_name: str, meta: Optional[Dict[str, object]]=None) -> float:
    """Επιστρέφει μέγεθος μοντέλου σε δισεκατομμύρια παραμέτρους από metadata ή από το όνομά του."""
    if isinstance(meta, dict):
        raw_b = meta.get('parameter_size_b')
        try:
            size_b = float(raw_b)
        except Exception:
            size_b = 0.0
        if size_b > 0:
            return max(0.0, min(size_b, 1000.0))
        try:
            size_bytes = float(meta.get('size_bytes') or 0)
        except Exception:
            size_bytes = 0.0
        if size_bytes > 0:
            return max(0.0, min(size_bytes / 1000000000.0, 1000.0))
    parsed = parse_parameter_size_to_billions(canonical_model_key(model_name))
    return max(0.0, min(parsed, 1000.0)) if parsed > 0 else 0.0

def get_model_context_tokens(meta: Optional[Dict[str, object]]=None) -> int:
    """Επιστρέφει μέγιστο context window σε tokens από metadata, ή 0 αν δεν είναι γνωστό."""
    if isinstance(meta, dict):
        try:
            ctx = int(meta.get('num_ctx_max') or 0)
        except Exception:
            ctx = 0
        if ctx > 0:
            return ctx
    return 0

def get_model_modified_ts(meta: Optional[Dict[str, object]]=None) -> float:
    """Επιστρέφει UNIX timestamp τελευταίας τροποποίησης μοντέλου από metadata, ή 0.0 αν δεν είναι διαθέσιμο."""
    if isinstance(meta, dict):
        try:
            raw_ts = float(meta.get('modified_ts') or 0)
        except Exception:
            raw_ts = 0.0
        if raw_ts > 0:
            return raw_ts
        raw_date = str(meta.get('modified_at') or '').strip()
        if raw_date:
            return parse_iso_datetime_to_timestamp(raw_date)
    return 0.0

def _clamp(value: float, low: float, high: float) -> float:
    """Επιστρέφει τιμή clipped στο εύρος [min_val, max_val]."""
    return max(low, min(value, high))

def _size_quality_strength(size_b: float) -> float:
    """Μετατρέπει μέγεθος μοντέλου (B παραμέτρων) σε normalized quality score."""
    import math
    if size_b <= 0:
        return 4.8
    normalized = math.log2(min(size_b, 1000.0) + 1.0) / math.log2(1001.0)
    return 3.8 + normalized * 6.2

def _size_speed_strength(size_b: float) -> float:
    """Μετατρέπει μέγεθος μοντέλου (B παραμέτρων) σε normalized speed score (αντίστροφη σχέση)."""
    import math
    if size_b <= 0:
        return 7.8
    normalized = math.log2(min(size_b, 1000.0) + 1.0) / math.log2(1001.0)
    return 9.8 - normalized * 7.2

def _context_strength(ctx_tokens: int) -> float:
    """Μετατρέπει context window (tokens) σε normalized strength score."""
    import math
    if ctx_tokens <= 0:
        return 3.2
    normalized = _clamp(math.log2(float(ctx_tokens)) / 18.0, 0.0, 1.08)
    bonus = 0.7 if ctx_tokens >= 200000 else 0.35 if ctx_tokens >= 128000 else 0.0
    return min(10.0, 3.6 + normalized * 6.0 + bonus)

def _freshness_strength(modified_ts: float) -> float:
    """Μετατρέπει UNIX timestamp τελευταίας τροποποίησης σε score που εκφράζει την ενημερότητα του μοντέλου."""
    if modified_ts <= 0:
        return 4.8
    age_days = max(0.0, (time.time() - modified_ts) / 86400.0)
    if age_days <= 21:
        return 10.0
    if age_days <= 45:
        return 9.4
    if age_days <= 90:
        return 8.6
    if age_days <= 180:
        return 7.6
    if age_days <= 365:
        return 6.3
    if age_days <= 540:
        return 5.2
    return 4.3

def _name_signal_bonus(model_name: str, criterion: str) -> float:
    """Υπολογίζει score bonus βάσει tokens στο όνομα μοντέλου που υποδηλώνουν εξειδίκευση (coder, thinking, -vl κ.α.)."""
    name = canonical_model_key(model_name)
    bonus = 0.0
    if criterion == 'coding':
        if any((token in name for token in ('coder', 'devstral', 'terminal', 'swe'))):
            bonus += 0.9
        if any((token in name for token in ('code', 'oss'))):
            bonus += 0.25
    elif criterion == 'reasoning':
        if any((token in name for token in ('thinking', 'reason', 'reasoning', 'r1'))):
            bonus += 0.95
        if any((token in name for token in ('deepseek', 'cogito'))):
            bonus += 0.2
    elif criterion == 'vision':
        if any((token in name for token in ('-vl', ':vl', 'vision', 'gemini', 'pixtral', 'llava'))):
            bonus += 0.95
    elif criterion == 'speed':
        if any((token in name for token in ('flash', 'nano', 'mini', 'small'))):
            bonus += 1.15
        if any((token in name for token in ('preview',))):
            bonus += 0.2
    elif criterion == 'overall':
        if any((token in name for token in ('thinking', 'coder', '-vl', ':vl', 'vision'))):
            bonus += 0.18
    return bonus

def score_model(model_name: str, meta: Optional[Dict[str, object]]=None, criterion: str='overall') -> float:
    """Βαθμολογεί ένα μοντέλο για συγκεκριμένο criterion συνδυάζοντας family prior, size, context, freshness, capabilities και name bonus."""
    criterion = (criterion or 'overall').strip().lower()
    if criterion not in _SCORING_CRITERIA:
        criterion = 'overall'
    profile = get_family_profile(model_name)
    base_prior = float(profile.get(criterion, _FAMILY_PRIOR_DEFAULTS.get(criterion, 7.0)))
    size_b = get_model_size_billions(model_name, meta)
    ctx = get_model_context_tokens(meta)
    modified_ts = get_model_modified_ts(meta)
    caps = get_model_capabilities(model_name, meta)
    size_quality = _size_quality_strength(size_b)
    size_speed = _size_speed_strength(size_b)
    context_strength = _context_strength(ctx)
    freshness = _freshness_strength(modified_ts)
    has_reasoning = 10.0 if 'reasoning' in caps else 0.0
    has_coding = 10.0 if 'coding' in caps else 0.0
    has_vision = 10.0 if 'vision' in caps else 0.0
    bonus = _name_signal_bonus(model_name, criterion)
    if criterion == 'coding':
        return base_prior * 0.56 + size_quality * 0.1 + context_strength * 0.1 + freshness * 0.05 + has_coding * 0.14 + has_reasoning * 0.04 + bonus
    if criterion == 'reasoning':
        return base_prior * 0.56 + size_quality * 0.11 + context_strength * 0.06 + freshness * 0.04 + has_reasoning * 0.13 + has_coding * 0.03 + bonus
    if criterion == 'context':
        if ctx > 0:
            return context_strength * 0.74 + base_prior * 0.14 + size_quality * 0.08 + freshness * 0.04
        return base_prior * 0.22 + size_quality * 0.12 + freshness * 0.06
    if criterion == 'vision':
        return base_prior * 0.58 + size_quality * 0.09 + context_strength * 0.06 + freshness * 0.04 + has_vision * 0.18 + has_reasoning * 0.02 + bonus
    if criterion == 'speed':
        return base_prior * 0.15 + size_speed * 0.6 + freshness * 0.1 + context_strength * 0.05 + (0.2 if size_b and size_b <= 24.0 else 0.0) + bonus
    if criterion == 'newest':
        return modified_ts if modified_ts > 0 else 0.0
    return base_prior * 0.56 + size_quality * 0.12 + context_strength * 0.06 + freshness * 0.05 + has_reasoning * 0.08 + has_coding * 0.06 + has_vision * 0.04 + bonus

def recommend_best_model(models: List[str], model_meta: Optional[Dict[str, Dict[str, object]]]=None, criterion: str='overall') -> str:
    """Επιστρέφει το μοντέλο με τη μεγαλύτερη βαθμολογία για το ζητούμενο criterion."""
    if not models:
        return ''
    scored = sorted(models, key=lambda model: score_model(model, (model_meta or {}).get(model, {}), criterion), reverse=True)
    return scored[0]

def wait_for_model_refresh(timeout: float=45.0, poll_interval: float=0.15) -> bool:
    """Μπλοκάρει polling μέχρι να ολοκληρωθεί refresh που τρέχει ήδη ή να λήξει το timeout."""
    deadline = time.time() + max(0.5, timeout)
    while time.time() < deadline:
        with REGISTRY.lock:
            in_progress = REGISTRY.refresh_in_progress
        if not in_progress:
            return True
        time.sleep(max(0.05, poll_interval))
    return False

def refresh_models(force: bool=False, wait_if_running: bool=True) -> None:
    """Ανακτά online τη λίστα cloud μοντέλων, αποθηκεύει cache στον δίσκο και ενημερώνει το global REGISTRY."""
    should_wait = False
    with REGISTRY.lock:
        if REGISTRY.refresh_in_progress:
            should_wait = True
        elif not force and REGISTRY.last_refresh_ts:
            if time.time() - REGISTRY.last_refresh_ts < MODEL_CACHE_SECONDS:
                return
        if not should_wait:
            REGISTRY.refresh_in_progress = True
    if should_wait:
        if wait_if_running:
            wait_for_model_refresh(timeout=60.0)
        return
    try:
        online_models, online_meta = fetch_official_cloud_catalog()
        refresh_ts = time.time()
        save_model_registry_cache_to_disk(online_models, online_meta, refresh_ts)
        with REGISTRY.lock:
            REGISTRY.models = list(online_models)
            REGISTRY.model_meta = copy.deepcopy(online_meta)
            REGISTRY.source = 'official-online'
            REGISTRY.last_error = ''
            REGISTRY.last_refresh_ts = refresh_ts
            REGISTRY.recommended_model = recommend_best_model(online_models, online_meta, 'overall')
    except Exception as exc:
        friendly_error = normalize_model_registry_refresh_error(exc)
        cached_models, cached_meta, cached_ts = load_model_registry_cache_from_disk()
        with REGISTRY.lock:
            if cached_models:
                REGISTRY.models = list(cached_models)
                REGISTRY.model_meta = copy.deepcopy(cached_meta)
                REGISTRY.source = 'stale-disk-cache'
                REGISTRY.last_error = friendly_error
                REGISTRY.last_refresh_ts = cached_ts or time.time()
                REGISTRY.recommended_model = recommend_best_model(cached_models, cached_meta, 'overall')
            else:
                REGISTRY.source = 'stale-online-cache' if REGISTRY.models else 'error'
                REGISTRY.last_error = friendly_error
                REGISTRY.last_refresh_ts = time.time()
                REGISTRY.recommended_model = recommend_best_model(REGISTRY.models, REGISTRY.model_meta, 'overall')
    finally:
        with REGISTRY.lock:
            REGISTRY.refresh_in_progress = False

def refresh_models_in_background(force: bool=False) -> None:
    """Εκκινεί μη-μπλοκαριστικό refresh μοντέλων σε background daemon thread."""

    def _runner() -> None:
        """Καλεί refresh_models χωρίς αναμονή για ήδη τρέχον refresh."""
        refresh_models(force=force, wait_if_running=False)
    threading.Thread(target=_runner, daemon=True, name='model-refresh-background').start()

def validate_python_code_block(code_text: str) -> Tuple[bool, str]:
    """Ελέγχει αν Python κώδικας parses χωρίς SyntaxError. Επιστρέφει (True, "") ή (False, περιγραφή σφάλματος)."""
    try:
        ast.parse(code_text, filename='<python_block>')
        return (True, '')
    except SyntaxError as exc:
        line_no = getattr(exc, 'lineno', None) or 0
        offset = getattr(exc, 'offset', None) or 0
        bad_line = ''
        lines = code_text.splitlines()
        if 1 <= line_no <= len(lines):
            bad_line = lines[line_no - 1]
        pointer = ' ' * max(offset - 1, 0) + '^' if offset else ''
        details = [f'SyntaxError στη γραμμή {line_no}: {exc.msg}']
        if bad_line:
            details.append(bad_line)
        if pointer:
            details.append(pointer)
        return (False, '\n'.join(details))
    except Exception as exc:
        return (False, f'Αποτυχία ελέγχου Python block: {exc}')

def resolve_python_for_generated_scripts() -> Tuple[Optional[List[str]], str]:
    """Εντοπίζει κατάλληλο Python interpreter για εκτέλεση παραγόμενων scripts, ειδικά σε frozen/packaged εκτελέσιμα."""
    if not getattr(sys, 'frozen', False):
        exe = os.path.normpath(sys.executable or 'python')
        return ([exe], 'sys.executable')
    candidates: List[Tuple[List[str], str]] = [(['py', '-3'], 'Python Launcher (py -3)'), (['python'], 'system python'), (['python3'], 'system python3')]
    creationflags = getattr(subprocess, 'CREATE_NO_WINDOW', 0) if os.name == 'nt' else 0
    for command, label in candidates:
        try:
            test = subprocess.run(command + ['-c', 'import sys; print(sys.executable)'], capture_output=True, text=True, timeout=8, creationflags=creationflags)
            if test.returncode == 0:
                return (command, label)
        except Exception:
            continue
    return (None, 'missing')

def launch_python_code_in_terminal(code_text: str, suggested_filename: str='') -> Tuple[bool, str]:
    """Αποθηκεύει Python block σε temp αρχείο και το ανοίγει σε νέο terminal window (Windows/macOS/Linux)."""
    code_text = str(code_text or '')
    if not code_text.strip():
        return (False, 'Το Python block είναι κενό.')
    if len(code_text) > 250000:
        return (False, 'Το Python block είναι υπερβολικά μεγάλο για εκτέλεση.')
    is_valid, validation_message = validate_python_code_block(code_text)
    if not is_valid:
        return (False, validation_message)
    exec_root_dir = os.path.join(tempfile.gettempdir(), 'ollama_cloud_chat_exec')
    os.makedirs(exec_root_dir, exist_ok=True)
    session_dir = tempfile.mkdtemp(prefix='run_', dir=exec_root_dir)
    requested_name = str(suggested_filename or '').strip() or suggest_python_filename(code_text)
    safe_name = sanitize_filename(requested_name)
    if not safe_name.lower().endswith('.py'):
        safe_name += '.py'
    script_name = safe_name
    script_path = os.path.join(session_dir, script_name)
    with open(script_path, 'w', encoding='utf-8', newline='\n') as f:
        f.write(code_text)
    python_cmd, python_source = resolve_python_for_generated_scripts()
    if not python_cmd:
        return (False, 'Δεν βρέθηκε εγκατεστημένος Python interpreter για το >Run.\nΣτο packaged .exe το sys.executable δείχνει το ίδιο το app και όχι python.exe.\nΕγκατάστησε Python ή χρησιμοποίησε πρώτα το 💾 Save και μετά τρέξε το .py χειροκίνητα.')
    launcher_stem = Path(safe_name).stem or 'generated_block'
    try:
        if os.name == 'nt':
            launcher_path = os.path.join(session_dir, f'run_{launcher_stem}.bat')
            command_line = subprocess.list2cmdline(python_cmd + [script_path])
            launcher_lines = ['@echo off', 'setlocal', 'chcp 65001>nul', f'title Python Code Block Runner - {safe_name}', command_line, 'echo.', 'echo Exit code: %ERRORLEVEL%', 'pause']
            with open(launcher_path, 'w', encoding='utf-8', newline='\r\n') as f:
                f.write('\r\n'.join(launcher_lines) + '\r\n')
            subprocess.Popen(['cmd.exe', '/k', launcher_path], creationflags=getattr(subprocess, 'CREATE_NEW_CONSOLE', 0))
            return (True, f'Το Python block αποθηκεύτηκε ως {safe_name} και εκτελείται σε νέο terminal με interpreter από {python_source}.')
        elif sys.platform == 'darwin':
            shell_cmd = ' '.join((shlex.quote(part) for part in python_cmd + [script_path]))
            shell_cmd += '; echo; echo Press Enter to close...; read'
            apple_script = f'tell application "Terminal" to do script {json.dumps(shell_cmd)}'
            subprocess.Popen(['osascript', '-e', apple_script])
        else:
            shell_cmd = ' '.join((shlex.quote(part) for part in python_cmd + [script_path]))
            shell_cmd += "; printf '\n\nPress Enter to close...'; read _"
            launched = False
            terminal_commands = [['x-terminal-emulator', '-e', 'bash', '-lc', shell_cmd], ['gnome-terminal', '--', 'bash', '-lc', shell_cmd], ['konsole', '-e', 'bash', '-lc', shell_cmd], ['xfce4-terminal', '-e', f'bash -lc {shlex.quote(shell_cmd)}']]
            for cmd in terminal_commands:
                try:
                    subprocess.Popen(cmd)
                    launched = True
                    break
                except FileNotFoundError:
                    continue
            if not launched:
                subprocess.Popen(python_cmd + [script_path])
    except Exception as exc:
        return (False, f'Αποτυχία ανοίγματος terminal: {exc}')
    return (True, f'Το Python block αποθηκεύτηκε ως {safe_name} και εκτελείται σε νέο terminal με interpreter από {python_source}.')

def _send_security_headers(handler: BaseHTTPRequestHandler) -> None:
    """Γράφει τα SECURITY_HEADERS στο HTTP response μέσω του handler."""
    for key, value in SECURITY_HEADERS.items():
        handler.send_header(key, value)

def json_response(handler: BaseHTTPRequestHandler, payload: Dict, status: int=200) -> None:
    """Στέλνει πλήρες JSON HTTP response με Content-Type, Content-Length και security headers."""
    data = json.dumps(payload, ensure_ascii=False).encode('utf-8')
    handler.send_response(status)
    handler.send_header('Content-Type', 'application/json; charset=utf-8')
    handler.send_header('Content-Length', str(len(data)))
    _send_security_headers(handler)
    handler.end_headers()
    handler.wfile.write(data)

def stream_json_line(handler: BaseHTTPRequestHandler, payload: Dict) -> None:
    """Γράφει ένα JSON-lines chunk στο streaming response. Εκτοξεύει BrokenPipeError αν ο client αποσυνδεθεί."""
    data = (json.dumps(payload, ensure_ascii=False) + '\n').encode('utf-8')
    try:
        handler.wfile.write(data)
        handler.wfile.flush()
    except (BrokenPipeError, ConnectionResetError, ConnectionAbortedError, OSError):
        raise BrokenPipeError('Client disconnected during stream')

def is_client_disconnect_error(exc: BaseException) -> bool:
    """Επιστρέφει True αν η εξαίρεση αντιστοιχεί σε disconnected client (broken pipe, reset, Windows errno)."""
    if isinstance(exc, (BrokenPipeError, ConnectionResetError, ConnectionAbortedError, TimeoutError)):
        return True
    if isinstance(exc, OSError):
        winerror = getattr(exc, 'winerror', None)
        errno_ = getattr(exc, 'errno', None)
        return winerror in {10053, 10054} or errno_ in {32, 54, 104, 107, 108}
    return False


def sanitize_download_filename(filename: str, fallback: str='assistant-response.pdf') -> str:
    """Καθαρίζει όνομα αρχείου λήψης ώστε να είναι ασφαλές για Content-Disposition και τοπική αποθήκευση."""
    raw_name = Path(str(filename or fallback)).name.strip().replace('\x00', '')
    if not raw_name:
        raw_name = fallback
    safe_name = re.sub(r'[^A-Za-z0-9._() \-]+', '-', raw_name)
    safe_name = re.sub(r'\s+', ' ', safe_name).strip(' .')
    if not safe_name:
        safe_name = fallback
    if not safe_name.lower().endswith('.pdf'):
        safe_name += '.pdf'
    return safe_name


def build_content_disposition_header(disposition: str, filename: str, fallback: str='download.bin') -> str:
    """Επιστρέφει ASCII-safe Content-Disposition header με UTF-8 filename* για συμβατότητα με unicode ονόματα."""
    raw_name = Path(str(filename or fallback)).name.strip().replace('\x00', '')
    if not raw_name:
        raw_name = fallback
    ascii_name = re.sub(r'[^A-Za-z0-9._() \-]+', '_', raw_name)
    ascii_name = re.sub(r'\s+', ' ', ascii_name).strip(' .')
    if not ascii_name:
        ascii_name = Path(fallback).name or 'download.bin'
    quoted_name = urllib.parse.quote(raw_name, safe='')
    return f"{disposition}; filename=\"{ascii_name}\"; filename*=UTF-8''{quoted_name}"

def _iter_headless_browser_candidates() -> List[Path]:
    """Συγκεντρώνει πιθανά εκτελέσιμα browsers που υποστηρίζουν headless print-to-pdf."""
    candidates: List[Path] = []
    which_names = [
        'msedge', 'msedge.exe',
        'microsoft-edge', 'microsoft-edge-stable',
        'chrome', 'chrome.exe',
        'google-chrome', 'google-chrome-stable',
        'chromium', 'chromium-browser',
        'brave', 'brave.exe', 'brave-browser',
    ]
    for name in which_names:
        resolved = shutil.which(name)
        if resolved:
            candidates.append(Path(resolved))

    if os.name == 'nt':
        program_files = os.environ.get('PROGRAMFILES', r'C:\Program Files')
        program_files_x86 = os.environ.get('PROGRAMFILES(X86)', r'C:\Program Files (x86)')
        local_app_data = os.environ.get('LOCALAPPDATA', '')
        windows_candidates = [
            Path(program_files) / 'Microsoft' / 'Edge' / 'Application' / 'msedge.exe',
            Path(program_files_x86) / 'Microsoft' / 'Edge' / 'Application' / 'msedge.exe',
            Path(program_files) / 'Google' / 'Chrome' / 'Application' / 'chrome.exe',
            Path(program_files_x86) / 'Google' / 'Chrome' / 'Application' / 'chrome.exe',
            Path(program_files) / 'BraveSoftware' / 'Brave-Browser' / 'Application' / 'brave.exe',
            Path(program_files_x86) / 'BraveSoftware' / 'Brave-Browser' / 'Application' / 'brave.exe',
        ]
        if local_app_data:
            windows_candidates.extend([
                Path(local_app_data) / 'Microsoft' / 'Edge' / 'Application' / 'msedge.exe',
                Path(local_app_data) / 'Google' / 'Chrome' / 'Application' / 'chrome.exe',
                Path(local_app_data) / 'BraveSoftware' / 'Brave-Browser' / 'Application' / 'brave.exe',
            ])
        candidates.extend(windows_candidates)
    elif sys.platform == 'darwin':
        candidates.extend([
            Path('/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge'),
            Path('/Applications/Google Chrome.app/Contents/MacOS/Google Chrome'),
            Path('/Applications/Chromium.app/Contents/MacOS/Chromium'),
            Path('/Applications/Brave Browser.app/Contents/MacOS/Brave Browser'),
        ])
    else:
        candidates.extend([
            Path('/usr/bin/microsoft-edge'),
            Path('/usr/bin/google-chrome'),
            Path('/usr/bin/chromium'),
            Path('/usr/bin/chromium-browser'),
            Path('/snap/bin/chromium'),
            Path('/usr/bin/brave-browser'),
        ])

    unique_paths: List[Path] = []
    seen: Set[str] = set()
    for candidate in candidates:
        try:
            normalized = str(candidate.expanduser().resolve(strict=False))
        except Exception:
            normalized = str(candidate)
        if normalized in seen:
            continue
        seen.add(normalized)
        unique_paths.append(candidate)
    return unique_paths

def _find_headless_pdf_browser() -> Optional[Path]:
    """Επιστρέφει τον πρώτο διαθέσιμο Chromium-based browser για server-side PDF export."""
    for candidate in _iter_headless_browser_candidates():
        try:
            if candidate.exists() and candidate.is_file():
                return candidate
        except OSError:
            continue
    return None

def _extract_primary_style_block(html_doc: str) -> str:
    """Εξάγει το κύριο <style> block από το index HTML ώστε να επαναχρησιμοποιηθεί στο printable export."""
    if not html_doc:
        return ''
    match = re.search(r'<style>\s*(.*?)\s*</style>', html_doc, flags=re.DOTALL | re.IGNORECASE)
    return match.group(1).strip() if match else ''

def _sanitize_mathjax_svg_cache_fragment(fragment: str) -> str:
    """Κρατά μόνο ασφαλές inline SVG cache markup της MathJax για να μη χάνονται glyphs στο PDF."""
    text = str(fragment or '').strip()
    if not text:
        return ''
    lower = text.lower()
    if '<script' in lower or '<iframe' in lower or '<object' in lower or '<embed' in lower:
        return ''
    if '<svg' not in lower or 'mjx' not in lower:
        return ''
    return text


def _polish_exported_pdf_with_pypdf(pdf_path: Path, document_title: str='Assistant response') -> bool:
    """Fallback polish με pypdf όταν δεν υπάρχει διαθέσιμο PyMuPDF: κόβει header/footer μέσω crop boxes και γράφει metadata."""
    try:
        from pypdf import PdfReader, PdfWriter
    except Exception:
        return False

    pdf_path = Path(pdf_path)
    if not pdf_path.exists() or pdf_path.stat().st_size <= 0:
        return False

    temp_output = pdf_path.with_name(pdf_path.stem + '_cropped.pdf')
    try:
        reader = PdfReader(str(pdf_path))
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)

        writer.add_metadata({
            '/Title': str(document_title or 'Assistant response'),
            '/Author': '',
            '/Creator': APP_TITLE,
            '/Producer': APP_TITLE,
            '/Subject': 'Assistant response export',
            '/Keywords': 'assistant,pdf,export,math,svg,markdown',
        })
        with temp_output.open('wb') as fh:
            writer.write(fh)
    except Exception:
        return False

    if temp_output.exists() and temp_output.stat().st_size > 1024:
        temp_output.replace(pdf_path)
        return True
    return False


def _polish_exported_pdf(pdf_path: Path, document_title: str='Assistant response') -> None:
    """Κάνει τελικό polish pass στο export PDF: αφαιρεί browser headers/footers, γράφει metadata και έχει pypdf fallback."""
    pdf_path = Path(pdf_path)
    if not pdf_path.exists() or pdf_path.stat().st_size <= 0:
        return

    polished_with_fitz = False
    try:
        import fitz  # PyMuPDF
    except Exception:
        fitz = None

    if fitz is not None:
        temp_output = pdf_path.with_name(pdf_path.stem + '_polished.pdf')
        doc = fitz.open(pdf_path)
        try:
            header_footer_needles = [
                'file:///',
                'assistant_export.html',
                str(document_title or '').strip(),
                '.pdf',
                'microsoft edge',
                'chrome',
            ]
            for page in doc:
                rect = page.rect
                top_h = min(56.0, max(38.0, rect.height * 0.060))
                bottom_h = min(44.0, max(30.0, rect.height * 0.048))
                redact_rects = [
                    fitz.Rect(rect.x0, rect.y0, rect.x1, rect.y0 + top_h),
                    fitz.Rect(rect.x0, rect.y1 - bottom_h, rect.x1, rect.y1),
                ]

                text_page = page.get_text('dict')
                for block in text_page.get('blocks', []):
                    if block.get('type') != 0:
                        continue
                    x0, y0, x1, y1 = block.get('bbox', [0, 0, 0, 0])
                    block_rect = fitz.Rect(x0, y0, x1, y1)
                    if block_rect.y1 <= rect.y0 + top_h + 10 or block_rect.y0 >= rect.y1 - bottom_h - 10:
                        text = []
                        for line in block.get('lines', []):
                            for span in line.get('spans', []):
                                text.append(str(span.get('text', '') or ''))
                        combined = ' '.join(text).strip().lower()
                        if combined:
                            if any(needle.lower() in combined for needle in header_footer_needles if needle):
                                redact_rects.append(block_rect + (-4, -3, 4, 3))
                            elif re.search(r'\b\d{1,2}/\d{1,2}/\d{2,4}\b', combined) or re.search(r'\b\d{1,2}:\d{2}\b', combined) or re.search(r'\b\d+\s*/\s*\d+\b', combined):
                                redact_rects.append(block_rect + (-4, -3, 4, 3))

                for redact_rect in redact_rects:
                    page.add_redact_annot(redact_rect, fill=(1, 1, 1))
                page.apply_redactions(
                    images=fitz.PDF_REDACT_IMAGE_NONE,
                    graphics=fitz.PDF_REDACT_LINE_ART_NONE,
                    text=fitz.PDF_REDACT_TEXT_REMOVE,
                )


            metadata = dict(doc.metadata or {})
            metadata.update({
                'title': str(document_title or 'Assistant response'),
                'author': '',
                'creator': APP_TITLE,
                'producer': APP_TITLE,
                'subject': 'Assistant response export',
                'keywords': 'assistant,pdf,export,math,svg,markdown',
            })
            doc.set_metadata(metadata)
            doc.save(temp_output, garbage=4, deflate=True)
            polished_with_fitz = temp_output.exists() and temp_output.stat().st_size > 1024
        finally:
            doc.close()

        if polished_with_fitz:
            temp_output.replace(pdf_path)

    if not polished_with_fitz:
        _polish_exported_pdf_with_pypdf(pdf_path, document_title=document_title)


def _build_assistant_pdf_document(html_fragment: str, theme: str='light', document_title: str='Assistant response', mathjax_svg_cache: str='') -> str:
    """Συνθέτει standalone printable HTML για headless browser export σε PDF.

    Για συνέπεια στην εκτύπωση, το PDF αποδίδεται πάντα με print-safe light theme,
    ανεξάρτητα από το active theme του UI. Έτσι αποφεύγονται σκούρα margins/containers
    όταν ο χρήστης βρίσκεται σε Dark Theme."""
    normalized_theme = 'light'
    base_css = _extract_primary_style_block(serve_index_html())
    prism_theme = 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism-solarizedlight.min.css'
    title_text = html.escape(str(document_title or 'Assistant response'))
    fragment = str(html_fragment or '').strip()
    mathjax_cache_fragment = str(mathjax_svg_cache or '').strip()
    cache_markup = f'<div class="mathjax-svg-cache" aria-hidden="true">{mathjax_cache_fragment}</div>' if mathjax_cache_fragment else ''
    extra_css = """
      @page {
        size: A4;
        margin: 18mm 12mm 18mm 12mm;
      }
      html, body {
        margin: 0 !important;
        min-height: auto !important;
      }
      html {
        background: #ffffff !important;
      }
      body {
        padding: 0 !important;
        background: #ffffff !important;
        color: #0f172a !important;
      }
      * {
        -webkit-print-color-adjust: exact !important;
        print-color-adjust: exact !important;
        box-sizing: border-box !important;
      }
      .app, .sidebar, .chat-panel, .chat-header, .messages-wrap, .messages, .composer, .footer-note, .top-actions {
        all: unset !important;
      }
      .print-shell {
        width: 100%;
        max-width: none;
        margin: 0;
        padding: 0;
      }
      .assistant-print-doc {
        display: block;
        width: 100%;
        margin: 0;
        padding: 0;
        background: #ffffff !important;
        color: #0f172a !important;
      }
      .assistant-print-body,
      .assistant-print-prompt-body {
        display: block;
        width: 100%;
        margin: 0;
        padding: 0;
        background: #ffffff !important;
        color: #0f172a !important;
        font-size: 11.25pt !important;
        line-height: 1.62 !important;
      }
      .assistant-print-body,
      .assistant-print-body *,
      .assistant-print-prompt-body,
      .assistant-print-prompt-body * {
        color: inherit;
      }
      .print-shell .assistant-export-cover {
        margin: 0 0 12px 0 !important;
      }
      .print-shell .assistant-export-prompt,
      .print-shell .assistant-export-answer {
        width: 100% !important;
        max-width: 100% !important;
        min-width: 0 !important;
        overflow: visible !important;
      }
      .print-shell .message-tools,
      .print-shell button,
      .print-shell .thinking-block,
      .print-shell details.thinking-block,
      .print-shell .thinking-summary,
      .print-shell .thinking-body {
        display: none !important;
      }
      .print-shell .msg,
      .print-shell .msg-head,
      .print-shell .msg-time,
      .print-shell .msg-role {
        all: unset !important;
      }
      .print-shell .msg-body,
      .print-shell .assistant-print-body,
      .print-shell .assistant-print-prompt-body,
      .print-shell .md-table-wrap,
      .print-shell .katex-display,
      .print-shell mjx-container,
      .print-shell mjx-container[display="true"],
      .print-shell pre,
      .print-shell .code-pre {
        overflow: hidden !important;
        max-width: 100% !important;
      }
      .print-shell .assistant-print-body > :first-child {
        margin-top: 0 !important;
      }
      .print-shell .assistant-print-body > :last-child {
        margin-bottom: 0 !important;
      }
      .print-shell .md-h1,
      .print-shell .md-h2,
      .print-shell .md-h3,
      .print-shell .md-h4,
      .print-shell .md-h5,
      .print-shell .md-h6 {
        color: #1d4ed8 !important;
        line-height: 1.25 !important;
        margin-top: 1.05em !important;
        margin-bottom: 0.35em !important;
        break-after: avoid-page !important;
        page-break-after: avoid !important;
      }
      .print-shell .md-p,
      .print-shell .md-list,
      .print-shell .md-bq,
      .print-shell .md-table-wrap,
      .print-shell .svg-render-card,
      .print-shell figure {
        break-inside: avoid-page;
        page-break-inside: avoid;
      }
      .print-shell .code-block {
        display: block !important;
        width: 100% !important;
        max-width: 100% !important;
        overflow: visible !important;
        break-inside: auto !important;
        page-break-inside: auto !important;
        box-decoration-break: clone !important;
        -webkit-box-decoration-break: clone !important;
      }
      .print-shell .code-toolbar {
        break-after: avoid-page !important;
        page-break-after: avoid !important;
      }
      .print-shell .md-p + .md-h1,
      .print-shell .md-p + .md-h2,
      .print-shell .md-p + .md-h3,
      .print-shell .md-list + .md-h1,
      .print-shell .md-list + .md-h2,
      .print-shell .md-list + .md-h3 {
        margin-top: 1.2em !important;
      }
      .print-shell .md-bq {
        background: #f8fafc !important;
        border-left-color: #60a5fa !important;
        color: #334155 !important;
      }
      .print-shell .code-pre,
      .print-shell pre {
        display: block !important;
        width: 100% !important;
        white-space: pre-wrap !important;
        overflow-wrap: anywhere !important;
        word-break: break-word !important;
        word-wrap: break-word !important;
        font-family: "Consolas", "Cascadia Code", "Fira Code", "Courier New", monospace !important;
        font-size: 10pt !important;
        line-height: 1.5 !important;
        background: #f8fafc !important;
        color: #0f172a !important;
        border: 1px solid #cbd5e1 !important;
        border-radius: 12px !important;
        box-shadow: none !important;
        padding: 14px 16px !important;
        overflow: visible !important;
        break-inside: auto !important;
        page-break-inside: auto !important;
        box-decoration-break: clone !important;
        -webkit-box-decoration-break: clone !important;
      }
      .print-shell code,
      .print-shell pre code,
      .print-shell .code-pre code {
        display: block !important;
        white-space: pre-wrap !important;
        overflow-wrap: anywhere !important;
        word-break: break-word !important;
        word-wrap: break-word !important;
        max-width: 100% !important;
        font-family: "Consolas", "Cascadia Code", "Fira Code", "Courier New", monospace !important;
      }
      .print-shell .md-table-wrap {
        width: 100% !important;
        overflow: visible !important;
        margin: 12px 0 !important;
      }
      .print-shell .md-table {
        width: 100% !important;
        min-width: 0 !important;
        table-layout: auto !important;
        border-collapse: separate !important;
        border-spacing: 0 !important;
        border: 1px solid #cbd5e1 !important;
        border-radius: 12px !important;
        overflow: hidden !important;
        background: #ffffff !important;
      }
      .print-shell .md-table thead {
        display: table-header-group !important;
      }
      .print-shell .md-table tr {
        break-inside: avoid-page;
        page-break-inside: avoid;
      }
      .print-shell .md-table th,
      .print-shell .md-table td {
        word-break: break-word !important;
        padding: 10px 14px !important;
        border-bottom: 1px solid #dbe4f0 !important;
        vertical-align: middle !important;
      }
      .print-shell .md-table th {
        background: #e2e8f0 !important;
        color: #0f172a !important;
        font-weight: 700 !important;
      }
      .print-shell .md-table tbody tr:nth-child(even) td {
        background: #f8fafc !important;
      }
      .print-shell .md-table tbody tr:last-child td {
        border-bottom: 0 !important;
      }
      .print-shell img,
      .print-shell svg,
      .print-shell canvas {
        max-width: 100% !important;
        height: auto !important;
      }
      .print-shell mjx-container,
      .print-shell mjx-container[display="true"] {
        overflow: visible !important;
        max-width: 100% !important;
      }
      .print-shell mjx-container svg,
      .print-shell .katex svg {
        overflow: visible !important;
        max-width: 100% !important;
        shape-rendering: geometricPrecision !important;
        text-rendering: geometricPrecision !important;
      }
      .print-shell img[data-pdf-rasterized="true"],
      .print-shell img[src^="data:image/png"] {
        image-rendering: auto !important;
        background: #ffffff !important;
      }
      .print-shell p,
      .print-shell li,
      .print-shell td,
      .print-shell th,
      .print-shell blockquote,
      .print-shell .assistant-export-subtitle,
      .print-shell .assistant-export-title,
      .print-shell .assistant-export-section-title,
      .print-shell .assistant-print-prompt-body,
      .print-shell .assistant-print-body,
      .print-shell .assistant-print-prompt-body *,
      .print-shell .assistant-print-body * {
        white-space: normal !important;
        overflow-wrap: anywhere !important;
        word-break: break-word !important;
      }
      .print-shell pre,
      .print-shell .code-pre,
      .print-shell code,
      .print-shell pre code,
      .print-shell .code-pre code {
        white-space: pre-wrap !important;
      }
      .print-shell svg,
      .print-shell .svg-render-card,
      .print-shell .svg-render-frame,
      .print-shell .svg-preview-wrap,
      .print-shell figure {
        break-inside: avoid-page !important;
        page-break-inside: avoid !important;
      }
      .print-shell .svg-render-card {
        background: #f8fafc !important;
        border: 1px solid #cbd5e1 !important;
        box-shadow: none !important;
      }
      .print-shell .svg-render-header {
        background: #0f172a !important;
        color: #ffffff !important;
      }
      .print-shell .svg-block,
      .print-shell .svg-preview-wrap,
      .print-shell .svg-print-figure,
      .print-shell .svg-print-figure img {
        background: #ffffff !important;
        box-shadow: none !important;
      }
      .print-shell .svg-block,
      .print-shell .svg-print-figure {
        border: 1px solid #dbe4f0 !important;
        border-radius: 16px !important;
        overflow: hidden !important;
      }
      .print-shell .svg-toolbar,
      .print-shell .svg-block [data-svg-source="1"],
      .print-shell .svg-block .code-block,
      .print-shell .svg-block .code-toolbar,
      .print-shell .svg-block pre,
      .print-shell .svg-block .code-pre {
        display: none !important;
      }
      .print-shell .svg-preview-wrap {
        padding: 12px !important;
        text-align: center !important;
      }
      .print-shell .svg-preview-image,
      .print-shell .svg-print-figure img,
      .print-shell .python-plot-render-print img {
        display: block !important;
        width: auto !important;
        max-width: 94% !important;
        max-height: 360px !important;
        height: auto !important;
        margin: 0 auto !important;
        border-radius: 0 !important;
        border: 0 !important;
      }
      .print-shell .svg-print-figure,
      .print-shell .python-plot-render-print {
        margin: 8px 0 !important;
        padding: 4px 0 !important;
        text-align: center !important;
        border: 0 !important;
        break-inside: avoid-page !important;
        page-break-inside: avoid !important;
      }
      .print-shell mjx-container[display="true"],
      .print-shell .katex-display {
        margin: 0.9em 0 !important;
        break-inside: avoid-page !important;
        page-break-inside: avoid !important;
      }
      .mathjax-svg-cache {
        position: absolute !important;
        left: -100000px !important;
        top: -100000px !important;
        width: 0 !important;
        height: 0 !important;
        overflow: hidden !important;
        visibility: hidden !important;
        pointer-events: none !important;
      }
      .mathjax-svg-cache svg {
        width: 0 !important;
        height: 0 !important;
        overflow: hidden !important;
      }
    """
    return f"""<!DOCTYPE html>
<html lang="el" data-theme="{normalized_theme}">
<head>
  <meta charset="utf-8" />
  <meta http-equiv="X-UA-Compatible" content="IE=edge" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{title_text}</title>
  <link rel="stylesheet" href="{prism_theme}" />
  <style>
{base_css}

{extra_css}
  </style>
</head>
<body>
  {cache_markup}
  <main class="print-shell">
    {fragment}
  </main>
</body>
</html>"""

def _safe_rmtree(path: Path, retries: int=8, delay: float=0.35) -> bool:
    """Διαγράφει φάκελο με retries ώστε να αντέχει σε προσωρινά locks του Chromium/Crashpad στα Windows."""
    target = Path(path)
    for attempt in range(max(1, retries)):
        try:
            if not target.exists():
                return True
            shutil.rmtree(target)
            return True
        except FileNotFoundError:
            return True
        except Exception as exc:
            if attempt >= retries - 1:
                log.warning('Αποτυχία cleanup temp dir %s: %s', target, exc)
                return False
            time.sleep(max(0.05, delay))
    return False


def _pdf_file_looks_valid(pdf_path: Path) -> bool:
    """Ελέγχει αν το παραγόμενο PDF φαίνεται ολοκληρωμένο και όχι απλώς προσωρινό/κενό αρχείο."""
    path = Path(pdf_path)
    try:
        if (not path.exists()) or (not path.is_file()):
            return False
        size = path.stat().st_size
        if size < 128:
            return False
        with path.open('rb') as handle:
            header = handle.read(5)
            if header != b'%PDF-':
                return False
            tail_read = min(size, 2048)
            handle.seek(max(0, size - tail_read))
            tail = handle.read(tail_read)
        return b'%%EOF' in tail
    except OSError:
        return False


def _wait_for_pdf_file_ready(pdf_path: Path, timeout_seconds: float=12.0, poll_interval: float=0.25) -> Tuple[bool, str]:
    """Περιμένει λίγο μετά το exit του browser γιατί σε Windows το PDF εμφανίζεται συχνά με μικρή καθυστέρηση."""
    path = Path(pdf_path)
    deadline = time.time() + max(0.5, timeout_seconds)
    last_size = -1
    stable_hits = 0
    while time.time() < deadline:
        try:
            if path.exists() and path.is_file():
                size = path.stat().st_size
                if size == last_size and size > 0:
                    stable_hits += 1
                else:
                    stable_hits = 0
                    last_size = size
                if stable_hits >= 1 and _pdf_file_looks_valid(path):
                    return (True, f'PDF έτοιμο ({size} bytes).')
            time.sleep(max(0.05, poll_interval))
        except OSError:
            time.sleep(max(0.05, poll_interval))
    if _pdf_file_looks_valid(path):
        try:
            return (True, f'PDF έτοιμο ({path.stat().st_size} bytes).')
        except OSError:
            return (True, 'PDF έτοιμο.')
    try:
        size = path.stat().st_size if path.exists() else 0
    except OSError:
        size = -1
    return (False, f'Το PDF δεν εμφανίστηκε ή δεν ολοκληρώθηκε έγκαιρα (exists={path.exists()}, size={size}).')


def _render_pdf_with_headless_browser(browser_path: Path, html_doc: str, output_pdf_path: Path) -> Tuple[bool, str]:
    """Τυπώνει standalone HTML σε PDF μέσω headless Chromium/Edge με ανθεκτικό έλεγχο ολοκλήρωσης και cleanup."""
    if not browser_path:
        return (False, 'Δεν βρέθηκε διαθέσιμος Chromium/Edge browser για headless PDF export.')

    output_pdf_path = Path(output_pdf_path).resolve()
    output_pdf_path.parent.mkdir(parents=True, exist_ok=True)
    creationflags = getattr(subprocess, 'CREATE_NO_WINDOW', 0)
    temp_dir_path = Path(tempfile.mkdtemp(prefix='ollama_chat_pdf_html_'))

    try:
        html_path = temp_dir_path / 'assistant_export.html'
        user_data_dir = temp_dir_path / 'browser_profile'
        user_data_dir.mkdir(parents=True, exist_ok=True)
        html_path.write_text(html_doc, encoding='utf-8')

        shared_args = [
            '--disable-gpu',
            '--hide-scrollbars',
            '--run-all-compositor-stages-before-draw',
            '--allow-file-access-from-files',
            '--disable-extensions',
            '--disable-default-apps',
            '--disable-sync',
            '--disable-background-networking',
            '--disable-component-update',
            '--disable-breakpad',
            '--disable-crash-reporter',
            '--metrics-recording-only',
            '--no-first-run',
            '--no-default-browser-check',
            '--print-to-pdf-no-header',
            '--no-pdf-header-footer',
            '--virtual-time-budget=20000',
            f'--user-data-dir={user_data_dir}',
            f'--print-to-pdf={output_pdf_path}',
            html_path.resolve().as_uri(),
        ]
        if os.name != 'nt':
            shared_args.insert(0, '--no-sandbox')

        variants = [
            [str(browser_path), '--headless=new', *shared_args],
            [str(browser_path), '--headless', *shared_args],
        ]

        errors: List[str] = []
        for command in variants:
            try:
                completed = subprocess.run(command, capture_output=True, text=True, timeout=90, creationflags=creationflags)
                ready, ready_detail = _wait_for_pdf_file_ready(output_pdf_path)
                if completed.returncode == 0 and ready:
                    return (True, ready_detail)

                stderr = (completed.stderr or '').strip()
                stdout = (completed.stdout or '').strip()
                try:
                    size = output_pdf_path.stat().st_size if output_pdf_path.exists() else 0
                except OSError:
                    size = -1
                detail_parts = [
                    f'Browser exit code {completed.returncode}',
                    f'pdf_exists={output_pdf_path.exists()}',
                    f'pdf_size={size}',
                    ready_detail,
                ]
                if stderr:
                    detail_parts.append(f'stderr: {stderr[:500]}')
                if stdout:
                    detail_parts.append(f'stdout: {stdout[:500]}')
                errors.append(' | '.join(part for part in detail_parts if part))
            except subprocess.TimeoutExpired:
                errors.append('Timeout κατά το headless print-to-pdf.')
            except Exception as exc:
                errors.append(str(exc))

        return (False, ' | '.join(part for part in errors if part) or 'Αποτυχία headless browser PDF export.')
    finally:
        _safe_rmtree(temp_dir_path)


def _decode_export_data_url(data_url: str) -> Tuple[bytes, str]:
    """Αποκωδικοποιεί data URLs που στέλνει το browser export pipeline για DOCX/PDF βοηθητικά assets."""
    text = str(data_url or '').strip()
    if not text.startswith('data:') or ',' not in text:
        return (b'', '')
    header, payload = text.split(',', 1)
    mime_type = header[5:].split(';', 1)[0].strip().lower() or 'application/octet-stream'
    try:
        if ';base64' in header.lower():
            return (base64.b64decode(payload), mime_type)
        return (urllib.parse.unquote_to_bytes(payload), mime_type)
    except Exception:
        return (b'', mime_type)


def _normalize_image_bytes_for_docx(image_bytes: bytes, mime_type: str='') -> bytes:
    """Μετατρέπει εικόνες σε μορφή που μπορεί να ενσωματώσει αξιόπιστα το python-docx."""
    raw = bytes(image_bytes or b'')
    if not raw:
        return b''
    normalized_mime = str(mime_type or '').strip().lower()
    try:
        if 'svg' in normalized_mime or raw.lstrip().startswith(b'<svg'):
            import cairosvg
            raw = cairosvg.svg2png(bytestring=raw)
            normalized_mime = 'image/png'
    except Exception:
        pass

    try:
        from PIL import Image
        with Image.open(io.BytesIO(raw)) as img:
            image_format = str(getattr(img, 'format', '') or '').upper()
            if image_format not in {'PNG', 'JPEG', 'JPG'}:
                converted = io.BytesIO()
                safe_image = img.convert('RGBA') if image_format not in {'JPEG', 'JPG'} else img.convert('RGB')
                safe_image.save(converted, format='PNG')
                return converted.getvalue()
    except Exception:
        pass
    return raw


def _apply_run_formatting_for_docx(run, state: Dict[str, object]) -> None:
    """Εφαρμόζει inline formatting σε run του Word."""
    font = run.font
    if state.get('bold'):
        run.bold = True
    if state.get('italic'):
        run.italic = True
    if state.get('underline'):
        run.underline = True
    if state.get('monospace'):
        font.name = 'Consolas'
    if state.get('subscript'):
        font.subscript = True
    if state.get('superscript'):
        font.superscript = True
    color_value = str(state.get('color') or '').strip().lstrip('#')
    if len(color_value) == 6:
        try:
            from docx.shared import RGBColor
            font.color.rgb = RGBColor.from_string(color_value.upper())
        except Exception:
            pass
    if state.get('link'):
        try:
            from docx.shared import RGBColor
            font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        except Exception:
            pass
        run.underline = True


def _set_paragraph_code_block_style(paragraph) -> None:
    """Δίνει απλή, καθαρή μορφή code block σε paragraph του DOCX."""
    try:
        from docx.shared import Pt
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        fmt = paragraph.paragraph_format
        fmt.left_indent = Pt(12)
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(6)
        fmt.line_spacing = 1.1

        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F6F8FA')
        p_pr.append(shd)
    except Exception:
        pass


def _node_has_docx_class(node, *class_names: str) -> bool:
    """Επιστρέφει True αν ο HTML κόμβος περιέχει κάποιο από τα ζητούμενα class names."""
    try:
        classes = {str(item).strip().lower() for item in (getattr(node, 'get', lambda *_: [])('class') or []) if str(item).strip()}
    except Exception:
        classes = set()
    wanted = {str(name).strip().lower() for name in class_names if str(name).strip()}
    return bool(classes & wanted)


def _extract_docx_preformatted_text(node) -> str:
    """Εξάγει κείμενο από highlighted pre/code block χωρίς να το σπάει token-by-token."""
    from bs4 import Tag

    if node is None:
        return ''

    target = node
    if isinstance(node, Tag):
        name = str(getattr(node, 'name', '') or '').lower()
        if name == 'pre':
            target = node.find('code') or node
        elif name != 'code':
            target = node.find('pre') or node.find('code') or node

    try:
        if isinstance(target, Tag):
            text = target.get_text('', strip=False)
        else:
            text = str(target or '')
    except Exception:
        text = str(target or '')

    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = text.replace('\u00a0', ' ')
    text = text.replace('\u200b', '').replace('\ufeff', '')
    return text.strip('\n')


def _extract_docx_code_language(node) -> str:
    """Ανιχνεύει τη γλώσσα ενός code block από classes/data attrs/toolbar labels."""
    from bs4 import Tag

    if node is None:
        return 'text'

    candidates = []

    def _collect_language_value(value) -> None:
        if not value:
            return
        if isinstance(value, (list, tuple, set)):
            for item in value:
                _collect_language_value(item)
            return
        text_value = str(value or '').strip()
        if text_value:
            candidates.append(text_value)

    if isinstance(node, Tag):
        _collect_language_value(node.get('data-language'))
        _collect_language_value(node.get('language'))
        _collect_language_value(node.get('class'))
        code_node = node.find('code') if node.name.lower() != 'code' else node
        if isinstance(code_node, Tag):
            _collect_language_value(code_node.get('data-language'))
            _collect_language_value(code_node.get('language'))
            _collect_language_value(code_node.get('class'))
        lang_label = node.find(class_=lambda value: value and 'code-lang' in ' '.join(value) if isinstance(value, (list, tuple)) else ('code-lang' in str(value)))
        if isinstance(lang_label, Tag):
            _collect_language_value(lang_label.get_text(' ', strip=True))

    for raw in candidates:
        text_value = str(raw or '').strip()
        if not text_value:
            continue
        for piece in re.split(r'\s+', text_value):
            match = re.search(r'language-([a-z0-9_+#.-]+)', piece, flags=re.IGNORECASE)
            if match:
                piece = match.group(1)
            piece = piece.strip().strip('.').strip().lower()
            if not piece:
                continue
            normalized = {
                'py': 'python',
                'python3': 'python',
                'js': 'javascript',
                'ts': 'typescript',
                'c++': 'cpp',
                'sh': 'bash',
                'shell': 'bash',
                'ps1': 'powershell',
                'plaintext': 'text',
                'plain': 'text',
                'txt': 'text',
                'yml': 'yaml',
                'md': 'markdown',
            }.get(piece, piece)
            if re.fullmatch(r'[a-z][a-z0-9_+#.-]{0,30}', normalized):
                return normalized
    return 'text'


def _merge_docx_styled_runs(runs):
    """Συγχωνεύει διαδοχικά runs με ίδιο style ώστε το DOCX να μένει ελαφρύ."""
    merged = []
    last_key = None
    for value, style in runs or []:
        if value is None or value == '':
            continue
        safe_style = dict(style or {})
        key = tuple(sorted(safe_style.items()))
        if merged and key == last_key:
            merged[-1] = (merged[-1][0] + value, merged[-1][1])
        else:
            merged.append((value, safe_style))
            last_key = key
    return merged


_DOCX_CODE_THEME = {
    'keyword': {'color': '7C3AED', 'bold': True},
    'builtin': {'color': '1D4ED8'},
    'string': {'color': '0F766E'},
    'comment': {'color': '64748B', 'italic': True},
    'number': {'color': 'DC2626'},
    'decorator': {'color': 'C2410C'},
    'function': {'color': '2563EB', 'bold': True},
    'class': {'color': '0F766E', 'bold': True},
    'constant': {'color': 'B45309', 'bold': True},
    'operator': {'color': '475569'},
}


_DOCX_PYTHON_KEYWORDS = {
    'False', 'None', 'True', 'and', 'as', 'assert', 'async', 'await', 'break', 'case',
    'class', 'continue', 'def', 'del', 'elif', 'else', 'except', 'finally', 'for',
    'from', 'global', 'if', 'import', 'in', 'is', 'lambda', 'match', 'nonlocal', 'not',
    'or', 'pass', 'raise', 'return', 'try', 'while', 'with', 'yield'
}


_DOCX_PYTHON_BUILTINS = {
    'abs', 'all', 'any', 'bool', 'bytes', 'callable', 'chr', 'dict', 'dir', 'enumerate',
    'filter', 'float', 'format', 'frozenset', 'getattr', 'hasattr', 'hash', 'hex', 'id',
    'int', 'isinstance', 'issubclass', 'iter', 'len', 'list', 'map', 'max', 'min', 'next',
    'object', 'open', 'ord', 'pow', 'print', 'range', 'repr', 'reversed', 'round', 'set',
    'slice', 'sorted', 'str', 'sum', 'super', 'tuple', 'type', 'zip', 'self', 'cls'
}


_DOCX_GENERIC_KEYWORDS = {
    'if', 'else', 'for', 'while', 'return', 'function', 'class', 'const', 'let', 'var',
    'import', 'export', 'from', 'try', 'catch', 'finally', 'new', 'public', 'private',
    'protected', 'static', 'async', 'await', 'break', 'continue', 'switch', 'case',
    'default', 'true', 'false', 'null', 'undefined', 'select', 'insert', 'update',
    'delete', 'create', 'drop', 'alter', 'where', 'join', 'group', 'order', 'by'
}


_DOCX_COMMENT_PREFIX = {
    'python': '#', 'bash': '#', 'shell': '#', 'powershell': '#', 'yaml': '#', 'yml': '#',
    'sql': '--', 'javascript': '//', 'typescript': '//', 'java': '//', 'c': '//',
    'cpp': '//', 'csharp': '//', 'go': '//', 'rust': '//', 'php': '//', 'swift': '//',
    'kotlin': '//', 'scala': '//', 'css': '/*', 'html': '<!--', 'xml': '<!--',
}


def _iter_docx_python_fallback_runs(code_text: str):
    """Απλό built-in syntax highlighter για Python όταν δεν υπάρχει Pygments."""
    token_re = re.compile(
        r"(?P<triple_double>(?:[furbFURB]{,3})\"\"\"[\s\S]*?\"\"\")"
        r"|(?P<triple_single>(?:[furbFURB]{,3})'''[\s\S]*?''')"
        r"|(?P<comment>\#.*?$)"
        r"|(?P<string>(?:[furbFURB]{,3})\"(?:\\.|[^\"\\])*\"|(?:[furbFURB]{,3})'(?:\\.|[^'\\])*')"
        r"|(?P<decorator>@[A-Za-z_][A-Za-z0-9_]*)"
        r"|(?P<number>\b(?:0x[0-9A-Fa-f]+|\d+(?:\.\d+)?(?:[eE][+-]?\d+)?)\b)"
        r"|(?P<identifier>\b[A-Za-z_][A-Za-z0-9_]*\b)"
        r"|(?P<operator>[+\-*/%=<>!&|^~:.,;()\[\]{}]+)"
        r"|(?P<space>\s+)"
        r"|(?P<other>.)",
        flags=re.MULTILINE,
    )

    runs = []
    expect_name = None
    for match in token_re.finditer(str(code_text or '')):
        kind = match.lastgroup or 'other'
        value = match.group(0)
        style = {}
        if kind in {'triple_double', 'triple_single', 'string'}:
            style = _DOCX_CODE_THEME['string']
        elif kind == 'comment':
            style = _DOCX_CODE_THEME['comment']
        elif kind == 'decorator':
            style = _DOCX_CODE_THEME['decorator']
        elif kind == 'number':
            style = _DOCX_CODE_THEME['number']
        elif kind == 'operator':
            style = _DOCX_CODE_THEME['operator']
            if value.strip() not in {'.'}:
                expect_name = None
        elif kind == 'identifier':
            if value in _DOCX_PYTHON_KEYWORDS:
                style = _DOCX_CODE_THEME['keyword']
                expect_name = 'function' if value == 'def' else ('class' if value == 'class' else None)
            elif expect_name == 'function':
                style = _DOCX_CODE_THEME['function']
                expect_name = None
            elif expect_name == 'class':
                style = _DOCX_CODE_THEME['class']
                expect_name = None
            elif value.isupper() and len(value) > 1:
                style = _DOCX_CODE_THEME['constant']
            elif value in _DOCX_PYTHON_BUILTINS:
                style = _DOCX_CODE_THEME['builtin']
            else:
                expect_name = None
        elif kind not in {'space'}:
            expect_name = None
        runs.append((value, dict(style or {})))
    return _merge_docx_styled_runs(runs)


def _iter_docx_generic_fallback_runs(code_text: str, language: str='text'):
    """Απλό fallback highlighter για JS/TS/SQL/Bash/YAML/κ.ά. χωρίς εξωτερική βιβλιοθήκη."""
    normalized_language = str(language or 'text').strip().lower()
    comment_prefix = _DOCX_COMMENT_PREFIX.get(normalized_language, '#')
    comment_re = {
        '/*': r'/\*[\s\S]*?\*/',
        '<!--': r'<!--[\s\S]*?-->',
        '--': r'--.*?$',
        '//': r'//.*?$',
        '#': r'\#.*?$',
    }.get(comment_prefix, r'\#.*?$')
    token_re = re.compile(
        rf"(?P<comment>{comment_re})"
        r"|(?P<string>\"(?:\\.|[^\"\\])*\"|'(?:\\.|[^'\\])*')"
        r"|(?P<number>\b(?:0x[0-9A-Fa-f]+|\d+(?:\.\d+)?(?:[eE][+-]?\d+)?)\b)"
        r"|(?P<identifier>\b[A-Za-z_][A-Za-z0-9_]*\b)"
        r"|(?P<operator>[+\-*/%=<>!&|^~:.,;()\[\]{}]+)"
        r"|(?P<space>\s+)"
        r"|(?P<other>.)",
        flags=re.MULTILINE,
    )

    runs = []
    for match in token_re.finditer(str(code_text or '')):
        kind = match.lastgroup or 'other'
        value = match.group(0)
        style = {}
        lowered = value.lower()
        if kind == 'comment':
            style = _DOCX_CODE_THEME['comment']
        elif kind == 'string':
            style = _DOCX_CODE_THEME['string']
        elif kind == 'number':
            style = _DOCX_CODE_THEME['number']
        elif kind == 'identifier':
            if lowered in _DOCX_GENERIC_KEYWORDS:
                style = _DOCX_CODE_THEME['keyword']
            elif value.isupper() and len(value) > 1:
                style = _DOCX_CODE_THEME['constant']
        elif kind == 'operator':
            style = _DOCX_CODE_THEME['operator']
        runs.append((value, dict(style or {})))
    return _merge_docx_styled_runs(runs)


def _iter_docx_syntax_runs(code_text: str, language: str='text'):
    """Επιστρέφει τμήματα κώδικα με style metadata για έγχρωμο DOCX syntax highlighting."""
    text_value = str(code_text or '').replace('\r\n', '\n').replace('\r', '\n')
    if not text_value:
        return []

    normalized_language = str(language or 'text').strip().lower() or 'text'
    aliases = {
        'text': 'text',
        'plain': 'text',
        'plaintext': 'text',
        'py': 'python',
        'python3': 'python',
        'js': 'javascript',
        'ts': 'typescript',
        'c++': 'cpp',
        'sh': 'bash',
        'shell': 'bash',
        'ps1': 'powershell',
        'yml': 'yaml',
    }
    normalized_language = aliases.get(normalized_language, normalized_language)

    try:
        from pygments import lex
        from pygments.lexers import TextLexer, get_lexer_by_name
        from pygments.styles import get_style_by_name
    except Exception:
        if normalized_language == 'python':
            return _iter_docx_python_fallback_runs(text_value)
        if normalized_language in {'javascript', 'typescript', 'json', 'sql', 'bash', 'powershell', 'yaml', 'html', 'xml', 'css', 'java', 'c', 'cpp', 'csharp', 'go', 'rust', 'php', 'swift', 'kotlin'}:
            return _iter_docx_generic_fallback_runs(text_value, language=normalized_language)
        return [(text_value, {})]

    try:
        lexer = TextLexer(stripnl=False) if normalized_language == 'text' else get_lexer_by_name(normalized_language, stripnl=False)
    except Exception:
        lexer = TextLexer(stripnl=False)

    try:
        style = get_style_by_name('friendly')
    except Exception:
        style = None

    runs = []
    last_style_key = None
    has_any_real_color = False
    for token_type, value in lex(text_value, lexer):
        if not value:
            continue
        current = token_type
        style_data = {}
        if style is not None:
            while current is not None:
                token_style = style.style_for_token(current)
                if token_style:
                    color = str(token_style.get('color') or '').strip()
                    if color:
                        style_data['color'] = color
                        has_any_real_color = True
                    if token_style.get('bold'):
                        style_data['bold'] = True
                    if token_style.get('italic'):
                        style_data['italic'] = True
                    if token_style.get('underline'):
                        style_data['underline'] = True
                    break
                current = getattr(current, 'parent', None)
        style_key = tuple(sorted(style_data.items()))
        if runs and style_key == last_style_key:
            runs[-1] = (runs[-1][0] + value, runs[-1][1])
        else:
            runs.append((value, style_data))
            last_style_key = style_key

    if not has_any_real_color:
        if normalized_language == 'python':
            return _iter_docx_python_fallback_runs(text_value)
        if normalized_language in {'javascript', 'typescript', 'json', 'sql', 'bash', 'powershell', 'yaml', 'html', 'xml', 'css', 'java', 'c', 'cpp', 'csharp', 'go', 'rust', 'php', 'swift', 'kotlin'}:
            return _iter_docx_generic_fallback_runs(text_value, language=normalized_language)

    return runs or [(text_value, {})]


def _append_syntax_highlighted_code_to_docx(paragraph, code_text: str, language: str='text') -> None:
    """Γράφει code block σε DOCX με Consolas και χρωματισμό token-level."""
    from docx.shared import Pt

    for chunk, style_state in _iter_docx_syntax_runs(code_text, language=language):
        if chunk == '':
            continue
        run = paragraph.add_run(chunk)
        state = {'monospace': True}
        state.update(style_state or {})
        _apply_run_formatting_for_docx(run, state)
        run.font.size = Pt(9.5)


def _add_picture_to_docx_run(run, image_bytes: bytes, max_width_inches: float=6.7):
    """Προσθέτει εικόνα σε run με ασφαλή κλιμάκωση.

    Πολύ μικρές inline εικόνες (π.χ. rasterized math glyphs) δεν πρέπει να
    διογκώνονται τεχνητά, γιατί στο DOCX δημιουργούν τεράστια κενά και χαλασμένη
    στοίχιση. Εδώ σεβόμαστε το intrinsic μέγεθος της εικόνας και απλώς το
    περιορίζουμε στο διαθέσιμο πλάτος."""
    if not image_bytes:
        return False
    try:
        from PIL import Image
        from docx.shared import Inches

        max_width = float(max_width_inches or 6.7)
        width_inches = min(max_width, 6.3)

        try:
            with Image.open(io.BytesIO(image_bytes)) as img:
                width_px = float(img.size[0] or 1)
                height_px = float(img.size[1] or 1)
                dpi = img.info.get('dpi', (96, 96))
                if isinstance(dpi, (tuple, list)) and dpi:
                    dpi_x = float(dpi[0] or 96.0)
                    dpi_y = float((dpi[1] if len(dpi) > 1 else dpi[0]) or 96.0)
                else:
                    dpi_x = float(dpi or 96.0)
                    dpi_y = float(dpi or 96.0)

                if dpi_x <= 1:
                    dpi_x = 96.0
                if dpi_y <= 1:
                    dpi_y = 96.0

                intrinsic_width = max(0.08, width_px / dpi_x)
                intrinsic_height = max(0.08, height_px / dpi_y)

                width_inches = min(max_width, intrinsic_width)

                max_height_inches = 8.4
                if intrinsic_height > 0 and width_inches > 0:
                    projected_height = width_inches * (height_px / max(width_px, 1.0))
                    if projected_height > max_height_inches:
                        width_inches = max(0.08, max_height_inches * (width_px / max(height_px, 1.0)))
        except Exception:
            width_inches = min(max_width, 6.3)

        run.add_picture(io.BytesIO(image_bytes), width=Inches(width_inches))
        return True
    except Exception:
        return False


def _normalize_docx_alt_text(node) -> str:
    """Καθαρίζει alt/aria-label κείμενο ώστε να μπορεί να χρησιμοποιηθεί ως fallback για math στο DOCX."""
    if node is None:
        return ''
    raw = str(
        node.get('alt', '')
        or node.get('aria-label', '')
        or node.get('title', '')
        or node.get('data-docx-alt', '')
        or ''
    )
    text = html.unescape(raw or '')
    text = re.sub(r'\s+', ' ', text).strip()
    generic = {'', 'math formula', 'formula', 'svg figure', 'image', 'picture', 'canvas'}
    return '' if text.lower() in generic else text


def _append_docx_text_runs(paragraph, node, state: Optional[Dict[str, object]]=None, max_width_inches: float=6.7) -> None:
    """Μετατρέπει inline HTML κόμβους σε runs του DOCX."""
    from bs4 import NavigableString, Tag

    current_state = dict(state or {})
    if node is None:
        return
    if isinstance(node, NavigableString):
        text = str(node)
        if text:
            run = paragraph.add_run(text)
            _apply_run_formatting_for_docx(run, current_state)
        return
    if not isinstance(node, Tag):
        return

    tag_name = node.name.lower()
    next_state = dict(current_state)
    if tag_name in {'strong', 'b'}:
        next_state['bold'] = True
    elif tag_name in {'em', 'i'}:
        next_state['italic'] = True
    elif tag_name == 'u':
        next_state['underline'] = True
    elif tag_name in {'code', 'kbd', 'samp'}:
        next_state['monospace'] = True
    elif tag_name == 'sub':
        next_state['subscript'] = True
    elif tag_name == 'sup':
        next_state['superscript'] = True
    elif tag_name == 'a':
        next_state['link'] = True

    if tag_name == 'br':
        paragraph.add_run().add_break()
        return

    classes = {str(item).strip().lower() for item in (node.get('class') or []) if str(item).strip()}
    is_docx_math = str(node.get('data-docx-math', '')).strip().lower() in {'1', 'true', 'yes'} or 'docx-math' in classes
    is_inline_math = str(node.get('data-docx-inline-math', '')).strip().lower() in {'1', 'true', 'yes'} or 'docx-inline-math' in classes

    if is_docx_math and tag_name != 'img':
        alt_text = _normalize_docx_alt_text(node) or re.sub(r'\s+', ' ', node.get_text(' ', strip=True) or '').strip()
        if alt_text:
            fallback = paragraph.add_run(alt_text)
            if is_inline_math:
                fallback.italic = True
            _apply_run_formatting_for_docx(fallback, next_state)
        return

    if tag_name == 'img':
        alt_text = _normalize_docx_alt_text(node)
        is_docx_math = str(node.get('data-docx-math', '')).strip().lower() in {'1', 'true', 'yes'} or 'docx-math' in classes
        is_inline_math = str(node.get('data-docx-inline-math', '')).strip().lower() in {'1', 'true', 'yes'} or 'docx-inline-math' in classes

        if is_docx_math and alt_text:
            fallback = paragraph.add_run(alt_text)
            if is_inline_math:
                fallback.italic = True
            _apply_run_formatting_for_docx(fallback, next_state)
            return

        image_bytes, mime_type = _decode_export_data_url(node.get('src', ''))
        if image_bytes:
            normalized = _normalize_image_bytes_for_docx(image_bytes, mime_type)
            run = paragraph.add_run()
            if not _add_picture_to_docx_run(run, normalized, max_width_inches=max_width_inches):
                if alt_text:
                    fallback = paragraph.add_run(f'[{alt_text}]')
                    _apply_run_formatting_for_docx(fallback, next_state)
        elif alt_text:
            fallback = paragraph.add_run(f'[{alt_text}]')
            _apply_run_formatting_for_docx(fallback, next_state)
        return

    for child in list(node.children):
        _append_docx_text_runs(paragraph, child, state=next_state, max_width_inches=max_width_inches)


def _set_docx_cell_text(cell, html_cell, max_width_inches: float=2.4) -> None:
    """Γεμίζει κελί Word table από HTML κελί."""
    from bs4 import Tag

    cell.text = ''
    first_paragraph = cell.paragraphs[0]
    first_paragraph.text = ''
    used_first = False

    def _new_paragraph_if_needed():
        nonlocal used_first
        if not used_first:
            used_first = True
            return first_paragraph
        return cell.add_paragraph()

    for child in list(getattr(html_cell, 'children', []) or []):
        if isinstance(child, Tag) and child.name and child.name.lower() in {'p', 'div', 'section', 'article'}:
            paragraph = _new_paragraph_if_needed()
            _append_docx_text_runs(paragraph, child, max_width_inches=max_width_inches)
        else:
            paragraph = _new_paragraph_if_needed()
            _append_docx_text_runs(paragraph, child, max_width_inches=max_width_inches)


def _append_html_to_docx(document, parent, node, list_level: int=0, list_type: Optional[str]=None, max_width_inches: float=6.7) -> None:
    """Μετατρέπει μπλοκ HTML του export fragment σε paragraphs/tables/εικόνες DOCX."""
    from bs4 import BeautifulSoup, NavigableString, Tag
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    if node is None:
        return
    if isinstance(node, BeautifulSoup):
        for child in list(node.children):
            _append_html_to_docx(document, parent, child, list_level=list_level, list_type=list_type, max_width_inches=max_width_inches)
        return
    if isinstance(node, NavigableString):
        if str(node).strip():
            paragraph = parent.add_paragraph()
            _append_docx_text_runs(paragraph, node, max_width_inches=max_width_inches)
        return
    if not isinstance(node, Tag):
        return

    tag_name = node.name.lower()
    if tag_name in {'style', 'script', 'noscript', 'template', 'meta', 'link', 'head', 'button', 'input', 'select', 'textarea'}:
        return

    if _node_has_docx_class(node, 'message-tools', 'code-toolbar', 'code-copy-btn', 'thinking-block'):
        return

    if _node_has_docx_class(node, 'code-block'):
        code_text = _extract_docx_preformatted_text(node)
        if not code_text.strip():
            return
        paragraph = parent.add_paragraph()
        _set_paragraph_code_block_style(paragraph)
        language = _extract_docx_code_language(node)
        _append_syntax_highlighted_code_to_docx(paragraph, code_text, language=language)
        return

    is_block_docx_math = str(node.get('data-docx-math', '')).strip().lower() in {'1', 'true', 'yes'} and str(node.get('data-docx-block-math', '')).strip().lower() in {'1', 'true', 'yes'}
    if is_block_docx_math:
        paragraph = parent.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _append_docx_text_runs(paragraph, node, state={'italic': True}, max_width_inches=max_width_inches)
        return

    block_like = {'article', 'section', 'div', 'main', 'header', 'footer', 'body', 'figure', 'figcaption'}
    heading_map = {'h1': 'Title', 'h2': 'Heading 1', 'h3': 'Heading 2', 'h4': 'Heading 3', 'h5': 'Heading 4', 'h6': 'Heading 5'}

    if tag_name in block_like:
        if tag_name == 'div' and not str(node.get_text(' ', strip=True) or '').strip():
            return
        if tag_name == 'div':
            block_children = [child for child in list(node.children) if isinstance(child, Tag) and child.name and child.name.lower() in {'p', 'div', 'section', 'article', 'ul', 'ol', 'pre', 'table', 'blockquote', 'img', 'figure', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}]
            if not block_children:
                paragraph = parent.add_paragraph()
                _append_docx_text_runs(paragraph, node, max_width_inches=max_width_inches)
                if paragraph.text.strip() or paragraph.runs:
                    return
                paragraph._element.getparent().remove(paragraph._element)
                return
        for child in list(node.children):
            _append_html_to_docx(document, parent, child, list_level=list_level, list_type=list_type, max_width_inches=max_width_inches)
        return

    if tag_name in heading_map:
        paragraph = parent.add_paragraph(style=heading_map[tag_name])
        _append_docx_text_runs(paragraph, node, max_width_inches=max_width_inches)
        if heading_map[tag_name] == 'Title':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return

    if tag_name == 'p':
        paragraph = parent.add_paragraph()
        _append_docx_text_runs(paragraph, node, max_width_inches=max_width_inches)
        return

    if tag_name in {'ul', 'ol'}:
        next_type = 'ol' if tag_name == 'ol' else 'ul'
        for child in list(node.children):
            _append_html_to_docx(document, parent, child, list_level=list_level + 1, list_type=next_type, max_width_inches=max_width_inches)
        return

    if tag_name == 'li':
        style_name = 'List Number' if list_type == 'ol' else 'List Bullet'
        if list_level > 1:
            suffix = min(5, list_level)
            candidate = f'{style_name} {suffix}'
            try:
                _ = document.styles[candidate]
                style_name = candidate
            except Exception:
                pass
        paragraph = parent.add_paragraph(style=style_name)
        inline_children = []
        nested_blocks = []
        for child in list(node.children):
            if isinstance(child, Tag) and child.name and child.name.lower() in {'ul', 'ol', 'div', 'p', 'pre', 'table', 'blockquote'}:
                nested_blocks.append(child)
            else:
                inline_children.append(child)
        for child in inline_children:
            _append_docx_text_runs(paragraph, child, max_width_inches=max_width_inches)
        for child in nested_blocks:
            _append_html_to_docx(document, parent, child, list_level=list_level, list_type=list_type, max_width_inches=max_width_inches)
        return

    if tag_name == 'pre' or _node_has_docx_class(node, 'code-pre'):
        paragraph = parent.add_paragraph()
        _set_paragraph_code_block_style(paragraph)
        code_text = _extract_docx_preformatted_text(node)
        if not code_text.strip():
            try:
                paragraph._element.getparent().remove(paragraph._element)
            except Exception:
                pass
            return
        language = _extract_docx_code_language(node)
        _append_syntax_highlighted_code_to_docx(paragraph, code_text, language=language)
        return

    if tag_name == 'blockquote':
        paragraph = parent.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        _append_docx_text_runs(paragraph, node, state={'italic': True}, max_width_inches=max_width_inches)
        return

    if tag_name == 'table':
        rows = []
        max_cols = 0
        for tr in node.find_all('tr'):
            cells = tr.find_all(['th', 'td'], recursive=False) or tr.find_all(['th', 'td'])
            if not cells:
                continue
            rows.append(cells)
            width = 0
            for cell in cells:
                try:
                    width += max(1, int(cell.get('colspan', 1)))
                except Exception:
                    width += 1
            max_cols = max(max_cols, width)
        if not rows or max_cols <= 0:
            return
        table = parent.add_table(rows=len(rows), cols=max_cols)
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        table.autofit = True
        for row_idx, cells in enumerate(rows):
            col_idx = 0
            for html_cell in cells:
                while col_idx < max_cols and table.cell(row_idx, col_idx).text:
                    col_idx += 1
                if col_idx >= max_cols:
                    break
                target_cell = table.cell(row_idx, col_idx)
                _set_docx_cell_text(target_cell, html_cell, max_width_inches=max(1.2, max_width_inches / max_cols))
                if html_cell.name.lower() == 'th' or row_idx == 0:
                    for paragraph in target_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    try:
                        from docx.oxml import OxmlElement
                        from docx.oxml.ns import qn
                        tc_pr = target_cell._tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), 'EAF2F8')
                        tc_pr.append(shd)
                    except Exception:
                        pass
                try:
                    colspan = max(1, int(html_cell.get('colspan', 1)))
                except Exception:
                    colspan = 1
                if colspan > 1 and col_idx + colspan - 1 < max_cols:
                    target_cell.merge(table.cell(row_idx, col_idx + colspan - 1))
                col_idx += colspan
        parent.add_paragraph()
        return

    if tag_name == 'img':
        alt_text = _normalize_docx_alt_text(node)
        is_docx_math = str(node.get('data-docx-math', '')).strip().lower() in {'1', 'true', 'yes'}

        if is_docx_math and alt_text:
            paragraph = parent.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if str(node.get('data-docx-block-math', '')).strip().lower() in {'1', 'true', 'yes'} else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(alt_text)
            run.italic = True
            return

        image_bytes, mime_type = _decode_export_data_url(node.get('src', ''))
        normalized = _normalize_image_bytes_for_docx(image_bytes, mime_type)
        if normalized:
            paragraph = parent.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            if not _add_picture_to_docx_run(run, normalized, max_width_inches=max_width_inches):
                if alt_text:
                    paragraph.add_run(f'[{alt_text}]')
        elif alt_text:
            paragraph = parent.add_paragraph()
            paragraph.add_run(f'[{alt_text}]')
        return

    if tag_name == 'hr':
        parent.add_paragraph('')
        return

    paragraph = parent.add_paragraph()
    _append_docx_text_runs(paragraph, node, max_width_inches=max_width_inches)


def _looks_like_export_css_leak(text: str) -> bool:
    """Ανιχνεύει CSS του export template που κατά λάθος πέρασε σαν κανονικό κείμενο."""
    sample = str(text or '').strip().lower()
    if not sample:
        return False
    markers = [
        '.assistant-export-',
        '.assistant-print-',
        'background: linear-gradient(',
        'text-overflow: ellipsis;',
        'page-break-inside: auto;',
        'box-decoration-break: clone;',
    ]
    hits = sum(1 for marker in markers if marker in sample)
    return hits >= 2


def _strip_export_css_leaks_from_soup(root) -> None:
    """Καθαρίζει CSS template leaks είτε ως <style> tags είτε ως plain-text paragraphs."""
    from bs4 import NavigableString, Tag

    if root is None:
        return

    for removable in list(root.find_all(['style', 'script', 'noscript', 'template', 'meta', 'link', 'head'])):
        try:
            removable.decompose()
        except Exception:
            pass

    # Αφαιρεί paragraphs/divs/pre που αποτελούνται αποκλειστικά από leaked CSS του exporter.
    for node in list(root.find_all(['p', 'div', 'pre'])):
        try:
            text = node.get_text('\n', strip=True)
        except Exception:
            text = ''
        if _looks_like_export_css_leak(text):
            try:
                node.decompose()
            except Exception:
                pass

    # Αν το CSS πέρασε σαν γυμνά text nodes στην αρχή του fragment, κόψε το preamble.
    pending = []
    started_real_content = False
    for child in list(getattr(root, 'children', []) or []):
        if isinstance(child, NavigableString):
            value = str(child or '')
            if not value.strip():
                continue
            if not started_real_content and _looks_like_export_css_leak(value):
                pending.append(child)
                continue
            started_real_content = True
            continue
        if isinstance(child, Tag):
            text = ''
            try:
                text = child.get_text('\n', strip=True)
            except Exception:
                text = ''
            if not started_real_content and _looks_like_export_css_leak(text):
                pending.append(child)
                continue
            started_real_content = True
    for node in pending:
        try:
            node.extract()
        except Exception:
            pass


def _build_assistant_docx_bytes(html_fragment: str, document_title: str='Assistant response') -> bytes:
    """Μετατρέπει το printable HTML fragment της απάντησης σε αρχείο DOCX χωρίς LibreOffice/soffice."""
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Cm, Pt

    fragment = str(html_fragment or '').strip()
    if not fragment:
        raise ValueError('Δεν δόθηκε printable HTML fragment για DOCX export.')

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    available_width_inches = max(4.5, float(section.page_width - section.left_margin - section.right_margin) / 914400.0)

    try:
        normal_style = document.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
    except Exception:
        pass
    try:
        title_style = document.styles['Title']
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(18)
    except Exception:
        pass

    core = document.core_properties
    core.title = str(document_title or 'Assistant response')
    core.author = ''
    core.subject = 'Assistant response export'
    core.keywords = 'assistant,docx,export,math,markdown'

    soup = BeautifulSoup(fragment, 'html.parser')
    for removable in soup.find_all(['style', 'script', 'noscript', 'template', 'meta', 'link', 'head']):
        try:
            removable.decompose()
        except Exception:
            pass
    for removable in soup.select('.mathjax-svg-cache'):
        try:
            removable.decompose()
        except Exception:
            pass

    root = soup.find(class_='assistant-print-doc') or soup
    _strip_export_css_leaks_from_soup(root)
    _append_html_to_docx(document, document, root, max_width_inches=available_width_inches)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def safe_read_json(handler: BaseHTTPRequestHandler) -> Dict:
    """Διαβάζει και αποκωδικοποιεί JSON body ασφαλώς με έλεγχο μεγέθους request."""
    try:
        content_length = int(handler.headers.get('Content-Length', '0'))
    except (ValueError, TypeError):
        content_length = 0
    if content_length > MAX_REQUEST_BODY_BYTES:
        log.warning('Απορρίφθηκε request %d bytes (όριο: %d)', content_length, MAX_REQUEST_BODY_BYTES)
        return {'__error__': 'request_too_large'}
    if content_length <= 0:
        return {}
    try:
        body = handler.rfile.read(content_length)
    except (OSError, ConnectionError) as exc:
        log.warning('Αποτυχία ανάγνωσης request body: %s', exc)
        return {}
    if not body:
        return {}
    try:
        return json.loads(body.decode('utf-8'))
    except (json.JSONDecodeError, UnicodeDecodeError):
        return {}

class QuietThreadingHTTPServer(ThreadingHTTPServer):
    """Ορίζει την κλάση QuietThreadingHTTPServer και συγκεντρώνει σε ένα σημείο τη σχετική κατάσταση και συμπεριφορά.

Ο διαχωρισμός αυτός βοηθά τον κώδικα να παραμένει οργανωμένος, επεκτάσιμος και ευκολότερος στη συντήρηση."""
    daemon_threads = True
    allow_reuse_address = True
    timeout = 60

    def handle_error(self, request, client_address) -> None:
        """Εξυπηρετεί HTTP σφάλματα επιστρέφοντας JSON αντί για HTML ώστε το frontend να τα χειριστεί σωστά."""
        _, exc, _ = sys.exc_info()
        if exc is not None and is_client_disconnect_error(exc):
            return
        super().handle_error(request, client_address)

def is_ollama_connection_refused(exc: object) -> bool:
    """Επιστρέφει True αν το σφάλμα αντιστοιχεί σε connection refused προς το Ollama API."""
    text = str(exc or '').lower()
    needles = ('10061', 'actively refused', 'connection refused', 'failed to establish a new connection', 'max retries exceeded')
    return any((n in text for n in needles))

def build_friendly_chat_error(exc: object) -> str:
    """Μετατρέπει τεχνικά σφάλματα chat API σε φιλικά μηνύματα με υπόδειξη επόμενης ενέργειας για τον χρήστη."""
    text = str(exc or '')
    lower = text.lower()
    if 'ollama_api_key' in lower or ('api key' in lower and ('missing' in lower or 'required' in lower)):
        return "Λείπει το Ollama Cloud API key για direct κλήσεις στο Ollama Cloud API. Βάλ'το στο πεδίο API Key του GUI ή αποθήκευσέ το στο settings αρχείο της εφαρμογής ή όρισέ το ως OLLAMA_API_KEY και ξανατρέξε την εφαρμογή."
    if is_ollama_connection_refused(exc):
        return 'Αποτυχία επικοινωνίας με το Ollama Cloud API. Έλεγξε τη σύνδεσή σου στο διαδίκτυο και δοκίμασε ξανά.'
    if '404' in text or 'not found' in lower or 'does not exist' in lower:
        return 'Το μοντέλο δεν βρέθηκε στο direct Ollama Cloud API. Στο direct mode τα επίσημα model names προέρχονται αποκλειστικά από το /api/tags και συχνά διαφέρουν από τα local *-cloud names. Κάνε Refresh Models και επίλεξε κάποιο από την επίσημη direct API λίστα.'
    if '401' in text or '403' in text or 'unauthorized' in lower or ('forbidden' in lower):
        return 'Το OLLAMA_API_KEY λείπει ή δεν είναι έγκυρο για το direct Ollama Cloud API. Έλεγξε το API key και δοκίμασε ξανά.'
    if 'network error' in lower or 'δικτυακό σφάλμα' in lower or 'name or service not known' in lower:
        return 'Δεν ήταν δυνατή η επικοινωνία με το Ollama Cloud API. Έλεγξε σύνδεση internet, firewall ή proxy.'
    if 'timeout' in lower or 'timed out' in lower or 'read timeout' in lower:
        return 'Το αίτημα προς το Ollama Cloud API έληξε (timeout). Το μοντέλο ίσως χρειάζεται περισσότερο χρόνο — δοκίμασε ξανά.'
    if 'context' in lower and ('length' in lower or 'window' in lower or 'exceed' in lower):
        return 'Το context window του μοντέλου ξεπεράστηκε. Δοκίμασε να καθαρίσεις το chat ή να μειώσεις το μέγεθος των συνημμένων.'
    return text or 'Άγνωστο σφάλμα επικοινωνίας με το Ollama Cloud API.'

def normalize_model_name(model: str) -> str:
    """Κανονικοποιεί ή καθαρίζει δεδομένα στο βήμα «normalize_model_name» ώστε οι επόμενες φάσεις να λαμβάνουν ασφαλή και συνεπή είσοδο.

Βασικά ορίσματα: model. Η απομόνωση αυτής της λογικής σε ξεχωριστή ρουτίνα μειώνει την επανάληψη και κάνει πιο εύκολο τον έλεγχο ή τη μελλοντική τροποποίηση."""
    model = (model or '').strip()
    if '/' in model and ':' in model:
        name_part, tag_part = model.split(':', 1)
        name_part = name_part.split('/')[-1]
        return f'{name_part}:{tag_part}'
    return model

def extract_chunk_content(chunk: object) -> str:
    """Εξάγει text περιεχόμενο από chunk του Ollama streaming API (message.content ή choices[0].delta.content)."""
    if chunk is None:
        return ''
    try:
        message = chunk.get('message', {}) if isinstance(chunk, dict) else getattr(chunk, 'message', None)
        if message is None:
            return str(getattr(chunk, 'content', '') or '')
        if isinstance(message, dict):
            return str(message.get('content', '') or '')
        return str(getattr(message, 'content', '') or '')
    except Exception:
        return ''

def extract_chunk_thinking(chunk: object) -> str:
    """Εξάγει thinking/reasoning περιεχόμενο από chunk του streaming API (message.thinking)."""
    if chunk is None:
        return ''
    try:
        message = chunk.get('message', {}) if isinstance(chunk, dict) else getattr(chunk, 'message', None)
        if message is None:
            return str(getattr(chunk, 'thinking', '') or '')
        if isinstance(message, dict):
            return str(message.get('thinking', '') or '')
        return str(getattr(message, 'thinking', '') or '')
    except Exception:
        return ''

def compose_display_assistant_text(content: str, thinking: str='') -> str:
    """Συνθέτει το τελικό text που εμφανίζεται στο UI συνδυάζοντας thinking block και main response."""
    safe_content = str(content or '')
    safe_thinking = str(thinking or '')
    if safe_thinking and safe_content:
        return f'<think>{safe_thinking}</think>\n\n{safe_content}'
    if safe_thinking:
        return f'<think>{safe_thinking}</think>'
    return safe_content
_INLINE_THINK_RE = re.compile('<think>.*?</think>\\s*', flags=re.IGNORECASE | re.DOTALL)

def strip_inline_think_blocks(text: str) -> str:
    """Αφαιρεί <think>...</think> blocks από streaming text που δεν εμφανίζονται ως ξεχωριστό thinking field."""
    return _INLINE_THINK_RE.sub('', str(text or '')).strip()

def is_gpt_oss_model(model: str) -> bool:
    """Επιστρέφει True αν το μοντέλο είναι GPT-OSS."""
    return 'gpt-oss' in str(model or '').strip().lower()

def is_qwen3_next_model(model: str) -> bool:
    """Επιστρέφει True αν το μοντέλο ανήκει στην Qwen3 series."""
    model_l = str(model or '').strip().lower()
    return 'qwen3-next' in model_l or 'qwen 3 next' in model_l

def is_qwen3_vl_model(model: str) -> bool:
    """Επιστρέφει True αν το μοντέλο ανήκει στην Qwen3-VL (vision) series."""
    model_l = str(model or '').strip().lower()
    return 'qwen3-vl' in model_l or 'qwen 3 vl' in model_l or 'qwen 3-vl' in model_l

def is_qwen3_coder_next_model(model: str) -> bool:
    """Επιστρέφει True αν το μοντέλο ανήκει στην Qwen3-Coder series."""
    model_l = str(model or '').strip().lower()
    return 'qwen3-coder-next' in model_l or 'qwen 3 coder next' in model_l

def is_reasoning_capable_model(model: str) -> bool:
    """Επιστρέφει True αν το μοντέλο υποστηρίζει extended reasoning / thinking."""
    model_l = str(model or '').strip().lower()
    thinking_hints = ('qwen3', 'deepseek-r1', 'deepseek-v3.1', 'reason', 'thinking', 'r1', 'gpt-oss')
    return any((token in model_l for token in thinking_hints))

def apply_qwen3_vl_nothink_workaround(messages: List[Dict], model: str, raw_mode: object) -> List[Dict]:
    """Εφαρμόζει workaround για Qwen3-VL που δεν αποδέχεται think=False αφαιρώντας το πεδίο από το payload."""
    mode = str(raw_mode or 'auto').strip().lower()
    if mode not in {'off', 'none'}:
        return messages
    if not is_qwen3_vl_model(model):
        return messages
    patched: List[Dict] = []
    last_user_index: Optional[int] = None
    for idx, item in enumerate(messages or []):
        if isinstance(item, dict):
            cloned = dict(item)
            patched.append(cloned)
            if str(cloned.get('role') or '').strip().lower() == 'user':
                last_user_index = idx
        else:
            patched.append(item)
    if last_user_index is None:
        return messages
    target = patched[last_user_index]
    content = str(target.get('content') or '')
    lower = content.lower()
    if '/no_think' not in lower and '/set nothink' not in lower:
        target['content'] = (content.rstrip() + '\n\n/no_think').strip()
    return patched

def resolve_think_mode(model: str, raw_mode: object) -> Optional[object]:
    """Καθορίζει την τιμή του πεδίου think για το request payload βάσει UI επιλογής και συμβατότητας μοντέλου."""
    mode = str(raw_mode or 'auto').strip().lower()
    model_l = str(model or '').strip().lower()
    if is_qwen3_coder_next_model(model):
        return None
    if mode in {'off', 'none'}:
        if is_gpt_oss_model(model):
            return 'low'
        if is_qwen3_next_model(model):
            return 'low'
        return False
    if mode == 'minimal':
        if is_gpt_oss_model(model):
            return 'low'
        if is_qwen3_next_model(model):
            return 'low'
        return True
    if mode in {'low', 'medium', 'high'}:
        if is_gpt_oss_model(model) or is_qwen3_next_model(model):
            return mode
        return True
    if mode == 'on':
        if is_gpt_oss_model(model):
            return 'medium'
        return True
    if mode != 'auto':
        log.warning('resolve_think_mode: άγνωστο mode %r — fallback σε auto', raw_mode)
    if is_gpt_oss_model(model):
        return 'medium'
    if any((token in model_l for token in ('qwen3-next',))):
        return True
    thinking_hints = ('qwen3', 'deepseek-r1', 'deepseek-v3.1', 'reason', 'thinking', 'r1')
    if any((token in model_l for token in thinking_hints)):
        return True
    return None

def iter_with_leading_chunk(first_chunk: object, iterator):
    """Επαναλήπτης που παράγει πρώτα ένα ήδη ληφθέν chunk και στη συνέχεια τα υπόλοιπα από τον αρχικό iterator."""
    yield first_chunk
    for chunk in iterator:
        yield chunk

def _build_think_fallback_candidates(model: str, think_value: Optional[object], raw_mode: object) -> List[Optional[object]]:
    """Χτίζει διατεταγμένη λίστα fallback τιμών think για επανεκκίνηση αν η αρχική τιμή απορριφθεί από το μοντέλο."""
    candidates: List[Optional[object]] = []
    mode = str(raw_mode or 'auto').strip().lower()

    def add(value: Optional[object]) -> None:
        """Προσθέτει μήνυμα στο ιστορικό συνομιλίας με αυτόματο trimming αν υπερβεί το MAX_HISTORY_MESSAGES."""
        if value not in candidates:
            candidates.append(value)
    add(think_value)
    if think_value is False:
        if is_qwen3_next_model(model):
            add('low')
            add(None)
        elif is_gpt_oss_model(model):
            add('low')
        else:
            add(None)
    elif think_value in {'low', 'medium', 'high'}:
        if is_qwen3_next_model(model):
            add(True)
            if mode == 'off':
                add(None)
        elif not is_gpt_oss_model(model):
            add(True)
    elif think_value is True:
        if is_qwen3_next_model(model):
            add('medium')
        elif is_gpt_oss_model(model):
            add('medium')
    elif think_value is None and is_reasoning_capable_model(model):
        add(True)
    return candidates

def _is_think_compat_error(exc: Exception) -> bool:
    """Επιστρέφει True αν το error message υποδηλώνει ότι το μοντέλο δεν αποδέχεται το πεδίο 'think'."""
    lower = str(exc).lower()
    return ' think ' in f' {lower} ' or 'invalid think value' in lower or 'reasoning_effort' in lower or ('reasoning effort' in lower)

def open_direct_cloud_chat_stream_with_fallback(*, model: str, messages: List[Dict], model_options: Optional[Dict], think_value: Optional[object], requested_mode: object) -> Tuple[object, Optional[object], List[str], bool]:
    """Ανοίγει chat stream με αυτόματο fallback αν το μοντέλο δεν αποδέχεται τις αρχικές παραμέτρους think."""
    candidates = _build_think_fallback_candidates(model, think_value, requested_mode)
    suppress_reasoning = str(requested_mode or '').strip().lower() in {'off', 'none'}
    warnings: List[str] = []
    first_error: Optional[Exception] = None
    last_error: Optional[Exception] = None
    for index, candidate in enumerate(candidates):
        stream_iter = direct_cloud_chat_stream(model=model, messages=messages, model_options=model_options if model_options else None, think_value=candidate)
        try:
            first_chunk = next(stream_iter)
            if index > 0:
                if candidate == 'low' and suppress_reasoning:
                    warnings.append(f"Το thinking mode για το μοντέλο {model} δεν δέχτηκε hard off από το Ollama Cloud API/model. Έγινε compatibility fallback σε think='low' και το reasoning trace αποκρύπτεται στο UI.")
                elif candidate is None:
                    warnings.append(f'Το thinking mode για το μοντέλο {model} δεν υποστηρίχθηκε πλήρως από το τρέχον Ollama Cloud API/model. Η απάντηση συνεχίζεται χωρίς ρητό think parameter.')
                else:
                    warnings.append(f'Το thinking mode για το μοντέλο {model} δεν υποστηρίχθηκε όπως ζητήθηκε και έγινε fallback σε think={candidate!r}.')
            return (iter_with_leading_chunk(first_chunk, stream_iter), candidate, warnings, suppress_reasoning)
        except StopIteration:
            if index > 0:
                warnings.append(f'Το thinking mode για το μοντέλο {model} χρειάστηκε fallback σε think={candidate!r}.')
            return (iter(()), candidate, warnings, suppress_reasoning)
        except Exception as exc:
            if first_error is None:
                first_error = exc
            last_error = exc
            if not _is_think_compat_error(exc):
                raise
            log.warning('Αποτυχία think compatibility για model=%s requested_mode=%r candidate=%r: %s', model, requested_mode, candidate, exc)
            continue
    raise last_error or first_error or RuntimeError('Αποτυχία εκκίνησης stream.')

def extract_token_stats(chunk: object) -> Optional[Dict]:
    """Εξάγει token usage στατιστικά από τελευταίο chunk streaming response."""
    try:
        if isinstance(chunk, dict):
            eval_count = chunk.get('eval_count')
            eval_duration = chunk.get('eval_duration')
            prompt_eval_count = chunk.get('prompt_eval_count')
            prompt_eval_duration = chunk.get('prompt_eval_duration')
            total_duration = chunk.get('total_duration')
            load_duration = chunk.get('load_duration')
        else:
            eval_count = getattr(chunk, 'eval_count', None)
            eval_duration = getattr(chunk, 'eval_duration', None)
            prompt_eval_count = getattr(chunk, 'prompt_eval_count', None)
            prompt_eval_duration = getattr(chunk, 'prompt_eval_duration', None)
            total_duration = getattr(chunk, 'total_duration', None)
            load_duration = getattr(chunk, 'load_duration', None)
        if not eval_count or not eval_duration or eval_duration <= 0:
            return None
        tokens_per_sec = round(eval_count / (eval_duration / 1000000000.0), 1)
        prompt_tokens_per_sec = None
        if prompt_eval_count and prompt_eval_duration and (prompt_eval_duration > 0):
            prompt_tokens_per_sec = round(prompt_eval_count / (prompt_eval_duration / 1000000000.0), 1)
        end_to_end_tokens_per_sec = None
        if total_duration and total_duration > 0:
            end_to_end_tokens_per_sec = round(eval_count / (total_duration / 1000000000.0), 1)
        return {'eval_count': eval_count, 'eval_duration': eval_duration, 'tokens_per_sec': tokens_per_sec, 'prompt_eval_count': prompt_eval_count, 'prompt_eval_duration': prompt_eval_duration, 'prompt_tokens_per_sec': prompt_tokens_per_sec, 'total_duration': total_duration, 'load_duration': load_duration, 'end_to_end_tokens_per_sec': end_to_end_tokens_per_sec}
    except Exception:
        pass
    return None

def get_effective_system_prompt(gui_system_prompt: str='', prompt_profile_id: str='', visualization_engine: str='auto') -> Tuple[str, str, str, str]:
    """Επιστρέφει το τελικό system prompt μαζί με metadata για prompt profile και visualization engine."""
    cleaned = (gui_system_prompt or '').strip()
    normalized_profile = normalize_prompt_profile_id(prompt_profile_id)
    normalized_engine = normalize_visualization_engine(visualization_engine)
    if cleaned:
        base_prompt = cleaned
        source = 'gui-custom'
    else:
        profile = get_prompt_profile(normalized_profile)
        base_prompt = str(profile.get('prompt', '') or '').strip() if profile else ''
        source = f'profile:{profile.get("id", normalized_profile)}' if profile else ''
    if not base_prompt:
        base_prompt, source = get_embedded_system_prompt()
    viz_instruction = build_visualization_engine_instruction(normalized_engine)
    effective_prompt = base_prompt.strip()
    if viz_instruction:
        effective_prompt = (effective_prompt + '\n\n' + viz_instruction).strip()
    return (effective_prompt, source, normalized_profile, normalized_engine)

def build_messages(system_prompt: str, session_messages: List[Dict]) -> List[Dict]:
    """Συνθέτει την πλήρη λίστα messages για το Ollama API, συμπεριλαμβάνοντας system prompt, ιστορικό και νέο user message."""
    messages: List[Dict] = []
    cleaned_system = (system_prompt or '').strip()
    if cleaned_system:
        messages.append({'role': 'system', 'content': cleaned_system})
    for item in session_messages:
        role = item.get('role', '').strip()
        content = item.get('content', '')
        if role in {'user', 'assistant'} and isinstance(content, str):
            msg: Dict = {'role': role, 'content': content}
            if role == 'assistant' and isinstance(item.get('thinking'), str) and item.get('thinking', '').strip():
                msg['thinking'] = item['thinking']
            if role == 'user' and item.get('images'):
                b64_images: List[str] = []
                for img in item['images']:
                    img_str = str(img or '').strip()
                    if not img_str:
                        continue
                    if len(img_str) > 260 or img_str.startswith('/') or (len(img_str) > 1 and img_str[1] == ':'):
                        try:
                            b64_images.append(base64.b64encode(Path(img_str).read_bytes()).decode('ascii'))
                        except Exception:
                            pass
                    else:
                        b64_images.append(img_str)
                if b64_images:
                    msg['images'] = b64_images
            messages.append(msg)
    return messages

def get_history_payload() -> List[Dict]:
    """Επιστρέφει thread-safe αντίγραφο του ιστορικού συνομιλίας για χρήση στο API call."""
    with SESSION.lock:
        return copy.deepcopy(SESSION.history)

def sanitize_filename(filename: str) -> str:
    """Κανονικοποιεί ή καθαρίζει δεδομένα στο βήμα «sanitize_filename» ώστε οι επόμενες φάσεις να λαμβάνουν ασφαλή και συνεπή είσοδο.

Βασικά ορίσματα: filename. Η απομόνωση αυτής της λογικής σε ξεχωριστή ρουτίνα μειώνει την επανάληψη και κάνει πιο εύκολο τον έλεγχο ή τη μελλοντική τροποποίηση."""
    basename = Path(filename).name
    cleaned = re.sub('[^\\w.\\-()\\[\\] ]+', '_', basename, flags=re.UNICODE).strip()
    cleaned = cleaned.lstrip('.')
    return cleaned or 'file'

def extract_original_generated_filename(stored_name: str) -> str:
    """Εξάγει το αρχικό suggested filename από μεταδεδομένα παραγόμενου αρχείου."""
    basename = Path(str(stored_name or '')).name
    match = re.match('^\\d{10,}_[0-9a-fA-F]{8}_(.+)$', basename)
    candidate = match.group(1) if match else basename
    safe_name = sanitize_filename(candidate)
    if safe_name and safe_name != 'file':
        return safe_name
    return 'generated_code.py'

def ensure_upload_dir() -> None:
    """Δημιουργεί τον φάκελο uploads αν δεν υπάρχει και επιστρέφει το Path του."""
    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

def ensure_generated_code_dir() -> None:
    """Δημιουργεί τον φάκελο generated_code αν δεν υπάρχει και επιστρέφει το Path του."""
    GENERATED_CODE_DIR.mkdir(parents=True, exist_ok=True)

def ensure_generated_media_dir() -> None:
    """Δημιουργεί τον κατάλογο που φιλοξενεί εικόνες/plots που παράγονται δυναμικά από Python code blocks."""
    GENERATED_MEDIA_DIR.mkdir(parents=True, exist_ok=True)

def extract_original_generated_media_filename(stored_name: str) -> str:
    """Ανακτά το αρχικό όνομα για αρχεία media που σερβίρονται αργότερα στο UI."""
    basename = Path(str(stored_name or '')).name
    match = re.match(r'^\d{10,}_[0-9a-fA-F]{8}_(.+)$', basename)
    candidate = match.group(1) if match else basename
    safe_name = sanitize_filename(candidate)
    if safe_name and safe_name != 'file':
        return safe_name
    return 'generated_plot.png'

def save_generated_media_file(source_path: Path, suggested_filename: str='') -> Dict[str, str]:
    """Αποθηκεύει plot/image artifact που παρήχθη από τον plot renderer και επιστρέφει metadata για το UI."""
    if not source_path.exists() or (not source_path.is_file()):
        raise ValueError('Το generated media αρχείο δεν βρέθηκε.')
    ensure_generated_media_dir()
    suffix = source_path.suffix.lower()
    if suffix not in {'.png', '.svg'}:
        suffix = '.png'
    requested_name = suggested_filename.strip() if suggested_filename else source_path.name
    safe_name = sanitize_filename(requested_name)
    if not safe_name.lower().endswith(suffix):
        safe_name = (Path(safe_name).stem or 'generated_plot') + suffix
    unique_name = f'{int(time.time() * 1000)}_{uuid.uuid4().hex[:8]}_{safe_name}'
    out_path = (GENERATED_MEDIA_DIR / unique_name).resolve()
    media_root = str(GENERATED_MEDIA_DIR.resolve()) + os.sep
    if not str(out_path).startswith(media_root):
        raise ValueError('Μη ασφαλές όνομα αρχείου για generated media attachment.')
    shutil.copy2(source_path, out_path)
    with SESSION.lock:
        SESSION.upload_paths.add(str(out_path))
    mime_type = 'image/svg+xml' if suffix == '.svg' else 'image/png'
    return {'name': safe_name, 'stored_name': unique_name, 'url': f'/generated-media/{urllib.parse.quote(unique_name)}', 'kind': 'image', 'mime_type': mime_type}


def _strip_python_code_fences(code_text: str) -> str:
    """Αφαιρεί fenced markdown περιτύλιγμα από Python code blocks όταν αυτό έχει διαρρεύσει στο render pipeline."""
    text = str(code_text or '').strip()
    if text.startswith('```'):
        lines = text.splitlines()
        if lines:
            lines = lines[1:]
        while lines and lines[-1].strip().startswith('```'):
            lines.pop()
        return '\n'.join(lines).strip('\n')
    return str(code_text or '')


def _join_broken_plot_lines(code_text: str) -> str:
    """Ενώνει γραμμές που έχουν σπάσει τεχνητά από το μοντέλο, π.χ. `axs` + `.text(...)` ή `vx` + `/2`."""
    lines = str(code_text or '').splitlines()
    if not lines:
        return str(code_text or '')
    joined: List[str] = []
    join_prefixes = ('.', ',', ')', ']', '}', '/', '*', '+', '-', '%')
    for raw_line in lines:
        line = raw_line.rstrip('\n')
        stripped = line.lstrip()
        if joined and stripped and stripped.startswith(join_prefixes):
            joined[-1] = joined[-1].rstrip() + stripped
            continue
        joined.append(line)
    return '\n'.join(joined)


def _comment_suspicious_plaintext_lines(code_text: str) -> str:
    """Μετατρέπει εμφανείς αφηγηματικές γραμμές σε σχόλια ώστε να μη σπάνε την εκτέλεση του plotting script."""
    out_lines: List[str] = []
    keywords = {
        'false', 'none', 'true', 'and', 'or', 'not', 'in', 'is', 'if', 'elif', 'else', 'for', 'while',
        'try', 'except', 'finally', 'with', 'as', 'def', 'class', 'return', 'yield', 'raise', 'import',
        'from', 'pass', 'break', 'continue', 'global', 'nonlocal', 'lambda', 'del', 'assert'
    }
    plain_text_pattern = re.compile(r'^[\s\u00A0\u1680\u2000-\u200A\u202F\u205F\u3000]*[A-Za-zΑ-Ωα-ωΆ-Ώά-ώ_][A-Za-zΑ-Ωα-ωΆ-Ώά-ώ0-9_\s,;:.!?”“"«»\-–—/\\]*$')
    for raw_line in str(code_text or '').splitlines():
        line = raw_line.rstrip('\n')
        stripped = line.strip()
        if (not stripped) or stripped.startswith('#'):
            out_lines.append(line)
            continue
        lowered = stripped.lower()
        looks_like_code = any(ch in stripped for ch in '=()[]{}:') or stripped.endswith('\\') or stripped.startswith(('plt', 'ax', 'fig', 'for ', 'if ', 'while ', 'def ', 'class ', 'return ', 'import ', 'from '))
        only_words = bool(plain_text_pattern.match(stripped))
        single_keyword = lowered in keywords
        single_identifier = bool(re.fullmatch(r'[A-Za-zΑ-Ωα-ωΆ-Ώά-ώ_][A-Za-zΑ-Ωα-ωΆ-Ώά-ώ0-9_]*', stripped))
        if only_words and (not looks_like_code) and (not single_keyword) and (not single_identifier):
            indent = line[:len(line) - len(line.lstrip(' '))]
            out_lines.append(f'{indent}# {stripped}')
            continue
        out_lines.append(line)
    return '\n'.join(out_lines)


def _repair_non_ascii_fstring_placeholders(code_text: str) -> str:
    """Αφαιρεί placeholders τύπου `{σταθερή}` από f-strings όταν πρόκειται προφανώς για plain text και όχι για έγκυρη Python έκφραση."""
    repaired_lines: List[str] = []
    fstring_hint = re.compile(r"(^|[^A-Za-z0-9_])(?:fr|rf|f)(['\"])|(^|[^A-Za-z0-9_])(?:fr|rf|f)(['\"]{3})", re.IGNORECASE)
    for line in str(code_text or '').splitlines():
        if fstring_hint.search(line):
            def _replace_placeholder(m: re.Match[str]) -> str:
                inner = str(m.group(1) or '').strip()
                if not inner:
                    return m.group(0)
                if all(ord(ch) < 128 for ch in inner) and re.fullmatch(r'[A-Za-z_][A-Za-z0-9_]*(?:\s*:[^{}]+)?', inner):
                    return m.group(0)
                return inner
            line = re.sub(r'\{([^{}\n]{1,120})\}', _replace_placeholder, line)
        repaired_lines.append(line)
    return '\n'.join(repaired_lines)

def _repair_plot_string_token(token_text: str) -> str:
    """Διορθώνει συνηθισμένα string issues σε matplotlib labels/text, όπως unsafe f-string placeholders και TeX backslashes."""
    text = str(token_text or '')
    match = re.match(r'(?i)^([rubf]*)(([\'\"]){{1,3}})(.*)\2$', text, re.DOTALL)
    if not match:
        return text
    prefix = match.group(1) or ''
    quote = match.group(2)
    body = match.group(4)
    if 'f' in prefix.lower():
        def _replace_placeholder(m: re.Match[str]) -> str:
            inner = str(m.group(1) or '').strip()
            if not inner:
                return m.group(0)
            if all(ord(ch) < 128 for ch in inner) and re.fullmatch(r'[A-Za-z_][A-Za-z0-9_]*(?:\s*:[^{}]+)?', inner):
                return m.group(0)
            return inner
        body = re.sub(r'\{([^{}\n]{1,120})\}', _replace_placeholder, body)
    needs_raw = ('\\' in body) and ('r' not in prefix.lower())
    fixed_prefix = prefix
    if needs_raw and ('b' not in prefix.lower()):
        fixed_prefix = ('r' + prefix) if prefix else 'r'
    return f'{fixed_prefix}{quote}{body}{quote}'


def repair_python_plot_code(code_text: str) -> str:
    """Εφαρμόζει ασφαλή heuristic repairs σε plotting code που παράχθηκε από LLM ώστε να γίνει πιο ανθεκτικό στο render."""
    repaired = _strip_python_code_fences(str(code_text or ''))
    repaired = repaired.replace('\r\n', '\n').replace('\r', '\n')
    repaired = repaired.replace('\u200b', '').replace('\ufeff', '')
    repaired = repaired.replace('“', '"').replace('”', '"').replace('‘', "'").replace('’', "'")
    repaired = _join_broken_plot_lines(repaired)
    repaired = _comment_suspicious_plaintext_lines(repaired)
    repaired = _repair_non_ascii_fstring_placeholders(repaired)
    try:
        tokens = []
        stream = io.StringIO(repaired)
        for token in tokenize.generate_tokens(stream.readline):
            if token.type == tokenize.STRING:
                replacement = _repair_plot_string_token(token.string)
                if replacement != token.string:
                    token = tokenize.TokenInfo(token.type, replacement, token.start, token.end, token.line)
            tokens.append(token)
        repaired = tokenize.untokenize(tokens)
    except Exception:
        pass
    return repaired

def validate_python_plot_code(code_text: str) -> Tuple[bool, str]:
    """Εκτελεί συντακτικό και βασικό safety validation πριν εκτελεστεί Python plotting block."""
    code_text = repair_python_plot_code(code_text)
    is_valid, validation_message = validate_python_code_block(code_text)
    if not is_valid:
        return (False, validation_message)
    try:
        tree = ast.parse(code_text, filename='<python_plot_block>')
    except SyntaxError as exc:
        return (False, f'Αποτυχία ανάλυσης plotting code: {exc}')
    aliases: Dict[str, str] = {}
    for node in ast.walk(tree):
        if isinstance(node, ast.Import):
            for alias in node.names:
                root = str(alias.name or '').split('.', 1)[0]
                if root not in SAFE_PLOT_IMPORT_ROOTS:
                    return (False, f'Το import "{alias.name}" δεν επιτρέπεται στον ασφαλή plot renderer.')
                aliases[alias.asname or root] = root
        elif isinstance(node, ast.ImportFrom):
            module_name = str(node.module or '')
            root = module_name.split('.', 1)[0] if module_name else ''
            if root not in SAFE_PLOT_IMPORT_ROOTS:
                return (False, f'Το import from "{module_name}" δεν επιτρέπεται στον ασφαλή plot renderer.')
            for alias in node.names:
                aliases[alias.asname or alias.name] = root
        elif isinstance(node, ast.Call):
            func = node.func
            if isinstance(func, ast.Name):
                func_name = str(func.id or '')
                if func_name in DISALLOWED_PLOT_CALL_NAMES:
                    return (False, f'Η κλήση "{func_name}(...)" δεν επιτρέπεται στον plot renderer.')
                mapped_root = aliases.get(func_name, func_name)
                if mapped_root in DISALLOWED_PLOT_ATTRIBUTE_ROOTS:
                    return (False, f'Η χρήση της βιβλιοθήκης "{mapped_root}" δεν επιτρέπεται στον plot renderer.')
            elif isinstance(func, ast.Attribute):
                root_name = ''
                value = func.value
                while isinstance(value, ast.Attribute):
                    value = value.value
                if isinstance(value, ast.Name):
                    root_name = str(value.id or '')
                mapped_root = aliases.get(root_name, root_name)
                if mapped_root in DISALLOWED_PLOT_ATTRIBUTE_ROOTS:
                    return (False, f'Η χρήση της βιβλιοθήκης "{mapped_root}" δεν επιτρέπεται στον plot renderer.')
    return (True, '')

def render_python_plot_to_generated_media(code_text: str, suggested_filename: str='') -> Tuple[bool, str, Optional[Dict[str, str]]]:
    """Τρέχει με ελεγχόμενο τρόπο matplotlib-based Python code και επιστρέφει generated image artifact."""
    code_text = repair_python_plot_code(str(code_text or '')).replace('\x00', '')
    if not code_text.strip():
        return (False, 'Το Python plotting block είναι κενό.', None)
    if len(code_text) > PLOT_RENDER_MAX_CODE_CHARS:
        return (False, 'Το Python plotting block είναι υπερβολικά μεγάλο για ασφαλή render εκτέλεση.', None)
    is_valid, validation_message = validate_python_plot_code(code_text)
    if not is_valid:
        return (False, validation_message, None)
    python_cmd, python_source = resolve_python_for_generated_scripts()
    if not python_cmd:
        return (False, 'Δεν βρέθηκε εγκατεστημένος Python interpreter για το plot renderer.', None)
    exec_root_dir = os.path.join(tempfile.gettempdir(), 'ollama_cloud_chat_exec')
    os.makedirs(exec_root_dir, exist_ok=True)
    session_dir = Path(tempfile.mkdtemp(prefix='plot_', dir=exec_root_dir))
    requested_name = str(suggested_filename or '').strip() or suggest_python_filename(code_text)
    safe_name = sanitize_filename(requested_name)
    if not safe_name.lower().endswith('.py'):
        safe_name += '.py'
    script_path = session_dir / safe_name
    script_path.write_text(code_text, encoding='utf-8', newline='\n')
    output_path = session_dir / ((Path(safe_name).stem or 'generated_plot') + '.png')
    runner_path = session_dir / '_plot_runner.py'
    runner_code = """import builtins
import json
import os
import re
import traceback
import warnings
import matplotlib
warnings.filterwarnings('ignore', category=SyntaxWarning)
matplotlib.use('Agg')
import matplotlib.pyplot as plt
OUTPUT_PATH = os.environ.get('OLLAMA_PLOT_OUTPUT', 'plot.png')
INPUT_PATH = os.environ.get('OLLAMA_PLOT_INPUT', '')

class _SafePlotNamespace(dict):
    def __missing__(self, key):
        if hasattr(builtins, key):
            return getattr(builtins, key)
        key_text = str(key or '')
        if any(ord(ch) > 127 for ch in key_text):
            return key_text
        raise KeyError(key)

def _patched_show(*args, **kwargs):
    figures = [plt.figure(num) for num in plt.get_fignums()]
    if figures:
        figures[-1].savefig(OUTPUT_PATH, dpi=170, bbox_inches='tight')

plt.show = _patched_show
namespace = _SafePlotNamespace({'__name__': '__main__', '__file__': INPUT_PATH, 'OUTPUT_PATH': OUTPUT_PATH, 'PLOT_OUTPUT_PATH': OUTPUT_PATH})
try:
    with open(INPUT_PATH, 'r', encoding='utf-8') as handle:
        source = handle.read()
    exec(compile(source, INPUT_PATH, 'exec'), namespace, namespace)
    if not os.path.exists(OUTPUT_PATH):
        figures = [plt.figure(num) for num in plt.get_fignums()]
        if figures:
            figures[-1].savefig(OUTPUT_PATH, dpi=170, bbox_inches='tight')
    if not os.path.exists(OUTPUT_PATH):
        raise RuntimeError('Το plotting script ολοκληρώθηκε αλλά δεν παρήγαγε matplotlib figure.')
    print(json.dumps({'ok': True, 'output_path': OUTPUT_PATH}))
except Exception as exc:
    print(json.dumps({'ok': False, 'error': str(exc), 'traceback': traceback.format_exc(limit=6)}))
    raise
"""
    runner_path.write_text(runner_code, encoding='utf-8', newline='\n')
    env = os.environ.copy()
    env['MPLBACKEND'] = 'Agg'
    env['OLLAMA_PLOT_INPUT'] = str(script_path)
    env['OLLAMA_PLOT_OUTPUT'] = str(output_path)
    creationflags = getattr(subprocess, 'CREATE_NO_WINDOW', 0) if os.name == 'nt' else 0
    try:
        proc = subprocess.run(python_cmd + [str(runner_path)], cwd=str(session_dir), env=env, capture_output=True, text=True, timeout=PLOT_RENDER_TIMEOUT_SECONDS, creationflags=creationflags)
    except subprocess.TimeoutExpired:
        shutil.rmtree(session_dir, ignore_errors=True)
        return (False, f'Λήξη χρόνου μετά από {PLOT_RENDER_TIMEOUT_SECONDS} δευτερόλεπτα κατά το render του plot.', None)
    except Exception as exc:
        shutil.rmtree(session_dir, ignore_errors=True)
        return (False, f'Αποτυχία εκτέλεσης plot renderer μέσω {python_source}: {exc}', None)
    stdout_text = str(proc.stdout or '').strip()
    stderr_text = str(proc.stderr or '').strip()
    if proc.returncode != 0 or (not output_path.exists()):
        details = (stderr_text or stdout_text or 'Το plotting script δεν παρήγαγε έγκυρο αποτέλεσμα.')[:PLOT_RENDER_MAX_STDOUT_CHARS]
        shutil.rmtree(session_dir, ignore_errors=True)
        return (False, details, None)
    try:
        suggested_media_name = f'{Path(safe_name).stem or "generated_plot"}.png'
        media_file = save_generated_media_file(output_path, suggested_filename=suggested_media_name)
    except Exception as exc:
        shutil.rmtree(session_dir, ignore_errors=True)
        return (False, f'Το plot δημιουργήθηκε αλλά απέτυχε η αποθήκευση του artifact: {exc}', None)
    shutil.rmtree(session_dir, ignore_errors=True)
    return (True, 'Το Python plot αποδόθηκε επιτυχώς.', media_file)

def suggest_python_filename(code_text: str) -> str:
    """Προτείνει περιγραφικό filename για παραγόμενο Python script βάσει περιεχομένου κώδικα."""
    cleaned = (code_text or '').strip()
    if not cleaned:
        return 'generated_code.py'
    patterns = ['^\\s*class\\s+([A-Za-z_][A-Za-z0-9_]*)', '^\\s*def\\s+([A-Za-z_][A-Za-z0-9_]*)', '^\\s*async\\s+def\\s+([A-Za-z_][A-Za-z0-9_]*)']
    for pattern in patterns:
        match = re.search(pattern, cleaned, flags=re.MULTILINE)
        if match:
            return sanitize_filename(match.group(1) + '.py')
    for line in cleaned.splitlines():
        stripped = line.strip()
        if stripped.startswith('#'):
            candidate = stripped.lstrip('#').strip()[:60]
            candidate = re.sub('\\s+', '_', candidate)
            candidate = sanitize_filename(candidate)
            if candidate and candidate != 'file':
                if not candidate.lower().endswith('.py'):
                    candidate += '.py'
                return candidate
            break
    return 'generated_code.py'

def save_generated_python_file(code_text: str, suggested_filename: str='') -> Dict[str, str]:
    """Αποθηκεύει παραγόμενο Python αρχείο στον GENERATED_CODE_DIR με sanitized filename."""
    cleaned = (code_text or '').replace('\r\n', '\n').replace('\r', '\n').strip() + '\n'
    if not cleaned.strip():
        raise ValueError('Δεν υπάρχει Python code για αποθήκευση.')
    ensure_generated_code_dir()
    requested_name = suggested_filename.strip() if suggested_filename else suggest_python_filename(cleaned)
    safe_name = sanitize_filename(requested_name)
    if not safe_name.lower().endswith('.py'):
        safe_name += '.py'
    unique_name = f'{int(time.time() * 1000)}_{uuid.uuid4().hex[:8]}_{safe_name}'
    out_path = (GENERATED_CODE_DIR / unique_name).resolve()
    generated_root = str(GENERATED_CODE_DIR.resolve()) + os.sep
    if not str(out_path).startswith(generated_root):
        raise ValueError('Μη ασφαλές όνομα αρχείου για generated Python attachment.')
    out_path.write_text(cleaned, encoding='utf-8')
    with SESSION.lock:
        SESSION.upload_paths.add(str(out_path))
    return {'name': safe_name, 'stored_name': unique_name, 'url': f'/generated-code/{urllib.parse.quote(unique_name)}', 'kind': 'file'}

def model_supports_images(model_name: str) -> bool:
    """Επιστρέφει True αν το μοντέλο υποστηρίζει vision/image input."""
    name = (model_name or '').lower()
    hints = ('vl', 'vision', 'gemini', 'gemma3', 'llava', 'minicpm-v', 'qwen2.5vl', 'qwen3-vl')
    return any((h in name for h in hints))

def save_uploaded_file(filename: str, data_base64: str) -> Path:
    """Αποθηκεύει ανεβασμένο αρχείο στον UPLOADS_DIR με επικύρωση μεγέθους και επέκτασης."""
    ensure_upload_dir()
    try:
        raw = base64.b64decode(data_base64, validate=True)
    except Exception as exc:
        raise ValueError(f'Μη έγκυρα base64 δεδομένα για το αρχείο: {filename}') from exc
    if len(raw) > MAX_UPLOAD_BYTES_PER_FILE:
        raise ValueError(f"Το αρχείο '{filename}' είναι πολύ μεγάλο. Μέγιστο: {MAX_UPLOAD_BYTES_PER_FILE // (1024 * 1024)} MB.")
    safe_name = sanitize_filename(filename)
    unique_name = f'{int(time.time() * 1000)}_{uuid.uuid4().hex[:8]}_{safe_name}'
    out_path = (UPLOADS_DIR / unique_name).resolve()
    uploads_resolved = str(UPLOADS_DIR.resolve()) + os.sep
    if not str(out_path).startswith(uploads_resolved):
        raise ValueError(f'Μη ασφαλές όνομα αρχείου: {filename}')
    out_path.write_bytes(raw)
    with SESSION.lock:
        SESSION.upload_paths.add(str(out_path))
    return out_path

def truncate_text(text: str, limit: int=MAX_TEXT_CHARS_PER_FILE) -> Tuple[str, bool]:
    """Κανονικοποιεί ή καθαρίζει δεδομένα στο βήμα «truncate_text» ώστε οι επόμενες φάσεις να λαμβάνουν ασφαλή και συνεπή είσοδο.

Βασικά ορίσματα: text, limit. Η απομόνωση αυτής της λογικής σε ξεχωριστή ρουτίνα μειώνει την επανάληψη και κάνει πιο εύκολο τον έλεγχο ή τη μελλοντική τροποποίηση."""
    cleaned = text.replace('\r\n', '\n').replace('\r', '\n').strip()
    if len(cleaned) <= limit:
        return (cleaned, False)
    return (cleaned[:limit].rstrip() + '\n\n...[TRUNCATED]...', True)

def extract_pdf_text(path: Path) -> Tuple[str, bool, str]:
    """Εξάγει κείμενο από PDF αρχείο χρησιμοποιώντας pymupdf ή pdfminer."""
    try:
        from pypdf import PdfReader
    except Exception:
        return ('', False, "Δεν βρέθηκε το optional package 'pypdf'. Εγκατάσταση: pip install pypdf")
    try:
        reader = PdfReader(str(path))
        parts: List[str] = []
        total_len = 0
        for page in reader.pages:
            extracted = page.extract_text() or ''
            if extracted:
                parts.append(extracted)
                total_len += len(extracted)
            if total_len >= MAX_TEXT_CHARS_PER_FILE:
                break
        joined = '\n\n'.join(parts).strip()
        if not joined:
            return ('', False, 'Δεν βρέθηκε εξαγώγιμο κείμενο στο PDF.')
        truncated_text, was_truncated = truncate_text(joined, MAX_TEXT_CHARS_PER_FILE)
        return (truncated_text, was_truncated, '')
    except Exception as exc:
        return ('', False, f'Αποτυχία ανάγνωσης PDF: {exc}')

def extract_text_for_context(path: Path) -> Tuple[str, bool, str]:
    """Εξάγει text περιεχόμενο από αρχείο για ένεση στο chat context, με truncation αν υπερβαίνει το MAX_TEXT_CHARS_PER_FILE."""
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        for enc in ('utf-8', 'utf-8-sig', 'cp1253', 'cp1252', 'latin-1'):
            try:
                content = path.read_text(encoding=enc)
                truncated_text, was_truncated = truncate_text(content, MAX_TEXT_CHARS_PER_FILE)
                return (truncated_text, was_truncated, '')
            except Exception:
                continue
        return ('', False, 'Το αρχείο δεν μπόρεσε να διαβαστεί ως κείμενο.')
    if suffix == '.pdf':
        return extract_pdf_text(path)
    return ('', False, 'Ο τύπος αρχείου δεν υποστηρίζεται για εξαγωγή κειμένου.')

def prepare_attachments(attachments: List[Dict], model_name: str) -> Tuple[List[Dict], List[str]]:
    """Επεξεργάζεται λίστα uploaded αρχείων και τα μετατρέπει σε attachment dicts κατάλληλα για το chat messages payload."""
    processed: List[Dict] = []
    warnings: List[str] = []
    if not attachments:
        return (processed, warnings)
    if len(attachments) > MAX_UPLOAD_FILES_PER_MESSAGE:
        raise ValueError(f'Μπορείς να στείλεις έως {MAX_UPLOAD_FILES_PER_MESSAGE} αρχεία ανά μήνυμα.')
    image_capable = model_supports_images(model_name)
    total_text_chars = 0
    for item in attachments:
        if not isinstance(item, dict):
            raise ValueError('Ένα από τα συνημμένα δεν είναι σε έγκυρη μορφή.')
        filename = str(item.get('name', 'file')).strip() or 'file'
        data_base64 = str(item.get('data_base64', '')).strip()
        mime_type = str(item.get('mime_type', '')).strip()
        if not data_base64:
            raise ValueError(f"Το αρχείο '{filename}' δεν περιέχει δεδομένα.")
        saved_path = save_uploaded_file(filename, data_base64)
        ext = saved_path.suffix.lower()
        entry: Dict = {'name': filename, 'path': str(saved_path), 'mime_type': mime_type, 'kind': 'other', 'text_excerpt': '', 'text_truncated': False, 'will_send_as_image': False, 'status_message': ''}
        if ext in IMAGE_EXTENSIONS:
            entry['kind'] = 'image'
            if image_capable:
                entry['will_send_as_image'] = True
            else:
                entry['status_message'] = 'Το αρχείο είναι εικόνα, αλλά το τρέχον μοντέλο δεν φαίνεται vision-capable.'
                warnings.append(f"Η εικόνα '{filename}' φορτώθηκε, αλλά για native image ανάλυση προτίμησε vision model όπως qwen3-vl.")
        else:
            entry['kind'] = 'document'
            text_excerpt, was_truncated, status_message = extract_text_for_context(saved_path)
            entry['text_excerpt'] = text_excerpt
            entry['text_truncated'] = was_truncated
            entry['status_message'] = status_message
            if text_excerpt:
                remaining = max(MAX_TOTAL_TEXT_CHARS_PER_MESSAGE - total_text_chars, 0)
                if remaining <= 0:
                    entry['text_excerpt'] = ''
                    entry['status_message'] = 'Το αρχείο φορτώθηκε, αλλά δεν μπήκε στο prompt λόγω ορίου context.'
                elif len(text_excerpt) > remaining:
                    trimmed, _ = truncate_text(text_excerpt, remaining)
                    entry['text_excerpt'] = trimmed
                    entry['text_truncated'] = True
                    total_text_chars += len(trimmed)
                else:
                    total_text_chars += len(text_excerpt)
        processed.append(entry)
    return (processed, warnings)

def build_user_message_content(user_text: str, processed_attachments: List[Dict]) -> str:
    """Χτίζει το content array ενός user message συνδυάζοντας κείμενο, εικόνες και documents."""
    parts: List[str] = [user_text.strip()]
    attachment_lines: List[str] = []
    context_blocks: List[str] = []
    for item in processed_attachments:
        if item['kind'] == 'image':
            if item['will_send_as_image']:
                attachment_lines.append(f"- Εικόνα: {item['name']} (θα σταλεί natively στο model)")
            else:
                note = item['status_message'] or 'Εικόνα χωρίς native vision υποστήριξη.'
                attachment_lines.append(f"- Εικόνα: {item['name']} ({note})")
        else:
            status_note = 'έτοιμο για context' if item['text_excerpt'] else item['status_message'] or 'χωρίς κείμενο'
            attachment_lines.append(f"- Αρχείο: {item['name']} ({status_note})")
            if item['text_excerpt']:
                context_blocks.append(f"### Αρχείο: {item['name']}\n{item['text_excerpt']}")
    if attachment_lines:
        parts.append('[Συνημμένα αρχεία]\n' + '\n'.join(attachment_lines))
    if context_blocks:
        parts.append('[Περιεχόμενο συνημμένων για χρήση ως context]\n\n' + '\n\n'.join(context_blocks))
    return '\n\n'.join((p for p in parts if p.strip()))

def serve_startup_html() -> str:
    """Επιστρέφει σελίδα εκκίνησης με SSE-based progress που ανακατευθύνει αυτόματα στο chat URL."""
    return f"""<!DOCTYPE html>\n<html lang="el">\n<head>\n  <meta charset="utf-8" />\n  <meta name="viewport" content="width=device-width, initial-scale=1" />\n  <title>Εκκίνηση — {html.escape(APP_TITLE)}</title>\n  <style>\n    *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}\n\n    body {{\n      min-height: 100vh;\n      background:\n        radial-gradient(circle at top left,  rgba(96,165,250,0.15), transparent 34%),\n        radial-gradient(circle at top right, rgba(94,234,212,0.12), transparent 26%),\n        linear-gradient(135deg, #0b1020, #101a35);\n      color: #e5eefc;\n      font-family: Consolas, "Cascadia Code", "Fira Code", monospace;\n      display: flex; align-items: center; justify-content: center;\n      flex-direction: column;\n      padding: 24px;\n    }}\n\n    .wrap {{ width: 680px; max-width: 100%; }}\n\n    .copyright {{\n      margin-top: 18px;\n      text-align: center;\n      color: #9fb0d1;\n      font-size: 0.92rem;\n      font-family: "Segoe UI", Inter, Arial, sans-serif;\n      letter-spacing: 0.5px;\n      padding-top: 16px;\n      border-top: 1px solid rgba(94,234,212,0.30);\n    }}\n\n    /* ── Header ── */\n    .header {{\n      text-align: center; margin-bottom: 28px;\n      animation: fade-down 0.5s ease-out;\n    }}\n    .logo   {{ font-size: 3.2rem; line-height: 1; margin-bottom: 10px; }}\n    .title  {{\n      font-size: 1.45rem; font-weight: 700; letter-spacing: 0.4px;\n      background: linear-gradient(135deg, #5eead4, #60a5fa);\n      -webkit-background-clip: text; -webkit-text-fill-color: transparent;\n    }}\n    .subtitle {{ color: #9fb0d1; font-size: 0.88rem; margin-top: 6px; }}\n\n    @keyframes fade-down {{\n      from {{ opacity: 0; transform: translateY(-14px); }}\n      to   {{ opacity: 1; transform: translateY(0); }}\n    }}\n\n    /* ── Terminal card ── */\n    .terminal {{\n      background: rgba(15,23,42,0.82);\n      border: 1px solid rgba(94,234,212,0.18);\n      border-radius: 20px;\n      backdrop-filter: blur(18px);\n      box-shadow: 0 24px 64px rgba(0,0,0,0.45);\n      overflow: hidden;\n      animation: fade-up 0.5s ease-out 0.1s both;\n    }}\n    @keyframes fade-up {{\n      from {{ opacity: 0; transform: translateY(14px); }}\n      to   {{ opacity: 1; transform: translateY(0); }}\n    }}\n\n    /* macOS-style titlebar */\n    .titlebar {{\n      display: flex; align-items: center; gap: 8px;\n      padding: 12px 18px;\n      background: rgba(15,23,42,0.98);\n      border-bottom: 1px solid rgba(94,234,212,0.10);\n    }}\n    .dot {{ width: 12px; height: 12px; border-radius: 50%; flex-shrink: 0; }}\n    .dot-r {{ background: #ff5f57; }}\n    .dot-y {{ background: #febc2e; }}\n    .dot-g {{ background: #28c840; }}\n    .bar-title {{\n      flex: 1; text-align: center;\n      color: #586e75; font-size: 0.8rem; letter-spacing: 0.3px;\n    }}\n\n    /* Log area */\n    .log-area {{\n      padding: 20px 22px;\n      min-height: 230px;\n      max-height: 380px;\n      overflow-y: auto;\n      scrollbar-width: thin;\n      scrollbar-color: rgba(94,234,212,0.2) transparent;\n    }}\n    .log-area::-webkit-scrollbar       {{ width: 5px; }}\n    .log-area::-webkit-scrollbar-thumb {{ background: rgba(94,234,212,0.22); border-radius: 99px; }}\n\n    .log-line {{\n      display: flex; gap: 14px;\n      padding: 2px 0; font-size: 0.875rem; line-height: 1.55;\n      animation: slide-in 0.22s ease-out;\n    }}\n    @keyframes slide-in {{\n      from {{ opacity: 0; transform: translateX(-10px); }}\n      to   {{ opacity: 1; transform: translateX(0); }}\n    }}\n\n    .log-t     {{ color: #586e75; flex-shrink: 0; width: 64px; }}\n    .log-lvl   {{ flex-shrink: 0; width: 56px; font-weight: 700; }}\n    .log-msg   {{ flex: 1; word-break: break-word; }}\n\n    .lvl-INFO    .log-lvl {{ color: #60a5fa; }}\n    .lvl-WARNING .log-lvl {{ color: #f59e0b; }}\n    .lvl-ERROR   .log-lvl {{ color: #f87171; }}\n    .lvl-READY   .log-lvl,\n    .lvl-READY   .log-msg {{ color: #34d399; }}\n    .lvl-READY   .log-msg {{ font-weight: 600; }}\n\n    /* Footer strip */\n    .footer {{\n      padding: 14px 22px;\n      border-top: 1px solid rgba(94,234,212,0.08);\n      background: rgba(15,23,42,0.55);\n    }}\n\n    /* Spinner */\n    .spinner {{\n      display: flex; align-items: center; gap: 10px;\n      color: #9fb0d1; font-size: 0.84rem;\n    }}\n    .dots {{ display: flex; gap: 5px; }}\n    .dots span {{\n      width: 7px; height: 7px; border-radius: 50%;\n      background: #5eead4; opacity: 0.25;\n      animation: dot-pulse 1.2s ease-in-out infinite;\n    }}\n    .dots span:nth-child(2) {{ animation-delay: 0.2s; }}\n    .dots span:nth-child(3) {{ animation-delay: 0.4s; }}\n    @keyframes dot-pulse {{\n      0%,80%,100% {{ opacity: 0.25; transform: scale(0.9); }}\n      40%          {{ opacity: 1;   transform: scale(1.1); }}\n    }}\n\n    /* Ready bar */\n    .ready-bar {{\n      display: none;\n      padding: 12px 16px;\n      background: rgba(52,211,153,0.12);\n      border: 1px solid rgba(52,211,153,0.28);\n      border-radius: 12px;\n      color: #34d399; font-size: 0.88rem; font-weight: 600;\n      text-align: center;\n    }}\n    .ready-bar a {{\n      color: #5eead4; cursor: pointer; text-decoration: underline;\n    }}\n    .progress-bar {{\n      height: 3px; background: rgba(52,211,153,0.15); border-radius: 99px;\n      margin-top: 10px; overflow: hidden;\n    }}\n    .progress-fill {{\n      height: 100%; width: 0;\n      background: linear-gradient(90deg, #5eead4, #60a5fa);\n      border-radius: 99px;\n      transition: width 1.6s linear;\n    }}\n  </style>\n</head>\n<body>\n  <div class="wrap">\n    <div class="header">\n      <div class="logo">☁️</div>\n      <div class="title">{html.escape(APP_TITLE)}</div>\n      <div class="subtitle">Εκκίνηση — παρακαλώ περίμενε…</div>\n    </div>\n\n    <div class="terminal">\n      <div class="titlebar">\n        <div class="dot dot-r"></div>\n        <div class="dot dot-y"></div>\n        <div class="dot dot-g"></div>\n        <div class="bar-title">startup log</div>\n      </div>\n\n      <div class="log-area" id="logArea"></div>\n\n      <div class="footer">\n        <div class="spinner" id="spinner">\n          <div class="dots"><span></span><span></span><span></span></div>\n          <span id="spinMsg">Αρχικοποίηση…</span>\n        </div>\n        <div class="ready-bar" id="readyBar">\n          ✅ Έτοιμο! Μεταβαίνεις αυτόματα…\n          &nbsp;<a onclick="goNow()">Πήγαινε τώρα</a>\n          <div class="progress-bar"><div class="progress-fill" id="progressFill"></div></div>\n        </div>\n      </div>\n    </div><!-- /.terminal -->\n\n    <div class="copyright">&copy; Ευάγγελος Πεφάνης</div>\n\n  </div><!-- /.wrap -->\n\n  <script>\n    var chatUrl  = null;\n    var logArea  = document.getElementById("logArea");\n    var spinner  = document.getElementById("spinner");\n    var spinMsg  = document.getElementById("spinMsg");\n    var readyBar = document.getElementById("readyBar");\n    var fillEl   = document.getElementById("progressFill");\n\n    var LEVEL_LABEL = {{\n      INFO: "INFO", WARNING: "WARN", ERROR: "ERR ", READY: "READY"\n    }};\n\n    function esc(s) {{\n      return String(s||"")\n        .replace(/&/g,"&amp;").replace(/</g,"&lt;").replace(/>/g,"&gt;");\n    }}\n\n    function addLine(ev) {{\n      var d = document.createElement("div");\n      d.className = "log-line lvl-" + ev.level;\n      d.innerHTML =\n        '<span class="log-t">'   + esc(ev.t)   + '</span>' +\n        '<span class="log-lvl">' + (LEVEL_LABEL[ev.level] || ev.level) + '</span>' +\n        '<span class="log-msg">' + esc(ev.msg)  + '</span>';\n      logArea.appendChild(d);\n      d.scrollIntoView({{ behavior: "smooth", block: "nearest" }});\n    }}\n\n    function goNow() {{\n      if (chatUrl) window.location.replace(chatUrl);\n    }}\n\n    var es = new EventSource("/startup-events");\n\n    es.onmessage = function(e) {{\n      var ev;\n      try {{ ev = JSON.parse(e.data); }} catch(_) {{ return; }}\n      addLine(ev);\n\n      if (ev.level !== "READY") {{\n        spinMsg.textContent = ev.msg.replace(/^[\\p{{Emoji}}\\s]+/u, "");\n      }}\n\n      if (ev.level === "READY") {{\n        chatUrl = ev.msg;\n        es.close();\n        spinner.style.display  = "none";\n        readyBar.style.display = "block";\n        // Animate progress bar → 100% then redirect\n        requestAnimationFrame(function() {{\n          fillEl.style.width = "100%";\n        }});\n        setTimeout(goNow, 1800);\n      }}\n    }};\n\n    es.onerror = function() {{\n      es.close();\n      spinMsg.textContent = "Αποτυχία SSE — ανανέωσε τη σελίδα.";\n      spinner.style.color = "#f59e0b";\n    }};\n  </script>\n</body>\n</html>"""

def serve_index_html() -> str:
    """Επιστρέφει το πλήρες index HTML της εφαρμογής με ενσωματωμένο SVG preview runtime."""
    system_prompt, _ = get_embedded_system_prompt()
    safe_prompt_json = json.dumps(system_prompt, ensure_ascii=False).replace('</', '<\\/')
    safe_prompt_profiles_json = json.dumps(get_prompt_profiles_catalog(), ensure_ascii=False).replace('</', '<\\/')
    safe_visualization_options_json = json.dumps(get_visualization_engine_options(), ensure_ascii=False).replace('</', '<\\/')
    accepted_types = ACCEPTED_FILE_TYPES
    html_doc = """<!DOCTYPE html>
<html lang="el" data-theme="dark">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>__APP_TITLE__</title>

  <!-- Prism.js — δύο themes: dark (prism-tomorrow) και light (prism-solarizedlight). -->
  <!-- Ενεργό theme εναλλάσσεται από JS στο applyTheme().                           -->
  <link id="prismDark"  rel="stylesheet"
        href="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism-tomorrow.min.css" />
  <link id="prismLight" rel="stylesheet" disabled
        href="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism-solarizedlight.min.css" />
  <link rel="stylesheet"
        href="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.css" />
  <script>
    window.MathJax = {
      loader: { load: ["[tex]/mhchem", "[tex]/physics", "[tex]/braket", "[tex]/cancel", "[tex]/bbox", "[tex]/mathtools"] },
      tex: {
        inlineMath: { "[+]": [["$", "$"]] },
        displayMath: [["$$", "$$"], ["\\[", "\\]"]],
        processEscapes: true,
        processEnvironments: true,
        packages: { "[+]": ["mhchem", "physics", "braket", "cancel", "bbox", "mathtools"] },
        tags: "ams",
        maxMacros: 1000
      },
      options: {
        skipHtmlTags: ["script", "noscript", "style", "textarea", "pre", "code"]
      },
      svg: {
        fontCache: "global"
      }
    };
  </script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/prism.min.js" defer></script>
  <script src="https://cdn.jsdelivr.net/npm/mathjax@4/tex-mml-svg.js" defer></script>
  <script src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.js" defer></script>
  <script src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/contrib/auto-render.min.js" defer></script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-python.min.js" defer></script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-javascript.min.js" defer></script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-typescript.min.js" defer></script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-bash.min.js" defer></script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-json.min.js" defer></script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-sql.min.js" defer></script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-css.min.js" defer></script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-markup.min.js" defer></script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-go.min.js" defer></script>
  <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-rust.min.js" defer></script>
  <!-- Inline: ενεργοποίηση σωστού Prism theme ΠΡΙΝ το πρώτο paint (αποφυγή flash) -->
  <script>
    (function () {
      var theme = "dark";
      try { theme = localStorage.getItem("ollama_chat_theme_v2") || "dark"; } catch (_) {}
      if (theme === "light") {
        var d = document.getElementById("prismDark");
        var l = document.getElementById("prismLight");
        if (d) d.disabled = true;
        if (l) l.disabled = false;
      }
    })();
  </script>

  <style>
    /* ── Variables (dark default) ── */
    :root {
      --bg1:       #0b1020;
      --bg2:       #101a35;
      --panel:     rgba(15, 23, 42, 0.78);
      --line:      rgba(148, 163, 184, 0.18);
      --text:      #e5eefc;
      --muted:     #9fb0d1;
      --accent:    #5eead4;
      --accent-2:  #60a5fa;
      --shadow:    0 20px 60px rgba(0,0,0,0.35);
      --radius:    22px;
      --radius-sm: 16px;
      --mono:      "Consolas", "Cascadia Code", "Fira Code", "Courier New", monospace;
      --sans:      "Segoe UI", Inter, Arial, sans-serif;
    }

    /* ── Reset ── */
    *, *::before, *::after { box-sizing: border-box; }
    html { color-scheme: dark; }
    html, body {
      margin: 0; min-height: 100%;
      font-family: var(--sans); color: var(--text);
      background:
        radial-gradient(circle at top left,  rgba(96,165,250,0.15), transparent 34%),
        radial-gradient(circle at top right, rgba(94,234,212,0.12), transparent 26%),
        linear-gradient(135deg, var(--bg1), var(--bg2));
      background-attachment: fixed;
    }
    body { padding: 22px; }
    ::selection { background: rgba(96,165,250,0.28); color: inherit; }

    /* ── Layout ── */
    .app {
      max-width: 1850px; margin: 0 auto;
      display: grid; grid-template-columns: 450px 1fr; gap: 22px;
    }
    .card {
      background: var(--panel); backdrop-filter: blur(18px);
      border: 1px solid var(--line); border-radius: var(--radius);
      box-shadow: var(--shadow);
    }

    /* ── Sidebar ── */
    .sidebar {
      padding: 18px;
      display: flex; flex-direction: column; gap: 16px;
      height: calc(100vh - 44px);
      position: sticky; top: 22px;
      overflow-y: auto;       /* scroll για μικρές οθόνες */
      scrollbar-width: thin;
    }
    .title { display: flex; align-items: center; justify-content: space-between; gap: 12px; }
    .title h1 { margin: 0; font-size: 1.22rem; line-height: 1.2; letter-spacing: 0.2px; }

    .pill {
      border-radius: 999px; padding: 7px 12px; font-size: 0.8rem; color: #dffcf6;
      background: linear-gradient(135deg, rgba(94,234,212,0.22), rgba(96,165,250,0.18));
      border: 1px solid rgba(94,234,212,0.2); white-space: nowrap;
    }

    .group {
      padding: 14px; border-radius: var(--radius-sm);
      background: rgba(15,23,42,0.52); border: 1px solid var(--line);
      flex-shrink: 0;
    }

    .label { display: block; margin-bottom: 8px; color: var(--muted); font-size: 0.92rem; font-weight: 600; }

    select, textarea, input[type="text"], input[type="file"] {
      width: 100%; border: 1px solid rgba(148,163,184,0.22);
      background: rgba(2,6,23,0.55); color: var(--text);
      border-radius: 14px; padding: 12px 14px; outline: none;
      font-family: var(--sans);
      transition: border-color 0.16s ease, box-shadow 0.16s ease;
    }
    textarea { resize: vertical; min-height: 90px; line-height: 1.45; }
    select:focus, textarea:focus, input:focus {
      border-color: rgba(96,165,250,0.55);
      box-shadow: 0 0 0 4px rgba(96,165,250,0.14);
    }
    input[type="file"] { padding: 10px; }

    .btn-row { display: grid; grid-template-columns: 1fr 1fr; gap: 10px; }

    button {
      border: 0; border-radius: 14px; padding: 12px 14px;
      cursor: pointer; font-weight: 700; font-size: 0.93rem;
      transition: transform 0.15s ease, filter 0.15s ease, opacity 0.15s ease;
    }
    button:hover:not(:disabled) { transform: translateY(-1px); filter: brightness(1.06); }
    button:disabled  { opacity: 0.52; cursor: not-allowed; transform: none; }
    .btn-full { width: 100%; }
    .primary  { color: #08111d; background: linear-gradient(135deg, var(--accent), var(--accent-2)); }
    .secondary {
      color: var(--text); background: rgba(51,65,85,0.85);
      border: 1px solid rgba(148,163,184,0.18);
    }

    #confirmThinkingProfileBtn {
      font-size: 0.82rem;
      font-weight: 400;
      padding: 9px 12px;
      letter-spacing: 0;
    }

    /* ── Chat panel ── */
    .chat-panel { height: calc(100vh - 44px); display: flex; flex-direction: column; overflow: hidden; }

    .chat-header {
      padding: 16px 26px; border-bottom: 1px solid var(--line);
      display: flex; flex-wrap: wrap; align-items: center;
      justify-content: space-between; gap: 10px;
      background: rgba(15,23,42,0.45); flex-shrink: 0;
    }
    .chat-header h2 { margin: 0; font-size: 1.08rem; }
    .header-left    { display: flex; flex-direction: column; gap: 3px; }
    .status-wrap    { display: flex; flex-wrap: wrap; gap: 8px; align-items: center; }

    .badge {
      border-radius: 999px; padding: 7px 12px; font-size: 0.82rem;
      background: rgba(51,65,85,0.8); border: 1px solid rgba(148,163,184,0.15);
      color: var(--muted);
    }
    .badge.ok   { color: #d1fae5; background: rgba(34,197,94,0.14);  border-color: rgba(34,197,94,0.24); }
    .badge.warn { color: #fef3c7; background: rgba(245,158,11,0.12); border-color: rgba(245,158,11,0.22); }
    .badge.err  { color: #fecaca; background: rgba(239,68,68,0.14);  border-color: rgba(239,68,68,0.24); }

    /* ── Messages ── */
    .messages-wrap { flex: 1; position: relative; overflow: hidden; }
    .messages {
      height: 100%; overflow-y: auto; padding: 22px;
      display: flex; flex-direction: column; gap: 14px;
      scroll-behavior: smooth;
    }
    .messages.streaming-active {
      scroll-behavior: auto;
    }

    .assistant-live-preview {
      display: block;
      width: 100%;
      padding: 2px 0;
      white-space: pre-wrap;
      word-break: break-word;
      overflow-wrap: anywhere;
      line-height: 1.62;
      font-family: var(--sans);
      color: inherit;
      min-height: 1.4em;
      contain: content;
    }
    .assistant-live-preview.live-code {
      font-family: var(--mono);
      font-size: 0.95em;
    }

    /* ── Realtime reasoning panel ── */
    .reasoning-panel {
      display: none;
      margin: 14px 22px 0;
      padding: 14px 16px;
      border-radius: 18px;
      border: 1px solid rgba(94,234,212,0.18);
      background: linear-gradient(180deg, rgba(14,26,50,0.72), rgba(12,22,40,0.62));
      box-shadow: inset 0 1px 0 rgba(255,255,255,0.03);
      flex-shrink: 0;
    }
    .reasoning-panel.visible { display: block; }
    .reasoning-panel.streaming {
      border-color: rgba(94,234,212,0.30);
      box-shadow: 0 0 0 1px rgba(94,234,212,0.08), 0 10px 30px rgba(0,0,0,0.14);
    }
    .reasoning-head {
      display: flex; align-items: flex-start; justify-content: space-between;
      gap: 12px; margin-bottom: 10px;
    }
    .reasoning-title-wrap { min-width: 0; }
    .reasoning-title {
      display: flex; align-items: center; gap: 10px;
      font-size: 0.95rem; font-weight: 800; letter-spacing: 0.2px;
    }
    .reasoning-meta {
      margin-top: 4px; color: var(--muted); font-size: 0.82rem;
    }
    .reasoning-toggle-btn {
      padding: 8px 12px; font-size: 0.8rem; white-space: nowrap;
    }
    .reasoning-body {
      margin: 0;
      max-height: 220px;
      overflow-y: auto;
      white-space: pre-wrap;
      word-break: break-word;
      font-family: var(--mono);
      font-size: 0.87rem;
      line-height: 1.56;
      color: #c9d9f3;
      background: rgba(2,6,23,0.34);
      border: 1px solid rgba(148,163,184,0.10);
      border-radius: 14px;
      padding: 12px 14px;
      scrollbar-width: thin;
      scrollbar-color: rgba(94,234,212,0.22) transparent;
    }
    .reasoning-body::-webkit-scrollbar { width: 6px; }
    .reasoning-body::-webkit-scrollbar-thumb {
      background: rgba(94,234,212,0.24); border-radius: 99px;
    }

    .empty-state {
      margin: auto; max-width: 820px; text-align: center;
      border: 1px dashed rgba(148,163,184,0.22); border-radius: 22px;
      padding: 34px 26px; color: var(--muted); background: rgba(15,23,42,0.28);
    }

    /* Scroll-to-bottom floating button */
    .scroll-to-bottom {
      position: absolute; bottom: 14px; right: 20px;
      background: var(--accent); color: #08111d;
      border: none; border-radius: 999px;
      padding: 10px 16px; font-weight: 800; font-size: 0.88rem;
      cursor: pointer; box-shadow: 0 6px 20px rgba(0,0,0,0.3);
      opacity: 0; pointer-events: none;
      transition: opacity 0.2s ease, transform 0.2s ease;
      z-index: 10;
    }
    .scroll-to-bottom.visible { opacity: 1; pointer-events: auto; }
    .scroll-to-bottom:hover   { transform: translateY(-2px); filter: brightness(1.08); }

    /* ── Messages ── */
    .msg {
      max-width: min(1020px, 95%); border-radius: 22px; padding: 14px 16px;
      border: 1px solid var(--line); box-shadow: 0 10px 30px rgba(0,0,0,0.18);
    }
    .msg.user      { align-self: flex-end;   background: linear-gradient(135deg, rgba(96,165,250,0.24), rgba(37,99,235,0.14)); }
    .msg.assistant { align-self: flex-start; background: rgba(15,23,42,0.75); max-width: 100%; width: 100%; }
    .msg.system    { align-self: center; background: rgba(51,65,85,0.6); color: var(--muted); max-width: 90%; }

    .msg-head {
      display: flex; align-items: center; justify-content: space-between;
      gap: 12px; margin-bottom: 10px; font-size: 0.84rem; color: var(--muted);
    }
    .msg-role { font-weight: 800; letter-spacing: 0.2px; }
    .msg-time { opacity: 0.8; white-space: nowrap; font-size: 0.8rem; }
    .msg-body { line-height: 1.62; font-size: 0.98rem; overflow-x: auto; }

    /* ── Markdown styles ── */
    .msg-body .md-h1,.msg-body .md-h2,.msg-body .md-h3,
    .msg-body .md-h4,.msg-body .md-h5,.msg-body .md-h6 {
      margin: 14px 0 6px; font-weight: 700; line-height: 1.3; color: var(--accent);
    }
    .msg-body .md-h1 { font-size: 1.35em; }
    .msg-body .md-h2 { font-size: 1.2em;  }
    .msg-body .md-h3 { font-size: 1.08em; }
    .msg-body .md-h4,.msg-body .md-h5,.msg-body .md-h6 { font-size: 1em; }
    .msg-body .md-p  { margin: 5px 0; }
    .msg-body .md-br { display: block; height: 5px; }
    .msg-body .md-hr { border: none; border-top: 1px solid var(--line); margin: 12px 0; }
    .msg-body .md-list { margin: 5px 0 5px 22px; padding: 0; }
    .msg-body .md-list li { margin: 3px 0; }
    .msg-body .md-bq {
      margin: 8px 0; padding: 10px 14px;
      border-left: 3px solid var(--accent);
      background: rgba(94,234,212,0.06);
      border-radius: 0 10px 10px 0; color: var(--muted);
    }
    .msg-body .md-bq p { margin: 0; }
    .msg-body .md-link { color: var(--accent-2); text-decoration: underline; }
    .msg-body .md-link:hover { opacity: 0.8; }
    .msg-body strong { color: #e2e8f0; font-weight: 700; }
    .msg-body em     { font-style: italic; color: #cbd5e1; }
    .msg-body del    { text-decoration: line-through; opacity: 0.7; }
    .msg-body .md-table-wrap {
      width: 100%;
      margin: 10px 0;
      overflow-x: auto;
      border-radius: 14px;
      border: 1px solid rgba(148,163,184,0.18);
      background: rgba(2,6,23,0.22);
    }
    .msg-body .md-table {
      width: max-content;
      min-width: 100%;
      border-collapse: collapse;
      font-size: 0.95rem;
    }
    .msg-body .md-table th,
    .msg-body .md-table td {
      padding: 10px 12px;
      border-bottom: 1px solid rgba(148,163,184,0.14);
      text-align: left;
      vertical-align: top;
    }
    .msg-body .md-table thead th {
      background: rgba(96,165,250,0.12);
      color: #e2e8f0;
      font-weight: 700;
    }
    .msg-body .md-table tbody tr:nth-child(even) {
      background: rgba(148,163,184,0.05);
    }
    .msg-body .md-table tbody tr:last-child td {
      border-bottom: none;
    }
    .msg-body .katex,
    .msg-body mjx-container {
      font-size: 1.04em;
      color: inherit;
    }
    .msg-body .katex-display,
    .msg-body mjx-container[display="true"] {
      margin: 0.85em 0;
      overflow-x: auto;
      overflow-y: hidden;
      padding: 4px 2px;
      display: block;
      max-width: 100%;
    }
    .msg-body mjx-container[display="true"] > svg {
      max-width: 100%;
      height: auto !important;
    }
    .msg-body mjx-container[display="false"] > svg {
      vertical-align: -0.18em;
    }

    /* ── Code blocks ── */
    .code-block {
      border: 1px solid rgba(148,163,184,0.18); border-radius: 16px;
      overflow: hidden; background: rgba(2,6,23,0.92);
      box-shadow: inset 0 1px 0 rgba(148,163,184,0.06); margin: 8px 0;
    }
    .code-toolbar {
      display: flex; align-items: center; justify-content: space-between;
      gap: 10px; padding: 9px 12px;
      background: rgba(15,23,42,0.98);
      border-bottom: 1px solid rgba(148,163,184,0.14);
    }
    .code-language {
      color: #cbd5e1; font-size: 0.78rem; font-weight: 700;
      letter-spacing: 0.3px; text-transform: uppercase;
    }
    .code-copy-btn {
      border: 1px solid rgba(148,163,184,0.18); background: rgba(30,41,59,0.9);
      color: var(--text); border-radius: 10px; padding: 6px 10px;
      font-size: 0.79rem; font-weight: 700; cursor: pointer;
    }
    .code-copy-btn:hover         { filter: brightness(1.1); }
    .code-copy-btn.done  { color: #d1fae5; border-color: rgba(34,197,94,0.24);  background: rgba(34,197,94,0.14); }
    .code-copy-btn.error { color: #fecaca; border-color: rgba(239,68,68,0.24);  background: rgba(239,68,68,0.14); }

    .code-pre {
      margin: 0 !important; padding: 16px 18px !important;
      overflow-x: auto; font-family: var(--mono) !important;
      font-size: 0.93rem !important; line-height: 1.65 !important;
      tab-size: 4; white-space: pre; text-align: left;
      background: transparent !important;
      scrollbar-width: thin;
      scrollbar-color: rgba(148,163,184,0.22) transparent;
    }
    .code-pre::-webkit-scrollbar       { height: 6px; }
    .code-pre::-webkit-scrollbar-track { background: transparent; }
    .code-pre::-webkit-scrollbar-thumb { background: rgba(148,163,184,0.28); border-radius: 99px; }
    .code-pre code[class*="language-"] {
      font-family: "Consolas", "Cascadia Code", "Fira Code", "Courier New", monospace !important;
      font-size: inherit !important;
      background: none !important;  /* prevent Prism dark background from leaking */
      font-variant-ligatures: none;
    }
    .msg.assistant .msg-body pre,
    .msg.assistant .msg-body pre *,
    .msg.assistant .msg-body code,
    .msg.assistant .msg-body code *,
    .assistant-print-body pre,
    .assistant-print-body pre *,
    .assistant-print-body code,
    .assistant-print-body code *,
    pre[class*="language-"],
    pre[class*="language-"] *,
    code[class*="language-"],
    code[class*="language-"] *,
    .code-pre,
    .code-pre *,
    .code-block pre,
    .code-block pre *,
    .code-block code,
    .code-block code * {
      font-family: "Consolas", "Cascadia Code", "Fira Code", "Courier New", monospace !important;
      font-variant-ligatures: none !important;
      font-feature-settings: "liga" 0, "calt" 0 !important;
    }
    .msg.assistant .msg-body pre,
    .assistant-print-body pre,
    .code-pre,
    .code-block pre {
      tab-size: 4 !important;
      -moz-tab-size: 4 !important;
    }

    .code-inline {
      display: inline-block; padding: 2px 6px; border-radius: 8px;
      background: rgba(30,41,59,0.92); border: 1px solid rgba(148,163,184,0.14);
      font-family: var(--mono); font-size: 0.91em; word-break: break-word;
    }

    /* ── Attachments ── */
    .attachment-list { margin-top: 10px; display: flex; flex-wrap: wrap; gap: 8px; }
    .attachment-chip {
      display: inline-flex; align-items: center; gap: 8px;
      padding: 6px 10px; border-radius: 999px;
      background: rgba(30,41,59,0.85); border: 1px solid rgba(148,163,184,0.18);
      font-size: 0.81rem; color: var(--text);
      text-decoration: none;
    }
    .attachment-chip.link { cursor: pointer; }
    .attachment-chip.link:hover { filter: brightness(1.08); }

    /* ── Thinking block (DeepSeek-R1, Qwen3, GLM-Z1 κ.ά.) ── */
    .thinking-block {
      margin: 8px 0 12px;
      border: 1px solid rgba(94,234,212,0.22);
      border-radius: 14px;
      overflow: hidden;
      background: rgba(14,26,50,0.55);
    }
    .thinking-summary {
      display: flex; align-items: center; gap: 8px;
      padding: 9px 14px; cursor: pointer;
      user-select: none; list-style: none;
      color: var(--muted); font-size: 0.83rem; font-weight: 600;
      background: rgba(94,234,212,0.06);
      border-bottom: 1px solid transparent;
      transition: background 0.15s;
    }
    .thinking-summary:hover { background: rgba(94,234,212,0.10); }
    details[open] .thinking-summary {
      border-bottom-color: rgba(94,234,212,0.15);
    }
    .thinking-icon { font-size: 0.95rem; }
    .thinking-label { flex: 1; }
    .thinking-chevron {
      font-size: 0.78rem; opacity: 0.6;
      transition: transform 0.2s;
    }
    details[open] .thinking-chevron { transform: rotate(90deg); }
    .thinking-body {
      padding: 12px 16px;
      font-size: 0.88rem; line-height: 1.58;
      color: #8da4c8;
      font-style: italic;
      white-space: pre-wrap;
      word-break: break-word;
      max-height: 340px;
      overflow-y: auto;
      scrollbar-width: thin;
      scrollbar-color: rgba(94,234,212,0.18) transparent;
    }
    .thinking-body::-webkit-scrollbar       { width: 4px; }
    .thinking-body::-webkit-scrollbar-thumb { background: rgba(94,234,212,0.22); border-radius: 99px; }

    /* Light theme thinking */
    html[data-theme="light"] .thinking-block {
      border-color: rgba(37,99,235,0.18);
      background: rgba(239,246,255,0.60);
    }
    html[data-theme="light"] .thinking-summary {
      background: rgba(37,99,235,0.06); color: #4a6080;
    }
    html[data-theme="light"] .thinking-summary:hover { background: rgba(37,99,235,0.10); }
    html[data-theme="light"] details[open] .thinking-summary {
      border-bottom-color: rgba(37,99,235,0.12);
    }
    html[data-theme="light"] .thinking-body { color: #5b7299; }
    .streaming-dots { display: inline-flex; gap: 5px; align-items: center; padding: 6px 0; }
    .streaming-dots span {
      width: 8px; height: 8px; border-radius: 50%;
      background: var(--accent); opacity: 0.25;
      animation: dot-pulse 1.2s ease-in-out infinite;
    }
    .streaming-dots span:nth-child(2) { animation-delay: 0.2s; }
    .streaming-dots span:nth-child(3) { animation-delay: 0.4s; }
    @keyframes dot-pulse {
      0%, 80%, 100% { opacity: 0.25; transform: scale(0.9); }
      40%            { opacity: 1;    transform: scale(1.1); }
    }

    /* ── Composer ── */
    .composer {
      border-top: 1px solid var(--line); padding: 16px 22px;
      background: rgba(15,23,42,0.4); flex-shrink: 0;
    }
    .composer textarea {
      min-height: 120px; height: 120px; max-height: 220px;
      font-size: 0.97rem; line-height: 1.5;
      font-family: var(--mono);
    }
    .composer-footer {
      margin-top: 10px; display: flex; flex-wrap: wrap; gap: 10px;
      align-items: center; justify-content: space-between;
    }
    .composer-left  { display: flex; gap: 12px; align-items: center; flex-wrap: wrap; }
    .composer-right { display: flex; gap: 10px; align-items: center; }
    .char-counter   { font-size: 0.81rem; color: var(--muted); }
    .char-counter.warn { color: #f59e0b; font-weight: 600; }
    .helper { color: var(--muted); font-size: 0.85rem; }

    /* ── Drag overlay ── */
    .drop-overlay {
      display: none; position: fixed; inset: 0;
      background: rgba(94,234,212,0.10); border: 3px dashed var(--accent);
      z-index: 9999; align-items: center; justify-content: center;
      font-size: 1.5rem; color: var(--accent); pointer-events: none;
      backdrop-filter: blur(2px);
    }
    .drop-overlay.active { display: flex; }

    /* ── Message tools ── */
    .message-tools { margin-top: 10px; display: flex; justify-content: flex-end; gap: 8px; }
    .tool-btn {
      padding: 7px 10px; border-radius: 10px;
      background: rgba(30,41,59,0.85); color: var(--text);
      font-size: 0.81rem; border: 1px solid rgba(148,163,184,0.16);
    }

    /* ── Model parameters panel ── */
    .param-row {
      display: flex; align-items: center; justify-content: space-between;
      gap: 10px; margin-bottom: 10px;
    }
    .param-row:last-child { margin-bottom: 0; }
    .param-label {
      font-size: 0.82rem; color: var(--muted); font-weight: 600;
      white-space: nowrap; min-width: 80px;
    }
    .param-value {
      font-size: 0.82rem; color: var(--accent); font-weight: 700;
      min-width: 36px; text-align: right; font-family: var(--mono);
    }
    input[type="range"] {
      flex: 1; height: 4px; border-radius: 4px; outline: none;
      background: rgba(148,163,184,0.25); padding: 0; border: none;
      cursor: pointer; accent-color: var(--accent);
    }
    input[type="range"]:focus { box-shadow: 0 0 0 3px rgba(94,234,212,0.18); }
    .param-seed-wrap {
      display: flex; gap: 8px; align-items: center;
    }
    .param-seed-wrap input[type="text"] {
      flex: 1; padding: 8px 10px; font-family: var(--mono);
      font-size: 0.82rem; border-radius: 10px;
    }

    /* ── Misc ── */
    .support-list { margin: 8px 0 0 18px; padding: 0; color: var(--muted); font-size: 0.81rem; line-height: 1.5; }
    .footer-note  { margin-top: auto; color: var(--muted); font-size: 0.84rem; line-height: 1.45; flex-shrink: 0; }
    .footer-note code {
      font-family: var(--mono); background: rgba(2,6,23,0.55);
      padding: 2px 7px; border-radius: 8px; font-size: 0.9em;
    }
    .muted { color: var(--muted); }
    .tiny  { font-size: 0.81rem; }

    /* ── Light theme ── */
    html[data-theme="light"] {
      color-scheme: light;
      --bg1:      #f5f8fd; --bg2:      #e7eef9;
      --panel:    rgba(255,255,255,0.94);
      --line:     rgba(37,53,84,0.10);
      --text:     #162338; --muted:    #5c6d86;
      --accent:   #2563eb; --accent-2: #06b6d4;
      --shadow:   0 22px 48px rgba(31,41,55,0.10);
    }
    html[data-theme="light"] body {
      background:
        radial-gradient(circle at top left,  rgba(37,99,235,0.12), transparent 36%),
        radial-gradient(circle at top right, rgba(6,182,212,0.10), transparent 28%),
        linear-gradient(135deg, var(--bg1), var(--bg2));
    }
    html[data-theme="light"] .card {
      background: linear-gradient(180deg, rgba(255,255,255,0.92), rgba(255,255,255,0.86));
      border-color: rgba(37,53,84,0.09);
    }
    html[data-theme="light"] .group {
      background: linear-gradient(180deg, rgba(248,251,255,0.96), rgba(242,247,253,0.90));
      border-color: rgba(37,53,84,0.08);
    }
    html[data-theme="light"] .pill { color: #0f3b57; }
    html[data-theme="light"] select,
    html[data-theme="light"] textarea,
    html[data-theme="light"] input {
      background: rgba(255,255,255,0.98); border-color: rgba(37,53,84,0.10); color: #162338;
    }
    html[data-theme="light"] .primary  { color: #fff; background: linear-gradient(135deg, #2563eb, #0ea5e9); }
    html[data-theme="light"] .secondary,
    html[data-theme="light"] .tool-btn,
    html[data-theme="light"] .attachment-chip {
      background: rgba(248,250,252,0.98); border-color: rgba(37,53,84,0.10); color: #1c2b40;
    }
    html[data-theme="light"] .chat-header,
    html[data-theme="light"] .composer {
      background: linear-gradient(180deg, rgba(255,255,255,0.88), rgba(247,250,254,0.92));
    }
    html[data-theme="light"] .reasoning-panel {
      background: linear-gradient(180deg, rgba(239,246,255,0.98), rgba(232,242,255,0.92));
      border-color: rgba(37,99,235,0.14);
    }
    html[data-theme="light"] .reasoning-panel.streaming {
      border-color: rgba(37,99,235,0.24);
      box-shadow: 0 0 0 1px rgba(37,99,235,0.05), 0 12px 28px rgba(37,99,235,0.06);
    }
    html[data-theme="light"] .reasoning-body {
      background: rgba(255,255,255,0.88);
      border-color: rgba(37,53,84,0.08);
      color: #365173;
    }
    html[data-theme="light"] .msg.user      { background: linear-gradient(135deg, rgba(37,99,235,0.12), rgba(6,182,212,0.10)); }
    html[data-theme="light"] .msg.assistant { background: linear-gradient(180deg, rgba(255,255,255,0.98), rgba(247,250,253,0.94)); }
    html[data-theme="light"] .msg.system    { background: rgba(242,246,252,0.98); color: #5b6c83; }
    html[data-theme="light"] .badge         { background: rgba(243,247,252,0.98); border-color: rgba(37,53,84,0.10); color: #5c6d86; }
    html[data-theme="light"] .badge.ok      { color: #166534; background: rgba(34,197,94,0.12);  border-color: rgba(34,197,94,0.20); }
    html[data-theme="light"] .badge.warn    { color: #92400e; background: rgba(245,158,11,0.12); border-color: rgba(245,158,11,0.20); }
    /* ═══════════════════════════════════════════════════════════════════════
       Light theme — code blocks: λευκό φόντο, πλήρης vivid token palette
       Base: Prism Solarized Light  +  custom overrides για μέγιστο χρώμα
       ═══════════════════════════════════════════════════════════════════════ */

    html[data-theme="light"] .code-block {
      background:   #eef6ff;               /* ανοιχτό γαλάζιο */
      border-color: rgba(37,99,235,0.14);
      box-shadow:   0 2px 12px rgba(37,99,235,0.07);
    }
    html[data-theme="light"] .code-toolbar {
      background:          #dbeafe;        /* blue-100 */
      border-bottom-color: rgba(37,99,235,0.12);
    }
    html[data-theme="light"] .code-language {
      color: #1e40af; letter-spacing: 0.3px;
    }
    html[data-theme="light"] .code-copy-btn {
      background:   #eff6ff;
      border-color: rgba(37,99,235,0.18);
      color:        #1e3a5f;
      box-shadow:   0 1px 3px rgba(37,99,235,0.06);
    }
    html[data-theme="light"] .code-copy-btn:hover {
      background: #dbeafe; filter: none;
    }
    html[data-theme="light"] .code-copy-btn.done  {
      color: #166534; background: rgba(34,197,94,0.13); border-color: rgba(34,197,94,0.25);
    }
    html[data-theme="light"] .code-copy-btn.error {
      color: #991b1b; background: rgba(239,68,68,0.10); border-color: rgba(239,68,68,0.25);
    }

    /* Βασικό χρώμα κειμένου */
    html[data-theme="light"] .code-pre,
    html[data-theme="light"] .code-pre code[class*="language-"] {
      color:      #1e293b;
      background: transparent !important;
    }

    html[data-theme="light"] .code-pre::-webkit-scrollbar-thumb {
      background: rgba(37,99,235,0.20);
    }

    /* ── Vivid token colours (Atom One Light inspired, more saturated) ── */

    /* Comments — readable gray, italic */
    html[data-theme="light"] .token.comment,
    html[data-theme="light"] .token.prolog,
    html[data-theme="light"] .token.doctype,
    html[data-theme="light"] .token.cdata {
      color: #93a1a1; font-style: italic;  /* Solarized base1 */
    }

    /* Keywords — vivid purple: def, class, import, if, for, return… */
    html[data-theme="light"] .token.keyword,
    html[data-theme="light"] .token.atrule,
    html[data-theme="light"] .token.rule {
      color: #7c3aed; font-weight: 600;   /* violet-600 */
    }

    /* Strings, chars, template literals — vivid green */
    html[data-theme="light"] .token.string,
    html[data-theme="light"] .token.char,
    html[data-theme="light"] .token.inserted,
    html[data-theme="light"] .token.attr-value {
      color: #16a34a;                      /* green-600 */
    }

    /* Numbers, booleans — amber/orange */
    html[data-theme="light"] .token.number,
    html[data-theme="light"] .token.boolean {
      color: #c2410c;                      /* orange-700 */
    }

    /* Functions, method names — vivid blue */
    html[data-theme="light"] .token.function,
    html[data-theme="light"] .token.function-variable {
      color: #1d4ed8;                      /* blue-700 */
    }

    /* Class names, types — pink/magenta */
    html[data-theme="light"] .token.class-name,
    html[data-theme="light"] .token.maybe-class-name,
    html[data-theme="light"] .token.builtin {
      color: #be185d;                      /* pink-700 */
    }

    /* Variables, parameters — teal */
    html[data-theme="light"] .token.variable,
    html[data-theme="light"] .token.parameter {
      color: #0f766e;                      /* teal-700 */
    }

    /* Properties — cyan/teal */
    html[data-theme="light"] .token.property {
      color: #0369a1;                      /* sky-700 */
    }

    /* Operators — dark teal */
    html[data-theme="light"] .token.operator,
    html[data-theme="light"] .token.entity,
    html[data-theme="light"] .token.url {
      color: #0891b2;                      /* cyan-600 */
    }

    /* HTML/XML tags — red */
    html[data-theme="light"] .token.tag,
    html[data-theme="light"] .token.deleted {
      color: #dc2626;                      /* red-600 */
    }

    /* HTML attribute names — purple */
    html[data-theme="light"] .token.attr-name {
      color: #9333ea;                      /* purple-600 */
    }

    /* Punctuation — dark neutral */
    html[data-theme="light"] .token.punctuation {
      color: #374151;                      /* gray-700 */
    }

    /* Regex — deep pink */
    html[data-theme="light"] .token.regex {
      color: #be185d;
    }

    /* Decorators / annotations — indigo */
    html[data-theme="light"] .token.decorator,
    html[data-theme="light"] .token.annotation {
      color: #4338ca; font-style: italic; /* indigo-700 */
    }

    /* Constants — amber */
    html[data-theme="light"] .token.constant,
    html[data-theme="light"] .token.symbol {
      color: #b45309;                      /* amber-700 */
    }

    html[data-theme="light"] .token.important,
    html[data-theme="light"] .token.bold   { font-weight: bold; }
    html[data-theme="light"] .token.italic { font-style: italic; }
    html[data-theme="light"] .token.namespace { opacity: 0.75; }

    /* JSON keys (property inside object) */
    html[data-theme="light"] .language-json .token.property {
      color: #1d4ed8;
    }

    /* SQL keywords */
    html[data-theme="light"] .language-sql .token.keyword {
      color: #7c3aed; font-weight: 700;
    }

    /* Bash builtins & variables */
    html[data-theme="light"] .language-bash .token.function { color: #1d4ed8; }
    html[data-theme="light"] .language-bash .token.variable { color: #0f766e; }

    /* ── Inline code & footer code ── */
    html[data-theme="light"] .code-inline {
      background:   #eef1f7;
      border-color: rgba(37,53,84,0.13);
      color:        #c7254e;
    }
    html[data-theme="light"] .footer-note code {
      background: rgba(229,237,248,0.88); color: #0550ae;
    }
    html[data-theme="light"] .msg-body .md-h1,
    html[data-theme="light"] .msg-body .md-h2,
    html[data-theme="light"] .msg-body .md-h3 { color: #1d4ed8; }
    html[data-theme="light"] .msg-body .md-bq { border-left-color: #2563eb; background: rgba(37,99,235,0.06); }
    html[data-theme="light"] .msg-body strong { color: #162338; }
    html[data-theme="light"] .msg-body em     { color: #334155; }
    html[data-theme="light"] .param-label  { color: var(--muted); }
    html[data-theme="light"] .param-value  { color: var(--accent); }
    html[data-theme="light"] input[type="range"] {
      background: rgba(37,53,84,0.15);
    }
    html[data-theme="light"] .scroll-to-bottom { box-shadow: 0 6px 20px rgba(37,53,84,0.2); }

    /* ── Help / User Guide Modal ── */
    .help-modal-backdrop {
      position: fixed;
      inset: 0;
      z-index: 120;
      display: none;
      align-items: center;
      justify-content: center;
      padding: 24px;
      background: rgba(2, 6, 23, 0.74);
      backdrop-filter: blur(10px);
    }
    .help-modal-backdrop.open { display: flex; }
    .help-modal {
      width: min(1020px, 100%);
      max-height: 88vh;
      overflow: hidden;
      display: flex;
      flex-direction: column;
      border-radius: 26px;
      border: 1px solid rgba(148,163,184,0.22);
      background:
        radial-gradient(circle at top right, rgba(96,165,250,0.14), transparent 30%),
        linear-gradient(180deg, rgba(15,23,42,0.96), rgba(2,6,23,0.98));
      box-shadow: 0 28px 80px rgba(0,0,0,0.45);
    }
    .help-modal-header {
      display: flex;
      align-items: flex-start;
      justify-content: space-between;
      gap: 16px;
      padding: 22px 24px 18px;
      border-bottom: 1px solid var(--line);
    }
    .help-modal-title h3 {
      margin: 0 0 6px;
      font-size: 1.28rem;
      color: var(--text);
      letter-spacing: 0.2px;
    }
    .help-modal-title .tiny {
      max-width: 760px;
      line-height: 1.55;
    }
    .help-modal-body {
      overflow: auto;
      padding: 22px 24px 26px;
      display: grid;
      gap: 18px;
    }
    .help-hero {
      padding: 18px 18px 16px;
      border-radius: 20px;
      border: 1px solid rgba(96,165,250,0.18);
      background: linear-gradient(135deg, rgba(30,41,59,0.82), rgba(15,23,42,0.92));
      box-shadow: inset 0 1px 0 rgba(255,255,255,0.04);
    }
    .help-hero strong {
      color: #cfe5ff;
    }
    .help-grid {
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 14px;
    }
    .help-card,
    .help-section {
      border-radius: 18px;
      border: 1px solid var(--line);
      background: rgba(15,23,42,0.68);
      padding: 16px 18px;
    }
    .help-card h4,
    .help-section h4 {
      margin: 0 0 10px;
      color: #dbeafe;
      font-size: 1rem;
      letter-spacing: 0.15px;
    }
    .help-card p,
    .help-section p {
      margin: 0;
      color: var(--text);
      line-height: 1.65;
    }
    .help-section ul,
    .help-card ul,
    .help-section ol {
      margin: 10px 0 0;
      padding-left: 20px;
      line-height: 1.7;
    }
    .help-section li,
    .help-card li {
      margin-bottom: 6px;
    }
    .help-kbd {
      display: inline-flex;
      align-items: center;
      justify-content: center;
      min-width: 34px;
      padding: 4px 8px;
      margin: 0 4px 4px 0;
      border-radius: 10px;
      border: 1px solid rgba(148,163,184,0.26);
      background: rgba(30,41,59,0.88);
      color: #e2e8f0;
      font-size: 0.82rem;
      font-weight: 700;
      box-shadow: inset 0 -2px 0 rgba(255,255,255,0.03);
    }
    .help-chip-row {
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      margin-top: 12px;
    }
    .help-chip {
      display: inline-flex;
      align-items: center;
      gap: 8px;
      padding: 8px 12px;
      border-radius: 999px;
      border: 1px solid rgba(96,165,250,0.2);
      background: rgba(59,130,246,0.08);
      color: #dbeafe;
      font-size: 0.85rem;
    }
    .help-steps {
      display: grid;
      gap: 12px;
      margin-top: 12px;
    }
    .help-step {
      border-radius: 16px;
      border: 1px solid rgba(148,163,184,0.16);
      background: rgba(15,23,42,0.44);
      padding: 14px 14px 12px;
    }
    .help-step strong {
      display: block;
      margin-bottom: 6px;
      color: #bfdbfe;
    }
    .help-two-col {
      display: grid;
      grid-template-columns: 1.15fr 0.85fr;
      gap: 14px;
    }
    .help-note {
      margin-top: 12px;
      padding: 12px 14px;
      border-radius: 14px;
      border: 1px solid rgba(250,204,21,0.18);
      background: rgba(250,204,21,0.06);
      color: #f8fafc;
      line-height: 1.65;
    }
    .help-footer-row {
      display: flex;
      flex-wrap: wrap;
      gap: 10px;
      align-items: center;
      justify-content: space-between;
      margin-top: 4px;
      color: var(--muted);
      font-size: 0.92rem;
    }
    html[data-theme="light"] .help-modal-backdrop {
      background: rgba(15, 23, 42, 0.42);
    }
    html[data-theme="light"] .help-modal {
      background:
        radial-gradient(circle at top right, rgba(59,130,246,0.12), transparent 28%),
        linear-gradient(180deg, rgba(255,255,255,0.98), rgba(248,250,252,0.99));
      box-shadow: 0 24px 70px rgba(15,23,42,0.18);
    }
    html[data-theme="light"] .help-card,
    html[data-theme="light"] .help-section,
    html[data-theme="light"] .help-step {
      background: rgba(255,255,255,0.88);
      border-color: rgba(37,53,84,0.12);
    }
    html[data-theme="light"] .help-hero {
      background: linear-gradient(135deg, rgba(239,246,255,0.98), rgba(248,250,252,0.98));
      border-color: rgba(37,99,235,0.14);
    }
    html[data-theme="light"] .help-card h4,
    html[data-theme="light"] .help-section h4,
    html[data-theme="light"] .help-step strong,
    html[data-theme="light"] .help-modal-title h3 {
      color: #12325b;
    }
    html[data-theme="light"] .help-hero strong,
    html[data-theme="light"] .help-chip {
      color: #12325b;
    }
    html[data-theme="light"] .help-chip {
      background: rgba(37,99,235,0.08);
      border-color: rgba(37,99,235,0.14);
    }
    html[data-theme="light"] .help-kbd {
      background: #eef4fb;
      color: #12325b;
      border-color: rgba(37,53,84,0.14);
    }

    /* ── Responsive ── */
    @media (max-width: 1100px) {
      .app { grid-template-columns: 1fr; }
      .sidebar { height: unset; position: static; max-height: none; overflow-y: visible; }
      .chat-panel { height: 75vh; }
      .msg { max-width: 100%; }
    }
    @media (max-width: 600px) {
      body { padding: 10px; }
      .app { gap: 12px; }
      .chat-panel { height: 65vh; }
      .composer textarea { min-height: 100px; height: 100px; }
      .help-modal-backdrop { padding: 12px; }
      .help-modal-header, .help-modal-body { padding-left: 16px; padding-right: 16px; }
      .help-grid, .help-two-col { grid-template-columns: 1fr; }
    }
  </style>
</head>
<body>

  <div class="drop-overlay" id="dropOverlay">📂 Άσε τα αρχεία εδώ</div>

  <div class="app">

    <!-- ── Sidebar ────────────────────────────────────────────────── -->
    <aside class="sidebar card">
      <div class="title">
        <h1>☁️ __APP_TITLE__</h1>
        <div class="pill">v4.0</div>
      </div>

      <div class="group">
        <label class="label" for="modelSelect">Μοντέλο Ollama</label>
        <input id="modelSearchInput" type="search" placeholder="Αναζήτηση μοντέλου..."
               title="Γράψε για φιλτράρισμα της λίστας μοντέλων" style="margin-bottom:8px;" />
        <select id="modelSelect" title="Επίλεξε cloud model"></select>
        <div class="tiny muted" style="margin-top:8px;">
          Η λίστα ανακτά τα ακριβή model names του official Ollama direct API catalog από το διαδίκτυο.
          Η αναζήτηση φιλτράρει τη λίστα χωρίς να επηρεάζει τη φόρτωση των μοντέλων.
        </div>
      </div>

      <div class="group">
        <label class="label" for="modelSortSelect">Ταξινόμηση μοντέλων</label>
        <select id="modelSortSelect" title="Επίλεξε κριτήριο αξιολόγησης για ταξινόμηση μοντέλων">
          <option value="overall">Overall / Καλύτερο συνολικά</option>
          <option value="coding">Coding / Προγραμματισμός</option>
          <option value="reasoning">Reasoning / Σκέψη</option>
          <option value="context">Long Context / Max Length</option>
          <option value="vision">Vision / Εικόνες</option>
          <option value="speed">Speed / Ταχύτητα</option>
          <option value="newest">Newest / Πιο νέο</option>
        </select>
        <div class="tiny muted" style="margin-top:8px;">
          Η κατάταξη γίνεται ευρετικά από metadata του direct catalog και model details, όπως context length,
          capabilities, μέγεθος μοντέλου και πρόσφατη ενημέρωση.
        </div>
      </div>

      <div class="group">
        <label class="label" for="thinkModeSelect">Thinking Mode</label>
        <select id="thinkModeSelect" title="Ρύθμιση thinking/reasoning mode του μοντέλου">
          <option value="auto">Auto</option>
          <option value="on" selected>On</option>
          <option value="off">Off</option>
          <option value="low">Low</option>
          <option value="medium">Medium</option>
          <option value="high">High</option>
        </select>
        <div class="tiny muted" id="thinkingSupportInfo" style="margin-top:8px;">
          Η λίστα επιλογών προσαρμόζεται αυτόματα ανάλογα με το επιλεγμένο μοντέλο και το thinking mode που υποστηρίζει.
        </div>
        <div class="btn-row" style="margin-top:8px;">
          <button class="secondary btn-full" id="confirmThinkingProfileBtn" title="Επιβεβαίωση του τρέχοντος thinking support profile για το επιλεγμένο μοντέλο">✅ Επιβεβαίωση Profile</button>
        </div>
      </div>

      <div class="group">
        <label class="label" for="ensembleModeSelect">Dual Model Ensemble</label>
        <select id="ensembleModeSelect" title="Επίλεξε λειτουργία dual model ensemble">
          <option value="off">Off</option>
          <option value="auto" selected>Auto</option>
          <option value="manual">Manual</option>
        </select>
        <input id="helperSearchInput" type="search" placeholder="Αναζήτηση helper model..."
               title="Γράψε για φιλτράρισμα της λίστας helper models" style="margin-top:8px; margin-bottom:8px;" />
        <select id="helperModelSelect" title="Επίλεξε βοηθητικό μοντέλο" disabled></select>
        <div class="tiny muted" id="ensembleModeInfo" style="margin-top:8px;">
          Off: μόνο το κύριο μοντέλο. Auto: αυτόματη επιλογή helper model. Manual: διαλέγεις εσύ το δεύτερο μοντέλο από τη λίστα.
        </div>
      </div>

      <div class="group">
        <label class="label" for="apiKeyInput">Ollama API Key</label>
        <input id="apiKeyInput" type="password" placeholder="ollama_..." autocomplete="off"
               title="API key για direct Ollama Cloud API" />
        <div class="btn-row" style="margin-top:10px;">
          <button class="primary" id="saveApiKeyBtn" title="Αποθήκευση API key σε αρχείο ρυθμίσεων">💾 Save Key</button>
          <button class="secondary" id="clearApiKeyBtn" title="Καθαρισμός αποθηκευμένου API key">🗑 Clear Key</button>
        </div>
        <div class="tiny muted" id="apiKeyInfo" style="margin-top:8px;">
          Το API key αποθηκεύεται σε τοπικό αρχείο ρυθμίσεων δίπλα στο .py.
        </div>
      </div>

      <div class="group">
        <label class="label" for="promptProfileSelect">Prompt Profile</label>
        <select id="promptProfileSelect" title="Επίλεξε preset system prompt profile για το studio"></select>
        <div class="btn-row" style="margin-top:10px;">
          <button class="primary" id="activatePromptProfileBtn" title="Επαναφορά του system prompt στο επιλεγμένο prompt profile">⚡ Apply Profile</button>
          <button class="secondary" id="savePromptSetupBtn" title="Αποθήκευση prompt profile, custom prompt και visualization engine στο settings file">💾 Save Setup</button>
        </div>
        <div class="tiny muted" id="promptProfileInfo" style="margin-top:8px;">
          Με την αλλαγή profile ενημερώνεται αυτόματα και το System Prompt στο UI.
        </div>
      </div>

      <div class="group">
        <label class="label" for="visualizationEngineSelect">Visualization Engine</label>
        <select id="visualizationEngineSelect" title="Επίλεξε πότε θα προτιμάται SVG ή Python / matplotlib"></select>
        <div class="tiny muted" id="visualizationEngineInfo" style="margin-top:8px;">
          Auto: έξυπνη επιλογή · SVG: διαγράμματα/σχήματα · Python Plot: πραγματικά data plots με matplotlib.
        </div>
      </div>

      <div class="group">
        <label class="label" for="systemPrompt">System Prompt</label>
        <textarea id="systemPrompt" title="System prompt της συνεδρίας">__SYSTEM_PROMPT__</textarea>
        <div class="tiny muted" style="margin-top:6px;">
          Η αλλαγή Prompt Profile ενημερώνει αυτόματα αυτό το πεδίο. Μπορείς έπειτα να κάνεις χειροκίνητο override και να το αποθηκεύσεις με Save Setup.
        </div>
      </div>

      <div class="group">
        <label class="label" for="fileInput">Αρχεία (ή drag & drop)</label>
        <input id="fileInput" type="file" multiple accept="__ACCEPTED_TYPES__"
               title="Επίλεξε αρχεία για context" />
        <div class="attachment-list" id="selectedFiles"></div>
        <div class="btn-row" style="margin-top:10px;">
          <button class="secondary" id="clearFilesBtn"    title="Αφαίρεση επιλεγμένων αρχείων">📎 Καθαρισμός</button>
          <button class="primary"   id="refreshModelsBtn" title="Ανάκτηση τελευταίας λίστας cloud models">🔄 Refresh Models</button>
        </div>
        <ul class="support-list">
          <li>Drag &amp; Drop υποστηρίζεται παντού.</li>
          <li>Εικόνες → natively σε vision models.</li>
          <li>TXT, PY, MD, JSON, CSV, PDF κ.ά. → context.</li>
        </ul>
      </div>

      <div class="group">
        <div class="label">⚙️ Παράμετροι μοντέλου</div>

        <div class="param-row">
          <span class="param-label" title="Δημιουργικότητα απάντησης (0=ντετερμινιστικό, 2=πολύ τυχαίο)">Temperature</span>
          <input type="range" id="paramTemp" min="0" max="2" step="0.05" value="0.8"
                 title="Temperature: 0 – 2" />
          <span class="param-value" id="paramTempVal">0.80</span>
        </div>

        <div class="param-row">
          <span class="param-label" title="Nucleus sampling: διατηρεί tokens που καλύπτουν το X% της πιθανότητας">Top-P</span>
          <input type="range" id="paramTopP" min="0.01" max="1" step="0.01" value="0.9"
                 title="Top-P: 0.01 – 1.00" />
          <span class="param-value" id="paramTopPVal">0.90</span>
        </div>

        <div class="param-row" style="align-items:flex-start; flex-direction:column; gap:6px;">
          <span class="param-label" title="Seed για αναπαραγώγιμα αποτελέσματα (-1 = τυχαίο)">Seed</span>
          <div class="param-seed-wrap" style="width:100%;">
            <input type="text" id="paramSeed" value="-1" placeholder="-1 (τυχαίο)"
                   title="Seed: αριθμός ≥ 0 για αναπαραγώγιμο, -1 για τυχαίο" />
            <button class="secondary" id="resetParamsBtn" style="padding:8px 10px; font-size:0.8rem; white-space:nowrap;"
                    title="Επαναφορά προεπιλεγμένων παραμέτρων">↺ Reset</button>
          </div>
        </div>

        <div class="param-row" style="align-items:flex-start; flex-direction:column; gap:6px;">
          <span class="param-label" title="Μέγιστο context / max length (num_ctx) για το επιλεγμένο μοντέλο">Max Length (num_ctx)</span>
          <div class="param-seed-wrap" style="width:100%;">
            <input type="number" id="paramNumCtx" min="256" step="256" value="" placeholder="κενό = auto / default"
                   title="Context window / max length για το επιλεγμένο μοντέλο" />
            <button class="secondary" id="clearNumCtxBtn" style="padding:8px 10px; font-size:0.8rem; white-space:nowrap;"
                    title="Καθαρισμός custom max length για το τρέχον μοντέλο">Auto</button>
          </div>
          <div class="tiny muted" id="paramNumCtxInfo">Ισχύει ξεχωριστά για κάθε μοντέλο.</div>
        </div>
      </div>

      <div class="group">
        <div class="label">Κατάσταση</div>
        <div id="modelInfo" class="muted tiny">Φόρτωση λίστας μοντέλων...</div>
      </div>

      <div class="btn-row">
        <button class="secondary" id="resetSystemPromptBtn" title="Επαναφορά default system prompt">🧠 Reset Prompt</button>
        <button class="secondary" id="clearChatBtn"         title="Καθαρισμός chat και uploads">🧹 Clear Chat</button>
      </div>
      <div class="btn-row">
        <button class="secondary" id="reloadSessionBtn"    title="Επαναφόρτωση ιστορικού από server">♻️ Reload Session</button>
        <button class="secondary" id="copySystemPromptBtn" title="Αντιγραφή system prompt στο clipboard">📋 Copy Prompt</button>
      </div>
      <div class="btn-row">
        <button class="secondary" id="exportChatBtn"  title="Εξαγωγή συνομιλίας ως Markdown αρχείο">💾 Export .md</button>
        <button class="secondary" id="autoScrollBtn"  title="Ενεργοποίηση/απενεργοποίηση αυτόματης κύλισης">📜 Auto-Scroll: ON</button>
      </div>

      <button class="secondary btn-full" id="helpGuideBtn" title="Προβολή πλήρων οδηγιών χρήσης της εφαρμογής">📘 Οδηγίες Χρήσης</button>
      <button class="secondary btn-full" id="themeToggleBtn" title="Εναλλαγή dark/light theme">☀️ Light Theme</button>

      <div class="footer-note">
        Direct Cloud API: βάλε το API key στο πεδίο του GUI και πάτησε <b>Save Key</b> ή όρισε <code>OLLAMA_API_KEY</code><br><br>
        Settings file: <code>ollama_cloud_chat_settings.json</code> δίπλα στο .py<br><br>
        Εκτέλεση: <code>python ollama_cloud_chat.py [--port N] [--no-browser]</code><br><br>
        <span style="display:block; margin-top:10px; padding-top:10px; border-top:1px solid var(--line); opacity:0.85; font-size:0.9rem; text-align:center; letter-spacing:0.3px;">&copy; Ευάγγελος Πεφάνης</span>
      </div>
    </aside>

    <!-- ── Chat panel ─────────────────────────────────────────────── -->
    <main class="chat-panel card">
      <div class="chat-header">
        <div class="header-left">
          <h2>Συνομιλία</h2>
          <div class="tiny muted">Enter αποστολή · Shift+Enter νέα γραμμή · Ctrl+Enter εναλλακτικό</div>
        </div>
        <div class="status-wrap">
          <span class="badge" id="backendBadge"       title="Κατάσταση direct Ollama Cloud API">⬤ Cloud API</span>
          <span class="badge" id="msgCountBadge">0 μηνύματα</span>
          <span class="badge" id="selectedModelBadge">Μοντέλο: -</span>
          <span class="badge" id="sourceBadge">Πηγή: -</span>
          <span class="badge" id="tokensPerSecBadge"  title="Ταχύτητα τελευταίας απάντησης" style="display:none;"></span>
          <span class="badge" id="streamBadge">Έτοιμο</span>
        </div>
      </div>

      <div class="reasoning-panel" id="reasoningPanel" aria-live="polite">
        <div class="reasoning-head">
          <div class="reasoning-title-wrap">
            <div class="reasoning-title">🧠 <span>Realtime βαθιά σκέψη</span></div>
            <div class="reasoning-meta" id="reasoningMeta">Αναμονή για thinking stream…</div>
          </div>
          <button class="secondary reasoning-toggle-btn" id="toggleReasoningBtn" title="Εμφάνιση ή απόκρυψη του panel σκέψης">🙈 Απόκρυψη</button>
        </div>
        <pre class="reasoning-body" id="reasoningContent"></pre>
      </div>

      <div class="messages-wrap">
        <div class="messages" id="messages">
          <div class="empty-state" id="emptyState">
            <h3 style="margin-top:0;">Έτοιμο για συνομιλία</h3>
            <div>Γράψε το δικό σου prompt και, αν θέλεις, πρόσθεσε αρχεία για context.</div>
            <div style="margin-top:10px;" class="tiny">
              Enter αποστολή · Shift+Enter νέα γραμμή · Ctrl+Enter εναλλακτικό
            </div>
          </div>
        </div>
        <button class="scroll-to-bottom" id="scrollToBottomBtn" title="Μετάβαση στο τέλος">↓ Τέλος</button>
      </div>

      <div class="composer">
        <textarea id="userInput" placeholder="Γράψε εδώ το user prompt σου…"
                  title="User prompt — Enter αποστολή, Shift+Enter νέα γραμμή, Ctrl+Enter εναλλακτικό"></textarea>
        <div class="composer-footer">
          <div class="composer-left">
            <span class="char-counter" id="charCounter">0 χαρ. / 0 λέξεις</span>
            <span class="helper" id="helperText">Enter · Shift+Enter · Ctrl+Enter</span>
          </div>
          <div class="composer-right">
            <button class="secondary" id="stopBtn" disabled title="Διακοπή streaming">⏹ Stop</button>
            <button class="primary"   id="sendBtn"          title="Αποστολή μηνύματος (Enter)">🚀 Αποστολή</button>
          </div>
        </div>
      </div>
    </main>

  </div><!-- /.app -->

  <div class="help-modal-backdrop" id="helpModal" aria-hidden="true">
    <div class="help-modal" role="dialog" aria-modal="true" aria-labelledby="helpModalTitle">
      <div class="help-modal-header">
        <div class="help-modal-title">
          <h3 id="helpModalTitle">📘 Οδηγίες Χρήσης του Ollama Cloud Chat Studio</h3>
          <div class="tiny muted">
            Αναλυτικός οδηγός για το τι κάνει κάθε περιοχή της εφαρμογής, πώς ξεκινάς γρήγορα,
            πώς χρησιμοποιείς τα μοντέλα, τα αρχεία, τα prompt profiles, το dual-model ensemble και τα εργαλεία του chat.
          </div>
        </div>
        <button class="secondary" id="closeHelpModalBtn" type="button" title="Κλείσιμο οδηγού χρήσης">✖ Κλείσιμο</button>
      </div>
      <div class="help-modal-body">
        <section class="help-hero">
          <p style="margin:0; line-height:1.75;">
            <strong>Η εφαρμογή είναι ένα ολοκληρωμένο browser-based studio συνομιλίας για Ollama Cloud models.</strong>
            Σου επιτρέπει να επιλέγεις μοντέλο, να ρυθμίζεις thinking mode, να χρησιμοποιείς δεύτερο helper model,
            να επισυνάπτεις αρχεία για context, να πειράζεις παραμέτρους παραγωγής, να διαχειρίζεσαι system prompt
            και να εξάγεις τη συνομιλία σου. Ο οδηγός αυτός είναι εδώ για να μπορείς να βρίσκεις όλα τα βασικά βήματα
            μέσα από το UI χωρίς να χρειάζεται εξωτερικό manual.
          </p>
          <div class="help-chip-row">
            <span class="help-chip">☁️ Direct Ollama Cloud API</span>
            <span class="help-chip">🧠 Thinking / Reasoning</span>
            <span class="help-chip">🤝 Dual-Model Ensemble</span>
            <span class="help-chip">📎 Attachments & Context</span>
            <span class="help-chip">💾 Export & Session Tools</span>
          </div>
        </section>

        <div class="help-grid">
          <section class="help-card">
            <h4>Γρήγορη εκκίνηση</h4>
            <div class="help-steps">
              <div class="help-step">
                <strong>1. Βάλε API key</strong>
                Συμπλήρωσε το πεδίο <b>Ollama API Key</b> και πάτησε <b>💾 Save Key</b> ώστε η εφαρμογή να καλεί το direct cloud API.
              </div>
              <div class="help-step">
                <strong>2. Επίλεξε μοντέλο</strong>
                Από το <b>Μοντέλο Ollama</b> διάλεξε το model που θέλεις. Μπορείς να γράψεις στο πεδίο αναζήτησης για γρήγορο φιλτράρισμα.
              </div>
              <div class="help-step">
                <strong>3. Ρύθμισε prompt / παραμέτρους</strong>
                Προαιρετικά διάλεξε Prompt Profile, Visualization Engine, Thinking Mode και παραμέτρους όπως Temperature, Top-P, Seed και Max Length.
              </div>
              <div class="help-step">
                <strong>4. Στείλε μήνυμα</strong>
                Γράψε το prompt σου στο κάτω πλαίσιο και πάτησε <b>🚀 Αποστολή</b> ή χρησιμοποίησε τα πλήκτρα <span class="help-kbd">Enter</span> / <span class="help-kbd">Ctrl+Enter</span>.
              </div>
            </div>
          </section>

          <section class="help-card">
            <h4>Βασικές συντομεύσεις πληκτρολογίου</h4>
            <p>
              <span class="help-kbd">Enter</span> Αποστολή μηνύματος<br>
              <span class="help-kbd">Shift+Enter</span> Νέα γραμμή στο prompt<br>
              <span class="help-kbd">Ctrl+Enter</span> Εναλλακτική αποστολή<br>
              <span class="help-kbd">Esc</span> Κλείσιμο αυτού του οδηγού
            </p>
            <div class="help-note">
              Κατά τη ροή απάντησης μπορείς να πατήσεις <b>⏹ Stop</b> για να διακόψεις το streaming. Αν έχεις ενεργό Auto-Scroll,
              η συνομιλία παραμένει αυτόματα στο τέλος καθώς έρχεται η απάντηση.
            </div>
          </section>
        </div>

        <div class="help-two-col">
          <section class="help-section">
            <h4>Αριστερό panel: τι κάνει κάθε ενότητα</h4>
            <ul>
              <li><b>Μοντέλο Ollama</b>: επιλέγεις το κύριο cloud model που θα απαντήσει.</li>
              <li><b>Ταξινόμηση μοντέλων</b>: αλλάζεις το heuristic ranking με βάση overall, coding, reasoning, long context, vision, speed ή newest.</li>
              <li><b>Thinking Mode</b>: ορίζεις αν το μοντέλο θα κάνει reasoning και σε ποια ένταση, όπου αυτό υποστηρίζεται.</li>
              <li><b>Dual Model Ensemble</b>: μπορείς να δουλέψεις μόνο με το κύριο μοντέλο ή με δεύτερο helper model σε Auto ή Manual mode.</li>
              <li><b>Ollama API Key</b>: αποθήκευση / καθαρισμός του API key στο τοπικό settings file.</li>
              <li><b>Prompt Profile</b>: φορτώνει έτοιμο στυλ system prompt για διαφορετικά use cases.</li>
              <li><b>Visualization Engine</b>: δίνει προτεραιότητα σε Auto, SVG ή Python Plot / matplotlib.</li>
              <li><b>System Prompt</b>: πλήρες prompt της συνεδρίας, το οποίο μπορείς να προσαρμόσεις και να αποθηκεύσεις.</li>
              <li><b>Αρχεία</b>: προσθήκη αρχείων με file picker ή drag &amp; drop για παροχή context στο μοντέλο.</li>
              <li><b>Παράμετροι μοντέλου</b>: Temperature, Top-P, Seed και Max Length (num_ctx).</li>
            </ul>
          </section>

          <section class="help-section">
            <h4>Chat panel: τι βλέπεις δεξιά</h4>
            <ul>
              <li><b>Κατάσταση / badges</b>: ενημέρωση για backend, πλήθος μηνυμάτων, επιλεγμένο μοντέλο, source και ταχύτητα.</li>
              <li><b>Realtime βαθιά σκέψη</b>: ειδικό panel που προβάλλει thinking/reasoning stream όταν υποστηρίζεται.</li>
              <li><b>Ιστορικό μηνυμάτων</b>: όλες οι απαντήσεις της συνομιλίας, με rendering Markdown, μαθηματικών, code blocks και SVG previews.</li>
              <li><b>Scroll to Bottom</b>: γρήγορη μετάβαση στο τέλος όταν έχεις ανέβει ψηλότερα στη συζήτηση.</li>
              <li><b>Composer</b>: το κάτω πλαίσιο όπου γράφεις prompt, βλέπεις μετρητή χαρακτήρων/λέξεων και στέλνεις μήνυμα.</li>
            </ul>
          </section>
        </div>

        <section class="help-section">
          <h4>Prompt Profiles και System Prompt</h4>
          <p>
            Τα <b>Prompt Profiles</b> είναι προκαθορισμένες στρατηγικές system prompt. Επίλεξε ένα profile και πάτησε
            <b>⚡ Apply Profile</b> για να γεμίσει αυτόματα το πεδίο <b>System Prompt</b>. Με το <b>💾 Save Setup</b>
            αποθηκεύεις το επιλεγμένο profile, τυχόν custom αλλαγές στο prompt και το Visualization Engine.
          </p>
          <ul>
            <li>Χρησιμοποίησε <b>🧠 Reset Prompt</b> για να επανέλθεις στο ενεργό profile.</li>
            <li>Χρησιμοποίησε <b>📋 Copy Prompt</b> για να αντιγράψεις το prompt στο clipboard.</li>
            <li>Πείραξε χειροκίνητα το System Prompt μόνο όταν θέλεις στοχευμένο override της συμπεριφοράς του μοντέλου.</li>
          </ul>
        </section>

        <section class="help-section">
          <h4>Thinking Mode και Realtime βαθιά σκέψη</h4>
          <p>
            Το <b>Thinking Mode</b> καθορίζει αν το μοντέλο θα χρησιμοποιήσει reasoning / thinking stream. Η εφαρμογή
            προσπαθεί να προσαρμόσει αυτόματα τις διαθέσιμες επιλογές ανάλογα με το επιλεγμένο μοντέλο.
          </p>
          <ul>
            <li><b>Auto</b>: η εφαρμογή και το μοντέλο αποφασίζουν δυναμικά το βέλτιστο mode.</li>
            <li><b>On / Off</b>: σταθερό άνοιγμα ή κλείσιμο reasoning, όπου επιτρέπεται.</li>
            <li><b>Low / Medium / High</b>: επίπεδα έντασης σκέψης όταν το model το υποστηρίζει ρητά.</li>
            <li>Το κουμπί <b>✅ Επιβεβαίωση Profile</b> σε βοηθά να επαληθεύσεις το τρέχον thinking support profile.</li>
            <li>Το panel <b>Realtime βαθιά σκέψη</b> μπορεί να εμφανιστεί ή να κρυφτεί από το κουμπί <b>🙈 Απόκρυψη</b>.</li>
          </ul>
        </section>

        <section class="help-section">
          <h4>Dual Model Ensemble: πότε και πώς να το χρησιμοποιείς</h4>
          <p>
            Η λειτουργία <b>Dual Model Ensemble</b> επιτρέπει στην εφαρμογή να συμβουλεύεται δεύτερο helper model πριν
            ολοκληρώσει την απάντηση το κύριο μοντέλο. Είναι χρήσιμη όταν θέλεις καλύτερο cross-check, επιπλέον reasoning,
            καλύτερη διαχείριση κώδικα ή υποβοήθηση για εικόνες/μεγάλο context.
          </p>
          <ul>
            <li><b>Off</b>: χρησιμοποιείται μόνο το κύριο μοντέλο.</li>
            <li><b>Auto</b>: η εφαρμογή επιλέγει αυτόματα helper model με βάση το task.</li>
            <li><b>Manual</b>: εσύ επιλέγεις ρητά το helper model από το dropdown.</li>
          </ul>
          <div class="help-note">
            Για γενική χρήση κράτησε το <b>Auto</b>. Για εξειδικευμένα σενάρια, π.χ. debugging ή code review, μπορείς να περάσεις σε <b>Manual</b>
            και να διαλέξεις δεύτερο μοντέλο που σου ταιριάζει περισσότερο.
          </div>
        </section>

        <div class="help-grid">
          <section class="help-section">
            <h4>Αρχεία, drag &amp; drop και context</h4>
            <ul>
              <li>Μπορείς να προσθέσεις πολλαπλά αρχεία από το <b>Αρχεία</b> ή με drag &amp; drop μέσα στο παράθυρο.</li>
              <li>Εικόνες αξιοποιούνται καλύτερα από vision-capable models.</li>
              <li>TXT, PY, MD, JSON, CSV, PDF και άλλα υποστηριζόμενα αρχεία μπαίνουν ως επιπλέον context.</li>
              <li>Με το <b>📎 Καθαρισμός</b> αφαιρείς τα προσωρινά επιλεγμένα αρχεία πριν από την αποστολή.</li>
            </ul>
          </section>

          <section class="help-section">
            <h4>Παράμετροι παραγωγής απάντησης</h4>
            <ul>
              <li><b>Temperature</b>: χαμηλά για πιο σταθερές απαντήσεις, ψηλά για πιο δημιουργικές.</li>
              <li><b>Top-P</b>: έλεγχος nucleus sampling.</li>
              <li><b>Seed</b>: ίδιο seed δίνει πιο αναπαραγώγιμες απαντήσεις.</li>
              <li><b>Max Length (num_ctx)</b>: custom context window ανά μοντέλο, όπου χρειάζεται.</li>
              <li><b>↺ Reset</b>: επαναφορά βασικών παραμέτρων στις προεπιλογές.</li>
              <li><b>Auto</b> δίπλα στο Max Length: καθαρίζει την custom τιμή και αφήνει το default behavior.</li>
            </ul>
          </section>
        </div>

        <section class="help-section">
          <h4>Κουμπιά χρησιμότητας και καθημερινή χρήση</h4>
          <ul>
            <li><b>🔄 Refresh Models</b>: ανανεώνει τη λίστα των διαθέσιμων official direct API models.</li>
            <li><b>🧹 Clear Chat</b>: καθαρίζει συνομιλία και uploads της τρέχουσας συνεδρίας.</li>
            <li><b>♻️ Reload Session</b>: ξαναφορτώνει ιστορικό από τον server.</li>
            <li><b>💾 Export .md</b>: εξαγωγή της συνομιλίας σε αρχείο Markdown.</li>
            <li><b>📜 Auto-Scroll</b>: ενεργοποιεί / απενεργοποιεί αυτόματη κύλιση προς τα κάτω.</li>
            <li><b>☀️ / 🌙 Theme</b>: εναλλαγή light και dark εμφάνισης.</li>
            <li><b>📘 Οδηγίες Χρήσης</b>: ανοίγει ξανά αυτό το αναλυτικό παράθυρο όποτε το χρειαστείς.</li>
          </ul>
        </section>

        <section class="help-section">
          <h4>Πρακτικές συμβουλές για καλύτερα αποτελέσματα</h4>
          <ul>
            <li>Δώσε σαφή prompt με ζητούμενα, πλαίσιο, μορφή απάντησης και γλώσσα εξόδου.</li>
            <li>Αν το task είναι τεχνικό ή επιστημονικό, επίλεξε κατάλληλο Prompt Profile πριν στείλεις το μήνυμα.</li>
            <li>Για κώδικα και debugging βοήθα το μοντέλο με σχετικά αρχεία, logs ή stack traces.</li>
            <li>Για διαγράμματα ή plots επίλεξε το σωστό <b>Visualization Engine</b>.</li>
            <li>Αν το μοντέλο παράγει υπερβολικά δημιουργικές απαντήσεις, μείωσε το <b>Temperature</b>.</li>
            <li>Αν δουλεύεις με μεγάλο υλικό, αξιοποίησε το <b>Max Length (num_ctx)</b> και κατάλληλα long-context models.</li>
          </ul>
        </section>

        <section class="help-section">
          <h4>Αν κάτι δεν λειτουργεί όπως περιμένεις</h4>
          <ol>
            <li>Έλεγξε ότι έχεις βάλει σωστό <b>Ollama API Key</b> και ότι πάτησες <b>Save Key</b>.</li>
            <li>Πάτησε <b>🔄 Refresh Models</b> για να φορτώσεις ξανά το catalog των μοντέλων.</li>
            <li>Δοκίμασε άλλο μοντέλο ή απενεργοποίησε προσωρινά το <b>Dual Model Ensemble</b>.</li>
            <li>Αν το prompt είναι περίπλοκο, κάν’ το πιο συγκεκριμένο ή χώρισέ το σε βήματα.</li>
            <li>Αν υπάρχουν πολλά αρχεία, αφαίρεσε όσα δεν είναι χρήσιμα για το τρέχον ερώτημα.</li>
            <li>Χρησιμοποίησε <b>🧹 Clear Chat</b> όταν θέλεις καθαρή νέα συνεδρία.</li>
          </ol>
        </section>

        <div class="help-footer-row">
          <span>Ο οδηγός αυτός είναι ενσωματωμένος στην εφαρμογή για άμεση on-screen βοήθεια.</span>
          <span>Χρήσιμος τόσο για πρώτο setup όσο και για καθημερινή χρήση.</span>
        </div>
      </div>
    </div>
  </div>

  <script>
    "use strict";

    // ── Constants ────────────────────────────────────────────────────────────
    const DEFAULT_SYSTEM_PROMPT  = __DEFAULT_SYSTEM_PROMPT_JSON__;
    const DEFAULT_PROMPT_PROFILES = __DEFAULT_PROMPT_PROFILES_JSON__;
    const DEFAULT_PROMPT_PROFILE_ID = __DEFAULT_PROMPT_PROFILE_ID_JSON__;
    const DEFAULT_VISUALIZATION_ENGINE = __DEFAULT_VISUALIZATION_ENGINE_JSON__;
    const VISUALIZATION_ENGINE_OPTIONS = __VISUALIZATION_ENGINE_OPTIONS_JSON__;
    const THEME_KEY              = "ollama_chat_theme_v2";
    const MODEL_KEY              = "ollama_chat_model_v2";
    const PARAMS_KEY             = "ollama_chat_params_v2";
    const THINK_MODE_KEY         = "ollama_chat_think_mode_v1";
    const THINK_PROFILE_CONFIRMATIONS_KEY = "ollama_chat_think_profile_confirmations_v1";
    const MODEL_SORT_KEY         = "ollama_chat_model_sort_v1";
    const ENSEMBLE_KEY           = "ollama_chat_ensemble_auto_v1";
    const ENSEMBLE_MODE_KEY      = "ollama_chat_ensemble_mode_v2";
    const ENSEMBLE_HELPER_KEY    = "ollama_chat_ensemble_helper_v1";
    const CHAR_WARN              = 8000;
    const DEFAULT_HELPER         = "Enter · Shift+Enter · Ctrl+Enter";
    const HEALTH_POLL_MS         = 15_000;  // polling interval για cloud API status
    const MODEL_REFRESH_POLL_MS  = 400;
    const BROWSER_SESSION_KEY    = "ollama_chat_browser_session_v1";
    const BROWSER_HEARTBEAT_MS   = 10_000;

    // ── App state ────────────────────────────────────────────────────────────
    const state = {
      isStreaming:            false,
      abortController:        null,
      currentAssistantNode:   null,
      currentThinkingText:    "",
      reasoningPanelVisible:  false,
      reasoningAutoOpen:      false,
      reasoningUserCollapsed: false,
      currentInlineThinkingOpen: true,
      reasoningStreamCompleted: false,
      models:                 [],
      selectedFiles:          [],
      theme:                  "dark",
      autoScroll:             true,
      chatHistory:            [],  // [{role, content, time}] — for export
      msgCount:               0,
      dragCounter:            0,   // reliable drag-leave detection
      lastTokensPerSec:       null,
      modelNumCtxByModel:     {},
      modelMaxNumCtxByModel:  {},
      modelMetaByModel:       {},
      modelDetailRequests:    {},
      currentThinkingProfile: null,
      confirmedThinkingProfilesByModel: {},
      helperModels:           [],
      ensembleMode:           "auto",
      promptProfiles:         Array.isArray(DEFAULT_PROMPT_PROFILES) ? DEFAULT_PROMPT_PROFILES : [],
      promptProfileMap:       {},
      activePromptProfile:    DEFAULT_PROMPT_PROFILE_ID || "scientific-technical",
      visualizationEngine:    DEFAULT_VISUALIZATION_ENGINE || "auto",
      modelSortCriterion:     "overall",
      lastModelsRefreshTs:    0,
      browserSessionId:       "",
      browserHeartbeatTimer:   null,
          };

    // ── DOM refs ─────────────────────────────────────────────────────────────
    const els = {
      modelSearchInput:     document.getElementById("modelSearchInput"),
      modelSelect:          document.getElementById("modelSelect"),
      modelSortSelect:      document.getElementById("modelSortSelect"),
      thinkModeSelect:      document.getElementById("thinkModeSelect"),
      thinkingSupportInfo:  document.getElementById("thinkingSupportInfo"),
      confirmThinkingProfileBtn: document.getElementById("confirmThinkingProfileBtn"),
      ensembleModeSelect:   document.getElementById("ensembleModeSelect"),
      helperSearchInput:    document.getElementById("helperSearchInput"),
      helperModelSelect:    document.getElementById("helperModelSelect"),
      ensembleModeInfo:     document.getElementById("ensembleModeInfo"),
      promptProfileSelect:  document.getElementById("promptProfileSelect"),
      activatePromptProfileBtn: document.getElementById("activatePromptProfileBtn"),
      savePromptSetupBtn:   document.getElementById("savePromptSetupBtn"),
      promptProfileInfo:    document.getElementById("promptProfileInfo"),
      visualizationEngineSelect: document.getElementById("visualizationEngineSelect"),
      visualizationEngineInfo: document.getElementById("visualizationEngineInfo"),
      systemPrompt:         document.getElementById("systemPrompt"),
      fileInput:            document.getElementById("fileInput"),
      selectedFiles:        document.getElementById("selectedFiles"),
      refreshModelsBtn:     document.getElementById("refreshModelsBtn"),
      clearFilesBtn:        document.getElementById("clearFilesBtn"),
      resetSystemPromptBtn: document.getElementById("resetSystemPromptBtn"),
      copySystemPromptBtn:  document.getElementById("copySystemPromptBtn"),
      exportChatBtn:        document.getElementById("exportChatBtn"),
      autoScrollBtn:        document.getElementById("autoScrollBtn"),
      helpGuideBtn:         document.getElementById("helpGuideBtn"),
      helpModal:            document.getElementById("helpModal"),
      closeHelpModalBtn:    document.getElementById("closeHelpModalBtn"),
      themeToggleBtn:       document.getElementById("themeToggleBtn"),
      clearChatBtn:         document.getElementById("clearChatBtn"),
      reloadSessionBtn:     document.getElementById("reloadSessionBtn"),
      modelInfo:            document.getElementById("modelInfo"),
      msgCountBadge:        document.getElementById("msgCountBadge"),
      selectedModelBadge:   document.getElementById("selectedModelBadge"),
      sourceBadge:          document.getElementById("sourceBadge"),
      streamBadge:          document.getElementById("streamBadge"),
      backendBadge:         document.getElementById("backendBadge"),
      tokensPerSecBadge:    document.getElementById("tokensPerSecBadge"),
      reasoningPanel:       document.getElementById("reasoningPanel"),
      reasoningMeta:        document.getElementById("reasoningMeta"),
      reasoningContent:     document.getElementById("reasoningContent"),
      toggleReasoningBtn:   document.getElementById("toggleReasoningBtn"),
      messages:             document.getElementById("messages"),
      userInput:            document.getElementById("userInput"),
      sendBtn:              document.getElementById("sendBtn"),
      stopBtn:              document.getElementById("stopBtn"),
      helperText:           document.getElementById("helperText"),
      charCounter:          document.getElementById("charCounter"),
      dropOverlay:          document.getElementById("dropOverlay"),
      scrollToBottomBtn:    document.getElementById("scrollToBottomBtn"),
      // Model parameters
      paramTemp:            document.getElementById("paramTemp"),
      paramTempVal:         document.getElementById("paramTempVal"),
      paramTopP:            document.getElementById("paramTopP"),
      paramTopPVal:         document.getElementById("paramTopPVal"),
      paramSeed:            document.getElementById("paramSeed"),
      paramNumCtx:          document.getElementById("paramNumCtx"),
      paramNumCtxInfo:      document.getElementById("paramNumCtxInfo"),
      clearNumCtxBtn:       document.getElementById("clearNumCtxBtn"),
      resetParamsBtn:       document.getElementById("resetParamsBtn"),
      apiKeyInput:          document.getElementById("apiKeyInput"),
      apiKeyInfo:           document.getElementById("apiKeyInfo"),
      saveApiKeyBtn:        document.getElementById("saveApiKeyBtn"),
      clearApiKeyBtn:       document.getElementById("clearApiKeyBtn"),
    };

    // ── Utilities ────────────────────────────────────────────────────────────

    function nowString() {
      return new Date().toLocaleTimeString("el-GR", {
        hour: "2-digit", minute: "2-digit", second: "2-digit"
      });
    }

    function escapeHtml(text) {
      return String(text || "")
        .replace(/&/g, "&amp;").replace(/</g, "&lt;").replace(/>/g, "&gt;")
        .replace(/"/g, "&quot;").replace(/'/g, "&#39;");
    }

    function setButtonFeedback(btn, successText, defaultText, cls = "done") {
      btn.textContent = successText;
      btn.classList.remove("done", "error");
      btn.classList.add(cls);
      setTimeout(() => { btn.textContent = defaultText; btn.classList.remove("done", "error"); }, 1500);
    }

    function countWords(text) {
      return text.trim() ? text.trim().split(/\\s+/).length : 0;
    }

    // ── Model parameters ─────────────────────────────────────────────────────

    const DEFAULT_PARAMS = { temperature: 0.8, top_p: 0.9, seed: -1 };

    function normalizeNumCtxValue(rawValue) {
      const text = String(rawValue ?? "").trim();
      if (!text) return null;
      const value = parseInt(text, 10);
      if (!Number.isFinite(value)) return null;
      if (value < 256 || value > 1048576) return null;
      return value;
    }

    function getSelectedModelKey() {
      return String(els.modelSelect.value || "").trim();
    }

    function syncNumCtxInputForSelectedModel() {
      const model = getSelectedModelKey();
      const saved = model ? state.modelNumCtxByModel[model] : null;
      const maxFromServer = model ? state.modelMaxNumCtxByModel[model] : null;
      const effective = saved != null ? saved : maxFromServer;
      if (els.paramNumCtx) {
        els.paramNumCtx.value = effective != null ? String(effective) : "";
        if (maxFromServer != null) {
          els.paramNumCtx.placeholder = String(maxFromServer);
        } else {
          els.paramNumCtx.placeholder = "ανάκτηση Max Length...";
        }
      }
      if (els.paramNumCtxInfo) {
        const label = model || "-";
        let shown = "Ανάκτηση metadata...";
        let modeLabel = "φόρτωση";
        if (saved != null) {
          shown = `${saved.toLocaleString("el-GR")} tokens`;
          modeLabel = "προσαρμοσμένο";
        } else if (maxFromServer != null) {
          shown = `${maxFromServer.toLocaleString("el-GR")} tokens`;
          modeLabel = "πραγματικό μέγιστο μοντέλου";
        }
        els.paramNumCtxInfo.textContent = `Μοντέλο: ${label} · Max Length: ${shown} · ${modeLabel}`;
      }
      if (model && maxFromServer == null) {
        ensureSelectedModelMeta(model);
      }
    }

    function getModelOptions() {
      const temperature = parseFloat(els.paramTemp.value);
      const top_p       = parseFloat(els.paramTopP.value);
      const seedRaw     = els.paramSeed.value.trim();
      const seed        = parseInt(seedRaw, 10);
      const numCtx      = normalizeNumCtxValue(els.paramNumCtx ? els.paramNumCtx.value : "");
      const opts = {};
      if (!isNaN(temperature)) opts.temperature = temperature;
      if (!isNaN(top_p))       opts.top_p       = top_p;
      if (!isNaN(seed) && seed >= 0) opts.seed  = seed;
      if (numCtx != null)      opts.num_ctx     = numCtx;
      return opts;
    }

    function saveParams() {
      try {
        localStorage.setItem(PARAMS_KEY, JSON.stringify({
          temperature: els.paramTemp.value,
          top_p:       els.paramTopP.value,
          seed:        els.paramSeed.value,
          num_ctx_by_model: state.modelNumCtxByModel,
        }));
      } catch (_) {}
    }

    function loadParams() {
      try {
        const saved = JSON.parse(localStorage.getItem(PARAMS_KEY) || "null");
        state.modelNumCtxByModel = (saved && saved.num_ctx_by_model && typeof saved.num_ctx_by_model === "object")
          ? saved.num_ctx_by_model
          : {};
        if (!saved) {
          syncNumCtxInputForSelectedModel();
          return;
        }
        if (saved.temperature != null) {
          els.paramTemp.value    = saved.temperature;
          els.paramTempVal.textContent = Number(saved.temperature).toFixed(2);
        }
        if (saved.top_p != null) {
          els.paramTopP.value    = saved.top_p;
          els.paramTopPVal.textContent = Number(saved.top_p).toFixed(2);
        }
        if (saved.seed != null) {
          els.paramSeed.value    = saved.seed;
        }
        syncNumCtxInputForSelectedModel();
      } catch (_) {
        state.modelNumCtxByModel = {};
        syncNumCtxInputForSelectedModel();
      }
    }

    function clearNumCtxForCurrentModel(showNotice = true) {
      const model = getSelectedModelKey();
      if (model && Object.prototype.hasOwnProperty.call(state.modelNumCtxByModel, model)) {
        delete state.modelNumCtxByModel[model];
      }
      syncNumCtxInputForSelectedModel();
      saveParams();
      if (showNotice && model) {
        renderSystemNotice(`Το Max Length του μοντέλου ${model} επανήλθε στο πραγματικό μέγιστο context του cloud tag.`);
      }
    }

    function updateCurrentModelNumCtxFromInput() {
      const model = getSelectedModelKey();
      if (!model || !els.paramNumCtx) return;
      const normalized = normalizeNumCtxValue(els.paramNumCtx.value);
      if (normalized == null) {
        delete state.modelNumCtxByModel[model];
        if (els.paramNumCtx.value.trim()) {
          els.paramNumCtx.value = "";
        }
      } else {
        state.modelNumCtxByModel[model] = normalized;
        els.paramNumCtx.value = String(normalized);
      }
      syncNumCtxInputForSelectedModel();
      saveParams();
    }

    function resetParams() {
      els.paramTemp.value          = DEFAULT_PARAMS.temperature;
      els.paramTempVal.textContent = DEFAULT_PARAMS.temperature.toFixed(2);
      els.paramTopP.value          = DEFAULT_PARAMS.top_p;
      els.paramTopPVal.textContent = DEFAULT_PARAMS.top_p.toFixed(2);
      els.paramSeed.value          = DEFAULT_PARAMS.seed;
      clearNumCtxForCurrentModel(false);
      saveParams();
      renderSystemNotice("Παράμετροι μοντέλου επαναφέρθηκαν στις προεπιλογές.");
    }

    // Live update labels as sliders move
    els.paramTemp.addEventListener("input", () => {
      els.paramTempVal.textContent = Number(els.paramTemp.value).toFixed(2);
      saveParams();
    });
    els.paramTopP.addEventListener("input", () => {
      els.paramTopPVal.textContent = Number(els.paramTopP.value).toFixed(2);
      saveParams();
    });
    els.paramSeed.addEventListener("change", saveParams);
    if (els.paramNumCtx) {
      els.paramNumCtx.addEventListener("change", updateCurrentModelNumCtxFromInput);
      els.paramNumCtx.addEventListener("blur", updateCurrentModelNumCtxFromInput);
    }
    if (els.clearNumCtxBtn) {
      els.clearNumCtxBtn.addEventListener("click", () => clearNumCtxForCurrentModel(true));
    }
    // ── Cloud API health polling ────────────────────────────────────────────────

    async function pollBackendHealth() {
      try {
        const resp = await fetch("/api/health");
        const data = await resp.json();
        if (data.cloud_api_configured) {
          els.backendBadge.textContent = "⬤ Cloud API: OK";
          els.backendBadge.className   = "badge ok";
          const keySource = data.api_key_source || "configured";
          els.backendBadge.title       = `Direct mode · API key source: ${keySource} · uptime ${Math.round(data.server_uptime_sec)}s`;
        } else {
          els.backendBadge.textContent = "⬤ Cloud API: KEY";
          els.backendBadge.className   = "badge err";
          els.backendBadge.title       = "Λείπει το Ollama Cloud API key από GUI/settings file ή OLLAMA_API_KEY";
        }
      } catch (_) {
        els.backendBadge.textContent = "⬤ Cloud API: ?";
        els.backendBadge.className   = "badge warn";
      }
    }

    function maskApiKey(value) {
      const key = String(value || "");
      if (!key) return "";
      if (key.length <= 10) return key[0] + "•".repeat(Math.max(0, key.length - 2)) + key.slice(-1);
      return key.slice(0, 4) + "•".repeat(Math.max(0, key.length - 8)) + key.slice(-4);
    }

    function setApiKeyInfo(message, tone = "") {
      if (!els.apiKeyInfo) return;
      els.apiKeyInfo.textContent = message;
      els.apiKeyInfo.className = `tiny ${tone === "error" ? "warn" : "muted"}`;
    }

    async function loadAppConfig() {
      try {
        const resp = await fetch("/api/app-config");
        const data = await resp.json();
        const key = String(data.ollama_api_key || "");
        if (Array.isArray(data.prompt_profiles) && data.prompt_profiles.length) {
          state.promptProfiles = data.prompt_profiles;
        }
        rebuildPromptProfileMap();
        state.activePromptProfile = String(data.active_prompt_profile || DEFAULT_PROMPT_PROFILE_ID || "scientific-technical");
        state.visualizationEngine = String(data.active_visualization_engine || DEFAULT_VISUALIZATION_ENGINE || "auto");
        populatePromptProfileSelect(state.activePromptProfile);
        populateVisualizationEngineSelect(state.visualizationEngine);
        const customPrompt = String(data.custom_system_prompt || "");
        if (els.systemPrompt) {
          els.systemPrompt.value = customPrompt || getPromptProfilePrompt(state.activePromptProfile) || DEFAULT_SYSTEM_PROMPT;
        }
        updatePromptProfileInfo();
        updateVisualizationEngineInfo();
        if (els.apiKeyInput) els.apiKeyInput.value = key;
        if (data.has_ollama_api_key) {
          const updated = data.updated_at ? ` · αποθήκευση ${data.updated_at}` : "";
          setApiKeyInfo(`API key φορτώθηκε από settings file (${maskApiKey(key)})${updated}`);
        } else {
          setApiKeyInfo("Δεν υπάρχει αποθηκευμένο API key στο settings file.");
        }
      } catch (_) {
        rebuildPromptProfileMap();
        populatePromptProfileSelect(state.activePromptProfile || DEFAULT_PROMPT_PROFILE_ID);
        populateVisualizationEngineSelect(state.visualizationEngine || DEFAULT_VISUALIZATION_ENGINE);
        if (els.systemPrompt && !String(els.systemPrompt.value || "").trim()) {
          els.systemPrompt.value = getPromptProfilePrompt(state.activePromptProfile) || DEFAULT_SYSTEM_PROMPT;
        }
        updatePromptProfileInfo();
        updateVisualizationEngineInfo();
        setApiKeyInfo("Αποτυχία φόρτωσης settings αρχείου.", "error");
      }
    }

    async function saveApiKey() {
      try {
        const key = String((els.apiKeyInput && els.apiKeyInput.value) || "").trim();
        const resp = await fetch("/api/app-config", {
          method: "POST",
          headers: { "Content-Type": "application/json" },
          body: JSON.stringify({
            ollama_api_key: key,
            active_prompt_profile: (els.promptProfileSelect && els.promptProfileSelect.value) || state.activePromptProfile || DEFAULT_PROMPT_PROFILE_ID,
            custom_system_prompt: (els.systemPrompt && els.systemPrompt.value) || "",
            active_visualization_engine: (els.visualizationEngineSelect && els.visualizationEngineSelect.value) || state.visualizationEngine || DEFAULT_VISUALIZATION_ENGINE,
          }),
        });
        const data = await resp.json().catch(() => ({}));
        if (!resp.ok) throw new Error(data.error || `HTTP ${resp.status}`);
        if (els.apiKeyInput) els.apiKeyInput.value = String((data.config && data.config.ollama_api_key) || key || "");
        state.activePromptProfile = String((data.config && data.config.active_prompt_profile) || state.activePromptProfile || DEFAULT_PROMPT_PROFILE_ID);
        state.visualizationEngine = String((data.config && data.config.active_visualization_engine) || state.visualizationEngine || DEFAULT_VISUALIZATION_ENGINE);
        if (els.promptProfileSelect) els.promptProfileSelect.value = state.activePromptProfile;
        if (els.visualizationEngineSelect) els.visualizationEngineSelect.value = state.visualizationEngine;
        updatePromptProfileInfo();
        updateVisualizationEngineInfo();
        const masked = maskApiKey((data.config && data.config.ollama_api_key) || key);
        const updated = data.config && data.config.updated_at ? ` · ${data.config.updated_at}` : "";
        setApiKeyInfo(key ? `API key αποθηκεύτηκε (${masked})${updated}` : `Το αποθηκευμένο API key καθαρίστηκε.${updated}`);
        await pollBackendHealth();
        if (key) renderSystemNotice("Αποθηκεύτηκε το Ollama API key στο settings file.");
        else renderSystemNotice("Καθαρίστηκε το αποθηκευμένο Ollama API key από το settings file.");
      } catch (err) {
        const msg = err && err.message ? err.message : String(err);
        setApiKeyInfo(`Σφάλμα αποθήκευσης API key: ${msg}`, "error");
        renderSystemNotice(`Σφάλμα αποθήκευσης API key: ${msg}`);
      }
    }

    async function clearApiKey() {
      if (els.apiKeyInput) els.apiKeyInput.value = "";
      await saveApiKey();
    }

    // ── Tokens/sec display ───────────────────────────────────────────────────

    function estimateLiveTokenCount(text) {
      const src = String(text || "");
      if (!src.trim()) return 0;
      const words = src.match(/[\\p{L}\\p{N}_]+/gu) || [];
      const punct = src.match(/[^\\s\\p{L}\\p{N}_]/gu) || [];
      return words.length + punct.length;
    }

    function showLiveTokenStats(currentText, streamStartMs, phaseLabel = "Generating") {
      if (!els.tokensPerSecBadge) return;
      const elapsedMs = Math.max(1, Date.now() - Number(streamStartMs || Date.now()));
      const elapsedSec = elapsedMs / 1000;
      if (elapsedSec < 0.08) return;

      const estimatedTokens = estimateLiveTokenCount(currentText);
      if (estimatedTokens <= 0) return;

      const liveTps = estimatedTokens / elapsedSec;
      els.tokensPerSecBadge.textContent = `⚡ ${liveTps.toFixed(1)} tok/s`;
      els.tokensPerSecBadge.className   = "badge";
      els.tokensPerSecBadge.title       = [
        `Live tok/s όσο γίνεται πιο κοντά στο πραγματικό κατά τη ροή`,
        `Phase: ${phaseLabel}`,
        `Estimated streamed tokens: ${estimatedTokens}`,
        `Elapsed: ${elapsedSec.toFixed(2)}s`,
        `Στο τέλος γίνεται reconcile με τα επίσημα eval_count / eval_duration του Ollama.`,
      ].join(" · ");
      els.tokensPerSecBadge.style.display = "";
    }

    function showTokenStats(tokenStats) {
      if (!tokenStats || !els.tokensPerSecBadge) return;
      const tps = Number(tokenStats.tokens_per_sec || 0);
      const tokens = Number(tokenStats.eval_count || 0);
      const promptTps = tokenStats.prompt_tokens_per_sec != null
        ? Number(tokenStats.prompt_tokens_per_sec)
        : null;
      const totalTps = tokenStats.end_to_end_tokens_per_sec != null
        ? Number(tokenStats.end_to_end_tokens_per_sec)
        : null;

      els.tokensPerSecBadge.textContent = `⚡ ${tps.toFixed(1)} tok/s`;
      els.tokensPerSecBadge.className   = "badge ok";
      els.tokensPerSecBadge.title       = [
        `Πραγματικό output speed: ${tps.toFixed(1)} tok/s`,
        `Output tokens: ${tokens}`,
        promptTps != null ? `Prompt speed: ${promptTps.toFixed(1)} tok/s` : null,
        totalTps != null ? `End-to-end speed: ${totalTps.toFixed(1)} tok/s` : null,
      ].filter(Boolean).join(" · ");
      els.tokensPerSecBadge.style.display = "";
      state.lastTokensPerSec = tps;
    }

    function hideTokenStats() {
      if (els.tokensPerSecBadge) els.tokensPerSecBadge.style.display = "none";
    }

    // ── Markdown renderer ─────────────────────────────────────────────────────

    /**
     * Inline markdown: escapes HTML first, then applies inline formatting.
     * Order matters: bold before italic to avoid greedy single-* matching.
     */
    function inlineMarkdown(text) {
      text = escapeHtml(text);
      text = text.replace(/\\*\\*(.+?)\\*\\*/g, "<strong>$1</strong>");
      text = text.replace(/__(.+?)__/g,     "<strong>$1</strong>");
      text = text.replace(/\\*([^*\\n]+?)\\*/g,            "<em>$1</em>");
      text = text.replace(/(?<!\\w)_([^_\\n]+?)_(?!\\w)/g, "<em>$1</em>");
      text = text.replace(/~~(.+?)~~/g,  "<del>$1</del>");
      text = text.replace(/`([^`\\n]+)`/g, '<code class="code-inline">$1</code>');
      text = text.replace(
        /\\[([^\\]]+)\\]\\((https?:\\/\\/[^\\)]+)\\)/g,
        '<a href="$2" target="_blank" rel="noopener noreferrer" class="md-link">$1</a>'
      );
      return text;
    }

    function splitMarkdownTableRow(line) {
      return String(line || "")
        .trim()
        .replace(/^\\|/, "")
        .replace(/\\|$/, "")
        .split("|")
        .map(cell => inlineMarkdown(cell.trim()));
    }

    function splitMarkdownTableAlignments(line) {
      return String(line || "")
        .trim()
        .replace(/^\\|/, "")
        .replace(/\\|$/, "")
        .split("|")
        .map(cell => {
          const part = cell.trim();
          if (/^:-{3,}:$/.test(part)) return "center";
          if (/^-{3,}:$/.test(part)) return "right";
          if (/^:-{3,}$/.test(part)) return "left";
          return "";
        });
    }

    function isMarkdownTableSeparator(line) {
      const stripped = String(line || "").trim();
      if (!stripped || !stripped.includes("|")) return false;
      const cells = stripped.replace(/^\\|/, "").replace(/\\|$/, "").split("|");
      return cells.length > 0 && cells.every(cell => /^:?-{3,}:?$/.test(cell.trim()));
    }

    function renderMarkdownTable(headerLine, separatorLine, bodyLines) {
      const headers = splitMarkdownTableRow(headerLine);
      const alignments = splitMarkdownTableAlignments(separatorLine);
      const rows = bodyLines.map(splitMarkdownTableRow);

      const renderCell = (tagName, value, align) => {
        const style = align ? ` style="text-align:${align}"` : "";
        return `<${tagName}${style}>${value || ""}</${tagName}>`;
      };

      const headHtml = `<thead><tr>${headers.map((cell, index) => renderCell("th", cell, alignments[index] || "")).join("")}</tr></thead>`;
      const bodyHtml = rows.length
        ? `<tbody>${rows.map(row => `<tr>${headers.map((_, index) => renderCell("td", row[index] || "", alignments[index] || "")).join("")}</tr>`).join("")}</tbody>`
        : "";

      return `<div class="md-table-wrap"><table class="md-table">${headHtml}${bodyHtml}</table></div>`;
    }

    let mathTypesetQueue = Promise.resolve();

    function mayContainScientificMarkup(text) {
      const source = String(text || "");
      if (!source) return false;
      return /(\\$\\$[\\s\\S]+?\\$\\$|\\\\\\[[\\s\\S]+?\\\\\\]|\\\\\\([\\s\\S]+?\\\\\\)|\\$[^$\\n][\\s\\S]*?\\$|\\\\(?:ce|pu|frac|dfrac|tfrac|sqrt|sum|prod|int|iint|iiint|oint|lim|log|ln|sin|cos|tan|alpha|beta|gamma|delta|epsilon|varepsilon|theta|lambda|mu|pi|sigma|omega|Omega|Delta|Gamma|Sigma|partial|nabla|vec|mathbf|mathbb|mathrm|mathcal|overline|underline|hat|bar|dot|ddot|times|cdot|pm|mp|neq|approx|sim|propto|leq|geq|ll|gg|to|rightarrow|leftarrow|leftrightarrow|Rightarrow|Leftarrow|Leftrightarrow|mapsto|implies|iff|land|lor|neg|oplus|otimes|forall|exists|infty|degree|angle|triangle|square|therefore|because|equiv|parallel|perp|notin|subset|supset|subseteq|supseteq|cup|cap|vdash|models|ohm|text)\\b)/.test(source);;;
    }

    function renderMathInElementSafe(root) {
      if (!root) return;

      if (window.MathJax && typeof window.MathJax.typesetPromise === "function") {
        mathTypesetQueue = mathTypesetQueue
          .catch(() => undefined)
          .then(() => window.MathJax.typesetPromise([root]))
          .catch((err) => {
            console.warn("MathJax render failed:", err);
          });
        return;
      }

      if (typeof window.renderMathInElement !== "function") return;

      try {
        window.renderMathInElement(root, {
          delimiters: [
            { left: "$$", right: "$$", display: true },
            { left: "\\\\[", right: "\\\\]", display: true },
            { left: "$", right: "$", display: false },
            { left: "\\\\(", right: "\\\\)", display: false },
          ],
          throwOnError: false,
          strict: "ignore",
          ignoredTags: ["script", "noscript", "style", "textarea", "pre", "code"],
        });
      } catch (err) {
        console.warn("KaTeX fallback render failed:", err);
      }
    }

    /**
     * Block markdown for a text segment (no code fences — those are handled separately).
     */
    function markdownToHtml(rawText) {
      const lines = rawText.split("\\n");
      const out   = [];
      let inUl = false, inOl = false, inBq = false;

      const closeUl  = () => { if (inUl) { out.push("</ul>");          inUl = false; } };
      const closeOl  = () => { if (inOl) { out.push("</ol>");          inOl = false; } };
      const closeBq  = () => { if (inBq) { out.push("</blockquote>"); inBq = false; } };
      const closeLists = () => { closeUl(); closeOl(); };

      for (let i = 0; i < lines.length; i += 1) {
        const line = lines[i];
        const nextLine = i + 1 < lines.length ? lines[i + 1] : "";

        if (line.includes("|") && isMarkdownTableSeparator(nextLine)) {
          closeLists(); closeBq();
          const bodyLines = [];
          let j = i + 2;
          while (j < lines.length) {
            const candidate = lines[j];
            if (!candidate.trim() || !candidate.includes("|")) break;
            bodyLines.push(candidate);
            j += 1;
          }
          out.push(renderMarkdownTable(line, nextLine, bodyLines));
          i = j - 1;
          continue;
        }

        // Heading  # … ######
        const hm = line.match(/^(#{1,6})\\s+(.*)/);
        if (hm) {
          closeLists(); closeBq();
          const lvl = hm[1].length;
          out.push(`<h${lvl} class="md-h${lvl}">${inlineMarkdown(hm[2])}</h${lvl}>`);
          continue;
        }

        // Horizontal rule  --- / *** / ___
        if (/^(\\s*[-*_]){3,}\\s*$/.test(line) && line.trim().length >= 3) {
          closeLists(); closeBq();
          out.push('<hr class="md-hr" />');
          continue;
        }

        // Blockquote  > ...
        const bm = line.match(/^>\\s?(.*)/);
        if (bm) {
          closeLists();
          if (!inBq) { out.push('<blockquote class="md-bq">'); inBq = true; }
          out.push(`<p>${inlineMarkdown(bm[1])}</p>`);
          continue;
        }
        closeBq();

        // Unordered list  - * +
        const um = line.match(/^[-*+]\\s+(.*)/);
        if (um) {
          closeOl();
          if (!inUl) { out.push('<ul class="md-list">'); inUl = true; }
          out.push(`<li>${inlineMarkdown(um[1])}</li>`);
          continue;
        }

        // Ordered list  1. 2. …
        const om = line.match(/^\\d+\\.\\s+(.*)/);
        if (om) {
          closeUl();
          if (!inOl) { out.push('<ol class="md-list">'); inOl = true; }
          out.push(`<li>${inlineMarkdown(om[1])}</li>`);
          continue;
        }

        closeLists();

        // Empty line
        if (!line.trim()) { out.push('<div class="md-br"></div>'); continue; }

        // Normal paragraph
        out.push(`<p class="md-p">${inlineMarkdown(line)}</p>`);
      }

      closeLists(); closeBq();
      return out.join("\\n");
    }

    // ── Message rendering ─────────────────────────────────────────────────────

    /**
     * Εξάγει <think>...</think> blocks, code fences και plain text.
     * Τύποι parts: "think" | "code" | "text"
     * Τα think blocks εμφανίζονται ως collapsible details element.
     */
    function parseMessageParts(sourceText) {
      const parts  = [];
      const text   = String(sourceText || "");
      let   cursor = 0;

      while (cursor < text.length) {
        // ── <think> block detection ──────────────────────────────────────────
        const thinkStart = text.indexOf("<think>", cursor);
        const fenceStart = text.indexOf("```",    cursor);

        // Determine which comes first
        const nextThink = thinkStart >= 0 ? thinkStart : Infinity;
        const nextFence = fenceStart >= 0 ? fenceStart : Infinity;

        if (nextThink === Infinity && nextFence === Infinity) {
          // Only plain text remains
          const rem = text.slice(cursor);
          if (rem) parts.push({ type: "text", content: rem });
          break;
        }

        if (nextThink < nextFence) {
          // <think> comes first
          const before = text.slice(cursor, thinkStart);
          if (before) parts.push({ type: "text", content: before });

          const thinkEnd = text.indexOf("</think>", thinkStart + 7);
          if (thinkEnd === -1) {
            // Still streaming — show partial thinking
            parts.push({ type: "think", content: text.slice(thinkStart + 7), complete: false });
            cursor = text.length;
          } else {
            parts.push({ type: "think", content: text.slice(thinkStart + 7, thinkEnd), complete: true });
            cursor = thinkEnd + 8; // length of "</think>"
          }
          continue;
        }

        // ── Code fence ───────────────────────────────────────────────────────
        const before = text.slice(cursor, fenceStart);
        if (before) parts.push({ type: "text", content: before });

        const afterFence = fenceStart + 3;
        const nlPos      = text.indexOf("\\n", afterFence);

        if (nlPos === -1) {
          parts.push({ type: "code", language: text.slice(afterFence).trim() || "text", content: "", complete: false });
          cursor = text.length;
          break;
        }

        const language  = text.slice(afterFence, nlPos).trim() || "text";
        const codeStart = nlPos + 1;
        const fenceEnd  = text.indexOf("```", codeStart);

        if (fenceEnd === -1) {
          parts.push({ type: "code", language, content: text.slice(codeStart), complete: false });
          cursor = text.length;
          break;
        }

        parts.push({ type: "code", language, content: text.slice(codeStart, fenceEnd), complete: true });
        cursor = fenceEnd + 3;
      }

      if (!parts.length) parts.push({ type: "text", content: "" });
      return parts;
    }

    // Language aliases → Prism class names
    const LANG_MAP = {
      py: "python", js: "javascript", ts: "typescript",
      sh: "bash", shell: "bash", zsh: "bash", fish: "bash",
      yml: "yaml", htm: "html", golang: "go",
    };

    function isPythonLanguage(language) {
      const normalizedLang = String(language || "").trim().toLowerCase();
      return normalizedLang === "python" || normalizedLang === "py";
    }

    function extractSuggestedPyFilenamesFromText(rawText) {
      const text = String(rawText || "");
      if (!text) return [];

      const parts = parseMessageParts(text);
      const textOnly = parts
        .filter(part => part.type === "text")
        .map(part => String(part.content || ""))
        .join("\\n");

      const results = [];
      const seen = new Set();
      const regex = /([A-Za-z0-9_][A-Za-z0-9._ -]{0,120}\\.py)\\b/gi;
      let match;
      while ((match = regex.exec(textOnly)) !== null) {
        const candidate = String(match[1] || "").trim().replace(/^['"`(\\[]+|['"`)\\],.:;!?]+$/g, "");
        if (!candidate) continue;
        const lower = candidate.toLowerCase();
        if (seen.has(lower)) continue;
        seen.add(lower);
        results.push(candidate);
      }
      return results;
    }

    function schedulePrismHighlight(codeNode, language = "", attempt = 0) {
      if (!codeNode || codeNode.nodeType !== 1) return;

      const normalizedLanguage = String(language || "").trim().toLowerCase();
      const prismLang = (LANG_MAP[normalizedLanguage] || normalizedLanguage || "text").toLowerCase();
      const maxAttempts = 20;

      const runHighlight = () => {
        if (!codeNode.isConnected) return;
        const prism = window.Prism;
        const grammar = prism && prism.languages ? prism.languages[prismLang] : null;

        if (prism && typeof prism.highlightElement === "function") {
          try {
            prism.highlightElement(codeNode);
          } catch (_) {}
        }

        const hasTokens = !!codeNode.querySelector(".token");
        if (!hasTokens && prism && grammar && typeof prism.highlight === "function") {
          try {
            codeNode.innerHTML = prism.highlight(codeNode.textContent || "", grammar, prismLang);
          } catch (_) {}
        }

        if (!codeNode.querySelector(".token") && attempt < maxAttempts) {
          window.setTimeout(() => schedulePrismHighlight(codeNode, prismLang, attempt + 1), Math.min(420, 70 + attempt * 18));
        }
      };

      if (document.readyState === "loading") {
        window.addEventListener("load", () => requestAnimationFrame(runHighlight), { once: true });
        return;
      }
      requestAnimationFrame(runHighlight);
    }

    function schedulePrismHighlightInContainer(container) {
      if (!container) return;
      requestAnimationFrame(() => {
        if (!container.isConnected) return;
        container.querySelectorAll(".code-pre code[class*='language-']").forEach((node) => {
          const match = String(node.className || "").match(/language-([a-z0-9_-]+)/i);
          schedulePrismHighlight(node, match ? match[1] : "text");
        });
      });
    }

    function createCodeBlock(language, code, suggestedFilename = "") {
      const wrapper  = document.createElement("div");
      wrapper.className = "code-block";

      const toolbar  = document.createElement("div");
      toolbar.className = "code-toolbar";

      const langNode = document.createElement("div");
      langNode.className   = "code-language";
      langNode.textContent = language || "text";

      const copyBtn  = document.createElement("button");
      copyBtn.type        = "button";
      copyBtn.className   = "code-copy-btn";
      copyBtn.textContent = "📋 Copy";
      copyBtn.title       = "Αντιγραφή κώδικα στο clipboard";
      copyBtn.addEventListener("click", async () => {
        try {
          await navigator.clipboard.writeText(code);
          setButtonFeedback(copyBtn, "✅ Copied", "📋 Copy");
        } catch {
          setButtonFeedback(copyBtn, "❌ Error", "📋 Copy", "error");
        }
      });

      toolbar.appendChild(langNode);

      const normalizedLang = String(language || "").trim().toLowerCase();
      if (isPythonLanguage(normalizedLang)) {
        const preferredFilename = String(suggestedFilename || "").trim();
        const defaultSaveLabel = preferredFilename ? `💾 ${preferredFilename}` : "💾 .py";
        const saveBtn = document.createElement("button");
        saveBtn.type = "button";
        saveBtn.className = "code-copy-btn";
        saveBtn.textContent = defaultSaveLabel;
        saveBtn.title = preferredFilename
          ? `Αποθήκευση του Python block ως ${preferredFilename}`
          : "Αποθήκευση του Python block σε επισυναπτόμενο αρχείο .py";
        saveBtn.addEventListener("click", async () => {
          try {
            const resp = await fetch("/api/export-python-block", {
              method: "POST",
              headers: { "Content-Type": "application/json" },
              body: JSON.stringify({ code, filename: preferredFilename }),
            });
            const data = await resp.json().catch(() => ({}));
            if (!resp.ok || !data.file || !data.file.url) {
              throw new Error(data.error || "Αποτυχία αποθήκευσης του Python block ως .py αρχείο.");
            }
            const messageWrapper = wrapper.closest('.msg');
            if (messageWrapper) appendGeneratedAttachmentToMessage(messageWrapper, data.file);
            triggerFileDownload(data.file.url, data.file.name || preferredFilename || 'generated_code.py');
            setButtonFeedback(saveBtn, "✅ Saved", defaultSaveLabel);
            renderSystemNotice(`Αποθηκεύτηκε το Python block ως αρχείο: ${data.file.name}`);
          } catch (err) {
            setButtonFeedback(saveBtn, "❌ Error", defaultSaveLabel, "error");
            renderSystemNotice(`Σφάλμα αποθήκευσης Python block: ${err && err.message ? err.message : String(err)}`);
          }
        });
        toolbar.appendChild(saveBtn);

        const defaultRunLabel = preferredFilename ? `▶ Run ${preferredFilename}` : "▶ Run";
        const runBtn = document.createElement("button");
        runBtn.type = "button";
        runBtn.className = "code-copy-btn";
        runBtn.textContent = defaultRunLabel;
        runBtn.title = preferredFilename
          ? `Άνοιγμα νέου terminal και εκτέλεση του Python block ως ${preferredFilename}`
          : "Άνοιγμα νέου terminal και εκτέλεση ολόκληρου του Python block";
        runBtn.addEventListener("click", async () => {
          const confirmed = window.confirm(
            preferredFilename
              ? `Να ανοίξει νέο terminal και να εκτελεστεί το Python block ως ${preferredFilename};`
              : "Να ανοίξει νέο terminal και να εκτελεστεί ολόκληρο το Python block;"
          );
          if (!confirmed) return;
          try {
            const resp = await fetch("/api/execute-python", {
              method: "POST",
              headers: { "Content-Type": "application/json" },
              body: JSON.stringify({ code, filename: preferredFilename }),
            });
            const data = await resp.json().catch(() => ({}));
            if (!resp.ok) {
              throw new Error(data.error || "Αποτυχία εκτέλεσης Python block.");
            }
            setButtonFeedback(
              runBtn,
              preferredFilename ? `▶ Running ${preferredFilename}` : "▶ Running",
              defaultRunLabel
            );
            renderSystemNotice(data.message || "Το Python block στάλθηκε για εκτέλεση σε νέο terminal.");
          } catch (err) {
            setButtonFeedback(runBtn, "❌ Error", defaultRunLabel, "error");
            renderSystemNotice(`Σφάλμα εκτέλεσης κώδικα: ${err && err.message ? err.message : String(err)}`);
          }
        });
        toolbar.appendChild(runBtn);

        if (looksLikePythonPlotCode(code) || String((els.visualizationEngineSelect && els.visualizationEngineSelect.value) || state.visualizationEngine || DEFAULT_VISUALIZATION_ENGINE || "auto") === "python-plot") {
          const plotBtn = document.createElement("button");
          plotBtn.type = "button";
          plotBtn.className = "code-copy-btn";
          plotBtn.textContent = "📈 Render Plot";
          plotBtn.title = preferredFilename
            ? `Ασφαλές render του Python plotting block ως ${preferredFilename}`
            : "Ασφαλές render του Python plotting block με matplotlib";
          plotBtn.addEventListener("click", async () => {
            setButtonFeedback(plotBtn, "⏳ Rendering…", "📈 Render Plot");
            await renderPythonPlotFromCode(code, preferredFilename, wrapper, plotBtn);
          });
          toolbar.appendChild(plotBtn);
        }
      }

      toolbar.appendChild(copyBtn);

      const pre      = document.createElement("pre");
      pre.className  = "code-pre";

      const codeNode = document.createElement("code");
      const prismLang = LANG_MAP[language.toLowerCase()] || language.toLowerCase();
      codeNode.className   = `language-${prismLang}`;
      codeNode.textContent = code;

      pre.appendChild(codeNode);
      wrapper.appendChild(toolbar);
      wrapper.appendChild(pre);

      schedulePrismHighlight(codeNode, prismLang);

      return wrapper;
    }

    function looksLikePythonPlotCode(code) {
      const src = String(code || "");
      if (!src.trim()) return false;
      const normalized = src.toLowerCase();
      return normalized.includes("matplotlib")
        || normalized.includes("plt.")
        || normalized.includes("ax.plot")
        || normalized.includes("ax.scatter")
        || normalized.includes("ax.bar")
        || normalized.includes("hist(")
        || normalized.includes("scatter(")
        || normalized.includes("bar(")
        || normalized.includes("plot(")
        || normalized.includes("dataframe.plot")
        || normalized.includes("series.plot");
    }

    function getPlotRenderContainer(wrapper) {
      let container = wrapper.querySelector(".python-plot-render");
      if (container) return container;
      container = document.createElement("div");
      container.className = "python-plot-render";
      container.style.marginTop = "12px";
      container.style.padding = "12px";
      container.style.border = "1px solid rgba(110, 231, 255, 0.16)";
      container.style.borderRadius = "14px";
      container.style.background = "rgba(15, 23, 42, 0.45)";
      wrapper.appendChild(container);
      return container;
    }

    function renderPythonPlotPreview(wrapper, plotFile) {
      if (!wrapper || !plotFile || !plotFile.url) return;
      const container = getPlotRenderContainer(wrapper);
      container.innerHTML = "";

      const title = document.createElement("div");
      title.textContent = `📈 Rendered Plot · ${plotFile.name || "generated_plot.png"}`;
      title.style.fontWeight = "700";
      title.style.marginBottom = "10px";
      title.style.color = "var(--text)";
      container.appendChild(title);

      const image = document.createElement("img");
      image.src = plotFile.url;
      image.alt = plotFile.name || "Rendered Python Plot";
      image.loading = "lazy";
      image.style.display = "block";
      image.style.width = "100%";
      image.style.maxWidth = "100%";
      image.style.height = "auto";
      image.style.borderRadius = "12px";
      image.style.background = "rgba(255,255,255,0.98)";
      image.style.border = "1px solid rgba(148,163,184,0.22)";
      container.appendChild(image);

      const actions = document.createElement("div");
      actions.style.display = "flex";
      actions.style.gap = "8px";
      actions.style.flexWrap = "wrap";
      actions.style.marginTop = "10px";

      const openLink = document.createElement("a");
      openLink.href = plotFile.url;
      openLink.target = "_blank";
      openLink.rel = "noopener noreferrer";
      openLink.className = "attachment-chip link";
      openLink.textContent = "🖼 Open Plot";
      actions.appendChild(openLink);

      const downloadBtn = document.createElement("button");
      downloadBtn.type = "button";
      downloadBtn.className = "code-copy-btn";
      downloadBtn.textContent = "💾 Download Plot";
      downloadBtn.addEventListener("click", () => {
        triggerFileDownload(plotFile.url, plotFile.name || "generated_plot.png");
      });
      actions.appendChild(downloadBtn);
      container.appendChild(actions);
    }

    async function renderPythonPlotFromCode(code, suggestedFilename, wrapper, triggerBtn) {
      try {
        const resp = await fetch("/api/render-python-plot", {
          method: "POST",
          headers: { "Content-Type": "application/json" },
          body: JSON.stringify({ code, filename: suggestedFilename || "" }),
        });
        const data = await resp.json().catch(() => ({}));
        if (!resp.ok || !data.file || !data.file.url) {
          throw new Error(data.error || "Αποτυχία render του Python plot.");
        }
        const messageWrapper = wrapper && wrapper.closest ? wrapper.closest(".msg") : null;
        if (messageWrapper) appendGeneratedAttachmentToMessage(messageWrapper, data.file);
        renderPythonPlotPreview(wrapper, data.file);
        if (triggerBtn) setButtonFeedback(triggerBtn, "✅ Plot Ready", "📈 Render Plot");
        renderSystemNotice(data.message || `Το plot ${data.file.name || "generated_plot.png"} δημιουργήθηκε.`);
      } catch (err) {
        if (triggerBtn) setButtonFeedback(triggerBtn, "❌ Plot Error", "📈 Render Plot", "error");
        renderSystemNotice(`Σφάλμα render plot: ${err && err.message ? err.message : String(err)}`);
      }
    }

    function createThinkingBlock(content, complete) {
      const details = document.createElement("details");
      details.className = "thinking-block";
      // Διατήρησε την επιλογή του χρήστη όσο γίνεται rerender κατά το streaming
      details.open = !complete ? state.currentInlineThinkingOpen !== false : false;

      const summary = document.createElement("summary");
      summary.className = "thinking-summary";

      const icon = document.createElement("span");
      icon.className = "thinking-icon"; icon.textContent = "🧠";

      const label = document.createElement("span");
      label.className = "thinking-label";
      const wordCount = content.trim() ? content.trim().split(/\\s+/).length : 0;
      label.textContent = complete
        ? `Βαθιά σκέψη · ${wordCount} λέξεις — κλικ για ${details.open ? "απόκρυψη" : "εμφάνιση"}`
        : `Σκέψη σε εξέλιξη… — κλικ για ${details.open ? "απόκρυψη" : "εμφάνιση"}`;

      const chev = document.createElement("span");
      chev.className = "thinking-chevron"; chev.textContent = "›";

      summary.appendChild(icon); summary.appendChild(label); summary.appendChild(chev);

      const body = document.createElement("div");
      body.className = "thinking-body"; body.textContent = content;

      details.appendChild(summary); details.appendChild(body);

      // Update label text when toggled
      details.addEventListener("toggle", () => {
        state.currentInlineThinkingOpen = details.open;
        if (complete) {
          label.textContent = `Βαθιά σκέψη · ${wordCount} λέξεις — κλικ για ${details.open ? "απόκρυψη" : "εμφάνιση"}`;
        } else {
          label.textContent = `Σκέψη σε εξέλιξη… — κλικ για ${details.open ? "απόκρυψη" : "εμφάνιση"}`;
        }
      });

      return details;
    }

    function extractStandaloneSvgBlocks(text) {
      const source = String(text || "");
      if (!source || source.indexOf("<svg") === -1 || source.indexOf("</svg>") === -1) return [source];

      const parts = [];
      const regex = /<svg\b[\\s\\S]*?<\\/svg>/ig;
      let lastIndex = 0;
      let match;
      while ((match = regex.exec(source)) !== null) {
        const start = match.index;
        const end = regex.lastIndex;
        if (start > lastIndex) {
          parts.push(source.slice(lastIndex, start));
        }
        parts.push(match[0]);
        lastIndex = end;
      }
      if (lastIndex < source.length) {
        parts.push(source.slice(lastIndex));
      }
      return parts.length ? parts : [source];
    }

    function createTextSegment(text, options = {}) {
      const source = String(text || "");
      const fragment = document.createDocumentFragment();
      const segments = extractStandaloneSvgBlocks(source);

      for (const segment of segments) {
        const piece = String(segment || "");
        const trimmed = piece.trim();
        if (!piece) continue;

        if (looksLikeSvgContent(trimmed)) {
          const svgNode = createSvgPreviewBlock(trimmed);
          if (svgNode) {
            fragment.appendChild(svgNode);
            continue;
          }
        }

        const div = document.createElement("div");
        div.innerHTML = markdownToHtml(piece);
        if (!options.skipMathRender && mayContainScientificMarkup(piece)) {
          renderMathInElementSafe(div);
        }
        if (div.childNodes.length === 0 && piece.trim()) {
          div.textContent = piece;
        }
        fragment.appendChild(div);
      }

      return fragment;
    }

    const liveMathRenderState = new WeakMap();

    function scheduleLiveScientificRender(container, sourceText, options = {}) {
      if (!container || !mayContainScientificMarkup(sourceText)) return;

      const minInterval = Math.max(0, Number(options.liveMathIntervalMs || 120));
      const forceNow = Boolean(options.forceMathPreview);
      const now = Date.now();
      let entry = liveMathRenderState.get(container);
      if (!entry) entry = { timer: 0, lastRun: 0 };

      const run = () => {
        entry.timer = 0;
        entry.lastRun = Date.now();
        try {
          renderMathInElementSafe(container);
        } catch (_) {}
        liveMathRenderState.set(container, entry);
      };

      const elapsed = now - Number(entry.lastRun || 0);
      if (forceNow || elapsed >= minInterval) {
        if (entry.timer) {
          clearTimeout(entry.timer);
          entry.timer = 0;
        }
        run();
        return;
      }

      if (!entry.timer) {
        entry.timer = window.setTimeout(run, Math.max(24, minInterval - elapsed));
      }
      liveMathRenderState.set(container, entry);
    }

    function composeDisplayContent(answerText = "", thinkingText = "") {
      const answer   = String(answerText || "");
      const thinking = String(thinkingText || "");
      if (thinking && answer) return `<think>${thinking}</think>\\n\\n${answer}`;
      if (thinking) return `<think>${thinking}</think>`;
      return answer;
    }

    function getThinkingStateFromRawContent(sourceText) {
      const parts = parseMessageParts(sourceText);
      const thinkParts = parts.filter(part => part.type === "think");
      return {
        text: thinkParts.map(part => part.content || "").join("\\n\\n"),
        complete: thinkParts.length ? thinkParts.every(part => part.complete !== false) : true,
      };
    }

    function applyReasoningPanelVisibility() {
      if (!els.reasoningPanel || !els.toggleReasoningBtn) return;
      const hasContent = Boolean(((els.reasoningContent && els.reasoningContent.textContent) || "").trim());
      const wantsVisible = state.reasoningPanelVisible || (state.reasoningAutoOpen && !state.reasoningUserCollapsed);
      const isVisible = hasContent && wantsVisible;
      els.reasoningPanel.classList.toggle("visible", isVisible);
      els.toggleReasoningBtn.textContent = isVisible ? "🙈 Απόκρυψη" : "👁 Εμφάνιση";
      els.toggleReasoningBtn.title = isVisible
        ? "Απόκρυψη του panel σκέψης"
        : "Εμφάνιση του panel σκέψης";
    }

    function resetReasoningPanel(hide = true) {
      state.currentThinkingText = "";
      state.reasoningAutoOpen = false;
      state.reasoningUserCollapsed = false;
      state.currentInlineThinkingOpen = true;
      state.reasoningStreamCompleted = false;
      if (hide) state.reasoningPanelVisible = false;
      if (els.reasoningContent) els.reasoningContent.textContent = "";
      if (els.reasoningMeta) els.reasoningMeta.textContent = "Αναμονή για thinking stream…";
      if (els.reasoningPanel) els.reasoningPanel.classList.remove("streaming");
      if (hide && els.reasoningPanel) els.reasoningPanel.classList.remove("visible");
      applyReasoningPanelVisibility();
    }

    function updateReasoningPanel(text, complete = false, streaming = false) {
      const safeText = String(text || "");
      state.currentThinkingText = safeText;

      if (!safeText.trim()) {
        resetReasoningPanel(true);
        return;
      }

      if (streaming && !complete && !state.reasoningUserCollapsed) {
        state.reasoningAutoOpen = true;
      }

      if (els.reasoningContent) {
        els.reasoningContent.textContent = safeText;
        els.reasoningContent.scrollTop = els.reasoningContent.scrollHeight;
      }

      const wordCount = safeText.trim() ? safeText.trim().split(/\\s+/).length : 0;
      if (els.reasoningMeta) {
        els.reasoningMeta.textContent = streaming && !complete
          ? `Σκέψη σε εξέλιξη… · ${wordCount} λέξεις`
          : `Ολοκληρωμένη σκέψη · ${wordCount} λέξεις`;
      }

      if (els.reasoningPanel) {
        els.reasoningPanel.classList.toggle("streaming", streaming && !complete);
      }
      applyReasoningPanelVisibility();

      if (complete) {
        setTimeout(() => {
          if (state.reasoningStreamCompleted) {
            state.reasoningAutoOpen = false;
            state.reasoningPanelVisible = false;
            state.reasoningUserCollapsed = false;
            applyReasoningPanelVisibility();
          }
        }, 350);
      }
    }

    const STREAM_RENDER_MIN_INTERVAL_MS = 95;

    function ensureAssistantStreamingPreviewNode(container) {
      if (!container) return null;
      let node = container.querySelector(".assistant-live-preview");
      if (!node) {
        container.innerHTML = "";
        node = document.createElement("div");
        node.className = "assistant-live-preview";
        container.appendChild(node);
      }
      return node;
    }

    function createStreamingCodeBlock(language, code) {
      const wrapper = document.createElement("div");
      wrapper.className = "code-block assistant-live-code-block";

      const toolbar = document.createElement("div");
      toolbar.className = "code-toolbar";

      const langLabel = document.createElement("span");
      langLabel.className = "code-lang";
      langLabel.textContent = String(language || "text").trim() || "text";
      toolbar.appendChild(langLabel);

      const status = document.createElement("span");
      status.className = "tiny muted";
      status.textContent = "streaming…";
      toolbar.appendChild(status);

      const pre = document.createElement("pre");
      pre.className = "code-pre";

      const codeNode = document.createElement("code");
      const normalizedLanguage = String(language || "").trim().toLowerCase();
      const prismLang = LANG_MAP[normalizedLanguage] || normalizedLanguage || "text";
      codeNode.className = `language-${prismLang}`;
      codeNode.textContent = String(code || "");

      pre.appendChild(codeNode);
      wrapper.appendChild(toolbar);
      wrapper.appendChild(pre);

      schedulePrismHighlight(codeNode, prismLang);

      return wrapper;
    }

    function updateAssistantStreamingPreview(container, content) {
      const sourceText = String(content || "");
      if (!container) return sourceText;
      container.dataset.rawContent = sourceText;
      renderMessageContent(container, sourceText, {
        skipMathRender: true,
        liveMathPreview: false,
        forceStreamingPreview: true,
      });
      return sourceText;
    }

    function flushAssistantStreamingRender(force = false) {
      if (!state.pendingAssistantRender || !state.currentAssistantNode) return "";

      const now = Date.now();
      const lastRun = Number(state.lastAssistantStreamRenderTs || 0);
      const elapsed = now - lastRun;
      if (!force && lastRun && elapsed < STREAM_RENDER_MIN_INTERVAL_MS) {
        const waitMs = Math.max(16, STREAM_RENDER_MIN_INTERVAL_MS - elapsed);
        if (!state.streamRenderTimer) {
          state.streamRenderTimer = window.setTimeout(() => {
            state.streamRenderTimer = 0;
            flushAssistantStreamingRender(true);
          }, waitMs);
        }
        return String((state.pendingAssistantRender && state.pendingAssistantRender.displayText) || "");
      }

      if (state.streamRenderTimer) {
        clearTimeout(state.streamRenderTimer);
        state.streamRenderTimer = 0;
      }

      state.lastAssistantStreamRenderTs = now;
      const pending = state.pendingAssistantRender;
      const hasSeparateThinking = Boolean(String(pending.separateThinkingText || "").trim());
      const displayText = hasSeparateThinking
        ? composeDisplayContent(pending.rawAnswerText, pending.separateThinkingText)
        : String(pending.rawAnswerText || "");

      pending.displayText = displayText;

      if (pending.streamFinished) {
        renderMessageContent(state.currentAssistantNode, displayText, {
          skipMathRender: false,
          liveMathPreview: false,
          forceMathPreview: true,
        });
      } else {
        updateAssistantStreamingPreview(state.currentAssistantNode, displayText);
      }

      if (hasSeparateThinking) {
        const thinkingComplete = state.reasoningStreamCompleted || pending.streamFinished;
        updateReasoningPanel(pending.separateThinkingText, thinkingComplete, !thinkingComplete);
      } else {
        const legacyThinking = getThinkingStateFromRawContent(displayText);
        if (legacyThinking.text.trim()) {
          updateReasoningPanel(legacyThinking.text, pending.streamFinished ? true : legacyThinking.complete, !pending.streamFinished && !legacyThinking.complete);
        } else if (!pending.streamFinished) {
          resetReasoningPanel(true);
        }
      }

      if (pending.streamFinished) {
        state.pendingAssistantRender = null;
      }
      return displayText;
    }

    function scheduleAssistantStreamingRender(force = false) {
      if (force) {
        return flushAssistantStreamingRender(true);
      }
      if (state.streamRenderRaf) {
        return String((state.pendingAssistantRender && state.pendingAssistantRender.displayText) || "");
      }
      state.streamRenderRaf = window.requestAnimationFrame(() => {
        state.streamRenderRaf = 0;
        flushAssistantStreamingRender(false);
      });
      return String((state.pendingAssistantRender && state.pendingAssistantRender.displayText) || "");
    }

    function renderAssistantStreamingView(rawAnswerText, separateThinkingText = "", streamFinished = false) {
      state.pendingAssistantRender = {
        rawAnswerText: String(rawAnswerText || ""),
        separateThinkingText: String(separateThinkingText || ""),
        streamFinished: Boolean(streamFinished),
        displayText: String((state.pendingAssistantRender && state.pendingAssistantRender.displayText) || ""),
      };
      return scheduleAssistantStreamingRender(streamFinished);
    }

    function renderMessageContent(container, content, options = {}) {
      const sourceText = String(content || "");
      container.innerHTML          = "";
      container.dataset.rawContent = sourceText;

      const frag  = document.createDocumentFragment();
      const parts = parseMessageParts(sourceText);
      const suggestedPyFilenames = extractSuggestedPyFilenamesFromText(sourceText);
      let pythonBlockIndex = 0;

      for (const part of parts) {
        if (part.type === "think") {
          frag.appendChild(createThinkingBlock(part.content || "", part.complete !== false));
        } else if (part.type === "code") {
          const rawCode = String(part.content || "");
          const normalizedLanguage = String(part.language || "").trim().toLowerCase();
          const isSvgBlock = isSvgLanguage(normalizedLanguage) || looksLikeSvgContent(rawCode.trim());
          const useStreamingPreview = !!options.forceStreamingPreview;

          if (isSvgBlock && !useStreamingPreview) {
            const svgNode = createSvgPreviewBlock(rawCode);
            if (svgNode) {
              frag.appendChild(svgNode);
            } else {
              const suggestedFilename = isPythonLanguage(part.language)
                ? (suggestedPyFilenames[pythonBlockIndex] || "")
                : "";
              frag.appendChild(createCodeBlock(part.language, rawCode, suggestedFilename));
              if (isPythonLanguage(part.language)) pythonBlockIndex += 1;
            }
          } else {
            const suggestedFilename = isPythonLanguage(part.language)
              ? (suggestedPyFilenames[pythonBlockIndex] || "")
              : "";
            frag.appendChild(
              useStreamingPreview
                ? createStreamingCodeBlock(part.language, rawCode)
                : createCodeBlock(part.language, rawCode, suggestedFilename)
            );
            if (isPythonLanguage(part.language)) pythonBlockIndex += 1;
          }
        } else {
          frag.appendChild(createTextSegment(part.content || "", options));
        }
      }
      container.appendChild(frag);
      schedulePrismHighlightInContainer(container);
      if (options.liveMathPreview) {
        scheduleLiveScientificRender(container, sourceText, options);
      }
      try {
        const ownerMessage = typeof container.closest === "function" ? container.closest(".msg") : null;
        if (ownerMessage && ownerMessage.classList && ownerMessage.classList.contains("assistant")) {
          syncAssistantPdfButtons();
        }
      } catch (_) {}
    }

    // ── Message creation ──────────────────────────────────────────────────────

    function createAttachmentChip(item) {
      const hasUrl = !!(item && item.url);
      const chip = document.createElement(hasUrl ? "a" : "div");
      chip.className = `attachment-chip${hasUrl ? " link" : ""}`;
      if (hasUrl) {
        chip.href = item.url;
        chip.target = "_blank";
        chip.rel = "noopener noreferrer";
        chip.download = item.name || "attachment";
        chip.title = `Άνοιγμα / λήψη: ${item.name || "attachment"}`;
      }
      const icon = item && item.kind === "image" ? "🖼" : "📄";
      chip.textContent = `${icon} ${(item && item.name) ? item.name : "attachment"}`;
      return chip;
    }

    function ensureMessageAttachmentList(messageWrapper) {
      let wrap = messageWrapper.querySelector('.attachment-list');
      if (!wrap) {
        wrap = document.createElement('div');
        wrap.className = 'attachment-list';
        messageWrapper.appendChild(wrap);
      }
      return wrap;
    }

    function appendGeneratedAttachmentToMessage(messageWrapper, item) {
      if (!messageWrapper || !item || !item.url) return;
      const wrap = ensureMessageAttachmentList(messageWrapper);
      const existing = Array.from(wrap.querySelectorAll('a.attachment-chip[href], .attachment-chip[data-name]'))
        .some(node => (node.getAttribute('href') || '') === item.url || (node.dataset && node.dataset.name) === item.name);
      if (existing) return;
      const chip = createAttachmentChip(item);
      if (chip.dataset) chip.dataset.name = item.name || '';
      wrap.appendChild(chip);
    }

    function triggerFileDownload(url, filename) {
      const a = document.createElement('a');
      a.href = url;
      a.download = filename || '';
      a.rel = 'noopener noreferrer';
      document.body.appendChild(a);
      a.click();
      a.remove();
    }

    function removeEmptyState() {
      const node = document.getElementById("emptyState");
      if (node) node.remove();
    }

    function scrollToBottom(force = false) {
      if (state.autoScroll || force) {
        const previousBehavior = els.messages.style.scrollBehavior;
        if (state.isStreaming) {
          els.messages.style.scrollBehavior = "auto";
        }
        els.messages.scrollTop = els.messages.scrollHeight;
        if (state.isStreaming) {
          els.messages.style.scrollBehavior = previousBehavior || "";
        }
      }
    }

    function updateScrollToBottomBtn() {
      const { scrollTop, scrollHeight, clientHeight } = els.messages;
      const atBottom = scrollHeight - scrollTop - clientHeight < 80;
      els.scrollToBottomBtn.classList.toggle("visible", !atBottom && !state.autoScroll);
    }

    function updateMsgCount() {
      state.msgCount = els.messages.querySelectorAll(".msg.user, .msg.assistant").length;
      els.msgCountBadge.textContent = `${state.msgCount} μηνύματα`;
    }

    function createStreamingPlaceholder() {
      const d = document.createElement("div");
      d.className = "streaming-dots";
      d.innerHTML = "<span></span><span></span><span></span>";
      return d;
    }

    function createMessage(role, content, attachments = []) {
      removeEmptyState();

      const wrapper = document.createElement("div");
      wrapper.className = `msg ${role}`;

      // Header
      const head     = document.createElement("div");
      head.className = "msg-head";
      const roleNode = document.createElement("div");
      roleNode.className   = "msg-role";
      roleNode.textContent = role === "user" ? "Εσύ" : role === "assistant" ? "Assistant" : "System";
      const timeNode = document.createElement("div");
      timeNode.className   = "msg-time";
      timeNode.textContent = nowString();
      head.appendChild(roleNode);
      head.appendChild(timeNode);

      // Body
      const body     = document.createElement("div");
      body.className = "msg-body";
      if (role === "assistant" && !content) {
        state.pendingAssistantRender = null;
        state.lastAssistantStreamRenderTs = 0;
        if (state.streamRenderTimer) { clearTimeout(state.streamRenderTimer); state.streamRenderTimer = 0; }
        if (state.streamRenderRaf) { cancelAnimationFrame(state.streamRenderRaf); state.streamRenderRaf = 0; }
        body.appendChild(createStreamingPlaceholder());
      } else {
        renderMessageContent(body, content || "");
      }

      wrapper.appendChild(head);
      wrapper.appendChild(body);

      // Attachments
      if (attachments && attachments.length) {
        const wrap = document.createElement("div");
        wrap.className = "attachment-list";
        for (const item of attachments) {
          wrap.appendChild(createAttachmentChip(item));
        }
        wrapper.appendChild(wrap);
      }

      // Copy button
      const tools   = document.createElement("div");
      tools.className = "message-tools";
      const copyBtn = document.createElement("button");
      copyBtn.type        = "button";
      copyBtn.className   = "tool-btn";
      copyBtn.textContent = "📋 Copy";
      copyBtn.title       = "Αντιγραφή μηνύματος";
      copyBtn.addEventListener("click", async () => {
        try {
          await navigator.clipboard.writeText(body.dataset.rawContent || body.textContent || "");
          setButtonFeedback(copyBtn, "✅ Copied", "📋 Copy");
        } catch {
          setButtonFeedback(copyBtn, "❌ Error", "📋 Copy", "error");
        }
      });
      tools.appendChild(copyBtn);
      wrapper.appendChild(tools);

      els.messages.appendChild(wrapper);
      scrollToBottom();
      updateMsgCount();

      return { wrapper, body };
    }

    function renderSystemNotice(text) {
      createMessage("system", text, []);
    }

    function renderEmptyState() {
      els.messages.innerHTML = `
        <div class="empty-state" id="emptyState">
          <h3 style="margin-top:0;">Έτοιμο για συνομιλία</h3>
          <div>Γράψε το δικό σου prompt και, αν θέλεις, πρόσθεσε αρχεία για context.</div>
          <div style="margin-top:10px;" class="tiny">
            Enter αποστολή · Shift+Enter νέα γραμμή · Ctrl+Enter εναλλακτικό
          </div>
        </div>`;
      state.msgCount = 0;
      els.msgCountBadge.textContent = "0 μηνύματα";
    }

    // ── File handling & Drag/Drop ─────────────────────────────────────────────

    function renderSelectedFiles() {
      els.selectedFiles.innerHTML = "";
      for (const file of state.selectedFiles) {
        const chip = document.createElement("div");
        chip.className   = "attachment-chip";
        chip.textContent = `📎 ${file.name}`;
        els.selectedFiles.appendChild(chip);
      }
    }

    function addFiles(fileList) {
      const existing = new Set(state.selectedFiles.map(f => `${f.name}|${f.size}`));
      for (const file of Array.from(fileList || [])) {
        const key = `${file.name}|${file.size}`;
        if (!existing.has(key)) { state.selectedFiles.push(file); existing.add(key); }
      }
      renderSelectedFiles();
    }

    async function readFileAsBase64(file) {
      return new Promise((resolve, reject) => {
        const reader  = new FileReader();
        reader.onload = () => {
          const parts = String(reader.result || "").split(",");
          resolve({
            name: file.name, size: file.size,
            mime_type: file.type || "",
            data_base64: parts.length > 1 ? parts[1] : "",
          });
        };
        reader.onerror = () => reject(new Error(`Αποτυχία ανάγνωσης: ${file.name}`));
        reader.readAsDataURL(file);
      });
    }

    async function collectFilesPayload() {
      const payload = [];
      for (const file of state.selectedFiles) payload.push(await readFileAsBase64(file));
      return payload;
    }

    // Reliable drag-enter / drag-leave with a counter (avoids child-element flicker)
    document.addEventListener("dragenter", (e) => {
      e.preventDefault();
      state.dragCounter++;
      els.dropOverlay.classList.add("active");
    });
    document.addEventListener("dragleave", () => {
      state.dragCounter--;
      if (state.dragCounter <= 0) {
        state.dragCounter = 0;
        els.dropOverlay.classList.remove("active");
      }
    });
    document.addEventListener("dragover", (e) => e.preventDefault());
    document.addEventListener("drop", (e) => {
      e.preventDefault();
      state.dragCounter = 0;
      els.dropOverlay.classList.remove("active");
      if (e.dataTransfer && e.dataTransfer.files.length) addFiles(e.dataTransfer.files);
    });

    // ── Model management ──────────────────────────────────────────────────────

    function getSavedModel(models) {
      try {
        const saved = localStorage.getItem(MODEL_KEY);
        return saved && models.includes(saved) ? saved : null;
      } catch { return null; }
    }

    function saveModel(model) {
      try { localStorage.setItem(MODEL_KEY, model); } catch (_) {}
    }

    function getSavedSortCriterion() {
      try {
        const saved = localStorage.getItem(MODEL_SORT_KEY);
        return saved || "overall";
      } catch { return "overall"; }
    }

    function saveSortCriterion(value) {
      try { localStorage.setItem(MODEL_SORT_KEY, value || "overall"); } catch (_) {}
    }

    function normalizeEnsembleMode(value) {
      const normalized = String(value || "auto").trim().toLowerCase();
      return ["off", "auto", "manual"].includes(normalized) ? normalized : "auto";
    }

    function getSavedEnsembleMode() {
      try {
        const explicit = localStorage.getItem(ENSEMBLE_MODE_KEY);
        if (explicit) return normalizeEnsembleMode(explicit);
      } catch (_) {}
      try {
        const legacy = localStorage.getItem(ENSEMBLE_KEY);
        if (legacy == null) return "auto";
        return legacy !== "0" ? "auto" : "off";
      } catch { return "auto"; }
    }

    function saveEnsembleMode(value) {
      const mode = normalizeEnsembleMode(value);
      try { localStorage.setItem(ENSEMBLE_MODE_KEY, mode); } catch (_) {}
      try { localStorage.setItem(ENSEMBLE_KEY, mode === "off" ? "0" : "1"); } catch (_) {}
    }

    function getSavedHelperModel(models) {
      try {
        const saved = localStorage.getItem(ENSEMBLE_HELPER_KEY);
        return saved && models.includes(saved) ? saved : null;
      } catch { return null; }
    }

    function saveHelperModel(model) {
      try {
        if (model) localStorage.setItem(ENSEMBLE_HELPER_KEY, model);
        else localStorage.removeItem(ENSEMBLE_HELPER_KEY);
      } catch (_) {}
    }

    function filterModelsBySearch(modelList, rawQuery) {
      const models = Array.isArray(modelList) ? modelList.filter(Boolean) : [];
      const query = String(rawQuery || "").trim().toLowerCase();
      if (!query) return models;
      const tokens = query.split(/\\s+/).filter(Boolean);
      return models.filter((model) => {
        const haystack = String(model || "").toLowerCase();
        return tokens.every((token) => haystack.includes(token));
      });
    }

    function parseParamSizeBillions(rawValue) {
      const textValue = String(rawValue || "").trim().toLowerCase();
      if (!textValue) return 0;
      const match = textValue.match(/(\\d+(?:\\.\\d+)?)\\s*([tbm])?/i);
      if (!match) return 0;
      const value = Number(match[1] || 0);
      const suffix = String(match[2] || "b").toLowerCase();
      if (!Number.isFinite(value) || value <= 0) return 0;
      if (suffix === "t") return value * 1000;
      if (suffix === "m") return value / 1000;
      return value;
    }

    function getModelMeta(model) {
      return (state.modelMetaByModel && state.modelMetaByModel[model] && typeof state.modelMetaByModel[model] === "object")
        ? state.modelMetaByModel[model]
        : {};
    }

    function getModelCapabilities(model) {
      const meta = getModelMeta(model);
      const caps = Array.isArray(meta.capabilities) ? meta.capabilities.map(x => String(x || "").toLowerCase()) : [];
      const name = String(model || "").toLowerCase();
      if (!caps.includes("vision") && ["vision", "-vl", ":vl", "gemini", "llava", "pixtral"].some(t => name.includes(t))) caps.push("vision");
      if (!caps.includes("coding") && ["coder", "code", "devstral"].some(t => name.includes(t))) caps.push("coding");
      if (!caps.includes("reasoning") && ["thinking", "reason", "r1", "gpt-oss", "qwen3.5", "kimi-k2", "deepseek", "cogito"].some(t => name.includes(t))) caps.push("reasoning");
      return Array.from(new Set(caps));
    }

    function getModelSizeBillions(model) {
      const meta = getModelMeta(model);
      const byParam = Number(meta.parameter_size_b || 0);
      if (Number.isFinite(byParam) && byParam > 0) return byParam;
      const byName = parseParamSizeBillions(model);
      if (Number.isFinite(byName) && byName > 0) return byName;
      const sizeBytes = Number(meta.size_bytes || 0);
      if (Number.isFinite(sizeBytes) && sizeBytes > 0) {
        return sizeBytes / 1_000_000_000;
      }
      return 0;
    }

    function getModelContextTokens(model) {
      const meta = getModelMeta(model);
      const numCtx = Number(meta.num_ctx_max || state.modelMaxNumCtxByModel[model] || 0);
      return Number.isFinite(numCtx) && numCtx > 0 ? numCtx : 0;
    }

    function getModelModifiedTs(model) {
      const meta = getModelMeta(model);
      const rawTs = Number(meta.modified_ts || 0);
      if (Number.isFinite(rawTs) && rawTs > 0) return rawTs;
      const rawDate = String(meta.modified_at || "").trim();
      if (!rawDate) return 0;
      const ts = Date.parse(rawDate);
      return Number.isFinite(ts) ? Math.trunc(ts / 1000) : 0;
    }

    const FAMILY_PRIOR_DEFAULTS = Object.freeze({
      overall: 7.50,
      coding: 7.20,
      reasoning: 7.35,
      context: 7.05,
      vision: 6.85,
      speed: 7.10,
    });

    const MODEL_FAMILY_PROFILES = Object.freeze([
      ["gemini-3-flash",   { overall: 9.12, coding: 8.62, reasoning: 8.82, context: 8.95, vision: 9.35, speed: 9.95 }],
      ["deepseek-v3.2",    { overall: 9.82, coding: 9.62, reasoning: 9.90, context: 9.14, vision: 6.20, speed: 4.70 }],
      ["deepseek-v3.1",    { overall: 9.72, coding: 9.54, reasoning: 9.80, context: 9.02, vision: 6.00, speed: 4.82 }],
      ["deepseek-r1",      { overall: 9.66, coding: 9.10, reasoning: 9.96, context: 8.40, vision: 5.25, speed: 3.86 }],
      ["qwen3.5",          { overall: 9.96, coding: 9.74, reasoning: 9.92, context: 9.52, vision: 9.84, speed: 4.42 }],
      ["qwen3-coder-next", { overall: 9.56, coding: 9.96, reasoning: 9.42, context: 9.00, vision: 5.35, speed: 5.06 }],
      ["qwen3-coder",      { overall: 9.68, coding: 9.98, reasoning: 9.60, context: 9.08, vision: 5.38, speed: 4.25 }],
      ["qwen3-vl",         { overall: 9.58, coding: 9.26, reasoning: 9.42, context: 9.02, vision: 9.98, speed: 4.52 }],
      ["qwen3-next",       { overall: 9.24, coding: 9.10, reasoning: 9.18, context: 8.86, vision: 7.92, speed: 5.12 }],
      ["kimi-k2-thinking", { overall: 9.62, coding: 9.22, reasoning: 9.90, context: 9.12, vision: 6.62, speed: 3.72 }],
      ["kimi-k2.5",        { overall: 9.76, coding: 9.40, reasoning: 9.78, context: 9.28, vision: 7.12, speed: 4.12 }],
      ["kimi-k2",          { overall: 9.58, coding: 9.22, reasoning: 9.62, context: 9.00, vision: 6.85, speed: 4.25 }],
      ["glm-5",            { overall: 9.60, coding: 9.36, reasoning: 9.56, context: 9.12, vision: 8.72, speed: 4.15 }],
      ["glm-4.7",          { overall: 9.46, coding: 9.18, reasoning: 9.44, context: 8.92, vision: 8.22, speed: 4.02 }],
      ["glm-4.6",          { overall: 9.38, coding: 9.12, reasoning: 9.36, context: 8.78, vision: 8.05, speed: 4.10 }],
      ["minimax-m2.7",     { overall: 9.42, coding: 9.06, reasoning: 9.40, context: 8.96, vision: 8.62, speed: 4.42 }],
      ["minimax-m2.5",     { overall: 9.34, coding: 8.98, reasoning: 9.32, context: 8.86, vision: 8.46, speed: 4.58 }],
      ["minimax-m2.1",     { overall: 9.22, coding: 8.92, reasoning: 9.20, context: 8.72, vision: 8.18, speed: 4.72 }],
      ["minimax-m2",       { overall: 9.14, coding: 8.88, reasoning: 9.10, context: 8.64, vision: 8.00, speed: 4.86 }],
      ["nemotron-3-super", { overall: 9.32, coding: 8.92, reasoning: 9.28, context: 8.96, vision: 7.42, speed: 4.82 }],
      ["nemotron-3-nano",  { overall: 8.76, coding: 8.36, reasoning: 8.68, context: 8.12, vision: 6.20, speed: 7.20 }],
      ["mistral-large-3",  { overall: 9.22, coding: 9.00, reasoning: 9.18, context: 8.82, vision: 7.82, speed: 4.32 }],
      ["devstral-small-2", { overall: 8.98, coding: 9.42, reasoning: 8.70, context: 8.52, vision: 5.02, speed: 6.62 }],
      ["devstral-2",       { overall: 9.18, coding: 9.74, reasoning: 8.96, context: 8.86, vision: 5.10, speed: 5.22 }],
      ["devstral",         { overall: 8.98, coding: 9.42, reasoning: 8.72, context: 8.52, vision: 5.05, speed: 6.20 }],
      ["gpt-oss",          { overall: 9.06, coding: 8.92, reasoning: 9.04, context: 8.62, vision: 5.10, speed: 5.90 }],
      ["cogito-2.1",       { overall: 9.28, coding: 9.08, reasoning: 9.42, context: 8.92, vision: 6.22, speed: 3.92 }],
      ["cogito",           { overall: 9.12, coding: 8.96, reasoning: 9.26, context: 8.76, vision: 6.02, speed: 4.20 }],
      ["gemini-3",         { overall: 9.74, coding: 9.42, reasoning: 9.72, context: 9.46, vision: 9.82, speed: 5.12 }],
      ["ministral-3",      { overall: 8.74, coding: 8.34, reasoning: 8.48, context: 8.22, vision: 6.22, speed: 8.24 }],
      ["ministral",        { overall: 8.62, coding: 8.22, reasoning: 8.36, context: 8.06, vision: 6.05, speed: 8.00 }],
      ["mistral-small",    { overall: 8.54, coding: 8.18, reasoning: 8.24, context: 7.96, vision: 5.92, speed: 8.20 }],
      ["gemma3",           { overall: 8.58, coding: 8.22, reasoning: 8.34, context: 7.82, vision: 8.12, speed: 7.50 }],
      ["rnj-1",            { overall: 8.32, coding: 7.96, reasoning: 8.16, context: 7.92, vision: 5.42, speed: 8.05 }],
      ["rnj",              { overall: 8.24, coding: 7.88, reasoning: 8.08, context: 7.84, vision: 5.32, speed: 8.10 }],
    ]);

    const MODEL_TRAIT_HINTS = Object.freeze([
      ["qwen3.5",          ["reasoning", "coding", "vision"]],
      ["qwen3-vl",         ["vision", "reasoning", "coding"]],
      ["qwen3-coder",      ["coding", "reasoning"]],
      ["qwen3-next",       ["reasoning", "coding", "vision"]],
      ["deepseek-v3.2",    ["reasoning", "coding"]],
      ["deepseek-v3.1",    ["reasoning", "coding"]],
      ["deepseek-r1",      ["reasoning"]],
      ["kimi-k2.5",        ["reasoning", "coding"]],
      ["kimi-k2",          ["reasoning", "coding"]],
      ["glm-5",            ["reasoning", "coding", "vision"]],
      ["glm-4",            ["reasoning", "coding", "vision"]],
      ["gemini-3",         ["reasoning", "coding", "vision"]],
      ["devstral",         ["coding"]],
      ["gpt-oss",          ["reasoning", "coding"]],
      ["nemotron-3-super", ["reasoning", "coding"]],
      ["nemotron-3-nano",  ["reasoning"]],
      ["mistral-large-3",  ["reasoning", "coding"]],
      ["cogito",           ["reasoning", "coding"]],
      ["ministral-3",      ["reasoning", "coding"]],
      ["gemma3",           ["reasoning", "coding", "vision"]],
    ]);

    function canonicalModelKey(model) {
      const raw = String(model || "").trim().toLowerCase();
      if (!raw) return "";
      const normalized = (raw.includes("/") && raw.includes(":"))
        ? `${raw.split(":", 1)[0].split("/").slice(-1)[0]}:${raw.slice(raw.indexOf(":") + 1)}`
        : raw;
      if (normalized.includes(":")) {
        const idx = normalized.indexOf(":");
        const family = normalized.slice(0, idx);
        let tag = normalized.slice(idx + 1);
        if (tag.endsWith("-cloud")) tag = tag.slice(0, -6);
        return `${family}:${tag}`.replace(/:+$/g, "");
      }
      return normalized.endsWith("-cloud") ? normalized.slice(0, -6) : normalized;
    }

    function modelMatchesPrefix(model, prefix) {
      const key = canonicalModelKey(model);
      const p = String(prefix || "").trim().toLowerCase();
      if (!key || !p) return false;
      return key.startsWith(p) || key.includes(p);
    }

    function getFamilyProfile(model) {
      for (const [prefix, profile] of MODEL_FAMILY_PROFILES) {
        if (modelMatchesPrefix(model, prefix)) return profile;
      }
      return FAMILY_PRIOR_DEFAULTS;
    }

    function getModelCapabilities(model) {
      const meta = getModelMeta(model);
      const caps = new Set(Array.isArray(meta.capabilities)
        ? meta.capabilities.map(x => String(x || "").toLowerCase()).filter(Boolean)
        : []);
      const key = canonicalModelKey(model);

      if (["vision", "-vl", ":vl", "gemini", "llava", "pixtral", "multimodal", "omni"].some(t => key.includes(t))) caps.add("vision");
      if (["coder", "code", "devstral", "claude-code", "swe", "terminal"].some(t => key.includes(t))) caps.add("coding");
      if (["thinking", "reason", "reasoning", "r1", "gpt-oss", "deepseek", "cogito", "kimi-k2", "glm-5", "glm-4.7", "glm-4.6"].some(t => key.includes(t))) caps.add("reasoning");
      for (const [prefix, hintedCaps] of MODEL_TRAIT_HINTS) {
        if (modelMatchesPrefix(key, prefix)) {
          hintedCaps.forEach(cap => caps.add(cap));
        }
      }
      caps.add("completion");
      return Array.from(caps);
    }

    function clamp(value, low, high) {
      return Math.max(low, Math.min(value, high));
    }

    function sizeQualityStrength(sizeB) {
      if (!Number.isFinite(sizeB) || sizeB <= 0) return 4.8;
      const normalized = Math.log2(Math.min(sizeB, 1000) + 1) / Math.log2(1001);
      return 3.8 + normalized * 6.2;
    }

    function sizeSpeedStrength(sizeB) {
      if (!Number.isFinite(sizeB) || sizeB <= 0) return 7.8;
      const normalized = Math.log2(Math.min(sizeB, 1000) + 1) / Math.log2(1001);
      return 9.8 - normalized * 7.2;
    }

    function contextStrength(ctx) {
      if (!Number.isFinite(ctx) || ctx <= 0) return 3.2;
      const normalized = clamp(Math.log2(ctx) / 18.0, 0.0, 1.08);
      const bonus = ctx >= 200000 ? 0.7 : (ctx >= 128000 ? 0.35 : 0);
      return Math.min(10.0, 3.6 + normalized * 6.0 + bonus);
    }

    function freshnessStrength(modifiedTs) {
      if (!Number.isFinite(modifiedTs) || modifiedTs <= 0) return 4.8;
      const ageDays = Math.max(0, (Date.now() / 1000 - modifiedTs) / 86400);
      if (ageDays <= 21) return 10.0;
      if (ageDays <= 45) return 9.4;
      if (ageDays <= 90) return 8.6;
      if (ageDays <= 180) return 7.6;
      if (ageDays <= 365) return 6.3;
      if (ageDays <= 540) return 5.2;
      return 4.3;
    }

    function nameSignalBonus(model, criterion) {
      const key = canonicalModelKey(model);
      let bonus = 0;
      if (criterion === "coding") {
        if (["coder", "devstral", "terminal", "swe"].some(t => key.includes(t))) bonus += 0.90;
        if (["code", "oss"].some(t => key.includes(t))) bonus += 0.25;
      } else if (criterion === "reasoning") {
        if (["thinking", "reason", "reasoning", "r1"].some(t => key.includes(t))) bonus += 0.95;
        if (["deepseek", "cogito"].some(t => key.includes(t))) bonus += 0.20;
      } else if (criterion === "vision") {
        if (["-vl", ":vl", "vision", "gemini", "pixtral", "llava"].some(t => key.includes(t))) bonus += 0.95;
      } else if (criterion === "speed") {
        if (["flash", "nano", "mini", "small"].some(t => key.includes(t))) bonus += 1.15;
        if (["preview"].some(t => key.includes(t))) bonus += 0.20;
      } else if (criterion === "overall") {
        if (["thinking", "coder", "-vl", ":vl", "vision"].some(t => key.includes(t))) bonus += 0.18;
      }
      return bonus;
    }

    function scoreModelForCriterion(model, criterion = "overall") {
      const chosenCriterion = ["overall", "coding", "reasoning", "context", "vision", "speed", "newest"].includes(String(criterion || "").toLowerCase())
        ? String(criterion || "overall").toLowerCase()
        : "overall";
      const profile = getFamilyProfile(model);
      const basePrior = Number(profile[chosenCriterion] || FAMILY_PRIOR_DEFAULTS[chosenCriterion] || 7.0);
      const sizeB = Math.max(0, Math.min(getModelSizeBillions(model) || 0, 1000));
      const ctx = Math.max(0, getModelContextTokens(model));
      const newest = getModelModifiedTs(model);
      const caps = getModelCapabilities(model);
      const sizeQuality = sizeQualityStrength(sizeB);
      const sizeSpeed = sizeSpeedStrength(sizeB);
      const ctxStrength = contextStrength(ctx);
      const freshness = freshnessStrength(newest);
      const hasReasoning = caps.includes("reasoning") ? 10 : 0;
      const hasCoding = caps.includes("coding") ? 10 : 0;
      const hasVision = caps.includes("vision") ? 10 : 0;
      const bonus = nameSignalBonus(model, chosenCriterion);

      switch (chosenCriterion) {
        case "coding":
          return basePrior * 0.56 + sizeQuality * 0.10 + ctxStrength * 0.10 + freshness * 0.05 + hasCoding * 0.14 + hasReasoning * 0.04 + bonus;
        case "reasoning":
          return basePrior * 0.56 + sizeQuality * 0.11 + ctxStrength * 0.06 + freshness * 0.04 + hasReasoning * 0.13 + hasCoding * 0.03 + bonus;
        case "context":
          return ctx > 0
            ? (ctxStrength * 0.74 + basePrior * 0.14 + sizeQuality * 0.08 + freshness * 0.04)
            : (basePrior * 0.22 + sizeQuality * 0.12 + freshness * 0.06);
        case "vision":
          return basePrior * 0.58 + sizeQuality * 0.09 + ctxStrength * 0.06 + freshness * 0.04 + hasVision * 0.18 + hasReasoning * 0.02 + bonus;
        case "speed":
          return basePrior * 0.15 + sizeSpeed * 0.60 + freshness * 0.10 + ctxStrength * 0.05 + ((sizeB > 0 && sizeB <= 24) ? 0.20 : 0) + bonus;
        case "newest":
          return newest > 0 ? newest : 0;
        case "overall":
        default:
          return basePrior * 0.56 + sizeQuality * 0.12 + ctxStrength * 0.06 + freshness * 0.05 + hasReasoning * 0.08 + hasCoding * 0.06 + hasVision * 0.04 + bonus;
      }
    }

    function sortModelsByCriterion(modelList, criterion = "overall") {
      const models = Array.isArray(modelList) ? [...modelList].filter(Boolean) : [];
      return models.sort((a, b) => {
        const diff = scoreModelForCriterion(b, criterion) - scoreModelForCriterion(a, criterion);
        if (Math.abs(diff) > 1e-9) return diff > 0 ? 1 : -1;
        return String(a).localeCompare(String(b), "en", { sensitivity: "base" });
      });
    }

    function getSortCriterionLabel(value) {
      const mapping = {
        overall: "Overall",
        coding: "Coding",
        reasoning: "Reasoning",
        context: "Long Context",
        vision: "Vision",
        speed: "Speed",
        newest: "Newest",
      };
      return mapping[value] || "Overall";
    }

    function getSavedThinkMode() {
      try {
        const saved = String(localStorage.getItem(THINK_MODE_KEY) || "").trim().toLowerCase();
        return ["auto", "on", "off", "low", "medium", "high"].includes(saved) ? saved : "on";
      } catch (_) {
        return "on";
      }
    }

    function saveThinkMode(value) {
      try {
        const normalized = String(value || "").trim().toLowerCase();
        if (normalized) localStorage.setItem(THINK_MODE_KEY, normalized);
      } catch (_) {}
    }

    function getThinkingSupportProfile(model) {
      const targetModel = String(model || "").trim();
      const key = canonicalModelKey(targetModel);
      const caps = getModelCapabilities(targetModel);
      const hasReasoning = caps.includes("reasoning");

      if (!targetModel) {
        return {
          profileKey: "empty",
          displayName: "Αναμονή μοντέλου",
          supportedModes: ["auto", "on", "off"],
          defaultMode: getSavedThinkMode(),
          exact: false,
          tone: "",
          note: "Επίλεξε μοντέλο για να προσαρμοστούν αυτόματα οι διαθέσιμες επιλογές Thinking Mode.",
        };
      }

      if (modelMatchesPrefix(key, "qwen3-coder-next")) {
        return {
          profileKey: "qwen3-coder-next",
          displayName: "Qwen3-Coder-Next",
          supportedModes: ["auto", "off"],
          defaultMode: "off",
          exact: true,
          tone: "",
          note: "Official non-thinking mode only. Το μοντέλο είναι γρήγορο coding model χωρίς ξεχωριστό reasoning trace.",
        };
      }

      if (modelMatchesPrefix(key, "gpt-oss")) {
        return {
          profileKey: "gpt-oss",
          displayName: "GPT-OSS",
          supportedModes: ["auto", "low", "medium", "high"],
          defaultMode: "medium",
          exact: true,
          tone: "",
          note: "Official reasoning effort only: low / medium / high. Το trace δεν απενεργοποιείται πλήρως.",
        };
      }

      if (modelMatchesPrefix(key, "qwen3-next")) {
        return {
          profileKey: "qwen3-next",
          displayName: "Qwen3-Next",
          supportedModes: ["auto", "on", "off", "low", "medium", "high"],
          defaultMode: "on",
          exact: false,
          tone: "warn",
          note: "Thinking-capable family. Στο Ollama Cloud το hard Off μπορεί να χρειάζεται compatibility fallback, οπότε το trace κρύβεται πλήρως στο UI όταν χρειαστεί.",
        };
      }

      if (modelMatchesPrefix(key, "deepseek-v3.1")) {
        return {
          profileKey: "deepseek-v3.1",
          displayName: "DeepSeek-V3.1",
          supportedModes: ["auto", "on", "off"],
          defaultMode: "on",
          exact: true,
          tone: "",
          note: "Official hybrid thinking / non-thinking model. Χρησιμοποιεί boolean thinking control (think=true/false).",
        };
      }

      if (modelMatchesPrefix(key, "deepseek-r1")) {
        return {
          profileKey: "deepseek-r1",
          displayName: "DeepSeek-R1",
          supportedModes: ["auto", "on", "off"],
          defaultMode: "on",
          exact: true,
          tone: "",
          note: "Official thinking model με boolean thinking control (think=true/false).",
        };
      }

      if (modelMatchesPrefix(key, "qwen3-vl")) {
        return {
          profileKey: "qwen3-vl",
          displayName: "Qwen3-VL",
          supportedModes: ["auto", "on", "off"],
          defaultMode: "on",
          exact: true,
          tone: "",
          note: "Thinking-capable vision model. Το Off στέλνει think=false και το app κρύβει τυχόν leaked trace αν το backend το επιστρέψει.",
        };
      }

      if (modelMatchesPrefix(key, "qwen3")) {
        return {
          profileKey: "qwen3",
          displayName: "Qwen3",
          supportedModes: ["auto", "on", "off"],
          defaultMode: "on",
          exact: true,
          tone: "",
          note: "Official thinking model με boolean thinking control (think=true/false).",
        };
      }

      if (hasReasoning) {
        return {
          profileKey: "generic-reasoning",
          displayName: "Reasoning-capable",
          supportedModes: ["auto", "on", "off"],
          defaultMode: "on",
          exact: false,
          tone: "warn",
          note: "Το μοντέλο φαίνεται reasoning-capable από metadata/όνομα. Εφαρμόζεται ασφαλές boolean mapping μέχρι να επιβεβαιωθεί πιο ειδικό profile.",
        };
      }

      return {
        profileKey: "non-thinking",
        displayName: "Non-thinking",
        supportedModes: ["auto", "off"],
        defaultMode: "off",
        exact: false,
        tone: "",
        note: "Δεν εντοπίστηκε official thinking support για το επιλεγμένο μοντέλο. Ενεργά παραμένουν μόνο τα ασφαλή modes Auto / Off.",
      };
    }

    function loadThinkingProfileConfirmations() {
      try {
        const raw = localStorage.getItem(THINK_PROFILE_CONFIRMATIONS_KEY);
        const parsed = raw ? JSON.parse(raw) : {};
        state.confirmedThinkingProfilesByModel = (parsed && typeof parsed === "object") ? parsed : {};
      } catch (_) {
        state.confirmedThinkingProfilesByModel = {};
      }
    }

    function saveThinkingProfileConfirmations() {
      try {
        localStorage.setItem(
          THINK_PROFILE_CONFIRMATIONS_KEY,
          JSON.stringify(state.confirmedThinkingProfilesByModel && typeof state.confirmedThinkingProfilesByModel === "object"
            ? state.confirmedThinkingProfilesByModel
            : {})
        );
      } catch (_) {}
    }

    function buildThinkingProfileSignature(profile) {
      const modes = Array.isArray(profile && profile.supportedModes) ? profile.supportedModes.join("|") : "";
      return `${String((profile && profile.profileKey) || "")}|${modes}`;
    }

    function getConfirmedThinkingProfileRecord(model) {
      const key = canonicalModelKey(model);
      const records = (state.confirmedThinkingProfilesByModel && typeof state.confirmedThinkingProfilesByModel === "object")
        ? state.confirmedThinkingProfilesByModel
        : {};
      return key ? (records[key] || null) : null;
    }

    function updateThinkingProfileConfirmButton(model, profile) {
      if (!els.confirmThinkingProfileBtn) return;
      const targetModel = String(model || "").trim();
      const confirmed = getConfirmedThinkingProfileRecord(targetModel);
      const currentSignature = buildThinkingProfileSignature(profile || {});
      const isConfirmed = !!(confirmed && confirmed.signature === currentSignature);
      els.confirmThinkingProfileBtn.disabled = !targetModel;
      els.confirmThinkingProfileBtn.textContent = isConfirmed ? "✅ Profile Επιβεβαιώθηκε" : "✅ Επιβεβαίωση Profile";
      els.confirmThinkingProfileBtn.title = targetModel
        ? (isConfirmed
            ? `Το profile για ${targetModel} έχει ήδη επιβεβαιωθεί.`
            : `Επιβεβαίωση του detected thinking profile για ${targetModel}`)
        : "Επίλεξε πρώτα μοντέλο για επιβεβαίωση profile.";
      els.confirmThinkingProfileBtn.dataset.confirmed = isConfirmed ? "1" : "0";
      els.confirmThinkingProfileBtn.classList.toggle("done", isConfirmed);
    }

    function confirmCurrentThinkingProfile() {
      const model = getSelectedModelKey();
      const profile = state.currentThinkingProfile || getThinkingSupportProfile(model);
      if (!model) {
        renderSystemNotice("Επίλεξε πρώτα μοντέλο για να επιβεβαιώσεις profile.");
        return;
      }
      const key = canonicalModelKey(model);
      state.confirmedThinkingProfilesByModel[key] = {
        model,
        profileKey: String((profile && profile.profileKey) || ""),
        displayName: String((profile && profile.displayName) || ""),
        signature: buildThinkingProfileSignature(profile || {}),
        confirmedAt: new Date().toISOString(),
      };
      saveThinkingProfileConfirmations();
      updateThinkingProfileConfirmButton(model, profile);
      applyThinkingModeSupportForModel(model, { preferredMode: (els.thinkModeSelect && els.thinkModeSelect.value) || getSavedThinkMode() });
      renderSystemNotice(`Επιβεβαιώθηκε το thinking profile για το μοντέλο ${model}.`);
    }

    function applyThinkingModeSupportForModel(model, options = {}) {
      const profile = getThinkingSupportProfile(model);
      state.currentThinkingProfile = profile;
      if (!els.thinkModeSelect) return profile;

      const labels = {
        auto: "Auto",
        on: "On",
        off: "Off",
        low: "Low",
        medium: "Medium",
        high: "High",
      };
      const allowed = new Set(Array.isArray(profile.supportedModes) ? profile.supportedModes : ["auto"]);
      const preferredRaw = Object.prototype.hasOwnProperty.call(options, "preferredMode")
        ? String(options.preferredMode || "")
        : String(els.thinkModeSelect.value || getSavedThinkMode() || profile.defaultMode || "auto");
      const preferred = preferredRaw.trim().toLowerCase();

      for (const opt of Array.from(els.thinkModeSelect.options || [])) {
        const value = String(opt.value || "").trim().toLowerCase();
        if (!opt.dataset.baseLabel) opt.dataset.baseLabel = opt.textContent;
        opt.textContent = opt.dataset.baseLabel;
        opt.disabled = !allowed.has(value);
      }

      let selected = preferred;
      if (!allowed.has(selected)) {
        selected = String(profile.defaultMode || "").trim().toLowerCase();
      }
      if (!allowed.has(selected)) {
        selected = allowed.has("auto") ? "auto" : (Array.from(allowed)[0] || "auto");
      }
      els.thinkModeSelect.value = selected;
      saveThinkMode(selected);

      const supportedPretty = Array.from(allowed).map(value => labels[value] || value).join(", ");
      const adjusted = preferred && selected !== preferred
        ? ` · προσαρμόστηκε σε <b>${labels[selected] || escapeHtml(selected)}</b>`
        : "";
      const confirmed = getConfirmedThinkingProfileRecord(model);
      const confirmedMatches = !!(confirmed && confirmed.signature === buildThinkingProfileSignature(profile));
      const confirmedHtml = confirmedMatches
        ? `<br><span class="tiny" style="display:inline-block; margin-top:4px; color:var(--ok);">✅ Επιβεβαιωμένο profile: ${escapeHtml((confirmed && confirmed.displayName) || (profile.displayName || "Thinking"))}</span>`
        : "";

      if (els.thinkingSupportInfo) {
        els.thinkingSupportInfo.innerHTML = `${profile.exact ? "Υποστήριξη" : "Υποστήριξη / συμβατότητα"} για <code>${escapeHtml(model || "-")}</code>: <b>${supportedPretty || "-"}</b>${adjusted}<br>${escapeHtml(profile.note || "")}${confirmedHtml}`;
        els.thinkingSupportInfo.className = `tiny ${profile.tone === "warn" ? "warn" : "muted"}`;
      }

      updateThinkingProfileConfirmButton(model, profile);
      els.thinkModeSelect.title = `${profile.displayName || "Thinking"} · ${profile.note || ""}`.trim();
      return profile;
    }

    function populateModelSelect(modelList, criterionRecommended = "", preferredModel = "", options = {}) {
      const allModels = Array.isArray(modelList) ? modelList.filter(Boolean) : [];
      const searchText = options && Object.prototype.hasOwnProperty.call(options, "searchText")
        ? String(options.searchText || "")
        : String((els.modelSearchInput && els.modelSearchInput.value) || "");
      const models = filterModelsBySearch(allModels, searchText);
      const preferWinner = !!(options && options.preferWinner);
      const allowSavedModel = options && Object.prototype.hasOwnProperty.call(options, "allowSavedModel")
        ? !!options.allowSavedModel
        : !preferWinner;
      els.modelSelect.innerHTML = "";

      if (!allModels.length) {
        const opt = document.createElement("option");
        opt.value = "";
        opt.textContent = "Δεν βρέθηκαν μοντέλα";
        els.modelSelect.appendChild(opt);
        return [];
      }

      if (!models.length) {
        const opt = document.createElement("option");
        opt.value = "";
        opt.textContent = "Καμία αντιστοίχιση στην αναζήτηση";
        els.modelSelect.appendChild(opt);
        return [];
      }

      for (const m of models) {
        const opt = document.createElement("option");
        opt.value = m;
        opt.textContent = (criterionRecommended && m === criterionRecommended) ? `★ ${m}` : m;
        els.modelSelect.appendChild(opt);
      }

      const saved = allowSavedModel ? getSavedModel(allModels) : null;
      const winner = (criterionRecommended && allModels.includes(criterionRecommended)) ? criterionRecommended : "";
      const currentValue = String(els.modelSelect.dataset.currentValue || els.modelSelect.value || "").trim();
      const preferred = (preferWinner ? winner : "")
        || (preferredModel && models.includes(preferredModel) ? preferredModel : "")
        || (currentValue && models.includes(currentValue) ? currentValue : "")
        || (saved && models.includes(saved) ? saved : "")
        || (winner && models.includes(winner) ? winner : "")
        || models[0] || "";

      els.modelSelect.value = preferred;
      els.modelSelect.dataset.currentValue = preferred;
      return models;
    }

    function rebuildModelSelect(preferredModel = "", preferWinner = false) {
      const criterion = (els.modelSortSelect && els.modelSortSelect.value) ? els.modelSortSelect.value : "overall";
      state.modelSortCriterion = criterion;
      const originalModels = Array.isArray(state.models) ? [...state.models] : [];
      const sortedModels = sortModelsByCriterion(originalModels, criterion);
      const criterionWinner = sortedModels[0] || "";
      const chosenModel = preferWinner ? "" : (preferredModel || els.modelSelect.value || "");
      populateModelSelect(sortedModels, criterionWinner, chosenModel, {
        preferWinner,
        allowSavedModel: !preferWinner,
      });
      updateHelperModelSelect();
      updateModelBadges();
      syncNumCtxInputForSelectedModel();
      return criterionWinner;
    }

    function updateHelperModelSelect(preferredModel = "") {
      if (!els.helperModelSelect) return [];
      const primaryModel = getSelectedModelKey();
      const allModels = sortModelsByCriterion(Array.isArray(state.models) ? [...state.models] : [], "overall")
        .filter((model) => model && model !== primaryModel);
      state.helperModels = allModels;
      els.helperModelSelect.innerHTML = "";

      if (!allModels.length) {
        const opt = document.createElement("option");
        opt.value = "";
        opt.textContent = "Δεν υπάρχουν helper models";
        els.helperModelSelect.appendChild(opt);
        return [];
      }

      const filtered = filterModelsBySearch(allModels, (els.helperSearchInput && els.helperSearchInput.value) || "");
      if (!filtered.length) {
        const opt = document.createElement("option");
        opt.value = "";
        opt.textContent = "Καμία αντιστοίχιση στην αναζήτηση";
        els.helperModelSelect.appendChild(opt);
        return [];
      }

      for (const model of filtered) {
        const opt = document.createElement("option");
        opt.value = model;
        opt.textContent = model;
        els.helperModelSelect.appendChild(opt);
      }

      const saved = getSavedHelperModel(allModels);
      const currentValue = String(els.helperModelSelect.dataset.currentValue || els.helperModelSelect.value || "").trim();
      const preferred = (preferredModel && filtered.includes(preferredModel) ? preferredModel : "")
        || (currentValue && filtered.includes(currentValue) ? currentValue : "")
        || (saved && filtered.includes(saved) ? saved : "")
        || filtered[0] || "";

      els.helperModelSelect.value = preferred;
      els.helperModelSelect.dataset.currentValue = preferred;
      return filtered;
    }

    function refreshHelperControls(preferredModel = "") {
      const mode = normalizeEnsembleMode((els.ensembleModeSelect && els.ensembleModeSelect.value) || state.ensembleMode || "auto");
      state.ensembleMode = mode;
      const manual = mode === "manual";
      if (els.helperSearchInput) els.helperSearchInput.disabled = !manual;
      if (els.helperModelSelect) els.helperModelSelect.disabled = !manual;
      updateHelperModelSelect(preferredModel);
      if (els.ensembleModeInfo) {
        els.ensembleModeInfo.textContent = manual
          ? "Manual: επίλεξε εσύ το δεύτερο helper model από τη λίστα. Το κύριο μοντέλο αποκλείεται αυτόματα."
          : mode === "auto"
            ? "Auto: το helper model επιλέγεται αυτόματα από το κύριο μοντέλο και το είδος του task."
            : "Off: χρησιμοποιείται μόνο το κύριο μοντέλο χωρίς helper.";
      }
      if (!state.isStreaming && els.helperText) {
        if (mode === "auto") {
          els.helperText.textContent = "🤝 Dual-model ensemble auto: ON";
        } else if (mode === "manual") {
          const helper = String((els.helperModelSelect && els.helperModelSelect.value) || "").trim();
          els.helperText.textContent = helper ? `🤝 Manual helper: ${helper}` : "🤝 Manual helper: επίλεξε δεύτερο μοντέλο";
        } else {
          els.helperText.textContent = DEFAULT_HELPER;
        }
      }
    }

    function updateModelBadges(data = null) {
      const model = els.modelSelect.value || "-";
      els.selectedModelBadge.textContent = `Μοντέλο: ${model}`;
      if (!data) return;

      const sourceText =
        data.source === "official-online" ? "Πηγή: official direct API catalog" :
        data.source === "stale-online-cache" ? "Πηγή: τελευταία επιτυχής official direct API λίστα" :
        data.source === "initializing" ? "Πηγή: αρχικοποίηση" :
        "Πηγή: σφάλμα online λίστας";

      els.sourceBadge.textContent = sourceText;
      els.sourceBadge.className   =
        data.source === "official-online" ? "badge ok" :
        data.last_error ? "badge warn" : "badge";
    }

    async function ensureSelectedModelMeta(model, force = false) {
      const targetModel = String(model || "").trim();
      if (!targetModel) return null;

      const existing = getModelMeta(targetModel);
      if (!force && existing && Number(existing.num_ctx_max || 0) >= 256 && existing.details_complete) {
        return existing;
      }
      if (!force && state.modelDetailRequests[targetModel]) {
        return state.modelDetailRequests[targetModel];
      }

      const request = (async () => {
        try {
          const query = `/api/model-details?model=${encodeURIComponent(targetModel)}${force ? "&force=1" : ""}`;
          const resp = await fetch(query);
          const data = await resp.json();
          if (resp.ok && data && data.meta && typeof data.meta === "object") {
            const merged = { ...(state.modelMetaByModel[targetModel] || {}), ...data.meta };
            state.modelMetaByModel[targetModel] = merged;
            const numCtxMax = Number(merged.num_ctx_max || 0);
            if (Number.isFinite(numCtxMax) && numCtxMax >= 256) {
              state.modelMaxNumCtxByModel[targetModel] = Math.trunc(numCtxMax);
            }
            if (targetModel === getSelectedModelKey()) {
              syncNumCtxInputForSelectedModel();
              applyThinkingModeSupportForModel(targetModel, { preferredMode: (els.thinkModeSelect && els.thinkModeSelect.value) || "" });
            }
            return merged;
          }
          return existing || null;
        } catch (_) {
          return existing || null;
        } finally {
          delete state.modelDetailRequests[targetModel];
        }
      })();

      state.modelDetailRequests[targetModel] = request;
      return request;
    }

    async function loadModels() {
      try {
        const currentSelection = getSelectedModelKey() || "";
        const resp = await fetch("/api/models");
        const data = await resp.json();

        const serverModels  = Array.isArray(data.models) ? data.models.filter(Boolean) : [];
        const modelMeta     = (data.model_meta && typeof data.model_meta === "object") ? data.model_meta : {};
        state.lastModelsRefreshTs = Number(data.last_refresh_ts || 0) || 0;
        state.modelMaxNumCtxByModel = {};
        state.modelMetaByModel = {};
        for (const [model, meta] of Object.entries(modelMeta)) {
          const safeMeta = (meta && typeof meta === "object") ? meta : {};
          state.modelMetaByModel[model] = safeMeta;
          const numCtxMax = Number((safeMeta && safeMeta.num_ctx_max) || 0);
          if (Number.isFinite(numCtxMax) && numCtxMax >= 256) {
            state.modelMaxNumCtxByModel[model] = Math.trunc(numCtxMax);
          }
        }

        state.models = Array.isArray(serverModels) ? [...serverModels] : [];
        const winner = rebuildModelSelect(currentSelection, !currentSelection);
        if (!currentSelection && winner) saveModel(winner);
        updateModelBadges(data);

        let info = `${state.models.length} μοντέλα Ollama Cloud/API · ταξινόμηση: καλύτερο → χειρότερο · πηγή: ${data.source || "-"}`;
        if (data.refresh_in_progress) info += " · ανανέωση σε εξέλιξη";
        if (winner && state.models.length) info += ` · 🏆 ${winner} (${getSortCriterionLabel(state.modelSortCriterion)})`;
        if (data.models_with_context != null) info += ` · ${data.models_with_context}/${state.models.length} με Max Length metadata`;
        if (data.last_error)  info += ` · ⚠ ${data.last_error}`;
        if (!serverModels.length) info += " · δεν βρέθηκε online official λίστα";
        els.modelInfo.textContent = info;
        syncNumCtxInputForSelectedModel();
        applyThinkingModeSupportForModel(getSelectedModelKey(), { preferredMode: getSavedThinkMode() });
        refreshHelperControls(getSavedHelperModel(Array.isArray(state.models) ? state.models : []) || "");
        ensureSelectedModelMeta(getSelectedModelKey());

      } catch (_) {
        state.modelMetaByModel = {};
        state.modelMaxNumCtxByModel = {};
        state.models = [];
        populateModelSelect([], "", "");
        updateHelperModelSelect();
        updateModelBadges({ source: "error", last_error: "Αποτυχία /api/models" });
        els.modelInfo.textContent = "Αποτυχία φόρτωσης official Ollama direct API models.";
        syncNumCtxInputForSelectedModel();
        applyThinkingModeSupportForModel("", { preferredMode: getSavedThinkMode() });
      }
    }

    async function waitForModelsRefresh(previousTs = 0) {
      const deadline = Date.now() + 45000;
      while (Date.now() < deadline) {
        try {
          const resp = await fetch("/api/models");
          const data = await resp.json();
          const ts = Number(data.last_refresh_ts || 0) || 0;
          if (!data.refresh_in_progress && (!previousTs || ts >= previousTs)) {
            return data;
          }
        } catch (_) {}
        await new Promise(resolve => setTimeout(resolve, MODEL_REFRESH_POLL_MS));
      }
      return null;
    }

    async function refreshModels(showNotice = false) {
      try {
        const previousTs = state.lastModelsRefreshTs || 0;
        els.modelInfo.textContent = "Ανανέωση όλων των διαθέσιμων official Ollama direct API models...";
        const resp = await fetch("/api/refresh-models", {
          method: "POST",
          headers: { "Content-Type": "application/json" },
          body: JSON.stringify({ force: true }),
        });
        try {
          const data = await resp.json();
          if (data && data.last_refresh_ts != null) {
            state.lastModelsRefreshTs = Number(data.last_refresh_ts || 0) || state.lastModelsRefreshTs;
          }
        } catch (_) {}
        const refreshed = await waitForModelsRefresh(previousTs);
        await loadModels();
        if (!refreshed) {
          updateModelBadges({ source: "error", last_error: "Η ανανέωση άργησε περισσότερο από το αναμενόμενο" });
          if (showNotice) renderSystemNotice("⚠ Η ανανέωση official direct API models άργησε περισσότερο από το αναμενόμενο. Εμφανίζονται τα πιο πρόσφατα διαθέσιμα δεδομένα.");
          return;
        }
        if (showNotice) renderSystemNotice(`Ανανεώθηκαν ${state.models.length} official Ollama direct API models.`);
      } catch (_) {
        updateModelBadges({ source: "error", last_error: "Αποτυχία ανανέωσης" });
        renderSystemNotice("Αποτυχία ανανέωσης official Ollama direct API models.");
      }
    }

    // ── Session management ────────────────────────────────────────────────────

    async function loadSession() {
      try {
        const resp = await fetch("/api/session");
        const data = await resp.json();

        els.messages.innerHTML = "";
        state.chatHistory      = [];
        resetReasoningPanel(true);

        if (!data.history || !data.history.length) { renderEmptyState(); return; }

        for (const item of data.history) {
          createMessage(item.role, item.content, item.attachments || []);
          state.chatHistory.push({ role: item.role, content: item.content, time: nowString() });
        }
      } catch (_) {
        renderSystemNotice("Αποτυχία φόρτωσης session.");
      }
    }

    async function clearChat() {
      let serverOk = false;
      let serverError = "";

      try {
        const resp = await fetch("/api/reset-chat", {
          method: "POST", headers: { "Content-Type": "application/json" }, body: "{}",
        });
        let data = {};
        try { data = await resp.json(); } catch (_) {}
        if (!resp.ok || data.ok === false) {
          throw new Error((data && data.error) || `HTTP ${resp.status}`);
        }
        serverOk = true;
      } catch (err) {
        serverError = (err && err.message) ? err.message : String(err || "Άγνωστο σφάλμα");
      }

      try {
        state.selectedFiles = [];
        state.chatHistory   = [];
        state.currentAssistantNode = null;
        state.currentThinkingText = "";
        state.reasoningStreamCompleted = false;
        resetReasoningPanel(true);
        renderSelectedFiles();
        renderEmptyState();
        if (serverOk) {
          renderSystemNotice("Το chat και τα προσωρινά αρχεία καθαρίστηκαν.");
        } else {
          renderSystemNotice(`Το chat καθαρίστηκε τοπικά, αλλά ο server δεν ολοκλήρωσε πλήρως τον καθαρισμό: ${serverError}`);
        }
      } catch (uiErr) {
        const uiMessage = (uiErr && uiErr.message) ? uiErr.message : String(uiErr || "Άγνωστο σφάλμα UI");
        renderSystemNotice(`Αποτυχία καθαρισμού chat: ${uiMessage}`);
      }
    }

    // ── Export chat ───────────────────────────────────────────────────────────

    function exportChat() {
      if (!state.chatHistory.length) {
        renderSystemNotice("Δεν υπάρχει ιστορικό για εξαγωγή."); return;
      }
      const model = els.modelSelect.value || "unknown";
      const date  = new Date().toLocaleString("el-GR");
      let md = `# Ollama Chat Export\\n\\n**Μοντέλο:** ${model}  \\n**Ημερομηνία:** ${date}\\n\\n---\\n\\n`;

      for (const item of state.chatHistory) {
        md += `## ${item.role === "user" ? "👤 Χρήστης" : "🤖 Assistant"}`;
        md += `  \\n*${item.time}*\\n\\n${item.content}\\n\\n---\\n\\n`;
      }

      const blob = new Blob([md], { type: "text/markdown;charset=utf-8" });
      const url  = URL.createObjectURL(blob);
      const a    = document.createElement("a");
      a.href     = url;
      a.download = `ollama-chat-${new Date().toISOString().slice(0, 10)}.md`;
      // Append to DOM, click, then remove — required for Firefox
      document.body.appendChild(a);
      a.click();
      document.body.removeChild(a);
      URL.revokeObjectURL(url);
    }

    // ── Streaming controls ────────────────────────────────────────────────────

    function setControlsDisabled(active) {
      const controls = [
        els.modelSearchInput, els.modelSelect, els.thinkModeSelect, els.confirmThinkingProfileBtn, els.ensembleModeSelect,
        els.helperSearchInput, els.helperModelSelect, els.promptProfileSelect, els.visualizationEngineSelect, els.systemPrompt, els.fileInput,
        els.refreshModelsBtn, els.clearFilesBtn, els.resetSystemPromptBtn,
        els.copySystemPromptBtn, els.themeToggleBtn, els.clearChatBtn,
        els.reloadSessionBtn, els.userInput, els.sendBtn,
        els.exportChatBtn, els.autoScrollBtn,
        els.paramTemp, els.paramTopP, els.paramSeed, els.paramNumCtx,
        els.clearNumCtxBtn, els.resetParamsBtn,
        els.apiKeyInput, els.saveApiKeyBtn, els.clearApiKeyBtn, els.activatePromptProfileBtn, els.savePromptSetupBtn,
      ];
      for (const el of controls) if (el) el.disabled = active;
      if (els.stopBtn) els.stopBtn.disabled = !active;
    }

    function finalizeStream(completionLabel = "") {
      state.isStreaming = false;
      setControlsDisabled(false);
      refreshHelperControls();
      if (els.stopBtn)     els.stopBtn.disabled     = true;
      if (els.streamBadge) {
        els.streamBadge.textContent = "Έτοιμο";
        els.streamBadge.className   = "badge ok";
        // Reset back to neutral after 4s
        setTimeout(() => {
          if (!state.isStreaming && els.streamBadge) {
            els.streamBadge.textContent = "Έτοιμο";
            els.streamBadge.className   = "badge";
          }
        }, 4000);
      }
      if (completionLabel && els.helperText) {
        els.helperText.textContent = completionLabel;
      }
    }

    function setStreamState(active, label = null) {
      state.isStreaming           = active;
      setControlsDisabled(active);
      if (els.messages && els.messages.classList) {
        els.messages.classList.toggle("streaming-active", !!active);
      }
      els.streamBadge.textContent = label || (active ? "Streaming..." : "Έτοιμο");
      els.streamBadge.className   = "badge" + (active ? " warn" : "");
    }

    function stopStreaming() {
      if (state.abortController) state.abortController.abort();
    }

    function schedulePageReloadAfterAnswer(delayMs = 350) {
      return;
    }

    // ── Send message ──────────────────────────────────────────────────────────

    async function sendMessage() {
      if (state.isStreaming) return;

      const userText     = els.userInput.value.trim();
      const model        = els.modelSelect.value;
      const thinkMode    = els.thinkModeSelect ? els.thinkModeSelect.value : "on";
      const systemPrompt = els.systemPrompt.value;

      if (!userText) { els.userInput.focus(); renderSystemNotice("Γράψε πρώτα το user prompt."); return; }
      if (!model)    { renderSystemNotice("Δεν υπάρχει επιλεγμένο μοντέλο."); return; }

      let attachmentsPayload = [];
      try {
        attachmentsPayload = await collectFilesPayload();
      } catch (err) {
        renderSystemNotice(`Σφάλμα αρχείων: ${err.message || err}`); return;
      }

      const optimisticAttachments = state.selectedFiles.map(f => ({
        name: f.name,
        kind: (f.type || "").startsWith("image/") ? "image" : "document",
      }));

      createMessage("user", userText, optimisticAttachments);
      state.chatHistory.push({ role: "user", content: userText, time: nowString() });

      els.userInput.value         = "";
      els.charCounter.textContent = "0 χαρ. / 0 λέξεις";
      els.charCounter.className   = "char-counter";
      state.selectedFiles         = [];
      els.fileInput.value         = "";
      renderSelectedFiles();

      const assistantMsg         = createMessage("assistant", "", []);
      assistantMsg.wrapper.dataset.exportModel = String(model || "").trim();
      assistantMsg.body.dataset.exportModel = String(model || "").trim();
      state.currentAssistantNode = assistantMsg.body;
      state.abortController      = new AbortController();
      resetReasoningPanel(true);
      state.reasoningStreamCompleted = false;

      const ensembleMode = normalizeEnsembleMode((els.ensembleModeSelect && els.ensembleModeSelect.value) || state.ensembleMode || "auto");
      const manualHelperModel = ensembleMode === "manual"
        ? String((els.helperModelSelect && els.helperModelSelect.value) || "").trim()
        : "";
      if (ensembleMode === "manual" && !manualHelperModel) {
        renderSystemNotice("Επίλεξε helper model από τη λίστα πριν στείλεις μήνυμα.");
        if (state.currentAssistantNode && state.currentAssistantNode.parentElement) {
          state.currentAssistantNode.parentElement.remove();
        }
        state.currentAssistantNode = null;
        state.abortController = null;
        return;
      }

      setStreamState(true, "Streaming...");
      els.helperText.textContent = `Αποστολή προς ${model}…`;
      saveModel(model);   // persist model selection
      if (manualHelperModel) saveHelperModel(manualHelperModel);
      let completionText = "";
      
      try {
        const response = await fetch("/api/chat", {
          method: "POST",
          headers: { "Content-Type": "application/json" },
          body: JSON.stringify({
            model, think_mode: thinkMode, system_prompt: systemPrompt,
            prompt_profile_id: (els.promptProfileSelect && els.promptProfileSelect.value) || state.activePromptProfile || DEFAULT_PROMPT_PROFILE_ID,
            visualization_engine: (els.visualizationEngineSelect && els.visualizationEngineSelect.value) || state.visualizationEngine || DEFAULT_VISUALIZATION_ENGINE,
            user_text: userText, attachments: attachmentsPayload,
            options: getModelOptions(),
            ensemble_mode: ensembleMode,
            ensemble_helper_model: manualHelperModel,
            ensemble_auto: ensembleMode === "auto",
          }),
          signal: state.abortController.signal,
        });

        if (!response.ok) {
          let serverError = `HTTP ${response.status}`;
          try { const d = await response.json(); if (d.error) serverError = d.error; } catch (_) {}
          throw new Error(serverError);
        }
        if (!response.body) throw new Error("Η ροή απάντησης δεν είναι διαθέσιμη.");

        const reader   = response.body.getReader();
        const decoder  = new TextDecoder("utf-8");
        let buffer       = "";
        let finalText    = "";
        let finalThinking = "";
        const streamStartMs = Date.now();

        while (true) {
          const { value, done } = await reader.read();
          if (done) break;

          buffer += decoder.decode(value, { stream: true });
          const lines = buffer.split("\\n");
          buffer = lines.pop() || "";

          for (const line of lines) {
            if (!line.trim()) continue;
            let payload = null;
            try { payload = JSON.parse(line); } catch { continue; }

            if (payload.type === "meta") {
              const exportPrimaryModel = String((payload.ensemble && payload.ensemble.primary_model) || model || "").trim();
              const exportHelperModel = String((payload.ensemble && payload.ensemble.helper_model) || "").trim();
              const exportModelLabel = exportPrimaryModel && exportHelperModel
                ? `${exportPrimaryModel} + ${exportHelperModel}`
                : (exportPrimaryModel || exportHelperModel || String(model || "").trim());
              if (assistantMsg && assistantMsg.wrapper && exportModelLabel) {
                assistantMsg.wrapper.dataset.exportModel = exportModelLabel;
              }
              if (assistantMsg && assistantMsg.body && exportModelLabel) {
                assistantMsg.body.dataset.exportModel = exportModelLabel;
              }
              if (payload.ensemble && payload.ensemble.enabled && payload.ensemble.helper_model && els.helperText) {
                const roleLabel = payload.ensemble.role_label || payload.ensemble.role || "helper";
                const reasonLabel = payload.ensemble.selection_reason ? ` · ${payload.ensemble.selection_reason}` : "";
                els.helperText.textContent = `🤝 Ensemble: ${payload.ensemble.primary_model || model} + ${payload.ensemble.helper_model} (${roleLabel}${reasonLabel})`;
              } else if (payload.ensemble && payload.ensemble.mode === "manual" && els.helperText) {
                els.helperText.textContent = "🤝 Manual helper ενεργό";
              }
              if (payload.warnings && payload.warnings.length) {
                renderSystemNotice(payload.warnings.join("\\n"));
              }
            } else if (payload.type === "thinking") {
              state.reasoningStreamCompleted = false;
              finalThinking += payload.content || "";
              renderAssistantStreamingView(finalText, finalThinking, false);
              els.streamBadge.textContent = "🧠 Thinking...";
              showLiveTokenStats(finalThinking, streamStartMs, "Thinking");
              scrollToBottom();
            } else if (payload.type === "thinking_done") {
              state.reasoningStreamCompleted = true;
              if (finalThinking.trim()) {
                updateReasoningPanel(finalThinking, true, false);
              }
            } else if (payload.type === "delta") {
              finalText += payload.content || "";
              renderAssistantStreamingView(finalText, finalThinking, false);
              if (!finalText.length && finalThinking.length) {
                els.streamBadge.textContent = "🧠 Thinking...";
                showLiveTokenStats(finalThinking, streamStartMs, "Thinking");
              } else {
                els.streamBadge.textContent = "✍️ Generating...";
                showLiveTokenStats(finalText, streamStartMs, "Generating");
              }
              scrollToBottom();
            } else if (payload.type === "error") {
              throw new Error(payload.error || "Άγνωστο σφάλμα.");
            } else if (payload.type === "done") {
              renderAssistantStreamingView(finalText, finalThinking, true);
              if (payload.elapsed_sec != null) {
                completionText = `✅ ${payload.elapsed_sec.toFixed(2)}s`;
              }
              if (payload.token_stats) {
                showTokenStats(payload.token_stats);
                const realTps = Number(payload.token_stats.tokens_per_sec || 0);
                if (payload.elapsed_sec != null) {
                  completionText += ` · ⚡ ${realTps.toFixed(1)} tok/s`;
                }
                els.streamBadge.textContent = `⚡ ${realTps.toFixed(1)} tok/s`;
                els.streamBadge.className = "badge ok";
              }
              finalizeStream(completionText);
            }
          }
        }

        const displayText = finalThinking.trim()
          ? composeDisplayContent(finalText, finalThinking)
          : finalText;

        if (!finalText.trim() && !finalThinking.trim()) {
          renderMessageContent(
            state.currentAssistantNode,
            "Δεν επιστράφηκε κείμενο. Έλεγξε το API key από το GUI/settings file και αν το μοντέλο είναι διαθέσιμο στο direct cloud catalog."
          );
          resetReasoningPanel(true);
        } else {
          renderAssistantStreamingView(finalText, finalThinking, true);
          state.chatHistory.push({ role: "assistant", content: displayText, time: nowString() });
        }


      } catch (err) {
        const errText = err && err.name === "AbortError"
          ? "Η ροή σταμάτησε από τον χρήστη."
          : `Σφάλμα: ${err && err.message ? err.message : String(err)}`;
        renderMessageContent(state.currentAssistantNode, errText);
        resetReasoningPanel(true);
        renderSystemNotice(errText);
      } finally {
        state.currentAssistantNode = null;
        state.abortController      = null;
        finalizeStream(completionText);
      }
    }

    function generateBrowserSessionId() {
      try {
        if (window.crypto && typeof window.crypto.randomUUID === "function") {
          return window.crypto.randomUUID();
        }
      } catch (_) {}
      return `browser_${Date.now()}_${Math.random().toString(16).slice(2)}`;
    }

    function getOrCreateBrowserSessionId() {
      try {
        const saved = sessionStorage.getItem(BROWSER_SESSION_KEY);
        if (saved) return saved;
        const created = generateBrowserSessionId();
        sessionStorage.setItem(BROWSER_SESSION_KEY, created);
        return created;
      } catch (_) {
        return generateBrowserSessionId();
      }
    }

    function postBrowserLifecycle(eventName, useBeacon = false) {
      const sessionId = String(state.browserSessionId || "").trim();
      if (!sessionId) return;
      const payload = JSON.stringify({ session_id: sessionId, event: eventName });
      if (useBeacon && navigator.sendBeacon) {
        try {
          const blob = new Blob([payload], { type: "application/json" });
          navigator.sendBeacon("/api/browser-session", blob);
          return;
        } catch (_) {}
      }
      fetch("/api/browser-session", {
        method: "POST",
        headers: { "Content-Type": "application/json" },
        body: payload,
        keepalive: true,
      }).catch(() => {});
    }

    function startBrowserHeartbeat() {
      if (state.browserHeartbeatTimer) {
        clearInterval(state.browserHeartbeatTimer);
      }
      state.browserHeartbeatTimer = setInterval(() => {
        postBrowserLifecycle("heartbeat", false);
      }, BROWSER_HEARTBEAT_MS);
    }

    function setupBrowserLifecycle() {
      state.browserSessionId = getOrCreateBrowserSessionId();
      postBrowserLifecycle("open", false);
      startBrowserHeartbeat();
      window.addEventListener("pageshow", () => postBrowserLifecycle("open", false));
      window.addEventListener("focus", () => postBrowserLifecycle("focus", false));
      let browserClosePosted = false;
      const postBrowserCloseOnce = () => {
        if (browserClosePosted) return;
        browserClosePosted = true;
        postBrowserLifecycle("close", true);
      };
      window.addEventListener("pagehide", (event) => {
        if (event && event.persisted) return;
        postBrowserCloseOnce();
      });
      window.addEventListener("beforeunload", postBrowserCloseOnce);
      document.addEventListener("visibilitychange", () => {
        if (document.visibilityState === "visible") {
          postBrowserLifecycle("visible", false);
        }
      });
    }

    // ── Theme ─────────────────────────────────────────────────────────────────

    function getSavedTheme() {
      try { const t = localStorage.getItem(THEME_KEY); return t === "light" ? "light" : "dark"; }
      catch { return "dark"; }
    }

    function switchPrismTheme(theme) {
      const dark  = document.getElementById("prismDark");
      const light = document.getElementById("prismLight");
      if (!dark || !light) return;
      if (theme === "light") {
        dark.disabled  = true;
        light.disabled = false;
      } else {
        light.disabled = true;
        dark.disabled  = false;
      }
      // Re-highlight ολόκληρη η σελίδα μετά το CSS swap
      schedulePrismHighlightInContainer(document);
    }

    function applyTheme(theme) {
      state.theme = theme === "light" ? "light" : "dark";
      document.documentElement.setAttribute("data-theme", state.theme);
      try { localStorage.setItem(THEME_KEY, state.theme); } catch (_) {}
      els.themeToggleBtn.textContent = state.theme === "light" ? "🌙 Dark Theme" : "☀️ Light Theme";
      switchPrismTheme(state.theme);
    }

    function toggleTheme() {
      applyTheme(state.theme === "light" ? "dark" : "light");
    }

    // ── Prompt profiles & visualization engine helpers ───────────────────────

    function rebuildPromptProfileMap() {
      state.promptProfileMap = {};
      for (const item of Array.isArray(state.promptProfiles) ? state.promptProfiles : []) {
        if (!item || !item.id) continue;
        state.promptProfileMap[String(item.id)] = item;
      }
    }

    function getPromptProfileById(profileId) {
      const fallbackId = state.activePromptProfile || DEFAULT_PROMPT_PROFILE_ID || "scientific-technical";
      const id = String(profileId || fallbackId || "").trim();
      return state.promptProfileMap[id] || state.promptProfileMap[fallbackId] || null;
    }

    function getPromptProfilePrompt(profileId) {
      const profile = getPromptProfileById(profileId);
      if (profile && typeof profile.prompt === "string" && profile.prompt.trim()) return profile.prompt;
      return DEFAULT_SYSTEM_PROMPT;
    }

    function populatePromptProfileSelect(selectedId = "") {
      if (!els.promptProfileSelect) return;
      const wanted = String(selectedId || state.activePromptProfile || DEFAULT_PROMPT_PROFILE_ID || "scientific-technical").trim();
      els.promptProfileSelect.innerHTML = "";
      for (const item of Array.isArray(state.promptProfiles) ? state.promptProfiles : []) {
        const option = document.createElement("option");
        option.value = item.id;
        option.textContent = item.label || item.id;
        els.promptProfileSelect.appendChild(option);
      }
      els.promptProfileSelect.value = wanted || DEFAULT_PROMPT_PROFILE_ID;
      state.activePromptProfile = els.promptProfileSelect.value || DEFAULT_PROMPT_PROFILE_ID;
      updatePromptProfileInfo();
    }

    function populateVisualizationEngineSelect(selectedEngine = "") {
      if (!els.visualizationEngineSelect) return;
      const wanted = String(selectedEngine || state.visualizationEngine || DEFAULT_VISUALIZATION_ENGINE || "auto").trim();
      els.visualizationEngineSelect.innerHTML = "";
      for (const item of Array.isArray(VISUALIZATION_ENGINE_OPTIONS) ? VISUALIZATION_ENGINE_OPTIONS : []) {
        const option = document.createElement("option");
        option.value = item.id;
        option.textContent = item.label || item.id;
        els.visualizationEngineSelect.appendChild(option);
      }
      els.visualizationEngineSelect.value = wanted || DEFAULT_VISUALIZATION_ENGINE || "auto";
      state.visualizationEngine = els.visualizationEngineSelect.value || DEFAULT_VISUALIZATION_ENGINE || "auto";
      updateVisualizationEngineInfo();
    }

    function updatePromptProfileInfo() {
      if (!els.promptProfileInfo) return;
      const profile = getPromptProfileById((els.promptProfileSelect && els.promptProfileSelect.value) || state.activePromptProfile);
      const promptText = String((els.systemPrompt && els.systemPrompt.value) || "").trim();
      const profilePrompt = String((profile && profile.prompt) || "").trim();
      const customState = promptText && profilePrompt && promptText !== profilePrompt
        ? " · υπάρχει custom override στο textarea"
        : (!promptText ? " · το textarea είναι κενό" : "");
      els.promptProfileInfo.textContent = profile
        ? `${profile.label || profile.id}: ${profile.description || ""}${customState}`
        : `Δεν βρέθηκε prompt profile.${customState}`;
    }

    function updateVisualizationEngineInfo() {
      if (!els.visualizationEngineInfo) return;
      const selected = String((els.visualizationEngineSelect && els.visualizationEngineSelect.value) || state.visualizationEngine || DEFAULT_VISUALIZATION_ENGINE || "auto").trim();
      const item = Array.isArray(VISUALIZATION_ENGINE_OPTIONS)
        ? VISUALIZATION_ENGINE_OPTIONS.find(entry => entry && entry.id === selected)
        : null;
      els.visualizationEngineInfo.textContent = item
        ? `${item.label || item.id}: ${item.description || ""}`
        : "Auto: έξυπνη επιλογή SVG ή Python plot ανά περίπτωση.";
    }

    function activatePromptProfile(notify = true) {
      const profile = getPromptProfileById((els.promptProfileSelect && els.promptProfileSelect.value) || state.activePromptProfile);
      if (!profile) {
        renderSystemNotice("Δεν βρέθηκε έγκυρο prompt profile.");
        return;
      }
      state.activePromptProfile = profile.id || DEFAULT_PROMPT_PROFILE_ID;
      if (els.promptProfileSelect) els.promptProfileSelect.value = state.activePromptProfile;
      if (els.systemPrompt) els.systemPrompt.value = String(profile.prompt || DEFAULT_SYSTEM_PROMPT || "");
      updatePromptProfileInfo();
      if (notify) renderSystemNotice(`Ενεργοποιήθηκε το prompt profile: ${profile.label || profile.id}`);
    }

    async function savePromptSetup() {
      try {
        const key = String((els.apiKeyInput && els.apiKeyInput.value) || "").trim();
        const payload = {
          ollama_api_key: key,
          active_prompt_profile: (els.promptProfileSelect && els.promptProfileSelect.value) || state.activePromptProfile || DEFAULT_PROMPT_PROFILE_ID,
          custom_system_prompt: (els.systemPrompt && els.systemPrompt.value) || "",
          active_visualization_engine: (els.visualizationEngineSelect && els.visualizationEngineSelect.value) || state.visualizationEngine || DEFAULT_VISUALIZATION_ENGINE,
        };
        const resp = await fetch("/api/app-config", {
          method: "POST",
          headers: { "Content-Type": "application/json" },
          body: JSON.stringify(payload),
        });
        const data = await resp.json().catch(() => ({}));
        if (!resp.ok) throw new Error(data.error || `HTTP ${resp.status}`);
        state.activePromptProfile = String((data.config && data.config.active_prompt_profile) || payload.active_prompt_profile || DEFAULT_PROMPT_PROFILE_ID);
        state.visualizationEngine = String((data.config && data.config.active_visualization_engine) || payload.active_visualization_engine || DEFAULT_VISUALIZATION_ENGINE);
        if (els.promptProfileSelect) els.promptProfileSelect.value = state.activePromptProfile;
        if (els.visualizationEngineSelect) els.visualizationEngineSelect.value = state.visualizationEngine;
        updatePromptProfileInfo();
        updateVisualizationEngineInfo();
        renderSystemNotice("Αποθηκεύτηκαν prompt profile, custom prompt και visualization engine στο settings file.");
      } catch (err) {
        renderSystemNotice(`Σφάλμα αποθήκευσης prompt setup: ${err && err.message ? err.message : String(err)}`);
      }
    }

    // ── System prompt helpers ─────────────────────────────────────────────────

    function resetSystemPrompt() {
      const profilePrompt = getPromptProfilePrompt((els.promptProfileSelect && els.promptProfileSelect.value) || state.activePromptProfile);
      els.systemPrompt.value = profilePrompt || DEFAULT_SYSTEM_PROMPT;
      updatePromptProfileInfo();
      renderSystemNotice("System prompt επανήλθε στο ενεργό prompt profile.");
    }

    async function copySystemPrompt() {
      try {
        await navigator.clipboard.writeText(els.systemPrompt.value || "");
        renderSystemNotice("System prompt αντιγράφηκε στο clipboard.");
      } catch {
        renderSystemNotice("Αποτυχία αντιγραφής.");
      }
    }

    // ── Help / User Guide modal ──────────────────────────────────────────────

    function isHelpModalOpen() {
      return Boolean(els.helpModal && els.helpModal.classList.contains("open"));
    }

    function openHelpModal() {
      if (!els.helpModal) return;
      els.helpModal.classList.add("open");
      els.helpModal.setAttribute("aria-hidden", "false");
      document.body.style.overflow = "hidden";
      window.setTimeout(() => {
        if (els.closeHelpModalBtn) els.closeHelpModalBtn.focus();
      }, 10);
    }

    function closeHelpModal({ restoreFocus = true } = {}) {
      if (!els.helpModal) return;
      els.helpModal.classList.remove("open");
      els.helpModal.setAttribute("aria-hidden", "true");
      document.body.style.overflow = "";
      if (restoreFocus && els.helpGuideBtn) {
        window.setTimeout(() => {
          try { els.helpGuideBtn.focus(); } catch (_) {}
        }, 10);
      }
    }

    // ── Auto-scroll ───────────────────────────────────────────────────────────

    function toggleAutoScroll() {
      state.autoScroll = !state.autoScroll;
      els.autoScrollBtn.textContent = `📜 Auto-Scroll: ${state.autoScroll ? "ON" : "OFF"}`;
      els.autoScrollBtn.style.opacity = state.autoScroll ? "1" : "0.6";
      updateScrollToBottomBtn();
      if (state.autoScroll) scrollToBottom(true);
    }

    // ── Char counter ──────────────────────────────────────────────────────────

    function updateCharCounter() {
      const text  = els.userInput.value;
      const chars = text.length;
      const words = countWords(text);
      els.charCounter.textContent = `${chars.toLocaleString("el-GR")} χαρ. / ${words} λέξεις`;
      els.charCounter.className   = `char-counter${chars > CHAR_WARN ? " warn" : ""}`;
    }

    // ── Event listeners ───────────────────────────────────────────────────────

    els.fileInput.addEventListener("change", (e) => addFiles(e.target.files));

    els.sendBtn.addEventListener("click",              sendMessage);
    els.stopBtn.addEventListener("click",              stopStreaming);
    if (els.activatePromptProfileBtn) els.activatePromptProfileBtn.addEventListener("click", () => activatePromptProfile(true));
    if (els.savePromptSetupBtn) els.savePromptSetupBtn.addEventListener("click", savePromptSetup);
    if (els.promptProfileSelect) els.promptProfileSelect.addEventListener("change", () => {
      state.activePromptProfile = els.promptProfileSelect.value || DEFAULT_PROMPT_PROFILE_ID;
      activatePromptProfile(false);
    });
    if (els.visualizationEngineSelect) els.visualizationEngineSelect.addEventListener("change", () => {
      state.visualizationEngine = els.visualizationEngineSelect.value || DEFAULT_VISUALIZATION_ENGINE;
      updateVisualizationEngineInfo();
    });
    if (els.systemPrompt) els.systemPrompt.addEventListener("input", updatePromptProfileInfo);
    els.resetSystemPromptBtn.addEventListener("click", resetSystemPrompt);
    els.copySystemPromptBtn.addEventListener("click",  copySystemPrompt);
    els.themeToggleBtn.addEventListener("click",       toggleTheme);
    els.exportChatBtn.addEventListener("click",        exportChat);
    els.autoScrollBtn.addEventListener("click",        toggleAutoScroll);
    if (els.helpGuideBtn) els.helpGuideBtn.addEventListener("click", openHelpModal);
    if (els.closeHelpModalBtn) els.closeHelpModalBtn.addEventListener("click", () => closeHelpModal());
    if (els.helpModal) {
      els.helpModal.addEventListener("click", (event) => {
        if (event.target === els.helpModal) closeHelpModal({ restoreFocus: false });
      });
    }
    document.addEventListener("keydown", (event) => {
      if (event.key === "Escape" && isHelpModalOpen()) {
        event.preventDefault();
        closeHelpModal();
      }
    });
    els.scrollToBottomBtn.addEventListener("click",    () => scrollToBottom(true));
    els.resetParamsBtn.addEventListener("click",       resetParams);
    els.clearFilesBtn.addEventListener("click", () => {
      state.selectedFiles = []; els.fileInput.value = ""; renderSelectedFiles();
    });
    els.clearChatBtn.addEventListener("click",    clearChat);
    els.reloadSessionBtn.addEventListener("click", loadSession);
    if (els.toggleReasoningBtn) {
      els.toggleReasoningBtn.addEventListener("click", () => {
        const hasContent = Boolean(((els.reasoningContent && els.reasoningContent.textContent) || "").trim());
        const currentlyVisible = hasContent && (state.reasoningPanelVisible || (state.reasoningAutoOpen && !state.reasoningUserCollapsed));
        if (currentlyVisible) {
          state.reasoningPanelVisible = false;
          state.reasoningAutoOpen = false;
          state.reasoningUserCollapsed = true;
        } else {
          state.reasoningPanelVisible = true;
          state.reasoningAutoOpen = false;
          state.reasoningUserCollapsed = false;
        }
        applyReasoningPanelVisibility();
      });
    }
    if (els.saveApiKeyBtn) els.saveApiKeyBtn.addEventListener("click", saveApiKey);
    if (els.clearApiKeyBtn) els.clearApiKeyBtn.addEventListener("click", clearApiKey);
    if (els.apiKeyInput) {
      els.apiKeyInput.addEventListener("keydown", (e) => {
        if (e.key === "Enter") { e.preventDefault(); saveApiKey(); }
      });
    }
    els.refreshModelsBtn.addEventListener("click", () => refreshModels(true));
    if (els.modelSortSelect) {
      els.modelSortSelect.addEventListener("change", () => {
        state.modelSortCriterion = els.modelSortSelect.value || "overall";
        saveSortCriterion(state.modelSortCriterion);
        const winner = rebuildModelSelect("", true);
        if (winner) saveModel(winner);
        applyThinkingModeSupportForModel(getSelectedModelKey(), { preferredMode: getSavedThinkMode() });
        let info = `${state.models.length} μοντέλα Ollama Cloud/API · ταξινόμηση: καλύτερο → χειρότερο`;
        if (winner) info += ` · 🏆 ${winner} (${getSortCriterionLabel(state.modelSortCriterion)})`;
        els.modelInfo.textContent = info;
      });
    }
    if (els.modelSearchInput) {
      els.modelSearchInput.addEventListener("input", () => {
        const previous = String(els.modelSelect.value || "").trim();
        populateModelSelect(sortModelsByCriterion(Array.isArray(state.models) ? [...state.models] : [], state.modelSortCriterion), "", previous, {
          preferWinner: false,
          allowSavedModel: true,
          searchText: els.modelSearchInput.value,
        });
        updateModelBadges();
      });
      els.modelSearchInput.addEventListener("keydown", (e) => {
        if (e.key === "Enter" && els.modelSelect && els.modelSelect.options.length > 0) {
          e.preventDefault();
          const first = els.modelSelect.options[0];
          if (first && first.value) {
            els.modelSelect.value = first.value;
            els.modelSelect.dataset.currentValue = first.value;
            saveModel(first.value);
            updateModelBadges();
            syncNumCtxInputForSelectedModel();
            applyThinkingModeSupportForModel(first.value, { preferredMode: getSavedThinkMode() });
            ensureSelectedModelMeta(first.value);
            refreshHelperControls();
          }
        }
      });
    }
    if (els.thinkModeSelect) {
      els.thinkModeSelect.addEventListener("change", () => {
        saveThinkMode(els.thinkModeSelect.value || "auto");
        applyThinkingModeSupportForModel(getSelectedModelKey(), { preferredMode: els.thinkModeSelect.value || "auto" });
      });
    }
    if (els.confirmThinkingProfileBtn) {
      els.confirmThinkingProfileBtn.addEventListener("click", confirmCurrentThinkingProfile);
    }
    if (els.ensembleModeSelect) {
      els.ensembleModeSelect.addEventListener("change", () => {
        const mode = normalizeEnsembleMode(els.ensembleModeSelect.value);
        els.ensembleModeSelect.value = mode;
        saveEnsembleMode(mode);
        refreshHelperControls();
      });
    }
    if (els.helperSearchInput) {
      els.helperSearchInput.addEventListener("input", () => {
        updateHelperModelSelect();
      });
    }
    if (els.helperModelSelect) {
      els.helperModelSelect.addEventListener("change", () => {
        const helper = String(els.helperModelSelect.value || "").trim();
        els.helperModelSelect.dataset.currentValue = helper;
        saveHelperModel(helper);
        if (!state.isStreaming && state.ensembleMode === "manual" && els.helperText) {
          els.helperText.textContent = helper ? `🤝 Manual helper: ${helper}` : "🤝 Manual helper: επίλεξε δεύτερο μοντέλο";
        }
      });
    }
    els.modelSelect.addEventListener("change", () => {
      const selectedModel = String(els.modelSelect.value || "").trim();
      els.modelSelect.dataset.currentValue = selectedModel;
      updateModelBadges();
      saveModel(selectedModel);
      syncNumCtxInputForSelectedModel();
      applyThinkingModeSupportForModel(selectedModel, { preferredMode: els.thinkModeSelect ? els.thinkModeSelect.value : getSavedThinkMode() });
      ensureSelectedModelMeta(selectedModel);
      refreshHelperControls();
      if (selectedModel && els.modelSearchInput) {
        els.modelSearchInput.blur();
      }
      if (els.modelSelect) {
        els.modelSelect.blur();
      }
    });

    // Char counter + focus reset
    els.userInput.addEventListener("input", updateCharCounter);
    els.userInput.addEventListener("focus", () => {
      if (!state.isStreaming) els.helperText.textContent = DEFAULT_HELPER;
    });

    // Keyboard shortcuts: Enter = send, Shift+Enter = newline, Ctrl+Enter = send
    els.userInput.addEventListener("keydown", (e) => {
      const send = (e.key === "Enter" && !e.shiftKey) || (e.key === "Enter" && e.ctrlKey);
      if (send) { e.preventDefault(); sendMessage(); }
    });

    // Scroll-to-bottom button visibility
    els.messages.addEventListener("scroll", updateScrollToBottomBtn);

    // ── Initialisation ────────────────────────────────────────────────────────

    applyTheme(getSavedTheme());
    if (els.modelSortSelect) {
      els.modelSortSelect.value = getSavedSortCriterion();
      state.modelSortCriterion = els.modelSortSelect.value || "overall";
    }
    state.ensembleMode = getSavedEnsembleMode();
    if (els.ensembleModeSelect) {
      els.ensembleModeSelect.value = state.ensembleMode;
    }
    rebuildPromptProfileMap();
    populatePromptProfileSelect(DEFAULT_PROMPT_PROFILE_ID);
    populateVisualizationEngineSelect(DEFAULT_VISUALIZATION_ENGINE);
    loadParams();
    loadThinkingProfileConfirmations();
    if (els.thinkModeSelect) {
      els.thinkModeSelect.value = getSavedThinkMode();
    }
    applyThinkingModeSupportForModel("", { preferredMode: getSavedThinkMode() });
    applyReasoningPanelVisibility();
    setupBrowserLifecycle();
    loadAppConfig();
    loadModels().finally(() => {
      refreshHelperControls(getSavedHelperModel(Array.isArray(state.models) ? state.models : []) || "");
    });
    loadSession();
    renderSystemNotice("🔄 Αυτόματη ανανέωση official direct API models σε εξέλιξη…");
    // Cloud API health — immediate check + polling
    pollBackendHealth();
    setInterval(pollBackendHealth, HEALTH_POLL_MS);

    // Αν η online απογραφή ολοκληρωθεί μετά το πρώτο paint, ενημέρωσε το dropdown αυτόματα και ενημέρωσε τον χρήστη.
    (function pollForOnlineModels() {
      let attempts = 0;
      let announced = false;
      const MAX_ATTEMPTS = 30; // 30 × 3s = 90s
      const timer = setInterval(async () => {
        attempts++;
        try {
          const resp = await fetch("/api/models");
          const data = await resp.json();
          if (data.source === "official-online") {
            clearInterval(timer);
            await loadModels();
            announced = true;
            renderSystemNotice(`✅ Αυτόματο Refresh Models: βρέθηκαν ${data.models?.length || 0} official direct API models.`);
          } else if (!data.refresh_in_progress && data.last_error) {
            clearInterval(timer);
            announced = true;
            await loadModels();
            renderSystemNotice(`❌ Αποτυχία αυτόματης ανανέωσης official direct API models: ${data.last_error}`);
          }
        } catch (_) {}
        if (attempts >= MAX_ATTEMPTS) {
          clearInterval(timer);
          if (!announced) {
            try {
              const resp = await fetch("/api/models");
              const data = await resp.json();
              await loadModels();
              if (data.refresh_in_progress) {
                renderSystemNotice("⚠ Η αυτόματη ανανέωση official direct API models συνεχίζεται περισσότερο από το αναμενόμενο. Εμφανίζονται προσωρινά τα πιο πρόσφατα διαθέσιμα δεδομένα.");
              } else if (Array.isArray(data.models) && data.models.length) {
                renderSystemNotice(`⚠ Η αυτόματη ανανέωση official direct API models δεν ολοκληρώθηκε online εγκαίρως. Φορτώθηκαν ${data.models.length} διαθέσιμα models από cache/τρέχουσα κατάσταση.`);
              } else {
                renderSystemNotice("⚠ Δεν ολοκληρώθηκε έγκαιρα η αυτόματη ανανέωση official direct API models.");
              }
            } catch (_) {
              renderSystemNotice("⚠ Δεν ολοκληρώθηκε έγκαιρα η αυτόματη ανανέωση official direct API models.");
            }
          }
        }
      }, 3000);
    })();
  </script>
</body>
</html>"""
    html_doc = html_doc.replace('__APP_TITLE__', html.escape(APP_TITLE)).replace('__SYSTEM_PROMPT__', html.escape(system_prompt)).replace('__DEFAULT_SYSTEM_PROMPT_JSON__', safe_prompt_json).replace('__DEFAULT_PROMPT_PROFILES_JSON__', safe_prompt_profiles_json).replace('__DEFAULT_PROMPT_PROFILE_ID_JSON__', json.dumps(DEFAULT_PROMPT_PROFILE_ID, ensure_ascii=False)).replace('__DEFAULT_VISUALIZATION_ENGINE_JSON__', json.dumps(DEFAULT_VISUALIZATION_ENGINE, ensure_ascii=False)).replace('__VISUALIZATION_ENGINE_OPTIONS_JSON__', safe_visualization_options_json).replace('__ACCEPTED_TYPES__', accepted_types)
    return html_doc

class AppHandler(BaseHTTPRequestHandler):
    """HTTP request handler της εφαρμογής.

Συγκεντρώνει τη λογική των GET και POST endpoints που εξυπηρετούν το frontend, τα uploads και το chat."""
    server_version = 'OllamaCloudChat/5.0'

    def log_message(self, format: str, *args) -> None:
        """Καταστέλλει τα ενδεικτικά log μηνύματα του BaseHTTPRequestHandler ώστε να μην γεμίζει η κονσόλα."""
        return

    def do_GET(self) -> None:
        """Δρομολογεί GET αιτήματα στον κατάλληλο handler βάσει path."""
        try:
            self._handle_GET()
        except Exception as exc:
            if is_client_disconnect_error(exc):
                return
            log.error('Unexpected GET error: %s', exc)
            try:
                json_response(self, {'error': 'Internal server error'}, status=500)
            except Exception:
                pass

    def _handle_GET(self) -> None:
        """Εξυπηρετεί GET endpoints: index, startup-events SSE, media, model-info και λοιπά static resources."""
        if self.path == '/' or self.path.startswith('/?'):
            body = serve_index_html().encode('utf-8')
            self.send_response(200)
            self.send_header('Content-Type', 'text/html; charset=utf-8')
            self.send_header('Content-Length', str(len(body)))
            self.send_header('Cache-Control', 'no-store')
            _send_security_headers(self)
            self.end_headers()
            self.wfile.write(body)
            return
        if self.path == '/startup':
            body = serve_startup_html().encode('utf-8')
            self.send_response(200)
            self.send_header('Content-Type', 'text/html; charset=utf-8')
            self.send_header('Content-Length', str(len(body)))
            self.send_header('Cache-Control', 'no-store')
            _send_security_headers(self)
            self.end_headers()
            self.wfile.write(body)
            return
        if self.path == '/startup-events':
            q = STARTUP.subscribe()
            self.send_response(200)
            self.send_header('Content-Type', 'text/event-stream; charset=utf-8')
            self.send_header('Cache-Control', 'no-cache')
            self.send_header('Connection', 'keep-alive')
            self.send_header('X-Accel-Buffering', 'no')
            _send_security_headers(self)
            self.end_headers()
            try:
                while True:
                    try:
                        event = q.get(timeout=20)
                        data = json.dumps(event, ensure_ascii=False)
                        self.wfile.write(f'data: {data}\n\n'.encode('utf-8'))
                        self.wfile.flush()
                        if event.get('level') == 'READY':
                            break
                    except _queue.Empty:
                        self.wfile.write(b': keepalive\n\n')
                        self.wfile.flush()
            except Exception:
                pass
            finally:
                STARTUP.unsubscribe(q)
            return
        if self.path.startswith('/generated-code/'):
            ensure_generated_code_dir()
            requested_name = urllib.parse.unquote(self.path.split('?', 1)[0][len('/generated-code/'):])
            safe_name = Path(requested_name).name
            file_path = (GENERATED_CODE_DIR / safe_name).resolve()
            generated_root = str(GENERATED_CODE_DIR.resolve()) + os.sep
            if not safe_name or not str(file_path).startswith(generated_root) or (not file_path.exists()) or (not file_path.is_file()):
                json_response(self, {'error': 'Το ζητούμενο generated .py αρχείο δεν βρέθηκε.'}, status=404)
                return
            raw = file_path.read_bytes()
            download_name = extract_original_generated_filename(safe_name)
            self.send_response(200)
            self.send_header('Content-Type', 'text/x-python; charset=utf-8')
            self.send_header('Content-Length', str(len(raw)))
            self.send_header('Content-Disposition', build_content_disposition_header('attachment', download_name, fallback='generated_code.py'))
            self.send_header('Cache-Control', 'no-store')
            _send_security_headers(self)
            self.end_headers()
            self.wfile.write(raw)
            return
        if self.path.startswith('/generated-media/'):
            ensure_generated_media_dir()
            requested_name = urllib.parse.unquote(self.path.split('?', 1)[0][len('/generated-media/'):])
            safe_name = Path(requested_name).name
            file_path = (GENERATED_MEDIA_DIR / safe_name).resolve()
            media_root = str(GENERATED_MEDIA_DIR.resolve()) + os.sep
            if not safe_name or not str(file_path).startswith(media_root) or (not file_path.exists()) or (not file_path.is_file()):
                json_response(self, {'error': 'Το ζητούμενο generated media αρχείο δεν βρέθηκε.'}, status=404)
                return
            raw = file_path.read_bytes()
            download_name = extract_original_generated_media_filename(safe_name)
            mime_type = 'image/svg+xml' if file_path.suffix.lower() == '.svg' else 'image/png'
            self.send_response(200)
            self.send_header('Content-Type', mime_type)
            self.send_header('Content-Length', str(len(raw)))
            self.send_header('Content-Disposition', build_content_disposition_header('inline', download_name, fallback='generated_plot.png'))
            self.send_header('Cache-Control', 'no-store')
            _send_security_headers(self)
            self.end_headers()
            self.wfile.write(raw)
            return
        if self.path == '/favicon.ico':
            self.send_response(204)
            self.end_headers()
            return
        if self.path == '/api/models':
            with REGISTRY.lock:
                no_models = not REGISTRY.models
                in_progress = REGISTRY.refresh_in_progress
            if no_models and in_progress:
                wait_for_model_refresh(timeout=60.0)
            elif no_models:
                refresh_models(force=True, wait_if_running=True)
            json_response(self, REGISTRY.as_dict())
            return
        if self.path.startswith('/api/model-details'):
            parsed = urllib.parse.urlsplit(self.path)
            query = urllib.parse.parse_qs(parsed.query or '')
            model = normalize_model_name(str((query.get('model') or [''])[0]).strip())
            force = str((query.get('force') or [''])[0]).strip().lower() in {'1', 'true', 'yes', 'on'}
            if not model:
                json_response(self, {'error': 'Δεν δόθηκε μοντέλο για metadata.'}, status=400)
                return
            meta = get_or_fetch_model_meta(model, force=force)
            if not meta:
                json_response(self, {'error': f'Δεν βρέθηκαν metadata για το μοντέλο {model}.'}, status=404)
                return
            json_response(self, {'model': model, 'meta': meta})
            return
        if self.path == '/api/session':
            json_response(self, {'history': get_history_payload()})
            return
        if self.path == '/api/app-config':
            payload = APP_CONFIG.as_public_dict()
            payload['api_key_source'] = get_ollama_api_key_source()
            payload['config_path'] = str(APP_CONFIG_FILE)
            payload['prompt_profiles'] = get_prompt_profiles_catalog()
            payload['visualization_engine_options'] = get_visualization_engine_options()
            json_response(self, payload)
            return
        if self.path == '/api/health':
            configured = is_direct_cloud_api_configured()
            json_response(self, {'status': 'ok' if configured else 'unavailable', 'mode': 'direct-cloud', 'cloud_api_configured': configured, 'api_key_source': get_ollama_api_key_source(), 'server_uptime_sec': round(time.time() - _SERVER_START_TIME, 1)}, status=200 if configured else 503)
            return
        json_response(self, {'error': 'Not found'}, status=404)

    def do_POST(self) -> None:
        """Δρομολογεί POST αιτήματα στον κατάλληλο handler βάσει path."""
        try:
            self._handle_POST()
        except Exception as exc:
            if is_client_disconnect_error(exc):
                return
            log.exception('Unexpected POST error: %s', exc)
            try:
                json_response(self, {'error': 'Internal server error'}, status=500)
            except Exception:
                pass

    def _handle_POST(self) -> None:
        """Εξυπηρετεί POST endpoints: chat streaming, session management, file upload, settings, PDF/DOCX export και plot rendering."""
        if self.path == '/api/browser-session':
            payload = safe_read_json(self)
            if payload.get('__error__') == 'request_too_large':
                json_response(self, {'error': 'Αίτημα πολύ μεγάλο.'}, status=413)
                return
            session_id = str(payload.get('session_id', '') or '').strip()[:128]
            event_name = str(payload.get('event', 'heartbeat') or 'heartbeat').strip().lower()
            if session_id:
                if event_name in {'open', 'heartbeat', 'visible', 'focus'}:
                    BROWSER_MONITOR.touch(session_id)
                elif event_name == 'close':
                    BROWSER_MONITOR.close(session_id)
            json_response(self, {'ok': True, 'active_sessions': BROWSER_MONITOR.active_count()})
            return
        if self.path == '/api/refresh-models':
            payload = safe_read_json(self)
            if payload.get('__error__') == 'request_too_large':
                json_response(self, {'error': 'Αίτημα πολύ μεγάλο.'}, status=413)
                return
            refresh_models(force=bool(payload.get('force', True)), wait_if_running=True)
            json_response(self, REGISTRY.as_dict())
            return
        if self.path == '/api/reset-chat':
            try:
                SESSION.reset()
                json_response(self, {'ok': True})
            except Exception as exc:
                json_response(self, {'error': f'Αποτυχία εκκαθάρισης session: {exc}'}, status=500)
            return
        if self.path == '/api/app-config':
            payload = safe_read_json(self)
            if payload.get('__error__') == 'request_too_large':
                json_response(self, {'error': 'Αίτημα πολύ μεγάλο.'}, status=413)
                return
            try:
                key = payload.get('ollama_api_key', None)
                active_prompt_profile = payload.get('active_prompt_profile', None)
                custom_system_prompt = payload.get('custom_system_prompt', None)
                active_visualization_engine = payload.get('active_visualization_engine', None)
                config = save_app_config_to_disk(key, active_prompt_profile=active_prompt_profile, custom_system_prompt=custom_system_prompt, active_visualization_engine=active_visualization_engine)
                json_response(self, {'ok': True, 'config': config.as_public_dict(), 'config_path': str(APP_CONFIG_FILE)})
            except Exception as exc:
                json_response(self, {'error': f'Αποτυχία αποθήκευσης settings file: {exc}'}, status=500)
            return
        if self.path == '/api/execute-python':
            payload = safe_read_json(self)
            if payload.get('__error__') == 'request_too_large':
                json_response(self, {'error': 'Αίτημα πολύ μεγάλο.'}, status=413)
                return
            code_text = str(payload.get('code', '') or '')
            suggested_filename = str(payload.get('filename', '') or '')
            ok, message = launch_python_code_in_terminal(code_text, suggested_filename=suggested_filename)
            json_response(self, {'ok': ok, 'message': message} if ok else {'error': message}, status=200 if ok else 400)
            return
        if self.path == '/api/export-python-block':
            payload = safe_read_json(self)
            if payload.get('__error__') == 'request_too_large':
                json_response(self, {'error': 'Αίτημα πολύ μεγάλο.'}, status=413)
                return
            code_text = str(payload.get('code', '') or '')
            suggested_filename = str(payload.get('filename', '') or '')
            try:
                file_info = save_generated_python_file(code_text, suggested_filename=suggested_filename)
                json_response(self, {'ok': True, 'file': file_info})
            except Exception as exc:
                json_response(self, {'error': f'Αποτυχία δημιουργίας .py αρχείου: {exc}'}, status=400)
            return
        if self.path == '/api/render-python-plot':
            payload = safe_read_json(self)
            if payload.get('__error__') == 'request_too_large':
                json_response(self, {'error': 'Αίτημα πολύ μεγάλο.'}, status=413)
                return
            code_text = str(payload.get('code', '') or '')
            suggested_filename = str(payload.get('filename', '') or '')
            ok, message, plot_file = render_python_plot_to_generated_media(code_text, suggested_filename=suggested_filename)
            if ok and plot_file:
                json_response(self, {'ok': True, 'message': message, 'file': plot_file})
            else:
                json_response(self, {'error': message or 'Αποτυχία render plot.'}, status=400)
            return

        if self.path == '/api/export-assistant-pdf':
            payload = safe_read_json(self)
            if payload.get('__error__') == 'request_too_large':
                json_response(self, {'error': 'Αίτημα πολύ μεγάλο.'}, status=413)
                return
            html_fragment = str(payload.get('html_fragment', '') or '')
            theme = str(payload.get('theme', 'light') or 'light').strip().lower()
            filename = sanitize_download_filename(str(payload.get('filename', 'assistant-response.pdf') or 'assistant-response.pdf'))
            mathjax_svg_cache = _sanitize_mathjax_svg_cache_fragment(payload.get('mathjax_svg_cache', ''))
            if not html_fragment.strip():
                json_response(self, {'error': 'Δεν δόθηκε printable HTML fragment για PDF export.'}, status=400)
                return
            browser_path = _find_headless_pdf_browser()
            if not browser_path:
                json_response(self, {'error': 'Δεν βρέθηκε εγκατεστημένος Microsoft Edge / Google Chrome / Chromium browser για server-side PDF export.'}, status=503)
                return
            try:
                document_title = Path(filename).stem or 'Assistant response'
                html_doc = _build_assistant_pdf_document(html_fragment, theme=theme, document_title=document_title, mathjax_svg_cache=mathjax_svg_cache)
                temp_dir = Path(tempfile.mkdtemp(prefix='ollama_chat_pdf_out_'))
                try:
                    pdf_path = temp_dir / filename
                    ok, detail = _render_pdf_with_headless_browser(browser_path, html_doc, pdf_path)
                    if (not ok) or (not _pdf_file_looks_valid(pdf_path)):
                        raise RuntimeError(detail or 'Αποτυχία headless browser export.')
                    _polish_exported_pdf(pdf_path, document_title=document_title)
                    if not _pdf_file_looks_valid(pdf_path):
                        raise RuntimeError('Το PDF δημιουργήθηκε αλλά απέτυχε το τελικό validation μετά το polish.')
                    raw = pdf_path.read_bytes()
                finally:
                    _safe_rmtree(temp_dir)
                self.send_response(200)
                self.send_header('Content-Type', 'application/pdf')
                self.send_header('Content-Length', str(len(raw)))
                self.send_header('Content-Disposition', build_content_disposition_header('attachment', filename, fallback='assistant-response.pdf'))
                self.send_header('Cache-Control', 'no-store')
                _send_security_headers(self)
                self.end_headers()
                self.wfile.write(raw)
            except Exception as exc:
                json_response(self, {'error': f'Αποτυχία δημιουργίας PDF: {exc}'}, status=500)
            return
        if self.path == '/api/export-assistant-docx':
            payload = safe_read_json(self)
            if payload.get('__error__') == 'request_too_large':
                json_response(self, {'error': 'Αίτημα πολύ μεγάλο.'}, status=413)
                return
            html_fragment = str(payload.get('html_fragment', '') or '')
            filename = sanitize_download_filename(str(payload.get('filename', 'assistant-response.docx') or 'assistant-response.docx'))
            if not html_fragment.strip():
                json_response(self, {'error': 'Δεν δόθηκε printable HTML fragment για Docx export.'}, status=400)
                return
            try:
                document_title = Path(filename).stem or 'Assistant response'
                raw = _build_assistant_docx_bytes(html_fragment, document_title=document_title)
                if not raw:
                    raise RuntimeError('Δεν δημιουργήθηκαν bytes DOCX.')
                self.send_response(200)
                self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                self.send_header('Content-Length', str(len(raw)))
                self.send_header('Content-Disposition', build_content_disposition_header('attachment', filename, fallback='assistant-response.docx'))
                self.send_header('Cache-Control', 'no-store')
                _send_security_headers(self)
                self.end_headers()
                self.wfile.write(raw)
            except Exception as exc:
                json_response(self, {'error': f'Αποτυχία δημιουργίας Docx: {exc}'}, status=500)
            return

        if self.path == '/api/chat':
            payload = safe_read_json(self)
            if payload.get('__error__') == 'request_too_large':
                json_response(self, {'error': 'Το αίτημα είναι πολύ μεγάλο. Μείωσε αριθμό ή μέγεθος αρχείων.'}, status=413)
                return
            model = normalize_model_name(str(payload.get('model', '')).strip())
            gui_system_prompt = str(payload.get('system_prompt', ''))
            prompt_profile_id = str(payload.get('prompt_profile_id', '') or '')
            visualization_engine = str(payload.get('visualization_engine', '') or '')
            think_mode = payload.get('think_mode', 'on')
            ensemble_mode_raw = str(payload.get('ensemble_mode', '') or '').strip().lower()
            if ensemble_mode_raw not in {'off', 'auto', 'manual'}:
                ensemble_mode_raw = 'auto' if bool(payload.get('ensemble_auto', True)) else 'off'
            ensemble_helper_model = normalize_model_name(str(payload.get('ensemble_helper_model', '') or '').strip())
            system_prompt, system_prompt_source, prompt_profile_id, visualization_engine = get_effective_system_prompt(gui_system_prompt, prompt_profile_id=prompt_profile_id, visualization_engine=visualization_engine)
            user_text = str(payload.get('user_text', '')).strip()
            attachments = payload.get('attachments', [])
            raw_opts = payload.get('options', {}) or {}
            model_options: Dict = {}
            try:
                temperature = float(raw_opts.get('temperature', -1))
                if 0.0 <= temperature <= 2.0:
                    model_options['temperature'] = temperature
            except (TypeError, ValueError):
                pass
            try:
                top_p = float(raw_opts.get('top_p', -1))
                if 0.0 < top_p <= 1.0:
                    model_options['top_p'] = top_p
            except (TypeError, ValueError):
                pass
            try:
                seed = int(raw_opts.get('seed', -1))
                if seed >= 0:
                    model_options['seed'] = seed
            except (TypeError, ValueError):
                pass
            try:
                num_ctx = int(raw_opts.get('num_ctx', 0))
                if 256 <= num_ctx <= 1048576:
                    model_options['num_ctx'] = num_ctx
            except (TypeError, ValueError):
                pass
            if not model:
                json_response(self, {'error': 'Δεν δόθηκε μοντέλο.'}, status=400)
                return
            if not user_text:
                json_response(self, {'error': 'Το user prompt πρέπει να δοθεί από εσένα.'}, status=400)
                return
            if attachments and (not isinstance(attachments, list)):
                json_response(self, {'error': 'Μη έγκυρα attachments.'}, status=400)
                return
            try:
                processed_attachments, warnings = prepare_attachments(attachments, model)
            except (ValueError, OSError, AttributeError, TypeError) as exc:
                json_response(self, {'error': str(exc)}, status=400)
                return
            self.send_response(200)
            self.send_header('Content-Type', 'application/x-ndjson; charset=utf-8')
            self.send_header('Cache-Control', 'no-cache')
            self.send_header('Connection', 'close')
            _send_security_headers(self)
            self.end_headers()
            start_time = time.time()
            try:
                visualization_hint = build_visualization_engine_user_hint(visualization_engine)
                effective_user_text = user_text.strip()
                if visualization_hint:
                    effective_user_text = (visualization_hint + '\\n\\n' + effective_user_text).strip()
                prepared_user_content = build_user_message_content(effective_user_text, processed_attachments)
                user_message: Dict = {'role': 'user', 'content': prepared_user_content}
                image_b64_list: List[str] = []
                for item in processed_attachments:
                    if item['kind'] == 'image' and item['will_send_as_image']:
                        try:
                            raw_bytes = Path(item['path']).read_bytes()
                            image_b64_list.append(base64.b64encode(raw_bytes).decode('ascii'))
                        except Exception as img_exc:
                            log.warning('Αδυναμία ανάγνωσης εικόνας %s: %s', item.get('name'), img_exc)
                if image_b64_list:
                    user_message['images'] = image_b64_list
                history_user_item: Dict = {'role': 'user', 'content': user_text, 'attachments': [{'name': item['name'], 'kind': item['kind']} for item in processed_attachments]}
                with SESSION.lock:
                    session_messages = list(SESSION.messages)
                    messages = build_messages(system_prompt, session_messages + [user_message])
                if not is_direct_cloud_api_configured():
                    raise RuntimeError("Λείπει το Ollama Cloud API key. Η εφαρμογή τρέχει μόνο σε direct Ollama Cloud API mode. Βάλ'το στο πεδίο API Key του GUI, στο settings file ή στο OLLAMA_API_KEY.")
                ensemble_info: Optional[Dict[str, object]] = None
                final_messages = list(messages)
                if ensemble_mode_raw == 'auto':
                    try:
                        ensemble_info = choose_auto_ensemble_helper(model, user_text, processed_attachments)
                    except Exception as ensemble_pick_exc:
                        log.warning('Αποτυχία επιλογής helper model για ensemble: %s', ensemble_pick_exc)
                        ensemble_info = None
                elif ensemble_mode_raw == 'manual':
                    try:
                        ensemble_info = choose_manual_ensemble_helper(model, ensemble_helper_model, user_text, processed_attachments)
                    except Exception as ensemble_pick_exc:
                        log.warning('Αποτυχία manual helper model για ensemble: %s', ensemble_pick_exc)
                        ensemble_info = None
                stream_json_line(self, {'type': 'meta', 'warnings': warnings, 'attachments': history_user_item['attachments'], 'system_prompt_source': system_prompt_source, 'ensemble': {'enabled': bool(ensemble_info), 'mode': ensemble_mode_raw, 'primary_model': model, 'helper_model': str((ensemble_info or {}).get('helper_model') or ''), 'criterion': str((ensemble_info or {}).get('criterion') or ''), 'role': str((ensemble_info or {}).get('role') or ''), 'role_label': str((ensemble_info or {}).get('role_label') or ''), 'selection_reason': str((ensemble_info or {}).get('selection_reason') or '')}})
                if ensemble_info:
                    helper_model = str(ensemble_info.get('helper_model') or '').strip()
                    helper_role = str(ensemble_info.get('role') or 'cross-checker').strip()
                    helper_meta = {}
                    with REGISTRY.lock:
                        helper_meta = copy.deepcopy(REGISTRY.model_meta.get(helper_model, {}))
                    helper_system = build_helper_system_prompt(model, helper_model, helper_role, dict(ensemble_info.get('traits') or {}))
                    helper_messages = insert_secondary_system_message(messages, helper_system)
                    helper_options = dict(model_options) if model_options else {}
                    if float(helper_options.get('temperature', 0.2) or 0.2) > 0.35:
                        helper_options['temperature'] = 0.2
                    helper_max_ctx = get_model_context_tokens(helper_meta)
                    try:
                        requested_ctx = int(helper_options.get('num_ctx') or 0)
                    except Exception:
                        requested_ctx = 0
                    if requested_ctx > 0 and helper_max_ctx > 0:
                        helper_options['num_ctx'] = min(requested_ctx, helper_max_ctx)
                    helper_num_predict = int(_ENSEMBLE_HELPER_MAX_TOKENS_BY_ROLE.get(helper_role, 160) or 160)
                    try:
                        current_num_predict = int(helper_options.get('num_predict') or 0)
                    except Exception:
                        current_num_predict = 0
                    if current_num_predict > 0:
                        helper_options['num_predict'] = min(current_num_predict, helper_num_predict)
                    else:
                        helper_options['num_predict'] = helper_num_predict
                    helper_timeout = int(_ENSEMBLE_HELPER_TIMEOUT_BY_ROLE.get(helper_role, 90) or 90)
                    stream_json_line(self, {'type': 'meta', 'status': f'🤝 Συμβουλεύομαι helper model {helper_model}…'})
                    helper_think_value = None
                    try:
                        helper_result = direct_cloud_chat_complete(model=helper_model, messages=helper_messages, model_options=helper_options if helper_options else None, think_value=helper_think_value, timeout=helper_timeout)
                        helper_text = str(helper_result.get('content') or '').strip()
                        if not helper_text:
                            helper_text = str(helper_result.get('thinking') or '').strip()
                        if helper_text:
                            final_messages = insert_secondary_system_message(messages, build_main_ensemble_guidance(helper_model, helper_role, helper_text))
                            stream_json_line(self, {'type': 'meta', 'status': f'✅ Έτοιμο το helper guidance από {helper_model}. Ξεκινά το κύριο μοντέλο…'})
                        else:
                            stream_json_line(self, {'type': 'meta', 'warnings': [f'Το βοηθητικό ensemble model {helper_model} δεν επέστρεψε guidance και η απάντηση συνεχίζεται μόνο με το κύριο μοντέλο.'], 'status': f'⚠️ Το helper model {helper_model} δεν έδωσε guidance. Συνεχίζω με το κύριο μοντέλο…'})
                    except Exception as helper_exc:
                        log.warning('Αποτυχία helper ensemble model %s: %s', helper_model, helper_exc)
                        stream_json_line(self, {'type': 'meta', 'warnings': [f'Το βοηθητικό ensemble model {helper_model} απέτυχε ({build_friendly_chat_error(helper_exc)}). Η απάντηση συνεχίζεται μόνο με το κύριο μοντέλο.'], 'status': f'⚠️ Παράλειψη helper model {helper_model}. Συνεχίζω με το κύριο μοντέλο…'})
                final_messages = apply_qwen3_vl_nothink_workaround(final_messages, model, think_mode)
                think_value = resolve_think_mode(model, think_mode)
                stream_json_line(self, {'type': 'meta', 'status': f'🧠 Το κύριο μοντέλο {model} ξεκινά να απαντά…'})
                response, _effective_think_value, compat_warnings, suppress_reasoning_output = open_direct_cloud_chat_stream_with_fallback(model=model, messages=final_messages, model_options=model_options if model_options else None, think_value=think_value, requested_mode=think_mode)
                if compat_warnings:
                    warnings = list(warnings) + list(compat_warnings)
                    stream_json_line(self, {'type': 'meta', 'warnings': compat_warnings})
                collected: List[str] = []
                collected_thinking: List[str] = []
                token_stats: Optional[Dict] = None
                thinking_started = False
                thinking_done_sent = False
                for chunk in response:
                    thinking_piece = extract_chunk_thinking(chunk)
                    if thinking_piece:
                        thinking_started = True
                        if not suppress_reasoning_output:
                            collected_thinking.append(thinking_piece)
                            stream_json_line(self, {'type': 'thinking', 'content': thinking_piece})
                    piece = extract_chunk_content(chunk)
                    if piece:
                        if thinking_started and (not thinking_done_sent):
                            if not suppress_reasoning_output:
                                stream_json_line(self, {'type': 'thinking_done'})
                            thinking_done_sent = True
                        collected.append(piece)
                        stream_json_line(self, {'type': 'delta', 'content': piece})
                    stats = extract_token_stats(chunk)
                    if stats:
                        token_stats = stats
                if collected_thinking and (not thinking_done_sent) and (not suppress_reasoning_output):
                    stream_json_line(self, {'type': 'thinking_done'})
                assistant_text = ''.join(collected).strip()
                assistant_text = strip_inline_think_blocks(assistant_text) if suppress_reasoning_output else assistant_text
                assistant_thinking = '' if suppress_reasoning_output else ''.join(collected_thinking).strip()
                if not assistant_text and (not assistant_thinking):
                    raise RuntimeError('Το μοντέλο δεν επέστρεψε περιεχόμενο. Έλεγξε το API key από το GUI/settings file και αν το μοντέλο είναι διαθέσιμο στο direct cloud catalog.')
                assistant_display_text = compose_display_assistant_text(assistant_text, assistant_thinking)
                elapsed = time.time() - start_time
                with SESSION.lock:
                    SESSION.messages.append(user_message)
                    SESSION.history.append(history_user_item)
                    SESSION.messages.append({'role': 'assistant', 'content': assistant_text, 'thinking': assistant_thinking})
                    SESSION.history.append({'role': 'assistant', 'content': assistant_display_text, 'attachments': []})
                    if len(SESSION.messages) > MAX_HISTORY_MESSAGES:
                        SESSION.messages[:] = SESSION.messages[-MAX_HISTORY_MESSAGES:]
                    if len(SESSION.history) > MAX_HISTORY_MESSAGES:
                        SESSION.history[:] = SESSION.history[-MAX_HISTORY_MESSAGES:]
                stream_json_line(self, {'type': 'done', 'elapsed_sec': elapsed, 'token_stats': token_stats})
            except Exception as exc:
                if is_client_disconnect_error(exc):
                    return
                friendly = build_friendly_chat_error(exc)
                log.error('Chat error: %s', exc)
                try:
                    stream_json_line(self, {'type': 'error', 'error': friendly})
                except Exception:
                    pass
            return
        json_response(self, {'error': 'Not found'}, status=404)

def parse_args() -> argparse.Namespace:
    """Αναλύει command-line arguments και επιστρέφει Namespace με port, API key, browser και λοιπές επιλογές."""
    parser = argparse.ArgumentParser(prog='ollama_cloud_chat', description=f'{APP_TITLE} — Web chat για Ollama cloud μοντέλα', formatter_class=argparse.ArgumentDefaultsHelpFormatter)
    parser.add_argument('--port', type=int, default=DEFAULT_PORT, help='Port του web server')
    parser.add_argument('--host', type=str, default=HOST, help='Host του web server')
    parser.add_argument('--no-browser', action='store_true', help='Μην ανοίξεις αυτόματα τον browser')
    parser.add_argument('--log-level', default='INFO', choices=['DEBUG', 'INFO', 'WARNING', 'ERROR'], help='Επίπεδο logging')
    parser.add_argument('--system-prompt-file', type=str, default='', metavar='FILE', help='Φόρτωση system prompt από εξωτερικό αρχείο .txt')
    return parser.parse_args()

def load_system_prompt_from_file(filepath: str) -> Optional[str]:
    """Φορτώνει system prompt από αρχείο κειμένου αν υπάρχει, αλλιώς επιστρέφει κενό string."""
    if not filepath:
        return None
    path = Path(filepath)
    if not path.exists():
        log.error('Το αρχείο system prompt δεν βρέθηκε: %s', filepath)
        return None
    try:
        content = path.read_text(encoding='utf-8').strip()
        if content:
            log.info('📄 System prompt φορτώθηκε από: %s (%d χαρ.)', filepath, len(content))
            return content
        log.warning('Το αρχείο system prompt είναι κενό: %s', filepath)
        return None
    except Exception as exc:
        log.error('Αποτυχία ανάγνωσης system prompt: %s', exc)
        return None

def open_browser_later(url: str, delay: float=0.9) -> None:
    """Ανοίγει τον browser σε daemon thread μετά μικρή καθυστέρηση ώστε ο server να είναι έτοιμος πρώτα."""

    def _worker() -> None:
        """Εκτελεί ασύγχρονη εργασία σε daemon thread."""
        time.sleep(delay)
        webbrowser.open(url, new=2)
    threading.Thread(target=_worker, daemon=True).start()

def _run_initialization(args: argparse.Namespace, port: int) -> None:
    """Εκτελεί πλήρη αρχικοποίηση εφαρμογής: δημιουργία φακέλων, φόρτωση cache, refresh μοντέλων, εκκίνηση server και watchdog."""
    global DEFAULT_SYSTEM_PROMPT
    url = f'http://{HOST}:{port}'
    if args.system_prompt_file:
        custom_prompt = load_system_prompt_from_file(args.system_prompt_file)
        if custom_prompt:
            DEFAULT_SYSTEM_PROMPT = custom_prompt
            slog('INFO', '📄 System prompt: %s (%d χαρ.)', args.system_prompt_file, len(custom_prompt))
        else:
            slog('WARNING', '⚠  Αποτυχία φόρτωσης system prompt — χρήση embedded.')
    else:
        slog('INFO', '📄 System prompt: embedded-in-code')
    slog('INFO', '📎 Αρχεία: drag & drop + file picker')
    slog('INFO', '☁️  Λειτουργία: direct Ollama Cloud API mode')
    if is_direct_cloud_api_configured():
        slog('INFO', '🔐 Ollama Cloud API key: βρέθηκε (στο .py ή στο περιβάλλον)')
    else:
        slog('WARNING', '⚠  Ollama Cloud API key: δεν βρέθηκε — βάλε το από το GUI ή στο %s ή όρισε OLLAMA_API_KEY', APP_CONFIG_FILE)
    ensure_upload_dir()
    cached_models, cached_meta, cached_ts = load_model_registry_cache_from_disk()
    with REGISTRY.lock:
        REGISTRY.models = list(cached_models)
        REGISTRY.model_meta = copy.deepcopy(cached_meta)
        REGISTRY.source = 'startup-disk-cache' if cached_models else 'initializing'
        REGISTRY.last_refresh_ts = float(cached_ts or 0.0)
        REGISTRY.last_error = ''
        REGISTRY.recommended_model = recommend_best_model(cached_models, cached_meta, 'overall') if cached_models else ''
        REGISTRY.refresh_in_progress = False
    slog('INFO', '🔄 Ανάκτηση όλων των διαθέσιμων official Ollama direct-cloud models…')
    slog('INFO', '✅ Server: %s', url)
    slog('INFO', '🛑 Ctrl+C για τερματισμό.')
    STARTUP.set_ready(url)
    refresh_models_in_background(force=True)


def _patch_svg_preview_in_index_html(html_doc: str) -> str:
    """Μετατρέπει fenced SVG blocks σε κανονικό preview εικόνας μέσα στο chat."""
    html_doc = str(html_doc or '')

    css_insert = '''    .svg-block {
      border: 1px solid rgba(148,163,184,0.18);
      border-radius: 16px;
      overflow: hidden;
      background: rgba(2,6,23,0.92);
      box-shadow: inset 0 1px 0 rgba(148,163,184,0.06);
      margin: 8px 0;
    }
    .svg-toolbar {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 10px;
      padding: 9px 12px;
      background: rgba(15,23,42,0.98);
      border-bottom: 1px solid rgba(148,163,184,0.14);
    }
    .svg-label {
      color: #cbd5e1;
      font-size: 0.78rem;
      font-weight: 700;
      letter-spacing: 0.3px;
      text-transform: uppercase;
    }
    .svg-actions {
      display: inline-flex;
      align-items: center;
      gap: 8px;
      flex-wrap: wrap;
      justify-content: flex-end;
    }
    .svg-preview-wrap {
      padding: 18px;
      background:
        linear-gradient(180deg, rgba(15,23,42,0.22), rgba(2,6,23,0.08)),
        repeating-linear-gradient(
          45deg,
          rgba(148,163,184,0.04) 0,
          rgba(148,163,184,0.04) 12px,
          rgba(255,255,255,0.02) 12px,
          rgba(255,255,255,0.02) 24px
        );
      overflow: auto;
      text-align: center;
    }
    .svg-preview-image {
      display: inline-block;
      width: auto;
      max-width: 100%;
      height: auto;
      max-height: none;
      border-radius: 12px;
      background: #ffffff;
      box-shadow: 0 10px 28px rgba(0,0,0,0.22);
    }
    html[data-theme="light"] .svg-preview-wrap {
      background:
        linear-gradient(180deg, rgba(255,255,255,0.88), rgba(241,245,249,0.95)),
        repeating-linear-gradient(
          45deg,
          rgba(15,23,42,0.03) 0,
          rgba(15,23,42,0.03) 12px,
          rgba(255,255,255,0.02) 12px,
          rgba(255,255,255,0.02) 24px
        );
    }

'''

    helper_insert = r'''    function looksLikeSvgContent(text) {
      const source = String(text || "").trim();
      if (!source) return false;
      return /^<svg\b[\s\S]*<\/svg>$/i.test(source);
    }

    function isSvgLanguage(language) {
      const normalized = String(language || "").trim().toLowerCase();
      return normalized === "svg" || normalized === "image/svg+xml";
    }

  
  function mapSvgMathScriptChars(value, kind) {
    const source = String(value || '');
    if (!source) return '';
    const subMap = {
      '0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅','6':'₆','7':'₇','8':'₈','9':'₉',
      '+':'₊','-':'₋','=':'₌','(':'₍',')':'₎',
      'a':'ₐ','e':'ₑ','h':'ₕ','i':'ᵢ','j':'ⱼ','k':'ₖ','l':'ₗ','m':'ₘ','n':'ₙ','o':'ₒ','p':'ₚ','r':'ᵣ','s':'ₛ','t':'ₜ','u':'ᵤ','v':'ᵥ','x':'ₓ',
      'β':'ᵦ','γ':'ᵧ','ρ':'ᵨ','φ':'ᵩ','χ':'ᵪ'
    };
    const supMap = {
      '0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹',
      '+':'⁺','-':'⁻','=':'⁼','(':'⁽',')':'⁾',
      'a':'ᵃ','b':'ᵇ','c':'ᶜ','d':'ᵈ','e':'ᵉ','f':'ᶠ','g':'ᵍ','h':'ʰ','i':'ⁱ','j':'ʲ','k':'ᵏ','l':'ˡ','m':'ᵐ','n':'ⁿ','o':'ᵒ','p':'ᵖ','r':'ʳ','s':'ˢ','t':'ᵗ','u':'ᵘ','v':'ᵛ','w':'ʷ','x':'ˣ','y':'ʸ','z':'ᶻ',
      'A':'ᴬ','B':'ᴮ','D':'ᴰ','E':'ᴱ','G':'ᴳ','H':'ᴴ','I':'ᴵ','J':'ᴶ','K':'ᴷ','L':'ᴸ','M':'ᴹ','N':'ᴺ','O':'ᴼ','P':'ᴾ','R':'ᴿ','T':'ᵀ','U':'ᵁ','V':'ⱽ','W':'ᵂ',
      'β':'ᵝ','γ':'ᵞ','δ':'ᵟ','θ':'ᶿ','ι':'ᶥ','Φ':'ᶲ','φ':'ᵠ','χ':'ᵡ'
    };
    const targetMap = kind === 'sub' ? subMap : supMap;
    let out = '';
    for (const ch of source) {
      if (Object.prototype.hasOwnProperty.call(targetMap, ch)) {
        out += targetMap[ch];
      } else {
        return '';
      }
    }
    return out;
  }

  function convertLatexLikeSvgText(value) {
    let text = String(value || '');
    if (!text) return '';

    const normalizeMathExpr = (expr) => {
      let s = String(expr || '');
      if (!s) return '';
      const replacements = [
        [/\\,/g, ' '],
        [/\\;/g, ' '],
        [/\\:/g, ' '],
        [/\\!/g, ''],
        [/~/g, ' '],
        [/\\cdot\b/g, '·'],
        [/\\times\b/g, '×'],
        [/\\pm\b/g, '±'],
        [/\\mp\b/g, '∓'],
        [/\\leq\b|\\le\b/g, '≤'],
        [/\\geq\b|\\ge\b/g, '≥'],
        [/\\neq\b/g, '≠'],
        [/\\approx\b/g, '≈'],
        [/\\propto\b/g, '∝'],
        [/\\infty\b/g, '∞'],
        [/\\rightarrow\b|\\to\b/g, '→'],
        [/\\leftarrow\b/g, '←'],
        [/\\leftrightarrow\b/g, '↔'],
        [/\\degree\b/g, '°'],
        [/\\circ\b/g, '°'],
        [/\\alpha\b/g, 'α'],
        [/\\beta\b/g, 'β'],
        [/\\gamma\b/g, 'γ'],
        [/\\delta\b/g, 'δ'],
        [/\\epsilon\b/g, 'ε'],
        [/\\varepsilon\b/g, 'ε'],
        [/\\zeta\b/g, 'ζ'],
        [/\\eta\b/g, 'η'],
        [/\\theta\b/g, 'θ'],
        [/\\vartheta\b/g, 'ϑ'],
        [/\\iota\b/g, 'ι'],
        [/\\kappa\b/g, 'κ'],
        [/\\lambda\b/g, 'λ'],
        [/\\mu\b/g, 'μ'],
        [/\\nu\b/g, 'ν'],
        [/\\xi\b/g, 'ξ'],
        [/\\pi\b/g, 'π'],
        [/\\varpi\b/g, 'ϖ'],
        [/\\rho\b/g, 'ρ'],
        [/\\varrho\b/g, 'ϱ'],
        [/\\sigma\b/g, 'σ'],
        [/\\varsigma\b/g, 'ς'],
        [/\\tau\b/g, 'τ'],
        [/\\upsilon\b/g, 'υ'],
        [/\\phi\b/g, 'φ'],
        [/\\varphi\b/g, 'ϕ'],
        [/\\chi\b/g, 'χ'],
        [/\\psi\b/g, 'ψ'],
        [/\\omega\b/g, 'ω'],
        [/\\Gamma\b/g, 'Γ'],
        [/\\Delta\b/g, 'Δ'],
        [/\\Theta\b/g, 'Θ'],
        [/\\Lambda\b/g, 'Λ'],
        [/\\Xi\b/g, 'Ξ'],
        [/\\Pi\b/g, 'Π'],
        [/\\Sigma\b/g, 'Σ'],
        [/\\Upsilon\b/g, 'Υ'],
        [/\\Phi\b/g, 'Φ'],
        [/\\Psi\b/g, 'Ψ'],
        [/\\Omega\b/g, 'Ω']
      ];
      replacements.forEach(([pattern, replacement]) => {
        s = s.replace(pattern, replacement);
      });

      for (let i = 0; i < 6; i += 1) {
        const next = s
          .replace(/\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}/g, '($1)/($2)')
          .replace(/\\sqrt\s*\{([^{}]+)\}/g, '√($1)');
        if (next === s) break;
        s = next;
      }

      s = s.replace(/([A-Za-zΑ-Ωα-ω0-9θΘΔΓΛΣΠΩμνρτφχψω])_\{([^{}]+)\}/g, (m, base, sub) => {
        const mapped = mapSvgMathScriptChars(sub, 'sub');
        return base + (mapped || ('_' + sub));
      });
      s = s.replace(/([A-Za-zΑ-Ωα-ω0-9θΘΔΓΛΣΠΩμνρτφχψω])_([A-Za-z0-9+\-=()βγρφχ]+)/g, (m, base, sub) => {
        const mapped = mapSvgMathScriptChars(sub, 'sub');
        return base + (mapped || ('_' + sub));
      });
      s = s.replace(/([A-Za-zΑ-Ωα-ω0-9θΘΔΓΛΣΠΩμνρτφχψω])\^\{([^{}]+)\}/g, (m, base, sup) => {
        const mapped = mapSvgMathScriptChars(sup, 'sup');
        return base + (mapped || ('^' + sup));
      });
      s = s.replace(/([A-Za-zΑ-Ωα-ω0-9θΘΔΓΛΣΠΩμνρτφχψω])\^([A-Za-z0-9+\-=()βγδθιΦφχ]+)/g, (m, base, sup) => {
        const mapped = mapSvgMathScriptChars(sup, 'sup');
        return base + (mapped || ('^' + sup));
      });

      s = s.replace(/[{}]/g, '');
      s = s.replace(/\\_/g, '_');
      s = s.replace(/\\\\/g, ' ');
      s = s.replace(/\s+/g, ' ').trim();
      return s;
    };

    text = text.replace(/\$\$?([\s\S]*?)\$\$?/g, (match, expr) => normalizeMathExpr(expr));
    text = normalizeMathExpr(text);
    return text;
  }

  function normalizeSvgMathTextNodes(root) {
    if (!root || !root.querySelectorAll) return;
    root.querySelectorAll('text, tspan, title, desc').forEach((node) => {
      const raw = node.textContent;
      if (!raw || raw.indexOf('$') === -1 && raw.indexOf('\\') === -1 && raw.indexOf('_') === -1 && raw.indexOf('^') === -1) return;
      const normalized = convertLatexLikeSvgText(raw);
      if (normalized && normalized !== raw) node.textContent = normalized;
    });
  }

  function sanitizeSvgMarkup(rawSvg) {
      const source = String(rawSvg || "").trim();
      if (!source) return "";

      try {
        const parser = new DOMParser();
        const doc = parser.parseFromString(source, "image/svg+xml");
        const root = doc.documentElement;
        if (!root || String(root.nodeName || "").toLowerCase() !== "svg") return "";
        if (doc.querySelector("parsererror")) return "";

        root.querySelectorAll("script, foreignObject, iframe, object, embed").forEach(node => node.remove());

        const allNodes = root.querySelectorAll("*");
        for (const node of allNodes) {
          for (const attr of Array.from(node.attributes || [])) {
            const attrName = String(attr.name || "");
            const attrValue = String(attr.value || "").trim();
            if (/^on/i.test(attrName)) {
              node.removeAttribute(attrName);
              continue;
            }
            if ((attrName === "href" || attrName === "xlink:href") && /^javascript:/i.test(attrValue)) {
              node.removeAttribute(attrName);
            }
          }
        }

        root.removeAttribute("onload");
        root.removeAttribute("onclick");
        root.setAttribute("preserveAspectRatio", root.getAttribute("preserveAspectRatio") || "xMidYMid meet");

        return new XMLSerializer().serializeToString(root);
      } catch (err) {
        console.warn("SVG sanitize failed:", err);
        return "";
      }
    }

    function svgMarkupToDataUrl(svgMarkup) {
      const source = String(svgMarkup || "").trim();
      if (!source) return "";
      return `data:image/svg+xml;charset=utf-8,${encodeURIComponent(source)}`;
    }

    function createSvgPreviewBlock(svgMarkup) {
      const sanitized = sanitizeSvgMarkup(svgMarkup);
      if (!sanitized) return null;

      const wrapper = document.createElement("div");
      wrapper.className = "svg-block";

      const toolbar = document.createElement("div");
      toolbar.className = "svg-toolbar";

      const label = document.createElement("div");
      label.className = "svg-label";
      label.textContent = "SVG Preview";

      const actions = document.createElement("div");
      actions.className = "svg-actions";

      const copyBtn = document.createElement("button");
      copyBtn.type = "button";
      copyBtn.className = "code-copy-btn";
      copyBtn.textContent = "📋 Copy SVG";
      copyBtn.title = "Αντιγραφή SVG source στο clipboard";
      copyBtn.addEventListener("click", async () => {
        try {
          await navigator.clipboard.writeText(sanitized);
          setButtonFeedback(copyBtn, "✅ Copied", "📋 Copy SVG");
        } catch {
          setButtonFeedback(copyBtn, "❌ Error", "📋 Copy SVG", "error");
        }
      });

      actions.appendChild(copyBtn);
      toolbar.appendChild(label);
      toolbar.appendChild(actions);

      const previewWrap = document.createElement("div");
      previewWrap.className = "svg-preview-wrap";

      const image = document.createElement("img");
      image.className = "svg-preview-image";
      image.loading = "lazy";
      image.alt = "SVG block diagram";
      image.src = svgMarkupToDataUrl(sanitized);

      previewWrap.appendChild(image);
      wrapper.appendChild(toolbar);
      wrapper.appendChild(previewWrap);
      return wrapper;
    }

'''

    old_render_loop = '''for (const part of parts) {
        if (part.type === "think") {
          frag.appendChild(createThinkingBlock(part.content || "", part.complete !== false));
        } else if (part.type === "code") {
          const suggestedFilename = isPythonLanguage(part.language)
            ? (suggestedPyFilenames[pythonBlockIndex] || "")
            : "";
          frag.appendChild(createCodeBlock(part.language, part.content || "", suggestedFilename));
          if (isPythonLanguage(part.language)) pythonBlockIndex += 1;
        } else {
          frag.appendChild(createTextSegment(part.content || ""));
        }
      }'''

    new_render_loop = '''for (const part of parts) {
        if (part.type === "think") {
          frag.appendChild(createThinkingBlock(part.content || "", part.complete !== false));
        } else if (part.type === "code") {
          const rawCode = part.content || "";
          const normalizedLanguage = String(part.language || "").trim().toLowerCase();
          const isSvgBlock = isSvgLanguage(normalizedLanguage) || looksLikeSvgContent(rawCode.trim());
          if (isSvgBlock) {
            const svgNode = createSvgPreviewBlock(rawCode);
            if (svgNode) {
              frag.appendChild(svgNode);
            } else {
              frag.appendChild(createCodeBlock(part.language, rawCode, ""));
            }
          } else {
            const suggestedFilename = isPythonLanguage(part.language)
              ? (suggestedPyFilenames[pythonBlockIndex] || "")
              : "";
            frag.appendChild(createCodeBlock(part.language, rawCode, suggestedFilename));
            if (isPythonLanguage(part.language)) pythonBlockIndex += 1;
          }
        } else {
          frag.appendChild(createTextSegment(part.content || ""));
        }
      }'''

    old_text_segment = '''function createTextSegment(text) {
      const div = document.createElement("div");
      div.innerHTML = markdownToHtml(text);
      if (mayContainScientificMarkup(text)) {
        renderMathInElementSafe(div);
      }
      return div;
    }'''

    new_text_segment = r'''function extractStandaloneSvgBlocks(text) {
      const source = String(text || "");
      if (!source || source.indexOf("<svg") === -1 || source.indexOf("</svg>") === -1) return [source];

      const parts = [];
      const regex = /<svg\b[\s\S]*?<\/svg>/ig;
      let lastIndex = 0;
      let match;
      while ((match = regex.exec(source)) !== null) {
        const start = match.index;
        const end = regex.lastIndex;
        if (start > lastIndex) {
          parts.push(source.slice(lastIndex, start));
        }
        parts.push(match[0]);
        lastIndex = end;
      }
      if (lastIndex < source.length) {
        parts.push(source.slice(lastIndex));
      }
      return parts.length ? parts : [source];
    }

    function createTextSegment(text) {
      const source = String(text || "");
      const fragment = document.createDocumentFragment();
      const segments = extractStandaloneSvgBlocks(source);

      for (const segment of segments) {
        const piece = String(segment || "");
        const trimmed = piece.trim();
        if (!piece) continue;

        if (looksLikeSvgContent(trimmed)) {
          const svgNode = createSvgPreviewBlock(trimmed);
          if (svgNode) {
            fragment.appendChild(svgNode);
            continue;
          }
        }

        const div = document.createElement("div");
        div.innerHTML = markdownToHtml(piece);
        if (mayContainScientificMarkup(piece)) {
          renderMathInElementSafe(div);
        }
        if (div.childNodes.length === 0 && piece.trim()) {
          div.textContent = piece;
        }
        fragment.appendChild(div);
      }

      return fragment;
    }'''

    if '.svg-block {' not in html_doc:
        html_doc = html_doc.replace('    .attachment-chip {\n', css_insert + '    .attachment-chip {\n', 1)

    if 'function looksLikeSvgContent(text)' not in html_doc:
        html_doc = html_doc.replace(
            '    function createCodeBlock(language, code, suggestedFilename = "") {\n',
            helper_insert + '    function createCodeBlock(language, code, suggestedFilename = "") {\n',
            1,
        )

    if old_render_loop in html_doc:
        html_doc = html_doc.replace(old_render_loop, new_render_loop, 1)

    if old_text_segment in html_doc:
        html_doc = html_doc.replace(old_text_segment, new_text_segment, 1)

    return html_doc

_original_serve_index_html = serve_index_html

def serve_index_html() -> str:
    """Επιστρέφει το index HTML με επιπλέον απόδοση SVG block diagrams ως εικόνες."""
    return _patch_svg_preview_in_index_html(_original_serve_index_html())

def main() -> None:
    """Σημείο εισόδου της εφαρμογής.

Η συνάρτηση εκτελεί τα βασικά βήματα αρχικοποίησης με σωστή σειρά και ενεργοποιεί τις επιμέρους ρουτίνες που χρειάζονται για να ξεκινήσει το πρόγραμμα."""
    global HOST
    args = parse_args()
    logging.getLogger().setLevel(getattr(logging, args.log_level))
    HOST = args.host
    sep = '=' * 68
    print(sep)
    print(f'  {APP_TITLE}')
    print(sep)
    port = find_free_port(HOST, start_port=args.port)
    server = QuietThreadingHTTPServer((HOST, port), AppHandler)
    BROWSER_MONITOR.attach_server(server)
    start_browser_session_watchdog()
    startup_url = f'http://{HOST}:{port}/startup'
    chat_url = f'http://{HOST}:{port}'
    print(f'  🌐 Server  : {chat_url}')
    print(f'  🚀 Startup : {startup_url}')
    if not args.no_browser:
        print('  🌍 Άνοιγμα browser…')
        open_browser_later(startup_url, delay=0.25)
    else:
        print('  🚫 --no-browser: ο browser δεν άνοιξε αυτόματα.')
    print('  🛑 Ctrl+C για τερματισμό.')
    print(sep)
    threading.Thread(target=_run_initialization, args=(args, port), daemon=True, name='startup-init').start()
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print('\n🛑 Τερματισμός εφαρμογής…')
    finally:
        server.server_close()
        try:
            SESSION.reset()
        except Exception:
            pass
        print('✅ Server έκλεισε.')

def _patch_pdf_export_in_index_html(html_doc: str) -> str:
    """Εμπλουτίζει το UI με εξαγωγή μεμονωμένης απάντησης Assistant σε PDF μέσω headless browser print engine."""
    if not isinstance(html_doc, str) or not html_doc:
        return html_doc

    css_anchor = '    .attachment-chip {\n'
    css_insert = '''    .pdf-export-host {
      position: fixed;
      left: -20000px;
      top: 0;
      width: 704px;
      padding: 0;
      margin: 0;
      pointer-events: none;
      z-index: -1;
      isolation: isolate;
    }
    .pdf-export-shell {
      width: 704px !important;
      max-width: 704px !important;
      margin: 0 !important;
    }

'''
    if '.pdf-export-host {' not in html_doc and css_anchor in html_doc:
        html_doc = html_doc.replace(css_anchor, css_insert + css_anchor, 1)

    helper_anchor = '    function createMessage(role, content, attachments = []) {\n'
    helper_insert = r'''    function sanitizePdfFilenamePart(value, fallback = "export") {
      const normalized = String(value || "")
        .normalize("NFKD")
        .replace(/[\u0300-\u036f]/g, "")
        .replace(/[^a-zA-Z0-9._-]+/g, "-")
        .replace(/-+/g, "-")
        .replace(/^-+|-+$/g, "");
      return normalized || fallback;
    }

    function getAssistantMessageExportModel(messageWrapper) {
      const body = messageWrapper ? messageWrapper.querySelector(".msg-body") : null;
      const messageModel = String(
        (messageWrapper && messageWrapper.dataset && messageWrapper.dataset.exportModel) ||
        (body && body.dataset && body.dataset.exportModel) ||
        (els.modelSelect && els.modelSelect.value) ||
        "assistant"
      ).trim();
      return messageModel || "assistant";
    }

    function buildAssistantPdfFilename(messageWrapper) {
      const modelName = sanitizePdfFilenamePart(getAssistantMessageExportModel(messageWrapper), "assistant");
      const isoStamp = new Date().toISOString().replace(/[:]/g, "-").replace(/\.\d+Z$/, "Z");
      return `assistant-response-${modelName}-${isoStamp}.pdf`;
    }

    function buildAssistantDocxFilename(messageWrapper) {
      const modelName = sanitizePdfFilenamePart(getAssistantMessageExportModel(messageWrapper), "assistant");
      const isoStamp = new Date().toISOString().replace(/[:]/g, "-").replace(/\.\d+Z$/, "Z");
      return `assistant-response-${modelName}-${isoStamp}.docx`;
    }

    function getPdfExportHost() {
      let host = document.getElementById("pdfExportHost");
      if (!host) {
        host = document.createElement("div");
        host.id = "pdfExportHost";
        host.className = "pdf-export-host";
        document.body.appendChild(host);
      }
      return host;
    }

    function getLastEligibleAssistantMessageForPdf() {
      if (!els || !els.messages || !els.messages.querySelectorAll) return null;
      const assistantMessages = Array.from(els.messages.querySelectorAll(".msg.assistant"));
      let lastEligible = null;
      for (const message of assistantMessages) {
        const body = message.querySelector(".msg-body");
        const rawContent = String((body && body.dataset && body.dataset.rawContent) || "").trim();
        const hasStreamingPlaceholder = Boolean(body && body.querySelector && body.querySelector(".streaming-placeholder"));
        if (rawContent && !hasStreamingPlaceholder) {
          lastEligible = message;
        }
      }
      return lastEligible;
    }

    function syncAssistantPdfButtons() {
      const scope = els && els.messages ? els.messages : document;
      if (!scope || !scope.querySelectorAll) return;
      const buttons = Array.from(scope.querySelectorAll(".pdf-export-btn, .docx-export-btn"));
      for (const button of buttons) {
        button.hidden = true;
        button.disabled = true;
        button.setAttribute("aria-hidden", "true");
      }
      const targetMessage = getLastEligibleAssistantMessageForPdf();
      if (!targetMessage) return;
      targetMessage.querySelectorAll(".pdf-export-btn, .docx-export-btn").forEach((button) => {
        button.hidden = false;
        button.disabled = false;
        button.removeAttribute("aria-hidden");
      });
    }

    async function waitForImagesInElement(root) {
      const images = Array.from((root && root.querySelectorAll) ? root.querySelectorAll("img") : []);
      if (!images.length) return;

      images.forEach((img) => {
        try {
          img.loading = "eager";
          img.decoding = "sync";
          img.setAttribute("loading", "eager");
          img.setAttribute("decoding", "sync");
          if ("fetchPriority" in img) img.fetchPriority = "high";
        } catch (_) {}
      });

      await Promise.all(images.map((img) => {
        if (img.complete && img.naturalWidth > 0) return Promise.resolve();
        return new Promise((resolve) => {
          const done = () => resolve();
          img.addEventListener("load", done, { once: true });
          img.addEventListener("error", done, { once: true });
          setTimeout(done, 4000);
        });
      }));
    }

    function collectMathJaxSvgGlobalCacheMarkup(root = document) {
      try {
        const scope = root && root.querySelectorAll ? root : document;
        const seen = new Set();
        const fragments = [];
        const candidates = Array.from(scope.querySelectorAll([
          'svg#MJX-SVG-global-cache',
          'svg[id*="MJX"]',
          'svg[style*="display: none"]',
          'svg[width="0"]',
          'svg[height="0"]',
          'defs[id*="MJX"]',
          'path[id^="MJX-"]',
          'g[id^="MJX-"]'
        ].join(', ')));

        for (const node of candidates) {
          const svg = node && node.tagName && String(node.tagName).toLowerCase() === 'svg'
            ? node
            : (node && typeof node.closest === 'function' ? node.closest('svg') : null);
          if (!svg) continue;
          const svgMarkup = String(svg.outerHTML || '');
          if (!svgMarkup) continue;
          const lower = svgMarkup.toLowerCase();
          const looksLikeMathJaxCache = lower.includes('mjx') && (lower.includes('<defs') || lower.includes('display: none') || lower.includes('width="0"') || lower.includes('height="0"'));
          if (!looksLikeMathJaxCache) continue;
          if (seen.has(svgMarkup)) continue;
          seen.add(svgMarkup);
          fragments.push(svgMarkup);
        }

        return fragments.join('\n');
      } catch (_) {
        return '';
      }
    }

    async function waitForMathRenderingToFinish() {
      try {
        if (window.MathJax && window.MathJax.startup && window.MathJax.startup.promise) {
          await window.MathJax.startup.promise.catch(() => undefined);
        }
      } catch (_) {}

      try {
        if (typeof mathTypesetQueue !== "undefined" && mathTypesetQueue && typeof mathTypesetQueue.then === "function") {
          await mathTypesetQueue.catch(() => undefined);
        }
      } catch (_) {}

      await new Promise((resolve) => requestAnimationFrame(() => requestAnimationFrame(resolve)));
    }

    async function blobToDataUrl(blob) {
      return await new Promise((resolve, reject) => {
        const reader = new FileReader();
        reader.onload = () => resolve(String(reader.result || ""));
        reader.onerror = () => reject(reader.error || new Error("FileReader failed"));
        reader.readAsDataURL(blob);
      });
    }

    async function tryFetchAsDataUrl(src) {
      const source = String(src || "").trim();
      if (!source || source.startsWith("data:")) return source;
      try {
        const response = await fetch(source);
        if (!response.ok) return source;
        const blob = await response.blob();
        return await blobToDataUrl(blob);
      } catch (_) {
        return source;
      }
    }

    async function inlineImagesAsDataUrls(root) {
      if (!root || !root.querySelectorAll) return;
      const images = Array.from(root.querySelectorAll("img"));
      for (const img of images) {
        try {
          const src = String(img.currentSrc || img.getAttribute("src") || "").trim();
          if (!src || src.startsWith("data:")) continue;
          const dataUrl = await tryFetchAsDataUrl(src);
          if (dataUrl && dataUrl.startsWith("data:")) {
            img.setAttribute("src", dataUrl);
            img.src = dataUrl;
          }
        } catch (_) {}
      }
      await waitForImagesInElement(root);
    }

    async function rasterizeSvgImagesForPdf(root) {
      if (!root || !root.querySelectorAll) return;
      const svgImages = Array.from(root.querySelectorAll('img[src^="data:image/svg+xml"], img[src*="image/svg+xml"]'));
      for (const img of svgImages) {
        try {
          const source = String(img.currentSrc || img.getAttribute('src') || '').trim();
          if (!source) continue;

          const rect = img.getBoundingClientRect();
          const targetWidth = Math.max(1, Math.round(rect.width || img.width || img.naturalWidth || 1));
          const targetHeight = Math.max(1, Math.round(rect.height || img.height || img.naturalHeight || 1));

          await new Promise((resolve) => {
            const worker = new Image();
            worker.decoding = 'sync';
            worker.onload = () => {
              try {
                const canvas = document.createElement('canvas');
                canvas.width = targetWidth;
                canvas.height = targetHeight;
                const ctx = canvas.getContext('2d', { alpha: false });
                if (!ctx) {
                  resolve();
                  return;
                }
                ctx.fillStyle = '#ffffff';
                ctx.fillRect(0, 0, canvas.width, canvas.height);
                ctx.drawImage(worker, 0, 0, canvas.width, canvas.height);
                const pngDataUrl = canvas.toDataURL('image/png');
                if (pngDataUrl && pngDataUrl.startsWith('data:image/png')) {
                  img.setAttribute('src', pngDataUrl);
                  img.src = pngDataUrl;
                  img.setAttribute('data-pdf-rasterized', 'true');
                }
              } catch (_) {}
              resolve();
            };
            worker.onerror = () => resolve();
            worker.src = source;
          });
        } catch (_) {}
      }
      await waitForImagesInElement(root);
    }

    function serializeSvgForExport(svgNode, width, height) {
      try {
        const clone = svgNode.cloneNode(true);
        clone.setAttribute('xmlns', 'http://www.w3.org/2000/svg');
        clone.setAttribute('xmlns:xlink', 'http://www.w3.org/1999/xlink');
        if (!clone.getAttribute('viewBox')) {
          const vbWidth = Math.max(1, Math.round(width || 1));
          const vbHeight = Math.max(1, Math.round(height || 1));
          clone.setAttribute('viewBox', `0 0 ${vbWidth} ${vbHeight}`);
        }
        clone.setAttribute('width', String(Math.max(1, Math.round(width || 1))));
        clone.setAttribute('height', String(Math.max(1, Math.round(height || 1))));
        return String(clone.outerHTML || '');
      } catch (_) {
        return '';
      }
    }

    async function rasterizeSvgMarkupToPngDataUrl(svgMarkup, width, height) {
      const markup = String(svgMarkup || '').trim();
      if (!markup) return '';
      return await new Promise((resolve) => {
        const worker = new Image();
        worker.decoding = 'sync';
        worker.onload = () => {
          try {
            const canvas = document.createElement('canvas');
            canvas.width = Math.max(1, Math.round(width || worker.width || 1));
            canvas.height = Math.max(1, Math.round(height || worker.height || 1));
            const ctx = canvas.getContext('2d', { alpha: false });
            if (!ctx) {
              resolve('');
              return;
            }
            ctx.fillStyle = '#ffffff';
            ctx.fillRect(0, 0, canvas.width, canvas.height);
            ctx.drawImage(worker, 0, 0, canvas.width, canvas.height);
            resolve(canvas.toDataURL('image/png'));
          } catch (_) {
            resolve('');
          }
        };
        worker.onerror = () => resolve('');
        worker.src = `data:image/svg+xml;charset=utf-8,${encodeURIComponent(markup)}`;
      });
    }

    async function replaceSvgNodeWithRasterImage(targetNode, svgNode, options = {}) {
      if (!targetNode || !svgNode || !targetNode.isConnected) return false;
      try {
        const rect = targetNode.getBoundingClientRect();
        const svgRect = svgNode.getBoundingClientRect();
        const width = Math.max(1, Math.round(rect.width || svgRect.width || svgNode.width?.baseVal?.value || 1));
        const height = Math.max(1, Math.round(rect.height || svgRect.height || svgNode.height?.baseVal?.value || 1));
        const markup = serializeSvgForExport(svgNode, width, height);
        const pngDataUrl = await rasterizeSvgMarkupToPngDataUrl(markup, width, height);
        if (!pngDataUrl || !pngDataUrl.startsWith('data:image/png')) return false;

        const img = document.createElement('img');
        img.alt = options.alt || targetNode.getAttribute('aria-label') || targetNode.getAttribute('alt') || 'Formula';
        img.loading = 'eager';
        img.decoding = 'sync';
        img.src = pngDataUrl;
        img.style.maxWidth = '100%';
        img.style.width = options.widthCss || `${width}px`;
        img.style.height = 'auto';
        img.style.display = options.display || 'inline-block';
        img.style.verticalAlign = options.verticalAlign || 'middle';

        if (options.attrs && typeof options.attrs === 'object') {
          Object.entries(options.attrs).forEach(([key, value]) => {
            try {
              img.setAttribute(String(key), String(value));
            } catch (_) {}
          });
        }

        targetNode.replaceWith(img);
        return true;
      } catch (_) {
        return false;
      }
    }

    const DOCX_SUPERSCRIPT_MAP = Object.freeze({
      '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
      '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
      '+': '⁺', '-': '⁻', '=': '⁼', '(': '⁽', ')': '⁾',
      'n': 'ⁿ', 'i': 'ⁱ'
    });

    const DOCX_SUBSCRIPT_MAP = Object.freeze({
      '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
      '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉',
      '+': '₊', '-': '₋', '=': '₌', '(': '₍', ')': '₎',
      'a': 'ₐ', 'e': 'ₑ', 'h': 'ₕ', 'i': 'ᵢ', 'j': 'ⱼ',
      'k': 'ₖ', 'l': 'ₗ', 'm': 'ₘ', 'n': 'ₙ', 'o': 'ₒ',
      'p': 'ₚ', 'r': 'ᵣ', 's': 'ₛ', 't': 'ₜ', 'u': 'ᵤ',
      'v': 'ᵥ', 'x': 'ₓ'
    });

    function convertSimpleMathScript(text, mapping, fallbackPrefix) {
      const value = String(text || '').trim();
      if (!value) return '';
      if (value.length > 12) return `${fallbackPrefix}(${value})`;
      let out = '';
      for (const ch of value) {
        if (Object.prototype.hasOwnProperty.call(mapping, ch)) {
          out += mapping[ch];
        } else {
          return `${fallbackPrefix}(${value})`;
        }
      }
      return out;
    }

    function cleanupDocxMathText(value) {
      return String(value || '')
        .replace(/\u2061/g, '')
        .replace(/\u200b/g, '')
        .replace(/\s+/g, ' ')
        .replace(/\s+([,.;:!?])/g, '$1')
        .replace(/\(\s+/g, '(')
        .replace(/\s+\)/g, ')')
        .trim();
    }

    function texToReadableMathText(value) {
      let text = String(value || '').trim();
      if (!text) return '';
      text = text
        .replace(/^\\\((.*)\\\)$/s, '$1')
        .replace(/^\\\[(.*)\\\]$/s, '$1')
        .replace(/^\$(.*)\$$/s, '$1');

      const replacements = [
        [/\s*\\cdot\s*/g, ' · '],
        [/\s*\\times\s*/g, ' × '],
        [/\s*\\div\s*/g, ' ÷ '],
        [/\\leq\b/g, '≤'],
        [/\\geq\b/g, '≥'],
        [/\\neq\b/g, '≠'],
        [/\\approx\b/g, '≈'],
        [/\\to\b/g, '→'],
        [/\\rightarrow\b/g, '→'],
        [/\\leftarrow\b/g, '←'],
        [/\\pm\b/g, '±'],
        [/\\mp\b/g, '∓'],
        [/\\degree\b/g, '°'],
        [/\\circ\b/g, '°'],
        [/\\infty\b/g, '∞'],
        [/\\alpha\b/g, 'α'],
        [/\\beta\b/g, 'β'],
        [/\\gamma\b/g, 'γ'],
        [/\\delta\b/g, 'δ'],
        [/\\epsilon\b/g, 'ε'],
        [/\\varepsilon\b/g, 'ε'],
        [/\\theta\b/g, 'θ'],
        [/\\vartheta\b/g, 'ϑ'],
        [/\\lambda\b/g, 'λ'],
        [/\\mu\b/g, 'μ'],
        [/\\pi\b/g, 'π'],
        [/\\rho\b/g, 'ρ'],
        [/\\sigma\b/g, 'σ'],
        [/\\tau\b/g, 'τ'],
        [/\\phi\b/g, 'φ'],
        [/\\varphi\b/g, 'φ'],
        [/\\omega\b/g, 'ω'],
        [/\\Delta\b/g, 'Δ'],
        [/\\Gamma\b/g, 'Γ'],
        [/\\Theta\b/g, 'Θ'],
        [/\\Lambda\b/g, 'Λ'],
        [/\\Pi\b/g, 'Π'],
        [/\\Sigma\b/g, 'Σ'],
        [/\\Phi\b/g, 'Φ'],
        [/\\Omega\b/g, 'Ω'],
        [/\\sin\b/g, 'sin'],
        [/\\cos\b/g, 'cos'],
        [/\\tan\b/g, 'tan'],
        [/\\cot\b/g, 'cot'],
        [/\\sec\b/g, 'sec'],
        [/\\csc\b/g, 'csc'],
        [/\\log\b/g, 'log'],
        [/\\ln\b/g, 'ln'],
      ];
      for (const [pattern, replacement] of replacements) {
        text = text.replace(pattern, replacement);
      }

      let previous = '';
      while (text !== previous) {
        previous = text;
        text = text.replace(/\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}/g, '($1)/($2)');
        text = text.replace(/\\sqrt\s*\{([^{}]+)\}/g, '√($1)');
        text = text.replace(/([A-Za-zΑ-Ωα-ω0-9\)\]])\s*\^\s*\{([^{}]+)\}/g, (_, base, exp) => `${base}${convertSimpleMathScript(exp, DOCX_SUPERSCRIPT_MAP, '^')}`);
        text = text.replace(/([A-Za-zΑ-Ωα-ω0-9\)\]])\s*_\s*\{([^{}]+)\}/g, (_, base, sub) => `${base}${convertSimpleMathScript(sub, DOCX_SUBSCRIPT_MAP, '_')}`);
        text = text.replace(/([A-Za-zΑ-Ωα-ω0-9\)\]])\s*\^\s*([A-Za-z0-9+\-=()])/g, (_, base, exp) => `${base}${convertSimpleMathScript(exp, DOCX_SUPERSCRIPT_MAP, '^')}`);
        text = text.replace(/([A-Za-zΑ-Ωα-ω0-9\)\]])\s*_\s*([A-Za-z0-9+\-=()])/g, (_, base, sub) => `${base}${convertSimpleMathScript(sub, DOCX_SUBSCRIPT_MAP, '_')}`);
        text = text.replace(/\\mathrm\s*\{([^{}]+)\}/g, '$1');
        text = text.replace(/\\text\s*\{([^{}]+)\}/g, '$1');
      }

      text = text
        .replace(/\\,/g, ' ')
        .replace(/\\!/g, '')
        .replace(/\\;/g, ' ')
        .replace(/\\:/g, ' ')
        .replace(/[{}]/g, '')
        .replace(/\\([A-Za-z]+)/g, '$1');

      return cleanupDocxMathText(text);
    }

    function serializeMathMlNodeToDocxText(node) {
      if (!node) return '';
      if (node.nodeType === Node.TEXT_NODE) {
        return cleanupDocxMathText(node.textContent);
      }
      if (node.nodeType !== Node.ELEMENT_NODE) return '';

      const tagName = String(node.localName || node.nodeName || '').toLowerCase();
      const children = Array.from(node.childNodes || []);
      const joinChildren = () => cleanupDocxMathText(children.map((child) => serializeMathMlNodeToDocxText(child)).join(''));

      if (tagName === 'annotation') {
        return texToReadableMathText(node.textContent || '');
      }
      if (['math', 'mrow', 'mstyle', 'semantics', 'mtd', 'mtr', 'mtable', 'mtext', 'mi', 'mn', 'mo'].includes(tagName)) {
        return joinChildren() || cleanupDocxMathText(node.textContent);
      }
      if (tagName === 'msub') {
        const base = serializeMathMlNodeToDocxText(children[0]);
        const sub = serializeMathMlNodeToDocxText(children[1]);
        return cleanupDocxMathText(`${base}${convertSimpleMathScript(sub, DOCX_SUBSCRIPT_MAP, '_')}`);
      }
      if (tagName === 'msup') {
        const base = serializeMathMlNodeToDocxText(children[0]);
        const sup = serializeMathMlNodeToDocxText(children[1]);
        return cleanupDocxMathText(`${base}${convertSimpleMathScript(sup, DOCX_SUPERSCRIPT_MAP, '^')}`);
      }
      if (tagName === 'msubsup') {
        const base = serializeMathMlNodeToDocxText(children[0]);
        const sub = serializeMathMlNodeToDocxText(children[1]);
        const sup = serializeMathMlNodeToDocxText(children[2]);
        return cleanupDocxMathText(`${base}${convertSimpleMathScript(sub, DOCX_SUBSCRIPT_MAP, '_')}${convertSimpleMathScript(sup, DOCX_SUPERSCRIPT_MAP, '^')}`);
      }
      if (tagName === 'mfrac') {
        const num = serializeMathMlNodeToDocxText(children[0]);
        const den = serializeMathMlNodeToDocxText(children[1]);
        return cleanupDocxMathText(`(${num})/(${den})`);
      }
      if (tagName === 'msqrt') {
        return cleanupDocxMathText(`√(${children.map((child) => serializeMathMlNodeToDocxText(child)).join('')})`);
      }
      if (tagName === 'mroot') {
        const base = serializeMathMlNodeToDocxText(children[0]);
        const degree = serializeMathMlNodeToDocxText(children[1]);
        return cleanupDocxMathText(`${base}^(1/${degree})`);
      }
      if (tagName === 'mfenced') {
        return cleanupDocxMathText(`(${children.map((child) => serializeMathMlNodeToDocxText(child)).join('')})`);
      }
      if (tagName === 'mover' || tagName === 'munder' || tagName === 'munderover') {
        return joinChildren();
      }
      return joinChildren() || cleanupDocxMathText(node.textContent);
    }

    function extractDocxMathAltText(node) {
      if (!node || !node.querySelector) return '';
      const directCandidates = [
        node.getAttribute('data-tex'),
        node.getAttribute('data-latex'),
        node.getAttribute('aria-label'),
        node.getAttribute('title'),
      ];
      for (const candidate of directCandidates) {
        const value = texToReadableMathText(candidate);
        if (value) return value;
      }

      const annotation = node.querySelector('annotation[encoding="application/x-tex"], annotation[encoding="application/tex"], annotation[encoding="TeX"], annotation[encoding="tex"]');
      if (annotation) {
        const value = texToReadableMathText(annotation.textContent || '');
        if (value) return value;
      }

      const mathNode = node.querySelector('math, mjx-assistive-mml, .MJX_Assistive_MathML');
      if (mathNode) {
        const value = serializeMathMlNodeToDocxText(mathNode);
        if (value) return value;
      }

      return texToReadableMathText(node.textContent || '');
    }

    function replaceMathNodeWithDocxText(targetNode, altText, isBlock) {
      if (!targetNode || !targetNode.isConnected) return false;
      const text = cleanupDocxMathText(altText);
      if (!text) return false;
      const replacement = document.createElement(isBlock ? 'div' : 'span');
      replacement.textContent = text;
      replacement.setAttribute('data-docx-math', '1');
      replacement.setAttribute('data-docx-inline-math', isBlock ? '0' : '1');
      replacement.setAttribute('data-docx-block-math', isBlock ? '1' : '0');
      replacement.setAttribute('data-docx-alt', text);
      replacement.className = isBlock ? 'docx-math docx-block-math' : 'docx-math docx-inline-math';
      if (isBlock) {
        replacement.style.display = 'block';
        replacement.style.textAlign = 'center';
        replacement.style.margin = '0.35em 0';
      } else {
        replacement.style.display = 'inline';
      }
      targetNode.replaceWith(replacement);
      return true;
    }

    async function rasterizeMathAndSvgForDocx(root) {
      if (!root || !root.querySelectorAll) return;

      const mathContainers = Array.from(root.querySelectorAll('mjx-container, .katex-display, .katex'));
      for (const node of mathContainers) {
        if (!node || !node.isConnected) continue;
        const isBlock = node.matches('mjx-container[display="true"], .katex-display');
        const altText = extractDocxMathAltText(node);
        if (replaceMathNodeWithDocxText(node, altText, isBlock)) {
          continue;
        }
        const svg = node.querySelector('svg');
        if (!svg) continue;
        await replaceSvgNodeWithRasterImage(node, svg, {
          display: isBlock ? 'block' : 'inline-block',
          alt: altText || 'Math formula',
          attrs: {
            'data-docx-math': '1',
            'data-docx-inline-math': isBlock ? '0' : '1',
            'data-docx-block-math': isBlock ? '1' : '0',
            'data-docx-alt': altText || 'Math formula',
            'class': isBlock ? 'docx-math docx-block-math' : 'docx-math docx-inline-math',
          },
        });
      }

      const svgNodes = Array.from(root.querySelectorAll('svg'));
      for (const svgNode of svgNodes) {
        if (!svgNode || !svgNode.isConnected) continue;
        if (svgNode.closest('.mathjax-svg-cache')) {
          try { svgNode.remove(); } catch (_) {}
          continue;
        }
        await replaceSvgNodeWithRasterImage(svgNode, svgNode, {
          display: 'block',
          widthCss: '100%',
          alt: 'SVG figure',
        });
      }

      // Για DOCX το Word δεν χειρίζεται αξιόπιστα embedded SVG data URLs.
      // Μετατρέπουμε και τα <img src="data:image/svg+xml..."> previews σε PNG
      // πριν σταλούν στον server, ώστε να μην εμφανίζεται fallback όπως
      // "[SVG block diagram]" στο τελικό .docx.
      await rasterizeSvgImagesForPdf(root);
      await waitForImagesInElement(root);
      await inlineImagesAsDataUrls(root);
    }

    function replaceCanvasElementsWithImages(root) {
      if (!root || !root.querySelectorAll) return;
      const canvases = Array.from(root.querySelectorAll("canvas"));
      for (const canvas of canvases) {
        try {
          const img = document.createElement("img");
          img.alt = canvas.getAttribute("aria-label") || canvas.getAttribute("alt") || "Canvas";
          img.loading = "eager";
          img.decoding = "sync";
          img.src = canvas.toDataURL("image/png");
          img.style.width = `${canvas.width || canvas.clientWidth || 1}px`;
          img.style.height = `${canvas.height || canvas.clientHeight || 1}px`;
          img.style.maxWidth = "100%";
          canvas.replaceWith(img);
        } catch (_) {}
      }
    }


function isLoosePipeRowText(text) {
  const value = String(text || "").replace(/\u00a0/g, " ").trim();
  return value.length >= 3 && value.startsWith("|") && value.endsWith("|") && value.includes("|");
}

function isLoosePipeDelimiterText(text) {
  const value = String(text || "").replace(/\u00a0/g, " ").trim();
  if (!isLoosePipeRowText(value)) return false;
  const inner = value.replace(/^\|/, "").replace(/\|$/, "").trim();
  if (!inner) return false;
  return inner.split("|").every((part) => /^\s*:?-{3,}:?\s*$/.test(part));
}

function splitLoosePipeRow(rowHtml) {
  const raw = String(rowHtml || "").trim().replace(/^\|/, "").replace(/\|$/, "");
  return raw.split("|").map((part) => part.trim());
}

function buildLoosePipeTableFromNodes(nodes) {
  if (!Array.isArray(nodes) || nodes.length < 2) return null;
  const textRows = nodes.map((node) => String(node.textContent || "").trim()).filter(Boolean);
  const htmlRows = nodes.map((node) => String(node.innerHTML || "").trim());
  if (textRows.length < 2) return null;
  const columnCount = Math.max(...htmlRows.map((row) => splitLoosePipeRow(row).length));
  if (columnCount < 2) return null;

  let headerCells = null;
  let dataStartIndex = 0;
  if (textRows.length >= 2 && isLoosePipeDelimiterText(textRows[1])) {
    headerCells = splitLoosePipeRow(htmlRows[0]);
    dataStartIndex = 2;
  }

  const wrap = document.createElement("div");
  wrap.className = "md-table-wrap normalized-pipe-table";
  const table = document.createElement("table");
  table.className = "md-table";
  wrap.appendChild(table);

  if (headerCells) {
    const thead = document.createElement("thead");
    const tr = document.createElement("tr");
    headerCells.forEach((cellHtml) => {
      const th = document.createElement("th");
      th.innerHTML = cellHtml || "&nbsp;";
      tr.appendChild(th);
    });
    thead.appendChild(tr);
    table.appendChild(thead);
  }

  const tbody = document.createElement("tbody");
  for (let i = dataStartIndex; i < htmlRows.length; i += 1) {
    const cells = splitLoosePipeRow(htmlRows[i]);
    if (!cells.length) continue;
    const tr = document.createElement("tr");
    for (let c = 0; c < columnCount; c += 1) {
      const td = document.createElement("td");
      td.innerHTML = (cells[c] || "").trim() || "&nbsp;";
      tr.appendChild(td);
    }
    tbody.appendChild(tr);
  }
  if (!tbody.children.length) return null;
  table.appendChild(tbody);
  return wrap;
}

function normalizeLoosePipeTables(root) {
  if (!root || !root.querySelector) return;
  const targets = Array.from(root.querySelectorAll(".assistant-print-prompt-body, .assistant-print-body"));
  if (!targets.length) targets.push(root);

  for (const body of targets) {
    const children = Array.from((body && body.children) || []);
    let group = [];

    const flush = () => {
      if (group.length < 2) {
        group = [];
        return;
      }
      const tableWrap = buildLoosePipeTableFromNodes(group);
      if (!tableWrap) {
        group = [];
        return;
      }
      const first = group[0];
      first.replaceWith(tableWrap);
      for (let i = 1; i < group.length; i += 1) {
        group[i].remove();
      }
      group = [];
    };

    for (const child of children) {
      const isEligible = child && child.matches && child.matches(".md-p, p, div, section, article");
      const text = String((child && child.textContent) || "").trim();
      if (isEligible && isLoosePipeRowText(text)) {
        group.push(child);
      } else {
        flush();
      }
    }
    flush();
  }
}

    function findRelatedUserPrompt(messageWrapper) {
      let node = messageWrapper ? messageWrapper.previousElementSibling : null;
      while (node) {
        if (node.classList && node.classList.contains("msg") && node.classList.contains("user")) {
          const body = node.querySelector(".msg-body");
          return String((body && body.dataset && body.dataset.rawContent) || (body && body.textContent) || "").trim();
        }
        node = node.previousElementSibling;
      }
      return "";
    }

    function normalizeCloneForPdfPrint(clone) {
      if (!clone) return;
      clone.classList.add("pdf-export-shell");
      clone.style.width = "704px";
      clone.style.maxWidth = "704px";
      clone.style.minWidth = "0";
      clone.style.margin = "0";

      clone.querySelectorAll("button, .message-tools, .streaming-placeholder").forEach((node) => node.remove());

      clone.querySelectorAll(".msg-body, .md-table-wrap, .katex-display, mjx-container, mjx-container[display='true'], pre, .code-pre").forEach((node) => {
        node.style.overflow = "visible";
        node.style.maxWidth = "100%";
      });

      clone.querySelectorAll("mjx-container, mjx-container[display='true']").forEach((node) => {
        node.style.overflow = "visible";
        node.style.maxWidth = "100%";
        node.style.display = node.getAttribute('display') === 'true' ? 'block' : 'inline-block';
      });

      clone.querySelectorAll("mjx-container svg, .katex svg").forEach((node) => {
        node.style.overflow = "visible";
        node.style.maxWidth = "100%";
      });

      clone.querySelectorAll(".md-table").forEach((node) => {
        node.style.width = "100%";
        node.style.minWidth = "0";
        node.style.tableLayout = "auto";
      });

      clone.querySelectorAll(".code-block").forEach((node) => {
        node.style.overflow = "visible";
        node.style.breakInside = "auto";
        node.style.pageBreakInside = "auto";
        node.style.boxDecorationBreak = "clone";
        node.style.webkitBoxDecorationBreak = "clone";
      });

      clone.querySelectorAll(".code-toolbar").forEach((node) => {
        node.style.breakAfter = "avoid-page";
        node.style.pageBreakAfter = "avoid";
      });

      clone.querySelectorAll("pre, .code-pre").forEach((node) => {
        node.style.overflow = "visible";
        node.style.breakInside = "auto";
        node.style.pageBreakInside = "auto";
        node.style.boxDecorationBreak = "clone";
        node.style.webkitBoxDecorationBreak = "clone";
      });

      clone.querySelectorAll("img, svg, canvas").forEach((node) => {
        node.style.maxWidth = "100%";
        if (node.style) node.style.height = "auto";
      });

      clone.querySelectorAll("img[src^='data:image/svg+xml'], img[data-pdf-rasterized='true'], .svg-preview-image").forEach((node) => {
        node.style.background = "#ffffff";
        node.style.display = "block";
      });
    }


    function replacePdfCloneCodeBlocksWithRenderedPlots(originalMessageWrapper, cloneRoot) {
      if (!originalMessageWrapper || !cloneRoot) return;
      const originalBlocks = Array.from(originalMessageWrapper.querySelectorAll('.msg-body .code-block'));
      const cloneBlocks = Array.from(cloneRoot.querySelectorAll('.assistant-print-body .code-block'));
      if (!originalBlocks.length || !cloneBlocks.length) return;
      const total = Math.min(originalBlocks.length, cloneBlocks.length);

      for (let i = 0; i < total; i += 1) {
        const originalBlock = originalBlocks[i];
        const cloneBlock = cloneBlocks[i];
        if (!originalBlock || !cloneBlock) continue;

        const renderedContainer = originalBlock.querySelector('.python-plot-render');
        const renderedImage = renderedContainer && renderedContainer.querySelector('img[src]');
        if (!renderedContainer || !renderedImage) continue;

        const figure = document.createElement('figure');
        figure.className = 'python-plot-render python-plot-render-print';
        figure.style.margin = '8px 0';
        figure.style.padding = '4px 0';
        figure.style.border = '0';
        figure.style.borderRadius = '0';
        figure.style.background = '#ffffff';
        figure.style.boxShadow = 'none';
        figure.style.textAlign = 'center';
        figure.style.breakInside = 'avoid';
        figure.style.pageBreakInside = 'avoid';

        const image = renderedImage.cloneNode(true);
        image.removeAttribute('loading');
        image.style.display = 'block';
        image.style.width = 'auto';
        image.style.maxWidth = '94%';
        image.style.maxHeight = '360px';
        image.style.height = 'auto';
        image.style.margin = '0 auto';
        image.style.borderRadius = '0';
        image.style.background = '#ffffff';
        image.style.border = '0';
        image.style.boxShadow = 'none';
        figure.appendChild(image);

        cloneBlock.replaceWith(figure);
      }

      cloneRoot.querySelectorAll('.assistant-print-body .python-plot-render-print figcaption, .assistant-print-body .python-plot-render-print a[href], .assistant-print-body .python-plot-render-print .plot-link, .assistant-print-body .python-plot-render-print .muted, .assistant-print-body .python-plot-render-print .secondary').forEach((node) => node.remove());
      cloneRoot.querySelectorAll('.assistant-print-body .python-plot-render-print').forEach((node) => {
        node.style.margin = '8px 0';
        node.style.padding = '4px 0';
        node.style.textAlign = 'center';
        Array.from(node.childNodes).forEach((child) => {
          if (child && child.nodeType === Node.TEXT_NODE && /Rendered Plot|generated-media|https?:\/\//i.test(String(child.textContent || ''))) {
            child.textContent = '';
          }
        });
      });
    }

    function replacePdfCloneSvgBlocksWithRenderedFigures(cloneRoot) {
      if (!cloneRoot || !cloneRoot.querySelectorAll) return;
      const svgBlocks = Array.from(cloneRoot.querySelectorAll('.assistant-print-body .svg-block, .assistant-print-prompt-body .svg-block'));
      for (const svgBlock of svgBlocks) {
        if (!svgBlock) continue;

        const previewImage = svgBlock.querySelector('.svg-preview-image, img[src]');
        if (!previewImage) {
          svgBlock.remove();
          continue;
        }

        const figure = document.createElement('figure');
        figure.className = 'svg-print-figure';
        figure.style.margin = '8px 0';
        figure.style.padding = '4px 0';
        figure.style.border = '0';
        figure.style.borderRadius = '0';
        figure.style.background = '#ffffff';
        figure.style.boxShadow = 'none';
        figure.style.textAlign = 'center';
        figure.style.breakInside = 'avoid';
        figure.style.pageBreakInside = 'avoid';

        const image = previewImage.cloneNode(true);
        image.removeAttribute('loading');
        image.style.display = 'block';
        image.style.width = 'auto';
        image.style.maxWidth = '94%';
        image.style.maxHeight = '360px';
        image.style.height = 'auto';
        image.style.margin = '0 auto';
        image.style.background = '#ffffff';
        image.style.borderRadius = '0';
        image.style.boxShadow = 'none';
        image.style.border = '0';
        figure.appendChild(image);

        svgBlock.replaceWith(figure);
      }
    }


    function extractDocxMathPlaceholders(source) {
      let text = String(source || '');
      const tokenMap = new Map();
      let counter = 0;

      const makeToken = (isBlock, rawValue) => {
        const readable = cleanupDocxMathText(texToReadableMathText(rawValue) || rawValue || '');
        if (!readable) return rawValue;
        const token = `DOCXMATH${isBlock ? 'BLOCK' : 'INLINE'}${counter += 1}TOKEN`;
        tokenMap.set(token, { text: readable, block: !!isBlock });
        return token;
      };

      text = text.replace(/(^|[\r\n])\s*\$\$([\s\S]+?)\$\$(?=\s*(?:$|[\r\n]))/g, (match, lead, expr) => `${lead}${makeToken(true, expr)}`);
      text = text.replace(/(^|[\r\n])\s*\\\[([\s\S]+?)\\\](?=\s*(?:$|[\r\n]))/g, (match, lead, expr) => `${lead}${makeToken(true, expr)}`);
      text = text.replace(/\\\(([\s\S]+?)\\\)/g, (match, expr) => makeToken(false, expr));
      text = text.replace(/(^|[^\\$])\$([^$\r\n]+?)\$/g, (match, lead, expr) => `${lead}${makeToken(false, expr)}`);

      return { text, tokenMap };
    }

    function createDocxMathElement(tokenData) {
      const isBlock = !!(tokenData && tokenData.block);
      const element = document.createElement(isBlock ? 'div' : 'span');
      const text = cleanupDocxMathText((tokenData && tokenData.text) || '');
      element.textContent = text;
      element.setAttribute('data-docx-math', '1');
      element.setAttribute('data-docx-inline-math', isBlock ? '0' : '1');
      element.setAttribute('data-docx-block-math', isBlock ? '1' : '0');
      element.setAttribute('data-docx-alt', text);
      element.className = isBlock ? 'docx-math docx-block-math' : 'docx-math docx-inline-math';
      if (isBlock) {
        element.style.display = 'block';
        element.style.textAlign = 'center';
        element.style.margin = '0.35em 0';
      } else {
        element.style.display = 'inline';
      }
      return element;
    }

    function replaceDocxMathPlaceholders(root, tokenMap) {
      if (!root || !tokenMap || !tokenMap.size) return;
      const tokenPattern = /DOCXMATH(?:BLOCK|INLINE)\d+TOKEN/g;
      const inlineAncestorSelector = 'a, strong, b, em, i, u, small, sup, sub, code, kbd, samp, label';
      const walker = document.createTreeWalker(root, NodeFilter.SHOW_TEXT, null);
      const textNodes = [];
      while (walker.nextNode()) {
        const current = walker.currentNode;
        if (current && tokenPattern.test(String(current.nodeValue || ''))) {
          textNodes.push(current);
        }
        tokenPattern.lastIndex = 0;
      }

      for (const textNode of textNodes) {
        if (!textNode || !textNode.parentNode) continue;
        const original = String(textNode.nodeValue || '');
        tokenPattern.lastIndex = 0;
        const matches = [...original.matchAll(tokenPattern)];
        if (!matches.length) continue;

        const parentEl = textNode.parentElement;
        if (
          matches.length === 1
          && matches[0].index === 0
          && matches[0][0].length === original.trim().length
          && parentEl
          && !parentEl.matches(inlineAncestorSelector)
        ) {
          const token = matches[0][0];
          const tokenData = tokenMap.get(token);
          if (tokenData && tokenData.block) {
            const replacement = createDocxMathElement(tokenData);
            if (parentEl.matches('p, .md-p')) {
              parentEl.replaceWith(replacement);
              continue;
            }
            textNode.parentNode.replaceChild(replacement, textNode);
            continue;
          }
        }

        const fragment = document.createDocumentFragment();
        let lastIndex = 0;
        for (const match of matches) {
          const token = match[0];
          const index = match.index || 0;
          if (index > lastIndex) {
            fragment.appendChild(document.createTextNode(original.slice(lastIndex, index)));
          }
          const tokenData = tokenMap.get(token);
          if (tokenData) {
            fragment.appendChild(createDocxMathElement({ text: tokenData.text, block: false }));
          } else {
            fragment.appendChild(document.createTextNode(token));
          }
          lastIndex = index + token.length;
        }
        if (lastIndex < original.length) {
          fragment.appendChild(document.createTextNode(original.slice(lastIndex)));
        }
        textNode.parentNode.replaceChild(fragment, textNode);
      }
    }

    function renderMessageContentForDocxExport(target, source) {
      if (!target) return;
      const prepared = extractDocxMathPlaceholders(source);
      renderMessageContent(target, prepared.text);
      replaceDocxMathPlaceholders(target, prepared.tokenMap);
    }


    async function buildDocxReadyClone(messageWrapper, rawContent, host, options = {}) {
      const article = document.createElement("article");
      article.className = "assistant-print-doc";

      const cover = document.createElement("header");
      cover.className = "assistant-export-cover";
      const modelDisplay = String((options && options.modelDisplay) || "").trim() || "Άγνωστο μοντέλο";
      cover.innerHTML = `
        <div class="assistant-export-eyebrow">Assistant Response</div>
        <h1 class="assistant-export-title">Μοντέλο : ${escapeHtml(modelDisplay)}</h1>
        <p class="assistant-export-subtitle">Ημερομηνία εξαγωγής: ${escapeHtml(new Date().toLocaleString("el-GR"))}</p>
      `;
      article.appendChild(cover);

      const promptText = String((options && options.userPrompt) || "").trim();
      if (promptText) {
        const promptSection = document.createElement("section");
        promptSection.className = "assistant-export-prompt";

        const promptTitle = document.createElement("h2");
        promptTitle.className = "assistant-export-section-title";
        promptTitle.textContent = "Prompt Χρήστη";
        promptSection.appendChild(promptTitle);

        const promptBody = document.createElement("div");
        promptBody.className = "msg-body assistant-print-prompt-body";
        promptSection.appendChild(promptBody);
        article.appendChild(promptSection);
        renderMessageContentForDocxExport(promptBody, promptText);
        promptBody.querySelectorAll(".thinking-block, details.thinking-block, mjx-container, .katex-display, .katex").forEach((node) => node.remove());
      }

      const answerSection = document.createElement("section");
      answerSection.className = "assistant-export-answer";
      const answerTitle = document.createElement("h2");
      answerTitle.className = "assistant-export-section-title";
      answerTitle.textContent = "Απάντηση Assistant";
      answerSection.appendChild(answerTitle);

      const body = document.createElement("div");
      body.className = "msg-body assistant-print-body";
      answerSection.appendChild(body);
      article.appendChild(answerSection);
      host.appendChild(article);

      renderMessageContentForDocxExport(body, rawContent || "");
      body.querySelectorAll(".thinking-block, details.thinking-block, mjx-container, .katex-display, .katex").forEach((node) => node.remove());
      normalizeLoosePipeTables(article);
      replacePdfCloneCodeBlocksWithRenderedPlots(messageWrapper, article);
      replacePdfCloneSvgBlocksWithRenderedFigures(article);

      replaceCanvasElementsWithImages(article);
      await waitForImagesInElement(article);
      await inlineImagesAsDataUrls(article);
      normalizeCloneForPdfPrint(article);
      await new Promise((resolve) => requestAnimationFrame(() => requestAnimationFrame(resolve)));
      return article;
    }

    async function buildPdfReadyClone(messageWrapper, rawContent, host, options = {}) {
      const article = document.createElement("article");
      article.className = "assistant-print-doc";
      const includeEmbeddedStyle = !options || options.includeEmbeddedStyle !== false;

      if (includeEmbeddedStyle) {
        const style = document.createElement("style");
        style.textContent = `
          .assistant-export-cover {
            margin: 0 0 18px 0;
            padding: 20px 22px;
            border: 1px solid #dbe4f0;
            border-radius: 18px;
            background: linear-gradient(180deg, #f8fbff 0%, #eef6ff 100%);
          }
          .assistant-export-eyebrow {
            margin: 0 0 8px 0;
            font-size: 10.5pt;
            font-weight: 700;
            letter-spacing: 0.08em;
            text-transform: uppercase;
            color: #2563eb;
          }
          .assistant-export-title {
            margin: 0;
            font-size: 11pt;
            line-height: 1.15;
            font-weight: 700;
            color: #0f172a;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
          }
          .assistant-export-subtitle {
            margin: 6px 0 0 0;
            font-size: 9.25pt;
            line-height: 1.4;
            color: #475569;
          }
          .assistant-export-prompt,
          .assistant-export-answer {
            margin: 0 0 16px 0;
            padding: 16px 18px;
            border: 1px solid #dbe4f0;
            border-radius: 16px;
            background: #ffffff;
            width: 100%;
            max-width: 100%;
            min-width: 0;
            overflow: visible;
          }
          .assistant-export-answer {
            padding-top: 18px;
          }
          .assistant-export-section-title {
            margin: 0 0 10px 0;
            font-size: 11pt;
            font-weight: 800;
            letter-spacing: 0.02em;
            color: #1e3a8a;
          }
          .assistant-print-prompt-body,
          .assistant-print-body {
            font-size: 11.25pt;
            line-height: 1.62;
            color: #0f172a;
            max-width: 100%;
            overflow-wrap: anywhere;
            word-break: break-word;
          }
          .assistant-print-prompt-body > :first-child,
          .assistant-print-body > :first-child {
            margin-top: 0 !important;
          }
          .assistant-print-prompt-body > :last-child,
          .assistant-print-body > :last-child {
            margin-bottom: 0 !important;
          }
          .assistant-print-body .code-block {
            overflow: visible;
            break-inside: auto;
            page-break-inside: auto;
            box-decoration-break: clone;
            -webkit-box-decoration-break: clone;
          }
          .assistant-print-body .code-toolbar {
            break-after: avoid-page;
            page-break-after: avoid;
          }
          .assistant-print-body pre,
          .assistant-print-body .code-pre {
            overflow: visible;
            break-inside: auto;
            page-break-inside: auto;
            box-decoration-break: clone;
            -webkit-box-decoration-break: clone;
          }
          .assistant-print-body .python-plot-render-print,
          .assistant-print-body .python-plot-render-print img {
            max-width: 100%;
          }
        `;
        article.appendChild(style);
      }

      const cover = document.createElement("header");
      cover.className = "assistant-export-cover";
      const modelDisplay = String((options && options.modelDisplay) || "").trim() || "Άγνωστο μοντέλο";
      cover.innerHTML = `
        <div class="assistant-export-eyebrow">Assistant Response</div>
        <h1 class="assistant-export-title">Μοντέλο : ${escapeHtml(modelDisplay)}</h1>
        <p class="assistant-export-subtitle">Ημερομηνία εξαγωγής: ${escapeHtml(new Date().toLocaleString("el-GR"))}</p>
      `;
      article.appendChild(cover);

      const promptText = String((options && options.userPrompt) || "").trim();
      if (promptText) {
        const promptSection = document.createElement("section");
        promptSection.className = "assistant-export-prompt";

        const promptTitle = document.createElement("h2");
        promptTitle.className = "assistant-export-section-title";
        promptTitle.textContent = "Prompt Χρήστη";
        promptSection.appendChild(promptTitle);

        const promptBody = document.createElement("div");
        promptBody.className = "msg-body assistant-print-prompt-body";
        promptSection.appendChild(promptBody);
        article.appendChild(promptSection);
        renderMessageContent(promptBody, promptText);
        promptBody.querySelectorAll(".thinking-block, details.thinking-block").forEach((node) => node.remove());
      }

      const answerSection = document.createElement("section");
      answerSection.className = "assistant-export-answer";
      const answerTitle = document.createElement("h2");
      answerTitle.className = "assistant-export-section-title";
      answerTitle.textContent = "Απάντηση Assistant";
      answerSection.appendChild(answerTitle);

      const body = document.createElement("div");
      body.className = "msg-body assistant-print-body";
      answerSection.appendChild(body);
      article.appendChild(answerSection);
      host.appendChild(article);

      renderMessageContent(body, rawContent || "");
      body.querySelectorAll(".thinking-block, details.thinking-block").forEach((node) => node.remove());
      normalizeLoosePipeTables(article);
      replacePdfCloneCodeBlocksWithRenderedPlots(messageWrapper, article);
      replacePdfCloneSvgBlocksWithRenderedFigures(article);

      await waitForMathRenderingToFinish();
      replaceCanvasElementsWithImages(article);
      await waitForImagesInElement(article);
      await inlineImagesAsDataUrls(article);
      await rasterizeSvgImagesForPdf(article);
      normalizeCloneForPdfPrint(article);
      await new Promise((resolve) => requestAnimationFrame(() => requestAnimationFrame(resolve)));
      return article;
    }

    async function downloadAssistantPdfFromServer(payload, filename) {
      const response = await fetch("/api/export-assistant-pdf", {
        method: "POST",
        headers: { "Content-Type": "application/json" },
        body: JSON.stringify(payload),
      });

      if (!response.ok) {
        let errorMessage = `HTTP ${response.status}`;
        try {
          const data = await response.json();
          if (data && data.error) errorMessage = data.error;
        } catch (_) {
          try {
            const text = await response.text();
            if (text) errorMessage = text;
          } catch (_) {}
        }
        throw new Error(errorMessage);
      }

      const blob = await response.blob();
      const objectUrl = URL.createObjectURL(blob);
      const anchor = document.createElement("a");
      anchor.href = objectUrl;
      anchor.download = filename || "assistant-response.pdf";
      document.body.appendChild(anchor);
      anchor.click();
      setTimeout(() => {
        URL.revokeObjectURL(objectUrl);
        anchor.remove();
      }, 1500);
    }

    async function downloadAssistantDocxFromServer(payload, filename) {
      const response = await fetch("/api/export-assistant-docx", {
        method: "POST",
        headers: { "Content-Type": "application/json" },
        body: JSON.stringify(payload),
      });

      if (!response.ok) {
        let errorMessage = `HTTP ${response.status}`;
        try {
          const data = await response.json();
          if (data && data.error) errorMessage = data.error;
        } catch (_) {
          try {
            const text = await response.text();
            if (text) errorMessage = text;
          } catch (_) {}
        }
        throw new Error(errorMessage);
      }

      const blob = await response.blob();
      const objectUrl = URL.createObjectURL(blob);
      const anchor = document.createElement("a");
      anchor.href = objectUrl;
      anchor.download = filename || "assistant-response.docx";
      document.body.appendChild(anchor);
      anchor.click();
      setTimeout(() => {
        URL.revokeObjectURL(objectUrl);
        anchor.remove();
      }, 1500);
    }

    async function exportAssistantMessageToPdf(messageWrapper, triggerBtn) {
      const defaultLabel = "📄 Export PDF";
      const safeFeedback = (label, mode) => {
        if (!triggerBtn) return;
        if (mode === "error") {
          setButtonFeedback(triggerBtn, label, defaultLabel, "error");
        } else {
          setButtonFeedback(triggerBtn, label, defaultLabel);
        }
      };

      try {
        if (!messageWrapper) {
          renderSystemNotice("Δεν βρέθηκε απάντηση για εξαγωγή PDF.");
          safeFeedback("❌ Error", "error");
          return;
        }

        const body = messageWrapper.querySelector(".msg-body");
        const rawContent = String((body && body.dataset && body.dataset.rawContent) || (body && body.textContent) || "").trim();
        if (!rawContent) {
          renderSystemNotice("Η απάντηση είναι κενή. Δεν υπάρχει περιεχόμενο για εξαγωγή PDF.");
          safeFeedback("❌ Error", "error");
          return;
        }

        if (triggerBtn) {
          triggerBtn.disabled = true;
          triggerBtn.textContent = "⏳ Export PDF...";
        }

        await waitForMathRenderingToFinish();

        if (document.fonts && document.fonts.ready) {
          try { await document.fonts.ready; } catch (_) {}
        }

        const filename = buildAssistantPdfFilename(messageWrapper);
        const userPrompt = findRelatedUserPrompt(messageWrapper);
        const modelDisplay = getAssistantMessageExportModel(messageWrapper);
        const host = getPdfExportHost();
        host.innerHTML = "";
        const clone = await buildPdfReadyClone(messageWrapper, rawContent, host, { userPrompt, modelDisplay });
        const htmlFragment = clone.outerHTML;
        const theme = String(document.documentElement.dataset.theme || "dark").toLowerCase() === "light" ? "light" : "dark";

        const mathjaxSvgCache = collectMathJaxSvgGlobalCacheMarkup(document);

        await downloadAssistantPdfFromServer({
          html_fragment: htmlFragment,
          theme,
          filename,
          mathjax_svg_cache: mathjaxSvgCache,
        }, filename);

        host.innerHTML = "";
        safeFeedback("✅ Saved");
        renderSystemNotice("Η απάντηση του Assistant εξήχθη σε PDF μέσω native headless browser print engine.");
      } catch (err) {
        const host = document.getElementById("pdfExportHost");
        if (host) host.innerHTML = "";
        safeFeedback("❌ Error", "error");
        renderSystemNotice(`Σφάλμα εξαγωγής PDF: ${err && err.message ? err.message : String(err)}`);
      } finally {
        if (triggerBtn) {
          triggerBtn.disabled = false;
        }
      }
    }

    async function exportAssistantMessageToDocx(messageWrapper, triggerBtn) {
      const defaultLabel = "📝 Export Docx";
      const safeFeedback = (label, mode) => {
        if (!triggerBtn) return;
        if (mode === "error") {
          setButtonFeedback(triggerBtn, label, defaultLabel, "error");
        } else {
          setButtonFeedback(triggerBtn, label, defaultLabel);
        }
      };

      try {
        if (!messageWrapper) {
          renderSystemNotice("Δεν βρέθηκε απάντηση για εξαγωγή Docx.");
          safeFeedback("❌ Error", "error");
          return;
        }

        const body = messageWrapper.querySelector(".msg-body");
        const rawContent = String((body && body.dataset && body.dataset.rawContent) || (body && body.textContent) || "").trim();
        if (!rawContent) {
          renderSystemNotice("Η απάντηση είναι κενή. Δεν υπάρχει περιεχόμενο για εξαγωγή Docx.");
          safeFeedback("❌ Error", "error");
          return;
        }

        if (triggerBtn) {
          triggerBtn.disabled = true;
          triggerBtn.textContent = "⏳ Export Docx...";
        }

        await waitForMathRenderingToFinish();
        if (document.fonts && document.fonts.ready) {
          try { await document.fonts.ready; } catch (_) {}
        }

        const filename = buildAssistantDocxFilename(messageWrapper);
        const userPrompt = findRelatedUserPrompt(messageWrapper);
        const modelDisplay = getAssistantMessageExportModel(messageWrapper);
        const host = getPdfExportHost();
        host.innerHTML = "";
        const clone = await buildDocxReadyClone(messageWrapper, rawContent, host, { userPrompt, modelDisplay });
        clone.querySelectorAll("style, script, noscript, template, meta, link, head").forEach((node) => { try { node.remove(); } catch (_) {} });
        await rasterizeMathAndSvgForDocx(clone);
        replaceCanvasElementsWithImages(clone);
        await waitForImagesInElement(clone);
        await inlineImagesAsDataUrls(clone);
        const htmlFragment = clone.outerHTML;

        await downloadAssistantDocxFromServer({
          html_fragment: htmlFragment,
          filename,
        }, filename);

        host.innerHTML = "";
        safeFeedback("✅ Saved");
        renderSystemNotice("Η απάντηση του Assistant εξήχθη σε Docx.");
      } catch (err) {
        const host = document.getElementById("pdfExportHost");
        if (host) host.innerHTML = "";
        safeFeedback("❌ Error", "error");
        renderSystemNotice(`Σφάλμα εξαγωγής Docx: ${err && err.message ? err.message : String(err)}`);
      } finally {
        if (triggerBtn) {
          triggerBtn.disabled = false;
        }
      }
    }

'''
    if 'function exportAssistantMessageToPdf(messageWrapper, triggerBtn)' not in html_doc and helper_anchor in html_doc:
        html_doc = html_doc.replace(helper_anchor, helper_insert + helper_anchor, 1)

    tools_anchor = r'''      const tools   = document.createElement("div");
      tools.className = "message-tools";
      const copyBtn = document.createElement("button");
'''
    tools_insert = r'''      const tools   = document.createElement("div");
      tools.className = "message-tools";

      if (role === "assistant") {
        const pdfBtn = document.createElement("button");
        pdfBtn.type = "button";
        pdfBtn.textContent = "📄 Export PDF";
        pdfBtn.title = "Εξαγωγή ολόκληρης της απάντησης του Assistant σε αρχείο PDF";
        pdfBtn.className = "tool-btn pdf-export-btn";
        pdfBtn.hidden = true;
        pdfBtn.disabled = true;
        pdfBtn.setAttribute("aria-hidden", "true");
        pdfBtn.addEventListener("click", async () => {
          await exportAssistantMessageToPdf(wrapper, pdfBtn);
        });
        tools.appendChild(pdfBtn);

        const docxBtn = document.createElement("button");
        docxBtn.type = "button";
        docxBtn.textContent = "📝 Export Docx";
        docxBtn.title = "Εξαγωγή ολόκληρης της απάντησης του Assistant σε αρχείο Microsoft Word (.docx)";
        docxBtn.className = "tool-btn docx-export-btn";
        docxBtn.hidden = true;
        docxBtn.disabled = true;
        docxBtn.setAttribute("aria-hidden", "true");
        docxBtn.addEventListener("click", async () => {
          await exportAssistantMessageToDocx(wrapper, docxBtn);
        });
        tools.appendChild(docxBtn);
      }

      const copyBtn = document.createElement("button");
'''
    if 'pdfBtn.textContent = "📄 Export PDF";' not in html_doc and tools_anchor in html_doc:
        html_doc = html_doc.replace(tools_anchor, tools_insert, 1)

    render_anchor = r'''    function renderMessageContent(container, content) {
      const sourceText = String(content || "");
      container.innerHTML          = "";
      container.dataset.rawContent = sourceText;

      const frag  = document.createDocumentFragment();
      const parts = parseMessageParts(sourceText);
      const suggestedPyFilenames = extractSuggestedPyFilenamesFromText(sourceText);
      let pythonBlockIndex = 0;

      for (const part of parts) {
        if (part.type === "think") {
          frag.appendChild(createThinkingBlock(part.content || "", part.complete !== false));
        } else if (part.type === "code") {
          const suggestedFilename = isPythonLanguage(part.language)
            ? (suggestedPyFilenames[pythonBlockIndex] || "")
            : "";
          frag.appendChild(createCodeBlock(part.language, part.content || "", suggestedFilename));
          if (isPythonLanguage(part.language)) pythonBlockIndex += 1;
        } else {
          frag.appendChild(createTextSegment(part.content || ""));
        }
      }
      container.appendChild(frag);
    }
'''
    render_insert = r'''    function renderMessageContent(container, content) {
      const sourceText = String(content || "");
      container.innerHTML          = "";
      container.dataset.rawContent = sourceText;

      const frag  = document.createDocumentFragment();
      const parts = parseMessageParts(sourceText);
      const suggestedPyFilenames = extractSuggestedPyFilenamesFromText(sourceText);
      let pythonBlockIndex = 0;

      for (const part of parts) {
        if (part.type === "think") {
          frag.appendChild(createThinkingBlock(part.content || "", part.complete !== false));
        } else if (part.type === "code") {
          const suggestedFilename = isPythonLanguage(part.language)
            ? (suggestedPyFilenames[pythonBlockIndex] || "")
            : "";
          frag.appendChild(createCodeBlock(part.language, part.content || "", suggestedFilename));
          if (isPythonLanguage(part.language)) pythonBlockIndex += 1;
        } else {
          frag.appendChild(createTextSegment(part.content || ""));
        }
      }
      container.appendChild(frag);
      try {
        const ownerMessage = typeof container.closest === "function" ? container.closest(".msg") : null;
        if (ownerMessage && ownerMessage.classList && ownerMessage.classList.contains("assistant")) {
          syncAssistantPdfButtons();
        }
      } catch (_) {}
    }
'''
    if 'syncAssistantPdfButtons();' not in html_doc and render_anchor in html_doc:
        html_doc = html_doc.replace(render_anchor, render_insert, 1)

    append_anchor = r'''      els.messages.appendChild(wrapper);
      scrollToBottom();
      updateMsgCount();
'''
    append_insert = r'''      els.messages.appendChild(wrapper);
      syncAssistantPdfButtons();
      scrollToBottom();
      updateMsgCount();
'''
    if 'els.messages.appendChild(wrapper);\n      syncAssistantPdfButtons();' not in html_doc and append_anchor in html_doc:
        html_doc = html_doc.replace(append_anchor, append_insert, 1)

    return html_doc

_previous_serve_index_html_with_svg = serve_index_html

def serve_index_html() -> str:
    """Επιστρέφει το index HTML με επιπλέον εξαγωγή απαντήσεων Assistant σε PDF."""
    return _patch_pdf_export_in_index_html(_previous_serve_index_html_with_svg())


def _patch_math_display_blocks_in_index_html(html_doc: str) -> str:
    """Διορθώνει τη διάσπαση πολυγραμμικών display-math blocks από το markdown renderer."""
    if not isinstance(html_doc, str) or not html_doc:
        return html_doc

    css_anchor = '    .attachment-chip {\n'
    css_insert = "    .math-display-raw {\n      display: block;\n      width: 100%;\n      overflow-x: auto;\n      padding: 2px 0;\n    }\n\n"
    if '.math-display-raw {' not in html_doc and css_anchor in html_doc:
        html_doc = html_doc.replace(css_anchor, css_insert + css_anchor, 1)

    helper_anchor = '    function markdownToHtml(rawText) {\n'
    helper_insert = r"""    function protectDisplayMathBlocks(rawText) {
      let text = String(rawText || "");
      const blocks = [];
      const patterns = [
        /\$\$[\s\S]+?\$\$/g,
        /\\\[[\s\S]+?\\\]/g,
      ];

      for (const pattern of patterns) {
        text = text.replace(pattern, (match) => {
          const token = `@@DISPLAY_MATH_${blocks.length}@@`;
          blocks.push(match);
          return `\n${token}\n`;
        });
      }

      return { text, blocks };
    }

"""
    if 'function protectDisplayMathBlocks(rawText)' not in html_doc and helper_anchor in html_doc:
        html_doc = html_doc.replace(helper_anchor, helper_insert + helper_anchor, 1)

    old_md = r"""    function markdownToHtml(rawText) {
      const lines = rawText.split("\n");
      const out   = [];
      let inUl = false, inOl = false, inBq = false;

      const closeUl  = () => { if (inUl) { out.push("</ul>");          inUl = false; } };
      const closeOl  = () => { if (inOl) { out.push("</ol>");          inOl = false; } };
      const closeBq  = () => { if (inBq) { out.push("</blockquote>"); inBq = false; } };
      const closeLists = () => { closeUl(); closeOl(); };

      for (let i = 0; i < lines.length; i += 1) {
        const line = lines[i];
        const nextLine = i + 1 < lines.length ? lines[i + 1] : "";

        if (line.includes("|") && isMarkdownTableSeparator(nextLine)) {
          closeLists(); closeBq();
          const bodyLines = [];
          let j = i + 2;
          while (j < lines.length) {
            const candidate = lines[j];
            if (!candidate.trim() || !candidate.includes("|")) break;
            bodyLines.push(candidate);
            j += 1;
          }
          out.push(renderMarkdownTable(line, nextLine, bodyLines));
          i = j - 1;
          continue;
        }

        // Heading  # … ######
        const hm = line.match(/^(#{1,6})\s+(.*)/);
        if (hm) {
          closeLists(); closeBq();
          const lvl = hm[1].length;
          out.push(`<h${lvl} class="md-h${lvl}">${inlineMarkdown(hm[2])}</h${lvl}>`);
          continue;
        }

        // Horizontal rule  --- / *** / ___
        if (/^(\s*[-*_]){3,}\s*$/.test(line) && line.trim().length >= 3) {
          closeLists(); closeBq();
          out.push('<hr class="md-hr" />');
          continue;
        }

        // Blockquote  > ...
        const bm = line.match(/^>\s?(.*)/);
        if (bm) {
          closeLists();
          if (!inBq) { out.push('<blockquote class="md-bq">'); inBq = true; }
          out.push(`<p>${inlineMarkdown(bm[1])}</p>`);
          continue;
        }
        closeBq();

        // Unordered list  - * +
        const um = line.match(/^[-*+]\s+(.*)/);
        if (um) {
          closeOl();
          if (!inUl) { out.push('<ul class="md-list">'); inUl = true; }
          out.push(`<li>${inlineMarkdown(um[1])}</li>`);
          continue;
        }

        // Ordered list  1. 2. …
        const om = line.match(/^\d+\.\s+(.*)/);
        if (om) {
          closeUl();
          if (!inOl) { out.push('<ol class="md-list">'); inOl = true; }
          out.push(`<li>${inlineMarkdown(om[1])}</li>`);
          continue;
        }

        closeLists();

        // Empty line
        if (!line.trim()) { out.push('<div class="md-br"></div>'); continue; }

        // Normal paragraph
        out.push(`<p class="md-p">${inlineMarkdown(line)}</p>`);
      }

      closeLists(); closeBq();
      return out.join("\n");
    }
"""

    new_md = r"""    function markdownToHtml(rawText) {
      const protectedMath = protectDisplayMathBlocks(rawText);
      const lines = protectedMath.text.split("\n");
      const out   = [];
      let inUl = false, inOl = false, inBq = false;

      const closeUl  = () => { if (inUl) { out.push("</ul>");          inUl = false; } };
      const closeOl  = () => { if (inOl) { out.push("</ol>");          inOl = false; } };
      const closeBq  = () => { if (inBq) { out.push("</blockquote>"); inBq = false; } };
      const closeLists = () => { closeUl(); closeOl(); };

      for (let i = 0; i < lines.length; i += 1) {
        const line = lines[i];
        const nextLine = i + 1 < lines.length ? lines[i + 1] : "";
        const mathMatch = line.trim().match(/^@@DISPLAY_MATH_(\d+)@@$/);

        if (mathMatch) {
          closeLists(); closeBq();
          const mathSource = protectedMath.blocks[Number(mathMatch[1])] || "";
          out.push(`<div class="math-display-raw">${escapeHtml(mathSource)}</div>`);
          continue;
        }

        if (line.includes("|") && isMarkdownTableSeparator(nextLine)) {
          closeLists(); closeBq();
          const bodyLines = [];
          let j = i + 2;
          while (j < lines.length) {
            const candidate = lines[j];
            if (!candidate.trim() || !candidate.includes("|")) break;
            bodyLines.push(candidate);
            j += 1;
          }
          out.push(renderMarkdownTable(line, nextLine, bodyLines));
          i = j - 1;
          continue;
        }

        // Heading  # … ######
        const hm = line.match(/^(#{1,6})\s+(.*)/);
        if (hm) {
          closeLists(); closeBq();
          const lvl = hm[1].length;
          out.push(`<h${lvl} class="md-h${lvl}">${inlineMarkdown(hm[2])}</h${lvl}>`);
          continue;
        }

        // Horizontal rule  --- / *** / ___
        if (/^(\s*[-*_]){3,}\s*$/.test(line) && line.trim().length >= 3) {
          closeLists(); closeBq();
          out.push('<hr class="md-hr" />');
          continue;
        }

        // Blockquote  > ...
        const bm = line.match(/^>\s?(.*)/);
        if (bm) {
          closeLists();
          if (!inBq) { out.push('<blockquote class="md-bq">'); inBq = true; }
          out.push(`<p>${inlineMarkdown(bm[1])}</p>`);
          continue;
        }
        closeBq();

        // Unordered list  - * +
        const um = line.match(/^[-*+]\s+(.*)/);
        if (um) {
          closeOl();
          if (!inUl) { out.push('<ul class="md-list">'); inUl = true; }
          out.push(`<li>${inlineMarkdown(um[1])}</li>`);
          continue;
        }

        // Ordered list  1. 2. …
        const om = line.match(/^\d+\.\s+(.*)/);
        if (om) {
          closeUl();
          if (!inOl) { out.push('<ol class="md-list">'); inOl = true; }
          out.push(`<li>${inlineMarkdown(om[1])}</li>`);
          continue;
        }

        closeLists();

        // Empty line
        if (!line.trim()) { out.push('<div class="md-br"></div>'); continue; }

        // Normal paragraph
        out.push(`<p class="md-p">${inlineMarkdown(line)}</p>`);
      }

      closeLists(); closeBq();
      return out.join("\n");
    }
"""

    if old_md in html_doc:
        html_doc = html_doc.replace(old_md, new_md, 1)

    return html_doc


_previous_serve_index_html_with_pdf = serve_index_html

def serve_index_html() -> str:
    """Επιστρέφει το index HTML με fix για display LaTeX blocks και σωστή σειρά patches."""
    return _patch_math_display_blocks_in_index_html(_previous_serve_index_html_with_pdf())


# ──────────────────────────────────────────────────────────────────────────────
# Runtime patch: local/offline math assets + robust delayed math rendering
# ──────────────────────────────────────────────────────────────────────────────

def _iter_runtime_asset_roots() -> List[Path]:
    """Επιστρέφει πιθανούς root φακέλους για static assets τόσο σε .py όσο και σε PyInstaller .exe."""
    roots: List[Path] = []
    candidates = [BASE_DIR]
    if getattr(sys, 'frozen', False):
        try:
            candidates.append(Path(sys.executable).resolve().parent)
        except Exception:
            pass
        meipass = getattr(sys, '_MEIPASS', None)
        if meipass:
            try:
                candidates.append(Path(meipass).resolve())
            except Exception:
                pass

    seen: Set[str] = set()
    for candidate in candidates:
        try:
            resolved = Path(candidate).resolve()
        except Exception:
            resolved = Path(candidate)
        key = str(resolved)
        if key not in seen:
            seen.add(key)
            roots.append(resolved)
    return roots


def _iter_runtime_asset_dirs() -> List[Path]:
    """Σαρώνει γνωστές τοποθεσίες όπου μπορεί να έχουν μπει web assets για το .exe."""
    dirs: List[Path] = []
    seen: Set[str] = set()
    candidate_names = ('web_assets', 'assets', 'static', 'vendor', 'third_party_assets')
    for root in _iter_runtime_asset_roots():
        for name in candidate_names:
            asset_dir = root / name
            try:
                if asset_dir.is_dir():
                    key = str(asset_dir.resolve())
                    if key not in seen:
                        seen.add(key)
                        dirs.append(asset_dir.resolve())
            except Exception:
                continue
    return dirs


def _sanitize_asset_relpath(relative_path: str) -> str:
    """Κανονικοποιεί σχετική διαδρομή asset ώστε να μην επιτρέπονται path traversal patterns."""
    raw = str(relative_path or '').replace('\\', '/').strip().strip('/')
    if not raw:
        return ''
    parts: List[str] = []
    for part in raw.split('/'):
        if part in {'', '.'}:
            continue
        if part == '..':
            return ''
        parts.append(part)
    return '/'.join(parts)


def _resolve_local_web_asset(relative_path: str) -> Optional[Path]:
    """Βρίσκει τοπικό asset αν υπάρχει. Υποστηρίζει και fallback αναζήτηση με filename."""
    safe_rel = _sanitize_asset_relpath(relative_path)
    if not safe_rel:
        return None

    asset_dirs = _iter_runtime_asset_dirs()
    if not asset_dirs:
        return None

    target_name = Path(safe_rel).name.lower()
    target_suffix = Path(safe_rel).suffix.lower()

    for asset_dir in asset_dirs:
        candidate = asset_dir / safe_rel
        try:
            if candidate.is_file():
                return candidate.resolve()
        except Exception:
            continue

    preferred_search_roots: List[Path] = []
    first_part = safe_rel.split('/', 1)[0].lower()
    for asset_dir in asset_dirs:
        scoped_dir = asset_dir / first_part
        try:
            if scoped_dir.is_dir():
                preferred_search_roots.append(scoped_dir.resolve())
        except Exception:
            pass
        preferred_search_roots.append(asset_dir)

    seen_roots: Set[str] = set()
    for search_root in preferred_search_roots:
        key = str(search_root)
        if key in seen_roots:
            continue
        seen_roots.add(key)
        try:
            for found in search_root.rglob(target_name):
                if not found.is_file():
                    continue
                if target_suffix and found.suffix.lower() != target_suffix:
                    continue
                return found.resolve()
        except Exception:
            continue

    return None


def _asset_content_type(asset_path: Path) -> str:
    """Δίνει Content-Type για JS/CSS/fonts/SVG assets."""
    suffix = str(Path(asset_path).suffix or '').lower()
    mapping = {
        '.css': 'text/css; charset=utf-8',
        '.js': 'application/javascript; charset=utf-8',
        '.mjs': 'application/javascript; charset=utf-8',
        '.json': 'application/json; charset=utf-8',
        '.map': 'application/json; charset=utf-8',
        '.svg': 'image/svg+xml',
        '.woff': 'font/woff',
        '.woff2': 'font/woff2',
        '.ttf': 'font/ttf',
        '.otf': 'font/otf',
        '.eot': 'application/vnd.ms-fontobject',
    }
    return mapping.get(suffix, 'application/octet-stream')


def _browser_asset_url(relative_path: str, cdn_url: str) -> str:
    """Επιστρέφει same-origin asset URL αν το αρχείο υπάρχει τοπικά, αλλιώς κρατά το CDN."""
    safe_rel = _sanitize_asset_relpath(relative_path)
    if safe_rel and _resolve_local_web_asset(safe_rel):
        return '/assets/' + urllib.parse.quote(safe_rel)
    return str(cdn_url or '')


def _pdf_asset_url(relative_path: str, cdn_url: str) -> str:
    """Επιστρέφει file:// asset URL για headless local PDF export όταν υπάρχει local asset."""
    safe_rel = _sanitize_asset_relpath(relative_path)
    asset_path = _resolve_local_web_asset(safe_rel)
    if asset_path:
        try:
            return asset_path.resolve().as_uri()
        except Exception:
            return str(asset_path)
    return str(cdn_url or '')


def _try_serve_local_asset(handler: BaseHTTPRequestHandler) -> bool:
    """Εξυπηρετεί τοπικά web assets από /assets/... ώστε το .exe να δουλεύει χωρίς υποχρεωτικό internet."""
    parsed = urllib.parse.urlsplit(getattr(handler, 'path', '') or '')
    if not parsed.path.startswith('/assets/'):
        return False

    rel_path = urllib.parse.unquote(parsed.path[len('/assets/'):])
    asset_path = _resolve_local_web_asset(rel_path)
    if not asset_path:
        body = b'Asset not found'
        handler.send_response(404)
        handler.send_header('Content-Type', 'text/plain; charset=utf-8')
        handler.send_header('Content-Length', str(len(body)))
        handler.send_header('Cache-Control', 'no-store')
        _send_security_headers(handler)
        handler.end_headers()
        handler.wfile.write(body)
        return True

    raw = asset_path.read_bytes()
    handler.send_response(200)
    handler.send_header('Content-Type', _asset_content_type(asset_path))
    handler.send_header('Content-Length', str(len(raw)))
    handler.send_header('Cache-Control', 'public, max-age=86400')
    _send_security_headers(handler)
    handler.end_headers()
    handler.wfile.write(raw)
    return True


_original_app_handler_handle_get = AppHandler._handle_GET

def _patched_app_handler_handle_get(self) -> None:
    """Προσθέτει endpoint για local bundled web assets πριν τρέξει η αρχική GET λογική."""
    if _try_serve_local_asset(self):
        return
    return _original_app_handler_handle_get(self)

AppHandler._handle_GET = _patched_app_handler_handle_get


def _patch_local_assets_and_math_fonts_in_index_html(html_doc: str) -> str:
    """Βελτιώνει το index HTML για .exe: local assets, math font fallbacks και delayed re-render."""
    html_doc = str(html_doc or '')
    if not html_doc:
        return html_doc

    replacements = {
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism-tomorrow.min.css':
            _browser_asset_url('prism/themes/prism-tomorrow.min.css', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism-tomorrow.min.css'),
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism-solarizedlight.min.css':
            _browser_asset_url('prism/themes/prism-solarizedlight.min.css', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism-solarizedlight.min.css'),
        'https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.css':
            _browser_asset_url('katex/katex.min.css', 'https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.css'),
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/prism.min.js':
            _browser_asset_url('prism/prism.min.js', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/prism.min.js'),
        'https://cdn.jsdelivr.net/npm/mathjax@4/tex-mml-svg.js':
            _browser_asset_url('mathjax/tex-mml-svg.js', 'https://cdn.jsdelivr.net/npm/mathjax@4/tex-mml-svg.js'),
        'https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.js':
            _browser_asset_url('katex/katex.min.js', 'https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.js'),
        'https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/contrib/auto-render.min.js':
            _browser_asset_url('katex/contrib/auto-render.min.js', 'https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/contrib/auto-render.min.js'),
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-python.min.js':
            _browser_asset_url('prism/components/prism-python.min.js', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-python.min.js'),
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-javascript.min.js':
            _browser_asset_url('prism/components/prism-javascript.min.js', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-javascript.min.js'),
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-typescript.min.js':
            _browser_asset_url('prism/components/prism-typescript.min.js', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-typescript.min.js'),
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-bash.min.js':
            _browser_asset_url('prism/components/prism-bash.min.js', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-bash.min.js'),
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-json.min.js':
            _browser_asset_url('prism/components/prism-json.min.js', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-json.min.js'),
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-sql.min.js':
            _browser_asset_url('prism/components/prism-sql.min.js', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-sql.min.js'),
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-css.min.js':
            _browser_asset_url('prism/components/prism-css.min.js', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-css.min.js'),
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-markup.min.js':
            _browser_asset_url('prism/components/prism-markup.min.js', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-markup.min.js'),
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-go.min.js':
            _browser_asset_url('prism/components/prism-go.min.js', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-go.min.js'),
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-rust.min.js':
            _browser_asset_url('prism/components/prism-rust.min.js', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-rust.min.js'),
    }
    for old, new in replacements.items():
        html_doc = html_doc.replace(old, new)

    html_doc = html_doc.replace(
        '--sans:      "Segoe UI", Inter, Arial, sans-serif;',
        '--sans:      "Segoe UI", "Segoe UI Symbol", "Cambria Math", "STIX Two Math", "Noto Sans Math", Inter, Arial, sans-serif;',
    )

    if 'text-rendering: geometricPrecision;' not in html_doc:
        font_css = '''
    .msg-body,
    .msg-body *,
    .reasoning-content,
    .reasoning-content *,
    .assistant-print-body,
    .assistant-print-body *,
    .md-p,
    .md-list,
    .md-bq,
    .md-table,
    .md-table th,
    .md-table td {
      font-family: var(--sans);
      text-rendering: geometricPrecision;
    }

    .msg-body mjx-container,
    .msg-body .katex,
    .msg-body math,
    .reasoning-content mjx-container,
    .reasoning-content .katex,
    .reasoning-content math,
    .assistant-print-body mjx-container,
    .assistant-print-body .katex,
    .assistant-print-body math {
      font-family: "Cambria Math", "STIX Two Math", "Noto Sans Math", "Segoe UI Symbol", var(--sans) !important;
    }

'''
        html_doc = html_doc.replace('    ::selection { background: rgba(96,165,250,0.28); color: inherit; }\n', '    ::selection { background: rgba(96,165,250,0.28); color: inherit; }\n' + font_css, 1)

    start_marker = 'function renderMathInElementSafe(root) {'
    end_marker = '\n\n    /**\n     * Block markdown for a text segment'
    start_idx = html_doc.find(start_marker)
    end_idx = html_doc.find(end_marker, start_idx if start_idx >= 0 else 0)
    if start_idx >= 0 and end_idx > start_idx:
        new_render_math = r'''function renderMathInElementSafe(root) {
      if (!window.__ollamaPendingMathRoots) window.__ollamaPendingMathRoots = new Set();
      const pendingMathRoots = window.__ollamaPendingMathRoots;

      function scheduleMathRenderRetry() {
        if (window.__ollamaMathRetryTimer) return;
        window.__ollamaMathRetryTimer = window.setTimeout(() => {
          window.__ollamaMathRetryTimer = 0;
          flushPendingMathRender();
        }, 180);
      }

      function queueMathJaxTypeset(nodes) {
        mathTypesetQueue = mathTypesetQueue
          .catch(() => undefined)
          .then(() => window.MathJax.typesetPromise(nodes))
          .catch((err) => {
            console.warn("MathJax render failed:", err);
            nodes.forEach((node) => {
              if (node) pendingMathRoots.add(node);
            });
            scheduleMathRenderRetry();
          });
      }

      function flushPendingMathRender() {
        const nodes = Array.from(pendingMathRoots).filter((node) => Boolean(node));
        if (!nodes.length) return;

        if (window.MathJax && typeof window.MathJax.typesetPromise === "function") {
          pendingMathRoots.clear();
          queueMathJaxTypeset(nodes);
          return;
        }

        if (typeof window.renderMathInElement === "function") {
          pendingMathRoots.clear();
          for (const node of nodes) {
            try {
              window.renderMathInElement(node, {
                delimiters: [
                  { left: "$$", right: "$$", display: true },
                  { left: "\\[", right: "\\]", display: true },
                  { left: "$", right: "$", display: false },
                  { left: "\\(", right: "\\)", display: false },
                ],
                throwOnError: false,
                strict: "ignore",
                ignoredTags: ["script", "noscript", "style", "textarea", "pre", "code"],
              });
            } catch (err) {
              console.warn("KaTeX fallback render failed:", err);
            }
          }
          return;
        }

        scheduleMathRenderRetry();
      }

      if (root) pendingMathRoots.add(root);

      if (!window.__ollamaMathLifecycleInstalled) {
        window.__ollamaMathLifecycleInstalled = true;
        document.addEventListener("readystatechange", () => {
          if (document.readyState === "interactive" || document.readyState === "complete") {
            flushPendingMathRender();
          }
        });
        window.addEventListener("load", () => {
          flushPendingMathRender();
          window.setTimeout(flushPendingMathRender, 350);
          window.setTimeout(flushPendingMathRender, 1200);
        });
        window.setInterval(() => {
          if (pendingMathRoots.size) flushPendingMathRender();
        }, 1500);
      }

      flushPendingMathRender();
    }'''
        html_doc = html_doc[:start_idx] + new_render_math + html_doc[end_idx:]

    return html_doc


_previous_build_assistant_pdf_document = _build_assistant_pdf_document

def _build_assistant_pdf_document(html_fragment: str, theme: str='light', document_title: str='Assistant response', mathjax_svg_cache: str='') -> str:
    """Override του printable HTML ώστε και το PDF export να χρησιμοποιεί local assets και math-friendly font fallback."""
    html_doc = _previous_build_assistant_pdf_document(
        html_fragment,
        theme=theme,
        document_title=document_title,
        mathjax_svg_cache=mathjax_svg_cache,
    )

    replacements = {
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism-solarizedlight.min.css':
            _pdf_asset_url('prism/themes/prism-solarizedlight.min.css', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism-solarizedlight.min.css'),
        'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism-tomorrow.min.css':
            _pdf_asset_url('prism/themes/prism-tomorrow.min.css', 'https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism-tomorrow.min.css'),
    }
    for old, new in replacements.items():
        html_doc = html_doc.replace(old, new)

    pdf_font_patch = '''
      .assistant-print-body,
      .assistant-print-body *,
      .print-shell .msg-body,
      .print-shell .msg-body *,
      .print-shell .md-p,
      .print-shell .md-list,
      .print-shell .md-bq,
      .print-shell .md-table th,
      .print-shell .md-table td {
        font-family: "Segoe UI", "Segoe UI Symbol", "Cambria Math", "STIX Two Math", "Noto Sans Math", Inter, Arial, sans-serif !important;
        text-rendering: geometricPrecision !important;
      }
      .print-shell mjx-container,
      .print-shell .katex,
      .print-shell math {
        font-family: "Cambria Math", "STIX Two Math", "Noto Sans Math", "Segoe UI Symbol", "Segoe UI", Inter, Arial, sans-serif !important;
      }
'''
    if 'text-rendering: geometricPrecision !important;' not in html_doc:
        html_doc = html_doc.replace('      .assistant-print-body,\n      .assistant-print-body * {\n        color: inherit;\n      }\n', '      .assistant-print-body,\n      .assistant-print-body * {\n        color: inherit;\n      }\n' + pdf_font_patch, 1)

    return html_doc


_previous_serve_index_html_with_local_math_assets = serve_index_html

def serve_index_html() -> str:
    """Επιστρέφει το τελικό index HTML με local assets, robust delayed math rendering και math font fallbacks."""
    return _patch_local_assets_and_math_fonts_in_index_html(_previous_serve_index_html_with_local_math_assets())




def _normalize_tex_to_pdf_fallback_text(value: str) -> str:
    """Μετατρέπει ωμό LaTeX σε αναγνώσιμο text fallback για PDF όταν αποτύχει το typesetting."""
    text = str(value or '').strip()
    if not text:
        return ''

    unwrap_patterns = [
        r'^\\\[(.*?)\\\]$',
        r'^\$\$(.*?)\$\$$',
        r'^\\\((.*?)\\\)$',
        r'^\$(.*?)\$$',
    ]

    for pattern in unwrap_patterns:
        match = re.match(pattern, text, flags=re.DOTALL)
        if match:
            text = str(match.group(1) or '').strip()
            break

    replacements = [
        (r'\\cdot\b', '·'),
        (r'\\times\b', '×'),
        (r'\\div\b', '÷'),
        (r'\\to\b', '→'),
        (r'\\rightarrow\b', '→'),
        (r'\\leftarrow\b', '←'),
        (r'\\leftrightarrow\b', '↔'),
        (r'\\Rightarrow\b', '⇒'),
        (r'\\Leftarrow\b', '⇐'),
        (r'\\Leftrightarrow\b', '⇔'),
        (r'\\geq\b', '≥'),
        (r'\\leq\b', '≤'),
        (r'\\neq\b', '≠'),
        (r'\\approx\b', '≈'),
        (r'\\sim\b', '∼'),
        (r'\\propto\b', '∝'),
        (r'\\pm\b', '±'),
        (r'\\mp\b', '∓'),
        (r'\\infty\b', '∞'),
        (r'\\partial\b', '∂'),
        (r'\\nabla\b', '∇'),
        (r'\\sum\b', '∑'),
        (r'\\prod\b', '∏'),
        (r'\\int\b', '∫'),
        (r'\\iint\b', '∬'),
        (r'\\iiint\b', '∭'),
        (r'\\oint\b', '∮'),
        (r'\\forall\b', '∀'),
        (r'\\exists\b', '∃'),
        (r'\\in\b', '∈'),
        (r'\\notin\b', '∉'),
        (r'\\subseteq\b', '⊆'),
        (r'\\supseteq\b', '⊇'),
        (r'\\subset\b', '⊂'),
        (r'\\supset\b', '⊃'),
        (r'\\cup\b', '∪'),
        (r'\\cap\b', '∩'),
        (r'\\land\b', '∧'),
        (r'\\lor\b', '∨'),
        (r'\\neg\b', '¬'),
        (r'\\oplus\b', '⊕'),
        (r'\\otimes\b', '⊗'),
        (r'\\parallel\b', '∥'),
        (r'\\perp\b', '⟂'),
        (r'\\lim\b', 'lim'),
        (r'\\log\b', 'log'),
        (r'\\ln\b', 'ln'),
        (r'\\sin\b', 'sin'),
        (r'\\cos\b', 'cos'),
        (r'\\tan\b', 'tan'),
        (r'\\alpha\b', 'α'),
        (r'\\beta\b', 'β'),
        (r'\\gamma\b', 'γ'),
        (r'\\delta\b', 'δ'),
        (r'\\epsilon\b', 'ε'),
        (r'\\varepsilon\b', 'ε'),
        (r'\\theta\b', 'θ'),
        (r'\\vartheta\b', 'ϑ'),
        (r'\\lambda\b', 'λ'),
        (r'\\mu\b', 'μ'),
        (r'\\pi\b', 'π'),
        (r'\\rho\b', 'ρ'),
        (r'\\sigma\b', 'σ'),
        (r'\\tau\b', 'τ'),
        (r'\\phi\b', 'φ'),
        (r'\\varphi\b', 'φ'),
        (r'\\omega\b', 'ω'),
        (r'\\Delta\b', 'Δ'),
        (r'\\Gamma\b', 'Γ'),
        (r'\\Theta\b', 'Θ'),
        (r'\\Lambda\b', 'Λ'),
        (r'\\Pi\b', 'Π'),
        (r'\\Sigma\b', 'Σ'),
        (r'\\Phi\b', 'Φ'),
        (r'\\Omega\b', 'Ω'),
    ]
    for pattern, replacement in replacements:
        text = re.sub(pattern, replacement, text)

    text = re.sub(r"([A-Za-zΑ-Ωα-ω])'", r"\1′", text)
    text = re.sub(r'\\lim\s*_\s*\{([^{}]+)\}', lambda m: f"lim {m.group(1)}", text)

    previous = None
    while previous != text:
        previous = text
        text = re.sub(r'\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}', r'(\1)/(\2)', text)
        text = re.sub(r'\\dfrac\s*\{([^{}]+)\}\s*\{([^{}]+)\}', r'(\1)/(\2)', text)
        text = re.sub(r'\\tfrac\s*\{([^{}]+)\}\s*\{([^{}]+)\}', r'(\1)/(\2)', text)
        text = re.sub(r'\\sqrt\s*\{([^{}]+)\}', r'√(\1)', text)
        text = re.sub(r'\\text\s*\{([^{}]+)\}', r'\1', text)
        text = re.sub(r'\\mathrm\s*\{([^{}]+)\}', r'\1', text)
        text = re.sub(r'\\mathbb\s*\{([^{}]+)\}', r'\1', text)
        text = re.sub(r'\\mathbf\s*\{([^{}]+)\}', r'\1', text)
        text = re.sub(r'\\vec\s*\{([^{}]+)\}', r'\1⃗', text)
        text = re.sub(r'\\overline\s*\{([^{}]+)\}', r'\1', text)
        text = re.sub(r'\\underline\s*\{([^{}]+)\}', r'\1', text)
        text = re.sub(r'\\hat\s*\{([^{}]+)\}', r'\1̂', text)
        text = re.sub(r'\\bar\s*\{([^{}]+)\}', r'\1̄', text)
        text = re.sub(r'\\dot\s*\{([^{}]+)\}', r'\1̇', text)
        text = re.sub(r'\\ddot\s*\{([^{}]+)\}', r'\1̈', text)

    text = text.replace('{', '').replace('}', '')
    text = text.replace('\\,', ' ')
    text = text.replace('\\;', ' ')
    text = text.replace('\\:', ' ')
    text = text.replace('\\!', '')
    text = re.sub(r'\\([A-Za-z]+)', r'\1', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text



def _patch_unrendered_pdf_math_fragment(html_fragment: str) -> str:
    """Αν μείνει ωμό LaTeX στο printable fragment, το μετατρέπει σε αναγνώσιμο math fallback."""
    fragment = str(html_fragment or '').strip()
    if not fragment or ('\\[' not in fragment and '$$' not in fragment and '\\(' not in fragment and '$' not in fragment):
        return fragment

    try:
        from bs4 import BeautifulSoup
    except Exception:
        return fragment

    try:
        soup = BeautifulSoup(fragment, 'html.parser')
    except Exception:
        return fragment

    selectors = [
        '.math-display-raw',
        '.assistant-print-body .md-p',
        '.assistant-print-body p',
        '.assistant-print-prompt-body .md-p',
        '.assistant-print-prompt-body p',
    ]
    seen_ids: Set[int] = set()
    for selector in selectors:
        for node in soup.select(selector):
            node_id = id(node)
            if node_id in seen_ids:
                continue
            seen_ids.add(node_id)
            try:
                if node.select_one('mjx-container, .katex, svg, math, code, pre, img'):
                    continue
            except Exception:
                continue
            raw_text = ' '.join(part for part in node.stripped_strings)
            if not raw_text:
                continue
            has_math_delimiters = ('\\[' in raw_text and '\\]' in raw_text) or ('$$' in raw_text) or ('\\(' in raw_text and '\\)' in raw_text)
            if not has_math_delimiters:
                continue
            fallback = _normalize_tex_to_pdf_fallback_text(raw_text)
            if not fallback:
                continue
            node.clear()
            node.append(fallback)
            try:
                existing_classes = list(node.get('class') or [])
                if 'pdf-math-fallback' not in existing_classes:
                    existing_classes.append('pdf-math-fallback')
                node['class'] = existing_classes
            except Exception:
                pass
            node['data-pdf-math-fallback'] = '1'
    return str(soup)


_previous_build_assistant_pdf_document_with_math_fallback = _build_assistant_pdf_document

def _build_assistant_pdf_document(html_fragment: str, theme: str='light', document_title: str='Assistant response', mathjax_svg_cache: str='') -> str:
    """Τελικό override για PDF: math fallback, local math assets και best-effort late typesetting."""
    safe_fragment = _patch_unrendered_pdf_math_fragment(html_fragment)
    html_doc = _previous_build_assistant_pdf_document_with_math_fallback(
        safe_fragment,
        theme=theme,
        document_title=document_title,
        mathjax_svg_cache=mathjax_svg_cache,
    )

    katex_css = _pdf_asset_url('katex/katex.min.css', 'https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.css')
    mathjax_js = _pdf_asset_url('mathjax/tex-mml-svg.js', 'https://cdn.jsdelivr.net/npm/mathjax@4/tex-mml-svg.js')
    katex_js = _pdf_asset_url('katex/katex.min.js', 'https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.js')
    katex_auto_js = _pdf_asset_url('katex/contrib/auto-render.min.js', 'https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/contrib/auto-render.min.js')

    pdf_math_css = """
      .print-shell .pdf-math-fallback,
      .print-shell [data-pdf-math-fallback="1"] {
        display: block !important;
        text-align: center !important;
        margin: 0.9em 0 !important;
        font-family: "Cambria Math", "STIX Two Math", "Noto Sans Math", "Segoe UI Symbol", "Segoe UI", Inter, Arial, sans-serif !important;
        font-size: 11.4pt !important;
        line-height: 1.55 !important;
        white-space: normal !important;
        word-break: break-word !important;
      }
      .print-shell mjx-container[display="true"],
      .print-shell .katex-display {
        text-align: center !important;
      }
"""
    if '.pdf-math-fallback' not in html_doc:
        html_doc = html_doc.replace('</style>', pdf_math_css + '\n  </style>', 1)

    math_head = f"""
  <link rel="stylesheet" href="{html.escape(katex_css, quote=True)}" />
  <script>
    window.MathJax = {{
      loader: {{ load: ["[tex]/mhchem", "[tex]/physics", "[tex]/braket", "[tex]/cancel", "[tex]/bbox", "[tex]/mathtools"] }},
      tex: {{
        inlineMath: [["$", "$"], ["\\(", "\\)"]],
        displayMath: [["$$", "$$"], ["\\[", "\\]"]],
        packages: {{ "[+]": ["mhchem", "physics", "braket", "cancel", "bbox", "mathtools"] }}
      }},
      svg: {{ fontCache: 'global' }},
      startup: {{ typeset: false }}
    }};
  </script>
  <script src="{html.escape(mathjax_js, quote=True)}" defer></script>
  <script src="{html.escape(katex_js, quote=True)}" defer></script>
  <script src="{html.escape(katex_auto_js, quote=True)}" defer></script>
"""
    if katex_css not in html_doc and '</head>' in html_doc:
        html_doc = html_doc.replace('</head>', math_head + '\n</head>', 1)

    math_runtime = r"""
<script>
(function () {
  function renderWithKatex(scope) {
    if (typeof window.renderMathInElement !== 'function') return;
    try {
      window.renderMathInElement(scope || document.body, {
        delimiters: [
          { left: '$$', right: '$$', display: true },
          { left: '\\[', right: '\\]', display: true },
          { left: '$', right: '$', display: false },
          { left: '\\(', right: '\\)', display: false }
        ],
        throwOnError: false,
        strict: 'ignore',
        ignoredTags: ['script', 'noscript', 'style', 'textarea', 'pre', 'code']
      });
    } catch (_) {}
  }

  async function renderPdfMath() {
    try {
      if (window.MathJax && typeof window.MathJax.typesetPromise === 'function') {
        await window.MathJax.typesetPromise([document.body]).catch(() => undefined);
      } else {
        renderWithKatex(document.body);
      }
    } catch (_) {
      renderWithKatex(document.body);
    }
  }

  function boot() {
    window.setTimeout(() => { renderPdfMath(); }, 40);
    window.setTimeout(() => { renderPdfMath(); }, 240);
    window.setTimeout(() => { renderPdfMath(); }, 900);
  }

  if (document.readyState === 'loading') {
    document.addEventListener('DOMContentLoaded', boot, { once: true });
  } else {
    boot();
  }
})();
</script>
"""
    if 'renderPdfMath()' not in html_doc and '</body>' in html_doc:
        html_doc = html_doc.replace('</body>', math_runtime + '\n</body>', 1)

    return html_doc

if __name__ == '__main__':
    main()


_previous_serve_index_html_with_svg_force_runtime = serve_index_html

def _patch_svg_force_runtime_in_index_html(html_doc: str) -> str:
    """Επιβάλλει runtime μετατροπή των SVG code blocks σε preview, ακόμη κι αν προηγούμενα string patches δεν έπιασαν."""
    html_doc = str(html_doc or '')
    marker = 'window.__svgRuntimePreviewFixInstalled'
    if marker in html_doc:
        return html_doc

    runtime_script = r"""
<script>
(function () {
  if (window.__svgRuntimePreviewFixInstalled) return;
  window.__svgRuntimePreviewFixInstalled = true;

  function looksLikeSvgContent(text) {
    const source = String(text || '').trim();
    return /^<svg\b[\s\S]*<\/svg>$/i.test(source);
  }

  function isSvgLanguage(language) {
    const normalized = String(language || '').trim().toLowerCase();
    return normalized === 'svg' || normalized === 'image/svg+xml' || normalized === 'xml+svg';
  }


  function mapSvgMathScriptChars(value, kind) {
    const source = String(value || '');
    if (!source) return '';
    const subMap = {
      '0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅','6':'₆','7':'₇','8':'₈','9':'₉',
      '+':'₊','-':'₋','=':'₌','(':'₍',')':'₎',
      'a':'ₐ','e':'ₑ','h':'ₕ','i':'ᵢ','j':'ⱼ','k':'ₖ','l':'ₗ','m':'ₘ','n':'ₙ','o':'ₒ','p':'ₚ','r':'ᵣ','s':'ₛ','t':'ₜ','u':'ᵤ','v':'ᵥ','x':'ₓ',
      'β':'ᵦ','γ':'ᵧ','ρ':'ᵨ','φ':'ᵩ','χ':'ᵪ'
    };
    const supMap = {
      '0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹',
      '+':'⁺','-':'⁻','=':'⁼','(':'⁽',')':'⁾',
      'a':'ᵃ','b':'ᵇ','c':'ᶜ','d':'ᵈ','e':'ᵉ','f':'ᶠ','g':'ᵍ','h':'ʰ','i':'ⁱ','j':'ʲ','k':'ᵏ','l':'ˡ','m':'ᵐ','n':'ⁿ','o':'ᵒ','p':'ᵖ','r':'ʳ','s':'ˢ','t':'ᵗ','u':'ᵘ','v':'ᵛ','w':'ʷ','x':'ˣ','y':'ʸ','z':'ᶻ',
      'A':'ᴬ','B':'ᴮ','D':'ᴰ','E':'ᴱ','G':'ᴳ','H':'ᴴ','I':'ᴵ','J':'ᴶ','K':'ᴷ','L':'ᴸ','M':'ᴹ','N':'ᴺ','O':'ᴼ','P':'ᴾ','R':'ᴿ','T':'ᵀ','U':'ᵁ','V':'ⱽ','W':'ᵂ',
      'β':'ᵝ','γ':'ᵞ','δ':'ᵟ','θ':'ᶿ','ι':'ᶥ','Φ':'ᶲ','φ':'ᵠ','χ':'ᵡ'
    };
    const targetMap = kind === 'sub' ? subMap : supMap;
    let out = '';
    for (const ch of source) {
      if (Object.prototype.hasOwnProperty.call(targetMap, ch)) {
        out += targetMap[ch];
      } else {
        return '';
      }
    }
    return out;
  }

  function convertLatexLikeSvgText(value) {
    let text = String(value || '');
    if (!text) return '';

    const normalizeMathExpr = (expr) => {
      let s = String(expr || '');
      if (!s) return '';
      const replacements = [
        [/\\,/g, ' '],
        [/\\;/g, ' '],
        [/\\:/g, ' '],
        [/\\!/g, ''],
        [/~/g, ' '],
        [/\\cdot\b/g, '·'],
        [/\\times\b/g, '×'],
        [/\\pm\b/g, '±'],
        [/\\mp\b/g, '∓'],
        [/\\leq\b|\\le\b/g, '≤'],
        [/\\geq\b|\\ge\b/g, '≥'],
        [/\\neq\b/g, '≠'],
        [/\\approx\b/g, '≈'],
        [/\\propto\b/g, '∝'],
        [/\\infty\b/g, '∞'],
        [/\\rightarrow\b|\\to\b/g, '→'],
        [/\\leftarrow\b/g, '←'],
        [/\\leftrightarrow\b/g, '↔'],
        [/\\degree\b/g, '°'],
        [/\\circ\b/g, '°'],
        [/\\alpha\b/g, 'α'],
        [/\\beta\b/g, 'β'],
        [/\\gamma\b/g, 'γ'],
        [/\\delta\b/g, 'δ'],
        [/\\epsilon\b/g, 'ε'],
        [/\\varepsilon\b/g, 'ε'],
        [/\\zeta\b/g, 'ζ'],
        [/\\eta\b/g, 'η'],
        [/\\theta\b/g, 'θ'],
        [/\\vartheta\b/g, 'ϑ'],
        [/\\iota\b/g, 'ι'],
        [/\\kappa\b/g, 'κ'],
        [/\\lambda\b/g, 'λ'],
        [/\\mu\b/g, 'μ'],
        [/\\nu\b/g, 'ν'],
        [/\\xi\b/g, 'ξ'],
        [/\\pi\b/g, 'π'],
        [/\\varpi\b/g, 'ϖ'],
        [/\\rho\b/g, 'ρ'],
        [/\\varrho\b/g, 'ϱ'],
        [/\\sigma\b/g, 'σ'],
        [/\\varsigma\b/g, 'ς'],
        [/\\tau\b/g, 'τ'],
        [/\\upsilon\b/g, 'υ'],
        [/\\phi\b/g, 'φ'],
        [/\\varphi\b/g, 'ϕ'],
        [/\\chi\b/g, 'χ'],
        [/\\psi\b/g, 'ψ'],
        [/\\omega\b/g, 'ω'],
        [/\\Gamma\b/g, 'Γ'],
        [/\\Delta\b/g, 'Δ'],
        [/\\Theta\b/g, 'Θ'],
        [/\\Lambda\b/g, 'Λ'],
        [/\\Xi\b/g, 'Ξ'],
        [/\\Pi\b/g, 'Π'],
        [/\\Sigma\b/g, 'Σ'],
        [/\\Upsilon\b/g, 'Υ'],
        [/\\Phi\b/g, 'Φ'],
        [/\\Psi\b/g, 'Ψ'],
        [/\\Omega\b/g, 'Ω']
      ];
      replacements.forEach(([pattern, replacement]) => {
        s = s.replace(pattern, replacement);
      });

      for (let i = 0; i < 6; i += 1) {
        const next = s
          .replace(/\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}/g, '($1)/($2)')
          .replace(/\\sqrt\s*\{([^{}]+)\}/g, '√($1)');
        if (next === s) break;
        s = next;
      }

      s = s.replace(/([A-Za-zΑ-Ωα-ω0-9θΘΔΓΛΣΠΩμνρτφχψω])_\{([^{}]+)\}/g, (m, base, sub) => {
        const mapped = mapSvgMathScriptChars(sub, 'sub');
        return base + (mapped || ('_' + sub));
      });
      s = s.replace(/([A-Za-zΑ-Ωα-ω0-9θΘΔΓΛΣΠΩμνρτφχψω])_([A-Za-z0-9+\-=()βγρφχ]+)/g, (m, base, sub) => {
        const mapped = mapSvgMathScriptChars(sub, 'sub');
        return base + (mapped || ('_' + sub));
      });
      s = s.replace(/([A-Za-zΑ-Ωα-ω0-9θΘΔΓΛΣΠΩμνρτφχψω])\^\{([^{}]+)\}/g, (m, base, sup) => {
        const mapped = mapSvgMathScriptChars(sup, 'sup');
        return base + (mapped || ('^' + sup));
      });
      s = s.replace(/([A-Za-zΑ-Ωα-ω0-9θΘΔΓΛΣΠΩμνρτφχψω])\^([A-Za-z0-9+\-=()βγδθιΦφχ]+)/g, (m, base, sup) => {
        const mapped = mapSvgMathScriptChars(sup, 'sup');
        return base + (mapped || ('^' + sup));
      });

      s = s.replace(/[{}]/g, '');
      s = s.replace(/\\_/g, '_');
      s = s.replace(/\\\\/g, ' ');
      s = s.replace(/\s+/g, ' ').trim();
      return s;
    };

    text = text.replace(/\$\$?([\s\S]*?)\$\$?/g, (match, expr) => normalizeMathExpr(expr));
    text = normalizeMathExpr(text);
    return text;
  }

  function normalizeSvgMathTextNodes(root) {
    if (!root || !root.querySelectorAll) return;
    root.querySelectorAll('text, tspan, title, desc').forEach((node) => {
      const raw = node.textContent;
      if (!raw || raw.indexOf('$') === -1 && raw.indexOf('\\') === -1 && raw.indexOf('_') === -1 && raw.indexOf('^') === -1) return;
      const normalized = convertLatexLikeSvgText(raw);
      if (normalized && normalized !== raw) node.textContent = normalized;
    });
  }

  function sanitizeSvgMarkup(rawSvg) {
    const source = String(rawSvg || '').trim();
    if (!source) return '';
    try {
      const parser = new DOMParser();
      const doc = parser.parseFromString(source, 'image/svg+xml');
      const root = doc.documentElement;
      if (!root || String(root.nodeName || '').toLowerCase() !== 'svg') return '';
      if (doc.querySelector('parsererror')) return '';
      root.querySelectorAll('script, foreignObject, iframe, object, embed').forEach((node) => node.remove());
      root.querySelectorAll('*').forEach((node) => {
        Array.from(node.attributes || []).forEach((attr) => {
          const name = String(attr.name || '');
          const value = String(attr.value || '').trim();
          if (/^on/i.test(name)) {
            node.removeAttribute(name);
            return;
          }
          if ((name === 'href' || name === 'xlink:href') && /^javascript:/i.test(value)) {
            node.removeAttribute(name);
          }
        });
      });
      root.removeAttribute('onload');
      root.removeAttribute('onclick');
      root.setAttribute('preserveAspectRatio', root.getAttribute('preserveAspectRatio') || 'xMidYMid meet');
      normalizeSvgMathTextNodes(root);
      return new XMLSerializer().serializeToString(root);
    } catch (err) {
      console.warn('SVG sanitize failed:', err);
      return '';
    }
  }

  function svgMarkupToDataUrl(svgMarkup) {
    const source = String(svgMarkup || '').trim();
    if (!source) return '';
    return 'data:image/svg+xml;charset=utf-8,' + encodeURIComponent(source);
  }

  function setButtonFeedbackSafe(button, activeText, defaultText, mode) {
    if (typeof window.setButtonFeedback === 'function') {
      window.setButtonFeedback(button, activeText, defaultText, mode);
      return;
    }
    const old = button.textContent;
    button.textContent = activeText;
    window.setTimeout(() => { button.textContent = defaultText || old; }, 1400);
  }

  function createSvgSourceCodeBlock(svgMarkup) {
    const sanitized = sanitizeSvgMarkup(svgMarkup);
    if (!sanitized) return null;

    const wrapper = document.createElement('div');
    wrapper.className = 'code-block';
    wrapper.dataset.svgSource = '1';
    wrapper.style.marginTop = '12px';

    const toolbar = document.createElement('div');
    toolbar.className = 'code-toolbar';

    const label = document.createElement('div');
    label.className = 'code-language';
    label.textContent = 'svg';

    const copyBtn = document.createElement('button');
    copyBtn.type = 'button';
    copyBtn.className = 'code-copy-btn';
    copyBtn.textContent = '📋 Copy SVG';
    copyBtn.title = 'Αντιγραφή SVG source στο clipboard';
    copyBtn.addEventListener('click', async () => {
      try {
        await navigator.clipboard.writeText(sanitized);
        setButtonFeedbackSafe(copyBtn, '✅ Copied', '📋 Copy SVG');
      } catch (_) {
        setButtonFeedbackSafe(copyBtn, '❌ Error', '📋 Copy SVG', 'error');
      }
    });

    toolbar.appendChild(label);
    toolbar.appendChild(copyBtn);

    const pre = document.createElement('pre');
    pre.className = 'code-pre';
    const code = document.createElement('code');
    code.className = 'language-svg';
    code.textContent = sanitized;
    pre.appendChild(code);

    wrapper.appendChild(toolbar);
    wrapper.appendChild(pre);

    try {
      if (window.Prism && typeof window.Prism.highlightElement === 'function') {
        window.Prism.highlightElement(code);
      }
    } catch (err) {
      console.warn('SVG source highlight failed:', err);
    }

    return wrapper;
  }

  function createSvgPreviewBlock(svgMarkup) {
    const sanitized = sanitizeSvgMarkup(svgMarkup);
    if (!sanitized) return null;

    const wrapper = document.createElement('div');
    wrapper.className = 'svg-block';
    wrapper.dataset.svgPreview = '1';

    const toolbar = document.createElement('div');
    toolbar.className = 'svg-toolbar';

    const label = document.createElement('div');
    label.className = 'svg-label';
    label.textContent = 'SVG Preview';

    const actions = document.createElement('div');
    actions.className = 'svg-actions';

    const copyBtn = document.createElement('button');
    copyBtn.type = 'button';
    copyBtn.className = 'code-copy-btn';
    copyBtn.textContent = '📋 Copy SVG';
    copyBtn.title = 'Αντιγραφή SVG source στο clipboard';
    copyBtn.addEventListener('click', async () => {
      try {
        await navigator.clipboard.writeText(sanitized);
        setButtonFeedbackSafe(copyBtn, '✅ Copied', '📋 Copy SVG');
      } catch (_) {
        setButtonFeedbackSafe(copyBtn, '❌ Error', '📋 Copy SVG', 'error');
      }
    });

    actions.appendChild(copyBtn);
    toolbar.appendChild(label);
    toolbar.appendChild(actions);

    const previewWrap = document.createElement('div');
    previewWrap.className = 'svg-preview-wrap';

    const image = document.createElement('img');
    image.className = 'svg-preview-image';
    image.loading = 'lazy';
    image.alt = 'SVG preview';
    image.src = svgMarkupToDataUrl(sanitized);

    previewWrap.appendChild(image);
    wrapper.appendChild(toolbar);
    wrapper.appendChild(previewWrap);

    const sourceBlock = createSvgSourceCodeBlock(sanitized);
    if (sourceBlock) wrapper.appendChild(sourceBlock);

    return wrapper;
  }

  function extractStandaloneSvgBlocks(text) {
    const source = String(text || '');
    if (!source || source.indexOf('<svg') === -1 || source.indexOf('</svg>') === -1) return [source];

    const parts = [];
    const regex = /<svg\b[\s\S]*?<\/svg>/ig;
    let lastIndex = 0;
    let match;
    while ((match = regex.exec(source)) !== null) {
      const start = match.index;
      const end = regex.lastIndex;
      if (start > lastIndex) parts.push(source.slice(lastIndex, start));
      parts.push(match[0]);
      lastIndex = end;
    }
    if (lastIndex < source.length) parts.push(source.slice(lastIndex));
    return parts.length ? parts : [source];
  }

  function createTextPieceNode(text) {
    const source = String(text || '');
    const div = document.createElement('div');
    if (typeof window.markdownToHtml === 'function') {
      div.innerHTML = window.markdownToHtml(source);
      try {
        if (typeof window.mayContainScientificMarkup === 'function' && typeof window.renderMathInElementSafe === 'function' && window.mayContainScientificMarkup(source)) {
          window.renderMathInElementSafe(div);
        }
      } catch (err) {
        console.warn('Scientific render failed for SVG-adjacent text:', err);
      }
      if (div.childNodes.length === 0 && source.trim()) div.textContent = source;
      return div;
    }
    div.style.whiteSpace = 'pre-wrap';
    div.textContent = source;
    return div;
  }

  function replaceMixedRawSvgNode(node) {
    if (!node) return false;
    const source = String(node.textContent || '');
    if (source.indexOf('<svg') === -1 || source.indexOf('</svg>') === -1) return false;
    const parts = extractStandaloneSvgBlocks(source);
    if (!parts.length || (parts.length === 1 && !looksLikeSvgContent(String(parts[0] || '').trim()))) return false;

    const fragment = document.createDocumentFragment();
    let replacedAnything = false;

    for (const part of parts) {
      const piece = String(part || '');
      const trimmed = piece.trim();
      if (!trimmed) continue;

      if (looksLikeSvgContent(trimmed)) {
        const preview = createSvgPreviewBlock(trimmed);
        if (preview) {
          fragment.appendChild(preview);
          replacedAnything = true;
          continue;
        }
      }

      fragment.appendChild(createTextPieceNode(piece));
    }

    if (!replacedAnything) return false;
    node.replaceWith(fragment);
    return true;
  }

  function upgradeSvgCodeBlocks(root) {
    const scope = root && root.querySelectorAll ? root : document;
    const codeNodes = Array.from(scope.querySelectorAll('.code-block code, pre code'));
    for (const codeNode of codeNodes) {
      if (!codeNode || codeNode.closest('[data-svg-preview="1"]') || codeNode.closest('[data-svg-source="1"]')) continue;
      const codeText = String(codeNode.textContent || '');
      const codeBlock = codeNode.closest('.code-block') || codeNode.closest('pre') || codeNode.parentElement;
      if (!codeBlock || codeBlock.dataset.svgUpgraded === '1') continue;
      const languageNode = codeBlock.querySelector('.code-language');
      const languageText = languageNode ? String(languageNode.textContent || '').trim() : '';
      const className = String(codeNode.className || '').toLowerCase();
      const isSvg = isSvgLanguage(languageText) || className.indexOf('language-svg') >= 0 || looksLikeSvgContent(codeText);
      if (!isSvg) continue;
      const preview = createSvgPreviewBlock(codeText);
      if (!preview) continue;
      codeBlock.dataset.svgUpgraded = '1';
      codeBlock.replaceWith(preview);
    }
  }

  function upgradeInlineRawSvg(root) {
    const scope = root && root.querySelectorAll ? root : document;
    const nodes = Array.from(scope.querySelectorAll('.msg.assistant .msg-body > div, .assistant-print-body > div, .assistant-print-prompt-body > div'));
    for (const node of nodes) {
      if (!node || node.closest('[data-svg-preview="1"]')) continue;
      if (node.querySelector && node.querySelector('.svg-block, img, svg')) continue;
      if (replaceMixedRawSvgNode(node)) continue;
      const text = String(node.textContent || '').trim();
      if (!looksLikeSvgContent(text)) continue;
      const preview = createSvgPreviewBlock(text);
      if (!preview) continue;
      node.replaceWith(preview);
    }
  }

  function runSvgUpgrade(root) {
    try { upgradeSvgCodeBlocks(root); } catch (err) { console.warn('SVG code-block upgrade failed:', err); }
    try { upgradeInlineRawSvg(root); } catch (err) { console.warn('SVG inline upgrade failed:', err); }
  }

  function installSvgObserver() {
    if (window.__svgPreviewObserverInstalled) return;
    window.__svgPreviewObserverInstalled = true;
    const observer = new MutationObserver((mutations) => {
      for (const mutation of mutations) {
        for (const node of Array.from(mutation.addedNodes || [])) {
          if (!node || node.nodeType !== 1) continue;
          runSvgUpgrade(node);
        }
      }
    });
    observer.observe(document.body || document.documentElement, { childList: true, subtree: true });
    window.__svgPreviewObserver = observer;
  }

  function bootstrapSvgPreviewFix() {
    runSvgUpgrade(document);
    window.setTimeout(() => runSvgUpgrade(document), 60);
    window.setTimeout(() => runSvgUpgrade(document), 220);
    window.setTimeout(() => runSvgUpgrade(document), 900);
    installSvgObserver();
  }

  if (document.readyState === 'loading') {
    document.addEventListener('DOMContentLoaded', bootstrapSvgPreviewFix, { once: true });
  } else {
    bootstrapSvgPreviewFix();
  }
})();
</script>
"""
    if '</body>' in html_doc:
        html_doc = html_doc.replace('</body>', runtime_script + '\n</body>', 1)
    else:
        html_doc += runtime_script
    return html_doc


def serve_index_html() -> str:
    """Τελικό index HTML με αναγκαστικό runtime SVG preview fix."""
    return _patch_svg_force_runtime_in_index_html(_previous_serve_index_html_with_svg_force_runtime())
